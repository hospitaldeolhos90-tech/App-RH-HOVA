import streamlit as st
import datetime
import time
import imaplib
import socket
import email
from email.header import decode_header
from email.utils import parsedate_to_datetime
import os
import smtplib
from email.mime.text import MIMEText
import re
import base64
import urllib.parse
import json
import io
import tempfile

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx2txt
except ImportError:
    docx2txt = None

try:
    import fitz
except ImportError:
    fitz = None

# ──────────────────────────────────────────
# CONFIGURAÇÕES
# ──────────────────────────────────────────
EMAIL_CONTA     = "rh@holhosvaledoaco.com.br"
SENHA_CONTA     = "20rhhova18"
IMAP_SERVER     = "email-ssl.com.br"
SMTP_SERVER     = "email-ssl.com.br"
SMTP_PORT       = 465
ENDERECO_HOVA   = "Rua Ponte Nova, 185 - Centro, Ipatinga/MG"
ARQUIVO_MEMORIA = "memoria_rh_hova.json"
ARQUIVO_LOCK    = "memoria_rh_hova.lock"

PALAVRAS_CV = [
    "curriculo", "currículo", "curriculum", "cv ", " cv", "vaga", "candidato",
    "candidatura", "emprego", "seleção", "selecao", "oportunidade",
    "recepcionista", "enfermagem", "faturamento", "administrativo", "aprendiz"
]

# ── Critérios de Triagem Inteligente ──────
# Cidades com alto custo de VT (>30km de Ipatinga)
# Cidades no raio de ~12km de Ipatinga — ACEITAR normalmente
CIDADES_RAIO_12KM = [
    "ipatinga",
    "coronel fabriciano","cel. fabriciano","cel fabriciano",
    "timóteo","timoteo",
    "santana do paraíso","santana do paraiso","santana paraiso",
]

# Outros estados e cidades distantes — REJEITAR com mensagem de localização
CIDADES_FORA = [
    # MG longe
    "belo horizonte","contagem","betim","governador valadares",
    "caratinga","manhuaçu","manhuacu","ponte nova","ubá","uba",
    "viçosa","vicosa","muriaé","muriae","juiz de fora","valadares",
    "uberlândia","uberlandia","montes claros","divinópolis","divinopolis",
    "sete lagoas","poços de caldas","pocos de caldas","itabira","ouro preto",
    # Outros estados
    "rio de janeiro","são paulo","sao paulo","vitória","vitoria",
    "brasília","brasilia","salvador","fortaleza","recife","manaus",
    "curitiba","porto alegre","belém","belem","goiânia","goiania",
    "maceió","maceio","natal","teresina","campo grande","cuiabá","cuiaba",
    "macapá","macapa","porto velho","boa vista","palmas","florianópolis",
    "florianopolis","aracaju","joão pessoa","joao pessoa","são luís","sao luis",
    "espírito santo","espirito santo"," es "," rj "," sp "," mg ",
    # Siglas de estado no texto
    "/es","/rj","/sp","/ba","/pe","/ce","/go","/pr","/rs","/sc",
]

def classificar_cidade(texto: str, cidade: str) -> str:
    """
    Retorna:
      'perto'  — cidade no raio 12km de Ipatinga → aceitar
      'longe'  — cidade fora do raio mas identificada → rejeitar c/ mensagem
      'outro_estado' — outro estado identificado → rejeitar c/ mensagem
      'indefinido' — não identificou cidade → triagem normal
    """
    t = f"{texto} {cidade}".lower()
    if any(c in t for c in CIDADES_RAIO_12KM):
        return 'perto'
    for c in CIDADES_FORA:
        if c in t:
            return 'longe'
    return 'indefinido'

def detectar_cidade_longe(texto: str, cidade: str) -> tuple[bool, str]:
    """Retrocompatibilidade — retorna (True, nome) se longe."""
    t = f"{texto} {cidade}".lower()
    for c in CIDADES_FORA:
        if c in t:
            return True, c.strip('/').strip().title()
    return False, ""

def detectar_cidade_perto(texto: str, cidade: str) -> bool:
    t = f"{texto} {cidade}".lower()
    return any(c in t for c in CIDADES_RAIO_12KM)

MSG_REJEICAO_CIDADE = (
    "Olá, tudo bem?\n\n"
    "Agradecemos muito seu interesse em fazer parte da equipe do "
    "Hospital de Olhos Vale do Aço!\n\n"
    "Após análise do seu currículo, verificamos que você reside fora da nossa "
    "região de atendimento. Nossas vagas são presenciais e destinadas a "
    "candidatos que residam em Ipatinga/MG ou cidades vizinhas no raio de "
    "aproximadamente 12km (Coronel Fabriciano, Timóteo e Santana do Paraíso).\n\n"
    "Infelizmente não temos condições de viabilizar o deslocamento de outras "
    "regiões ou estados no momento.\n\n"
    "Agradecemos sua compreensão e desejamos muito sucesso na sua busca!\n\n"
    "Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
)

# Vagas abertas no momento — edite conforme o post divulgado
VAGAS_ABERTAS = [
    "RECEPCAO E ATENDIMENTO",
    "TECNICO E ENFERMAGEM",
    "ADMINISTRATIVO",
    "FATURAMENTO",
    "JOVEM APRENDIZ",
    "TRIAGEM GERAL",
]

# ── EPTOM — Instituição de Jovem Aprendiz ──
EMAIL_EPTOM       = "eptom.aprendiz@gmail.com"
EMAIL_EPTOM_RESP  = "eptom@eptom.webnode.com.br"  # e-mail de retorno para a EPTOM

# Dados fixos do Hospital para preencher a ficha automaticamente
FICHA_EMPRESA_DADOS = {
    "empresa":        "Hospital de Olhos Vale do Aço",
    "cnpj":           "05.011.179/0001-04",
    "resp_setor":     "Josiane",
    "tel_resp_setor": "(31) 9 8860-0023",
    "resp_contrato":  "Paula",
    "email_contrato": "rh@holhosvaledoaco.com.br",
    "horario":        "A definir pela EPTOM",
    "salario":        "R$ 761,55",
}

def detectar_primeiro_emprego(texto: str) -> bool:
    """Retorna True se o currículo indica candidato sem experiência."""
    if not texto: return False
    tl = texto.lower()
    sem_exp = [
        "primeiro emprego","sem experiência","sem experiencia",
        "nenhuma experiência","nenhuma experiencia",
        "não possuo experiência","nao possuo experiencia",
        "estudante","recém formado","recem formado","recém-formado"
    ]
    tem_exp = [
        "empresa","cargo","período","periodo","experiência profissional",
        "experiencia profissional","histórico profissional","atuação"
    ]
    # Se tem palavras de sem-exp OU não tem nenhuma seção de experiência
    if any(p in tl for p in sem_exp): return True
    if not any(p in tl for p in tem_exp): return True
    return False

def detectar_cidade_longe(texto: str, cidade: str) -> tuple[bool, str]:
    """Retorna (True, nome_cidade) se candidato mora longe."""
    t = f"{texto} {cidade}".lower()
    for c in CIDADES_LONGE:
        if c in t: return True, c.title()
    return False, ""

def detectar_cidade_perto(texto: str, cidade: str) -> bool:
    t = f"{texto} {cidade}".lower()
    return any(c in t for c in CIDADES_PERTO)

MESES_NOMES = {
    1:"Janeiro", 2:"Fevereiro", 3:"Marco", 4:"Abril",
    5:"Maio",    6:"Junho",     7:"Julho", 8:"Agosto",
    9:"Setembro",10:"Outubro", 11:"Novembro", 12:"Dezembro"
}

# ──────────────────────────────────────────
# PÁGINA
# ──────────────────────────────────────────
st.set_page_config(
    page_title="HOVA | Seleção de Talentos",
    layout="wide",
    initial_sidebar_state="auto"
)

# Sidebar: abre no desktop, recolhida mas acessível no mobile
st.markdown("""
<script>
(function() {
    function isMobile() { return window.innerWidth <= 768; }

    function gerenciarSidebar() {
        if (!isMobile()) {
            // Desktop — forçar abertura
            var btn = document.querySelector(
                'button[data-testid="collapsedControl"],' +
                'button[aria-label="Open sidebar"],' +
                'button[aria-label="Abrir barra lateral"]'
            );
            if (btn) btn.click();

            // Desktop — esconder botão de fechar dentro da sidebar
            document.querySelectorAll(
                '[data-testid="stSidebarCollapseButton"],' +
                '[data-testid="stSidebarNavCollapseButton"]'
            ).forEach(function(b) {
                var l = (b.getAttribute("aria-label") || "").toLowerCase();
                if (l.includes("close") || l.includes("fechar") || l.includes("collapse"))
                    b.style.display = "none";
            });
        }
        // Mobile — NÃO esconder nenhum botão, deixar o usuário abrir/fechar livremente
    }

    setTimeout(gerenciarSidebar, 300);
    setTimeout(gerenciarSidebar, 1000);
    setInterval(gerenciarSidebar, 5000);
    window.addEventListener('resize', gerenciarSidebar);
})();
</script>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────
# CSS — design moderno, verde petróleo + branco
# ──────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

/* ── Reset ── */
#MainMenu {visibility:hidden;}
footer    {visibility:hidden;}
header    {visibility:hidden;}

/* ── Base ── */
html, body, .stApp, [class*="css"] {
    font-family: 'Inter', 'Segoe UI', system-ui, -apple-system, sans-serif !important;
}
.stApp { background: #EAECEF; }

/* ═══════════════════════════════════════
   SIDEBAR — verde petróleo escuro, sempre visível
═══════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background: #003329 !important;
    border-right: none !important;
    min-width: 240px !important;
    max-width: 280px !important;
}
section[data-testid="stSidebar"] > div { padding-top: 0 !important; }

/* Esconder o botão de fechar dentro da sidebar */
section[data-testid="stSidebar"] button[aria-label*="Close"],
section[data-testid="stSidebar"] button[aria-label*="close"],
section[data-testid="stSidebar"] button[aria-label*="Fechar"],
section[data-testid="stSidebar"] button[aria-label*="Collapse"],
section[data-testid="stSidebar"] button[aria-label*="collapse"],
[data-testid="stSidebarNavCollapseButton"],
[data-testid="stSidebarCollapseButton"] {
    display: none !important;
    visibility: hidden !important;
    opacity: 0 !important;
    pointer-events: none !important;
}

/* Botão de reabrir (fora da sidebar, quando colapsada) — mantemos visível */
/* mas forçamos o estado via JS */

/* Sidebar no mobile — permite recolher normalmente */
@media (max-width: 768px) {
    section[data-testid="stSidebar"] {
        min-width: unset !important;
        max-width: unset !important;
    }
    /* Garantir que o botão de abrir sidebar fique visível no mobile */
    button[data-testid="collapsedControl"] {
        display: flex !important;
        visibility: visible !important;
        opacity: 1 !important;
        pointer-events: auto !important;
        position: fixed !important;
        top: 12px !important;
        left: 8px !important;
        z-index: 999 !important;
        background: #004D40 !important;
        color: #FFFFFF !important;
        border-radius: 8px !important;
        width: 40px !important;
        height: 40px !important;
        border: none !important;
        box-shadow: 0 2px 8px rgba(0,77,64,0.3) !important;
    }
}

/* Todos os textos da sidebar — branco */
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div,
section[data-testid="stSidebar"] .stMarkdown {
    color: rgba(255,255,255,0.85) !important;
}
section[data-testid="stSidebar"] small,
section[data-testid="stSidebar"] .stCaption { color: rgba(255,255,255,0.45) !important; }

/* ── Radio buttons na sidebar — FIX VERMELHO ── */
section[data-testid="stSidebar"] .stRadio > div { gap: 2px !important; }
section[data-testid="stSidebar"] .stRadio label {
    background: transparent !important;
    border-radius: 6px !important;
    padding: 8px 10px !important;
    cursor: pointer !important;
    transition: background 0.15s !important;
    color: rgba(255,255,255,0.75) !important;
    font-size: 12px !important;
    font-weight: 500 !important;
}
section[data-testid="stSidebar"] .stRadio label:hover {
    background: rgba(255,255,255,0.08) !important;
    color: #FFFFFF !important;
}
/* O círculo do radio — forçar verde petróleo */
section[data-testid="stSidebar"] input[type="radio"] {
    accent-color: #26A69A !important;
}
/* Label selecionado */
section[data-testid="stSidebar"] .stRadio label[data-baseweb="radio"]:has(input:checked),
section[data-testid="stSidebar"] input[type="radio"]:checked + div {
    color: #FFFFFF !important;
}
/* Workaround direto para o ponto vermelho do radio */
section[data-testid="stSidebar"] [data-baseweb="radio"] [data-checked="true"] div,
section[data-testid="stSidebar"] [data-baseweb="radio"] div[role="radio"][aria-checked="true"] {
    border-color: #26A69A !important;
    background-color: #26A69A !important;
}
section[data-testid="stSidebar"] [data-baseweb="radio"] div[role="radio"] {
    border-color: rgba(255,255,255,0.3) !important;
    background-color: transparent !important;
}

/* ── Select/Slider na sidebar ── */
section[data-testid="stSidebar"] .stSelectbox > div > div,
section[data-testid="stSidebar"] .stSelectbox select {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    color: #FFFFFF !important;
    border-radius: 7px !important;
}
section[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] div[role="slider"] {
    background: #26A69A !important;
    border-color: #26A69A !important;
}
section[data-testid="stSidebar"] .stSlider div[data-testid="stThumbValue"] {
    color: #FFFFFF !important;
}

/* ── Input na sidebar ── */
section[data-testid="stSidebar"] input {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    color: #FFFFFF !important;
    border-radius: 7px !important;
}
section[data-testid="stSidebar"] input::placeholder {
    color: rgba(255,255,255,0.35) !important;
}
section[data-testid="stSidebar"] input:focus {
    border-color: #26A69A !important;
    box-shadow: 0 0 0 2px rgba(38,166,154,0.25) !important;
}

/* ── Divisor na sidebar ── */
section[data-testid="stSidebar"] hr {
    border-color: rgba(255,255,255,0.1) !important;
}

/* ── Botões na sidebar ── */
section[data-testid="stSidebar"] div[data-testid="stButton"] button {
    background: rgba(255,255,255,0.08) !important;
    color: rgba(255,255,255,0.8) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    height: 40px !important;
    font-size: 11px !important;
}
section[data-testid="stSidebar"] div[data-testid="stButton"] button:hover {
    background: rgba(255,255,255,0.15) !important;
    color: #FFFFFF !important;
    transform: none !important;
}
section[data-testid="stSidebar"] div[data-testid="stButton"] button[kind="primary"] {
    background: #26A69A !important;
    color: #FFFFFF !important;
    border: none !important;
}
section[data-testid="stSidebar"] div[data-testid="stButton"] button[kind="primary"]:hover {
    background: #00897B !important;
}

/* ═══════════════════════════════════════
   ANIMAÇÕES
═══════════════════════════════════════ */
@keyframes fadeUp {
    from { opacity:0; transform:translateY(14px); }
    to   { opacity:1; transform:translateY(0); }
}
@keyframes scaleIn {
    from { opacity:0; transform:scale(0.97); }
    to   { opacity:1; transform:scale(1); }
}

/* ═══════════════════════════════════════
   HERO CARD — maior, mais impactante
═══════════════════════════════════════ */
.hero-card {
    background: #FFFFFF;
    border-radius: 20px;
    padding: 44px 52px;
    margin-bottom: 24px;
    display: flex;
    justify-content: space-between;
    align-items: center;
    box-shadow: 0 4px 24px rgba(0,0,0,0.07);
    border: 1px solid #E2E6EA;
    animation: scaleIn 0.4s ease forwards;
    position: relative;
    overflow: hidden;
}
/* Linha decorativa lateral verde */
.hero-card::before {
    content: '';
    position: absolute;
    left: 0; top: 0; bottom: 0;
    width: 5px;
    background: linear-gradient(180deg, #004D40, #26A69A);
    border-radius: 20px 0 0 20px;
}
.hero-marca {
    font-size: 11px;
    font-weight: 600;
    color: #26A69A;
    letter-spacing: 3px;
    text-transform: uppercase;
    margin-bottom: 10px;
}
.hero-titulo {
    font-size: 52px;
    font-weight: 900;
    color: #0D1B2A;
    letter-spacing: -2px;
    line-height: 0.95;
    text-transform: uppercase;
}
.hero-titulo span {
    color: #004D40;
    display: block;
}
.hero-sub {
    font-size: 11px;
    color: #9AA5B4;
    letter-spacing: 2.5px;
    text-transform: uppercase;
    margin-top: 12px;
    font-weight: 500;
}
/* Stats — mais elegantes */
.hero-stats { display:flex; gap:10px; }
.stat-box {
    background: #F5F7FA;
    border: 1px solid #E2E6EA;
    border-radius: 14px;
    padding: 18px 24px;
    text-align: center;
    min-width: 88px;
    transition: transform 0.2s;
}
.stat-box:hover { transform: translateY(-2px); }
.stat-box .n {
    font-size: 32px;
    font-weight: 900;
    color: #004D40;
    line-height: 1;
    display: block;
    letter-spacing: -1px;
}
.stat-box .l {
    font-size: 9px;
    color: #9AA5B4;
    text-transform: uppercase;
    letter-spacing: 2px;
    margin-top: 6px;
    display: block;
    font-weight: 600;
}

/* ═══════════════════════════════════════
   BOTÕES PRINCIPAIS — verde petróleo
═══════════════════════════════════════ */
div[data-testid="stButton"] button {
    height: 48px !important;
    border-radius: 9px !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 1.5px !important;
    transition: all 0.18s ease !important;
    border: none !important;
    cursor: pointer !important;
    font-family: 'Inter', sans-serif !important;
}
div[data-testid="stButton"] button[kind="primary"] {
    background: #004D40 !important;
    color: #FFFFFF !important;
    box-shadow: 0 2px 10px rgba(0,77,64,0.25) !important;
}
div[data-testid="stButton"] button[kind="primary"]:hover {
    background: #003329 !important;
    box-shadow: 0 5px 18px rgba(0,77,64,0.35) !important;
    transform: translateY(-1px) !important;
}
div[data-testid="stButton"] button[kind="secondary"] {
    background: #FFFFFF !important;
    color: #004D40 !important;
    border: 1.5px solid #B2DFDB !important;
}
div[data-testid="stButton"] button[kind="secondary"]:hover {
    background: #F0FAF8 !important;
    border-color: #004D40 !important;
    color: #003329 !important;
}
div[data-testid="stButton"] button[kind="secondary"]:focus {
    outline: none !important;
    box-shadow: 0 0 0 3px rgba(0,77,64,0.12) !important;
}

/* ═══════════════════════════════════════
   ABAS
═══════════════════════════════════════ */
div[data-testid="stTabs"] {
    background: #FFFFFF;
    border-radius: 14px;
    padding: 0 16px;
    box-shadow: 0 1px 6px rgba(0,0,0,0.06);
    margin-bottom: 22px;
    border: 1px solid #E2E6EA;
}
div[data-testid="stTabs"] button[data-baseweb="tab"] {
    font-size: 10.5px !important;
    font-weight: 700 !important;
    padding: 16px 13px !important;
    color: #9AA5B4 !important;
    border-bottom: 2.5px solid transparent !important;
    background: transparent !important;
    letter-spacing: 1px;
    text-transform: uppercase;
    font-family: 'Inter', sans-serif !important;
}
div[data-testid="stTabs"] button[aria-selected="true"] {
    color: #004D40 !important;
    border-bottom: 2.5px solid #004D40 !important;
}

/* ═══════════════════════════════════════
   CARD CANDIDATO
═══════════════════════════════════════ */
.card-cand {
    background: #FFFFFF;
    border: 1px solid #E2E6EA;
    border-radius: 20px;
    padding: 48px 56px;
    margin-bottom: 24px;
    box-shadow: 0 4px 20px rgba(0,0,0,0.05);
    animation: fadeUp 0.3s ease forwards;
    text-align: center;
}
.cand-nome {
    font-size: 28px;
    font-weight: 800;
    color: #0D1B2A;
    text-transform: uppercase;
    letter-spacing: -0.5px;
    line-height: 1.15;
    margin-bottom: 8px;
}
.cand-info { font-size: 13px; color: #9AA5B4; margin-bottom: 18px; line-height: 1.6; }

/* ═══════════════════════════════════════
   TAGS
═══════════════════════════════════════ */
.tag {
    display: inline-block;
    padding: 5px 13px;
    border-radius: 20px;
    font-size: 10px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    margin: 3px;
}
.tag-verde  { background:#E6F4F1; color:#004D40; border:1px solid #B2DFDB; }
.tag-cinza  { background:#F1F3F5; color:#4A5568; border:1px solid #DDE1E7; }
.tag-azul   { background:#EBF4FF; color:#1A56DB; border:1px solid #C3D9F7; }
.tag-manual { background:#FFF8EC; color:#92540A; border:1px solid #F6D860; }

/* ═══════════════════════════════════════
   RESUMO CURRÍCULO
═══════════════════════════════════════ */
.cv-resumo {
    background: #F8FAFB;
    border: 1px solid #E2E6EA;
    border-radius: 12px;
    padding: 24px 28px;
    font-size: 13px;
    color: #4A5568;
    line-height: 1.85;
    text-align: left;
    margin-top: 20px;
}

/* ═══════════════════════════════════════
   AVATAR
═══════════════════════════════════════ */
.avatar {
    width: 82px; height: 82px;
    border-radius: 50%;
    background: linear-gradient(145deg, #004D40, #26A69A);
    color: #FFFFFF;
    display: flex; justify-content: center; align-items: center;
    font-size: 24px; font-weight: 800;
    margin: 0 auto 22px auto;
    box-shadow: 0 4px 16px rgba(0,77,64,0.25);
    letter-spacing: 1px;
}
.avatar-img {
    width: 88px; height: 88px;
    border-radius: 50%; object-fit: cover;
    border: 3px solid #B2DFDB;
    margin: 0 auto 22px auto; display: block;
    box-shadow: 0 4px 14px rgba(0,0,0,0.1);
}

/* ═══════════════════════════════════════
   CARDS DE ETAPAS
═══════════════════════════════════════ */
.card-agendado {
    background: #FFFFFF;
    border: 1px solid #E2E6EA;
    border-radius: 14px;
    padding: 22px 26px;
    margin-bottom: 12px;
    border-left: 4px solid #004D40;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    transition: box-shadow 0.2s;
}
.card-agendado:hover { box-shadow: 0 4px 16px rgba(0,0,0,0.08); }
.card-contratado {
    background: linear-gradient(135deg, #F2FAF8, #E8F5F2);
    border: 1px solid #B2DFDB;
    border-radius: 14px;
    padding: 22px 26px;
    margin-bottom: 12px;
    border-left: 4px solid #004D40;
}
.card-pendente {
    background: #FFFDF7;
    border: 1px solid #E8D9A6;
    border-radius: 14px;
    padding: 18px 22px;
    margin-bottom: 10px;
    border-left: 4px solid #B7791F;
}
.card-alerta {
    background: #FFFAF8;
    border: 1px solid #E5BCBC;
    border-radius: 14px;
    padding: 18px 22px;
    margin-bottom: 10px;
    border-left: 4px solid #9B2C2C;
}

/* ═══════════════════════════════════════
   FORMULÁRIO AGENDAMENTO
═══════════════════════════════════════ */
.form-sched {
    background: #F5FFFE;
    border: 1.5px solid #B2DFDB;
    border-radius: 16px;
    padding: 32px 36px;
    margin-bottom: 22px;
    box-shadow: 0 2px 10px rgba(0,77,64,0.06);
}

/* ═══════════════════════════════════════
   NOTIFICAÇÕES
═══════════════════════════════════════ */
.notif {
    border-radius: 9px;
    padding: 14px 20px;
    font-size: 13px;
    font-weight: 600;
    text-align: center;
    margin-bottom: 16px;
    letter-spacing: 0.2px;
}
.notif-ok   { background:#EBF8F4; border:1px solid #9DCFBF; color:#00382E; }
.notif-info { background:#EBF4FF; border:1px solid #93C5FD; color:#1E40AF; }
.notif-warn { background:#FFFBEB; border:1px solid #D4A853; color:#92540A; }

/* ═══════════════════════════════════════
   ESTADO VAZIO
═══════════════════════════════════════ */
.empty { text-align:center; padding:72px 20px; }
.empty .e-title {
    font-size: 11px; font-weight: 700;
    color: #CDD5DF; letter-spacing: 3px; text-transform: uppercase;
}
.empty .e-sub { font-size: 12px; color: #DDE1E7; margin-top: 8px; }

/* ═══════════════════════════════════════
   INPUTS PRINCIPAIS
═══════════════════════════════════════ */
div[data-testid="stTextInput"] input {
    border-radius: 8px !important;
    border: 1.5px solid #D1D8E0 !important;
    font-size: 13px !important;
    color: #0D1B2A !important;
    padding: 10px 14px !important;
    background: #FFFFFF !important;
    font-family: 'Inter', sans-serif !important;
}
div[data-testid="stTextInput"] input:focus {
    border-color: #004D40 !important;
    box-shadow: 0 0 0 3px rgba(0,77,64,0.09) !important;
}
div[data-testid="stTextInput"] input::placeholder { color: #B0BAC8 !important; }

/* ── Barra de pesquisa — destaque ── */
div[data-testid="stTextInput"]:has(input[aria-label="busca_global"]) input,
div[data-testid="stTextInput"] input[id*="busca_global"] {
    height: 48px !important;
    font-size: 14px !important;
    padding-left: 18px !important;
}

/* ═══════════════════════════════════════
   BOTÃO WHATSAPP
═══════════════════════════════════════ */
.wa-btn {
    display: block;
    background: #00A884;
    color: #FFFFFF !important;
    text-decoration: none;
    border-radius: 9px;
    padding: 13px 18px;
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    text-align: center;
    margin-bottom: 10px;
    transition: background 0.18s;
    box-shadow: 0 2px 8px rgba(0,168,132,0.2);
}
.wa-btn:hover { background: #007F65; color:#FFFFFF !important; }

/* ═══════════════════════════════════════
   LABEL SIDEBAR
═══════════════════════════════════════ */
.sb-label {
    font-size: 9px;
    font-weight: 700;
    color: rgba(255,255,255,0.45);
    letter-spacing: 2px;
    text-transform: uppercase;
    margin-bottom: 10px;
    display: block;
}

/* ═══════════════════════════════════════
   MÓDULO FUNCIONÁRIOS
═══════════════════════════════════════ */
.func-grid {
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
    gap: 16px;
    margin-bottom: 24px;
}
.func-card {
    background: #FFFFFF;
    border: 1px solid #E2E6EA;
    border-radius: 16px;
    padding: 24px 16px 18px;
    text-align: center;
    box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    transition: transform 0.18s, box-shadow 0.18s;
    animation: fadeUp 0.3s ease forwards;
}
.func-card:hover {
    transform: translateY(-3px);
    box-shadow: 0 6px 20px rgba(0,77,64,0.12);
}
.func-avatar {
    width: 80px; height: 80px;
    border-radius: 50%;
    background: linear-gradient(145deg, #004D40, #26A69A);
    color: #FFF;
    display: flex; justify-content: center; align-items: center;
    font-size: 26px; font-weight: 800;
    margin: 0 auto 14px auto;
    box-shadow: 0 4px 14px rgba(0,77,64,0.25);
}
.func-avatar-img {
    width: 80px; height: 80px;
    border-radius: 50%; object-fit: cover;
    border: 3px solid #B2DFDB;
    margin: 0 auto 14px auto; display: block;
    box-shadow: 0 4px 14px rgba(0,0,0,0.1);
}
.func-nome {
    font-size: 13px; font-weight: 800;
    color: #0D1B2A; text-transform: uppercase;
    letter-spacing: 0.3px; line-height: 1.3;
    margin-bottom: 4px;
}
.func-cargo {
    font-size: 11px; color: #004D40;
    font-weight: 600; letter-spacing: 0.5px;
    text-transform: uppercase; margin-bottom: 4px;
}
.func-data {
    font-size: 10px; color: #9AA5B4;
    margin-bottom: 14px;
}
.dossie-card {
    background: #FFFFFF;
    border: 1px solid #E2E6EA;
    border-radius: 16px;
    padding: 28px 32px;
    margin-bottom: 20px;
    box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    animation: fadeUp 0.3s ease;
}
.ex-func-card {
    background: #F8FAFB;
    border: 1px solid #E2E6EA;
    border-radius: 10px;
    padding: 14px 18px;
    margin-bottom: 8px;
    border-left: 3px solid #9AA5B4;
}

/* ── Abas internas do módulo ── */
.mod-tab-btn {
    display: inline-block;
    padding: 8px 20px;
    border-radius: 20px;
    font-size: 11px; font-weight: 700;
    letter-spacing: 1px; text-transform: uppercase;
    cursor: pointer; margin-right: 8px; margin-bottom: 16px;
    border: 1.5px solid #E2E6EA;
    background: #FFFFFF; color: #9AA5B4;
}


/* ═══════════════════════════════════════
   MOBILE — elementos exclusivos
═══════════════════════════════════════ */
.mobile-header { display:none; }
.mobile-chips  { display:none; }
.mobile-bottom-nav { display:none; }
.desktop-only  { display:block; }

.mobile-stat { background:rgba(255,255,255,0.12); border-radius:8px; padding:5px 8px; text-align:center; min-width:52px; }
.ms-n { display:block; color:#fff; font-size:16px; font-weight:700; line-height:1; }
.ms-l { display:block; color:rgba(255,255,255,0.5); font-size:8px; letter-spacing:1px; margin-top:2px; }

.chip { display:inline-block; padding:5px 14px; border-radius:20px; font-size:11px; font-weight:600; margin-right:6px; cursor:pointer; border:0.5px solid #E2E6EA; background:#F5F7FA; color:#4A5568; white-space:nowrap; }
.chip.active { background:#004D40; color:#fff; border-color:#004D40; }

.bnav-item { display:flex; flex-direction:column; align-items:center; gap:3px; padding:4px 10px; color:#9AA5B4; font-size:9px; font-weight:600; letter-spacing:0.5px; text-transform:uppercase; text-decoration:none; }
.bnav-item.active { color:#004D40; }

@media (max-width: 768px) {
    /* Mostrar mobile / esconder desktop */
    .mobile-header     { display:flex !important; background:#003329; padding:12px 16px; border-radius:0 0 14px 14px; align-items:center; justify-content:space-between; margin-bottom:10px; position:sticky; top:0; z-index:100; box-shadow:0 2px 12px rgba(0,51,41,0.3); }
    .mobile-chips      { display:flex !important; overflow-x:auto; white-space:nowrap; padding:8px 12px; background:#fff; border-bottom:0.5px solid #E2E6EA; margin-bottom:6px; scrollbar-width:none; }
    .mobile-chips::-webkit-scrollbar { display:none; }
    .mobile-bottom-nav { display:flex !important; position:fixed; bottom:0; left:0; right:0; background:#fff; border-top:0.5px solid #E2E6EA; padding:8px 0 14px; z-index:200; justify-content:space-around; box-shadow:0 -2px 16px rgba(0,0,0,0.08); }
    .desktop-only      { display:none !important; }

    /* Padding p/ não sobrepor bottom nav */
    .stApp { padding-bottom: 76px !important; }

    /* Botão menu flutuante */
    button[data-testid="collapsedControl"] {
        display:flex !important; visibility:visible !important; opacity:1 !important;
        pointer-events:auto !important; position:fixed !important;
        bottom:80px !important; right:14px !important; top:auto !important; left:auto !important;
        z-index:199 !important; background:#004D40 !important; color:#fff !important;
        border-radius:50% !important; width:44px !important; height:44px !important;
        border:none !important; box-shadow:0 4px 14px rgba(0,77,64,0.35) !important;
    }

    /* Sidebar */
    section[data-testid="stSidebar"] { min-width:unset !important; max-width:unset !important; }

    /* Componentes */
    .card-cand   { padding:20px 12px !important; }
    .cand-nome   { font-size:18px !important; }
    .cv-resumo   { font-size:12px !important; padding:12px !important; }
    .hova-card   { padding:14px 8px 10px !important; }
    .hova-card-nome { font-size:11px !important; }
    .hova-card-cargo-bar { font-size:10px !important; padding:5px 6px !important; }
    .dossie-header { padding:18px 12px !important; }
    .form-sched  { padding:16px 12px !important; }
    .notif       { font-size:12px !important; padding:10px 12px !important; }
    .avatar      { width:68px !important; height:68px !important; font-size:20px !important; }
    div[data-testid="stButton"] button { height:52px !important; font-size:11px !important; }
    div[data-testid="stTabs"] { overflow-x:auto !important; -webkit-overflow-scrolling:touch; }
    div[data-testid="stTabs"] button[data-baseweb="tab"] { padding:12px 9px !important; font-size:9.5px !important; white-space:nowrap; }
}
@media (max-width: 480px) {
    .cand-nome { font-size:16px !important; }
    .ms-n      { font-size:14px !important; }
}
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────
# PERSISTÊNCIA
# ──────────────────────────────────────────
def _serial(obj):
    if isinstance(obj, (datetime.date, datetime.datetime)): return obj.isoformat()
    if isinstance(obj, datetime.time): return obj.strftime('%H:%M:%S')
    if isinstance(obj, bytes): return base64.b64encode(obj).decode('utf-8')
    raise TypeError(f"Não serializável: {type(obj)}")

# ──────────────────────────────────────────
# PERSISTÊNCIA — SUPABASE (cliente oficial)
# ──────────────────────────────────────────
@st.cache_resource
def _get_supabase_client():
    """Cria cliente Supabase uma única vez e reutiliza."""
    try:
        from supabase import create_client
        url = st.secrets["supabase"]["url"].rstrip("/")
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception as e:
        st.error(f"Erro ao conectar ao Supabase: {e}")
        return None

def _sb_get() -> dict:
    """Lê o registro principal do Supabase."""
    try:
        sb = _get_supabase_client()
        if not sb:
            return {}
        res = sb.table("hova_dados").select("dados").eq("id", "principal").execute()
        if res.data:
            return res.data[0]["dados"] or {}
        return {}
    except Exception as e:
        st.warning(f"Supabase leitura: {e}")
        return {}

def _sb_set(dados: dict):
    """Grava/atualiza o registro principal no Supabase."""
    try:
        sb = _get_supabase_client()
        if not sb:
            return
        payload_str = json.dumps(dados, default=_serial, ensure_ascii=False)
        payload_obj = json.loads(payload_str)
        sb.table("hova_dados").upsert({
            "id":         "principal",
            "dados":      payload_obj,
            "updated_at": datetime.datetime.utcnow().isoformat(),
        }).execute()
    except Exception as e:
        st.warning(f"Supabase gravação: {e}")

def _sb_backup_automatico(dados: dict):
    """
    Salva um snapshot diário automático no Supabase.
    Chave: backup_YYYY-MM-DD — sobrescreve se já existir hoje.
    Mantém os últimos 30 dias automaticamente.
    """
    try:
        sb = _get_supabase_client()
        if not sb:
            return
        hoje = datetime.date.today().isoformat()
        payload_str = json.dumps(dados, default=_serial, ensure_ascii=False)
        payload_obj = json.loads(payload_str)
        sb.table("hova_dados").upsert({
            "id":         f"backup_{hoje}",
            "dados":      payload_obj,
            "updated_at": datetime.datetime.utcnow().isoformat(),
        }).execute()
        # Limpar backups com mais de 30 dias
        corte = (datetime.date.today() - datetime.timedelta(days=30)).isoformat()
        sb.table("hova_dados").delete().lt(
            "id", f"backup_{corte}"
        ).like("id", "backup_%").execute()
    except Exception:
        pass  # backup é silencioso — não interrompe o fluxo

def salvar_json():
    """Serializa o estado, salva no Supabase e aciona backup diário automático."""
    try:
        dados = {
            "aguardando":      st.session_state.aguardando_retorno,
            "agendados":       st.session_state.agendados,
            "contratados":     st.session_state.contratados,
            "ex_funcionarios": st.session_state.ex_funcionarios,
            "favoritos":       st.session_state.favoritos,
        }
        _sb_set(dados)
        # Backup diário: só dispara uma vez por dia por sessão
        hoje = datetime.date.today().isoformat()
        if st.session_state.get('_ultimo_backup') != hoje:
            _sb_backup_automatico(dados)
            st.session_state['_ultimo_backup'] = hoje
    except Exception as e:
        st.warning(f"Aviso ao salvar: {e}")

def _fix_datas(lista):
    for c in lista:
        for k in ['data_entrevista','data_inicio_contrato','data_inicio_experiencia',
                  'data_desligamento']:
            if k in c and isinstance(c[k], str):
                try: c[k] = datetime.date.fromisoformat(c[k])
                except: c[k] = None
        for k in ['hora_entrevista','opcao_1','opcao_2','opcao_3','hora_inicio_contrato']:
            if k in c and isinstance(c[k], str):
                try: c[k] = datetime.time.fromisoformat(c[k])
                except: c[k] = None
        for k in ['arquivo_bytes','foto']:
            if k in c and c[k]:
                try: c[k] = base64.b64decode(c[k])
                except: c[k] = None
        if 'documentos' in c and isinstance(c['documentos'], dict):
            for nome_doc, val in c['documentos'].items():
                if val and isinstance(val, str):
                    try: c['documentos'][nome_doc] = base64.b64decode(val)
                    except: c['documentos'][nome_doc] = None
    return lista

def carregar_json():
    """Carrega dados do Supabase. Verifica updated_at para sincronizar múltiplos usuários."""
    try:
        sb = _get_supabase_client()
        if not sb:
            st.session_state._carregado = True
            return

        # Verificar timestamp remoto
        res_ts = sb.table("hova_dados").select("updated_at").eq("id","principal").execute()
        remote_ts = res_ts.data[0]["updated_at"] if res_ts.data else ""
    except Exception:
        remote_ts = ""

    last_ts = st.session_state.get('_sb_ts', '')
    if st.session_state.get('_carregado') and remote_ts == last_ts and last_ts:
        return

    try:
        d = _sb_get()
        if d:
            st.session_state.aguardando_retorno = _fix_datas(d.get("aguardando",      []))
            st.session_state.agendados          = _fix_datas(d.get("agendados",       []))
            st.session_state.contratados        = _fix_datas(d.get("contratados",     []))
            st.session_state.ex_funcionarios    = _fix_datas(d.get("ex_funcionarios", []))
            st.session_state.favoritos          = _fix_datas(d.get("favoritos",       []))
            proc = set()
            for lst in [st.session_state.aguardando_retorno,
                        st.session_state.agendados,
                        st.session_state.contratados]:
                for c in lst:
                    if c.get('email'): proc.add(c['email'].lower().strip())
            st.session_state._processados = proc
            st.session_state._sb_ts       = remote_ts
    except Exception as e:
        st.warning(f"Aviso ao carregar: {e}")
    finally:
        st.session_state._carregado = True

# ──────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────
_def = {
    'cvs': [], 'agendados': [], 'contratados': [],
    'aguardando_retorno': [], 'cvs_antigos': [],
    'ex_funcionarios': [],
    'favoritos': [],
    'historico_emails': set(),
    'candidato_foco': None, 'contratar_foco': None,
    'perfil_foco': None,
    'nao_contratar_foco': None,
    'rejeitar_foco': None,
    'editar_agendado': None,
    'pular_idx': {},
    'fav_idx': 0,
    'sync_msg': None, 'sync_logs': [],
    'executar_sync': False, 'limite_sync': 30,
    'alertas_processados': set(),
    '_processados': set(),
}
for k, v in _def.items():
    if k not in st.session_state:
        st.session_state[k] = v

carregar_json()

# ──────────────────────────────────────────
# UTILITÁRIOS
# ──────────────────────────────────────────
def send_email(dest, assunto, corpo):
    try:
        m = MIMEText(corpo, 'plain', 'utf-8')
        m['Subject'] = assunto
        m['From']    = EMAIL_CONTA
        m['To']      = dest
        m['Bcc']     = EMAIL_CONTA
        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as s:
            s.login(EMAIL_CONTA, SENHA_CONTA)
            s.send_message(m)
        return True
    except: return False

def horario_disponivel(data: datetime.date, hora: datetime.time) -> bool:
    """Retorna True se o slot data+hora não está ocupado em agendados."""
    return not any(
        a.get('data_entrevista') == data and a.get('hora_entrevista') == hora
        for a in st.session_state.agendados
    )

def horarios_livres(data: datetime.date,
                    opcoes: list[datetime.time]) -> list[datetime.time]:
    """Retorna quais horários das opções ainda estão livres."""
    return [h for h in opcoes if h and horario_disponivel(data, h)]
    try:
        m = MIMEText(corpo, 'plain', 'utf-8')
        m['Subject'] = assunto
        m['From']    = EMAIL_CONTA
        m['To']      = dest
        m['Bcc']     = EMAIL_CONTA
        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as s:
            s.login(EMAIL_CONTA, SENHA_CONTA)
            s.send_message(m)
        return True
    except: return False

def iniciais(nome):
    p = nome.strip().split()
    if len(p) >= 2: return f"{p[0][0]}{p[-1][0]}".upper()
    return nome[:2].upper() if nome else "CV"

def resumo(texto: str) -> str:
    """
    Extração inteligente e gratuita — sem API, sem tokens.
    Foca no que o RH quer ver: endereço, experiências (empresa+cargo+período) e formação.
    """
    if not texto or len(texto.strip()) < 20:
        return "<i style='color:#9AA5B4;'>Nenhum texto extraído do documento.</i>"

    # Normalizar
    texto = re.sub(r'\r\n|\r', '\n', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto).strip()
    linhas = [l.strip() for l in texto.splitlines()]
    tl = texto.lower()

    html = []

    # ── 1. CIDADE / ENDEREÇO ─────────────────────────────────────
    cidades_vale = ["ipatinga", "coronel fabriciano", "timóteo", "timoteo",
                    "santana do paraíso", "santana do paraiso", "belo horizonte",
                    "governador valadares", "caratinga"]
    cidade_enc = ""
    for linha in linhas[:25]:   # endereço geralmente nas primeiras linhas
        ll = linha.lower()
        for cid in cidades_vale:
            if cid in ll:
                cidade_enc = linha
                break
        if cidade_enc:
            break

    # Fallback: procurar padrão "Cidade/UF" ou "Cidade - UF"
    if not cidade_enc:
        m = re.search(r'([A-ZÀ-Ú][a-zA-ZÀ-ú\s]+(?:do|de|da)?\s*[A-ZÀ-Ú][a-zA-ZÀ-ú]+)\s*/\s*MG', texto)
        if m:
            cidade_enc = m.group(0)

    # ── 2. TELEFONE ──────────────────────────────────────────────
    tel_enc = ""
    m_tel = re.search(r'\(?\d{2}\)?\s?(?:9\s?)?\d{4,5}[\s\-]?\d{4}', texto)
    if m_tel:
        tel_enc = re.sub(r'\s+', ' ', m_tel.group(0)).strip()

    # ── 3. EXPERIÊNCIAS PROFISSIONAIS ────────────────────────────
    # Localizar seção de experiência
    padroes_exp = [
        r'experi[eê]nci[as]\s+profissionais?',
        r'hist[oó]rico\s+profissional',
        r'atua[çc][aã]o\s+profissional',
        r'experi[eê]ncia',
    ]
    ini_exp = -1
    for p in padroes_exp:
        m = re.search(p, tl)
        if m:
            ini_exp = m.start()
            break

    # Localizar seção de formação para delimitar o bloco de experiência
    padroes_form = [
        r'forma[çc][aã]o\s+acad[eê]mica',
        r'forma[çc][aã]o\s+escolar',
        r'escolaridade',
        r'forma[çc][aã]o',
        r'instru[çc][aã]o',
    ]
    ini_form = len(texto)
    for p in padroes_form:
        m = re.search(p, tl)
        if m and m.start() > ini_exp:
            ini_form = m.start()
            break

    experiencias = []
    if ini_exp != -1:
        bloco_exp = texto[ini_exp:ini_form]
        linhas_exp = [l.strip() for l in bloco_exp.splitlines() if l.strip()]

        # Padrões de período: 01/2020 - 12/2022 | jan/2020 a dez/2022 | 2018-2020
        pat_periodo = re.compile(
            r'(?:'
            r'\d{2}/\d{4}\s*[–\-a]\s*\d{2}/\d{4}'      # 01/2020 - 12/2022
            r'|(?:jan|fev|mar|abr|mai|jun|jul|ago|set|out|nov|dez)[a-z]*/\d{4}'
            r'.*?(?:jan|fev|mar|abr|mai|jun|jul|ago|set|out|nov|dez)[a-z]*/\d{4}'
            r'|\d{2}/\d{2}/\d{4}\s*[–\-a_]\s*\d{2}/\d{2}/\d{4}'  # 17/07/2018_11/05/2020
            r'|\d{4}\s*[–\-]\s*(?:\d{4}|atual|presente|atualmente)'  # 2018 - atual
            r')',
            re.IGNORECASE
        )

        # Palavras que indicam cargo
        pat_cargo = re.compile(
            r'\b(auxiliar|assistente|analista|técnico|tecnico|enfermeiro|recepcionista|'
            r'atendente|operador|coordenador|supervisor|gerente|diretor|ajudante|'
            r'colaborador|agente|consultor|vendedor|caixa|balconista|estoquista|'
            r'faturista|secretária|secretario|motorista|porteiro|vigilante)\b',
            re.IGNORECASE
        )

        exp_atual: dict = {}
        for linha in linhas_exp[1:]:  # pular o título da seção
            ll_lower = linha.lower()

            # Ignorar títulos de seção
            if any(re.search(p, ll_lower) for p in padroes_exp + padroes_form):
                continue

            # Detectar período
            m_per = pat_periodo.search(linha)
            if m_per:
                if exp_atual:
                    experiencias.append(exp_atual)
                exp_atual = {"periodo": re.sub(r'\s+', ' ', linha).strip(),
                             "empresa": "", "cargo": ""}
                continue

            # Detectar cargo
            if pat_cargo.search(linha) and len(linha) < 80:
                if exp_atual:
                    if not exp_atual.get("cargo"):
                        exp_atual["cargo"] = linha
                    elif not exp_atual.get("empresa"):
                        exp_atual["empresa"] = linha
                continue

            # Linha com "EMPRESA:" ou "CARGO:" explícitos
            if re.match(r'empresa\s*:', ll_lower):
                val = re.sub(r'(?i)empresa\s*:\s*', '', linha).strip()
                if exp_atual:
                    exp_atual["empresa"] = val
                continue
            if re.match(r'cargo\s*:', ll_lower):
                val = re.sub(r'(?i)cargo\s*:\s*', '', linha).strip()
                if exp_atual:
                    exp_atual["cargo"] = val
                continue
            if re.match(r'per[ií]odo\s*:', ll_lower):
                val = re.sub(r'(?i)per[ií]odo\s*:\s*', '', linha).strip()
                if exp_atual and not exp_atual.get("periodo"):
                    exp_atual["periodo"] = val
                continue

            # Linha que parece nome de empresa (maiúsculas, sem ser título)
            if (linha.isupper() or re.match(r'^[A-ZÀ-Ú][A-Za-zÀ-ú\s\.\-&]+(?:LTDA|S\.A\.|ME|EIRELI|EPP)?$', linha)) \
                    and len(linha) > 3 and len(linha) < 70:
                if exp_atual and not exp_atual.get("empresa"):
                    exp_atual["empresa"] = linha

        if exp_atual and (exp_atual.get("empresa") or exp_atual.get("cargo") or exp_atual.get("periodo")):
            experiencias.append(exp_atual)

        # Limitar a 3 experiências
        experiencias = experiencias[:3]

    # ── 4. FORMAÇÃO ──────────────────────────────────────────────
    formacao_enc = ""
    graus = [
        r'ensino\s+m[eé]dio\s+completo', r'ensino\s+m[eé]dio\s+incompleto',
        r'ensino\s+fundamental\s+completo',
        r'gradua[çc][aã]o\s+em\s+[\w\s]+',
        r'p[oó]s[\s\-]gradua[çc][aã]o\s+em\s+[\w\s]+',
        r't[eé]cnico\s+em\s+[\w\s]+',
        r'superior\s+completo', r'superior\s+incompleto',
        r'faculdade\s+de\s+[\w\s]+',
    ]
    for p in graus:
        m = re.search(p, tl)
        if m:
            # Pegar a linha completa onde foi encontrado
            inicio_linha = tl.rfind('\n', 0, m.start()) + 1
            fim_linha    = tl.find('\n', m.end())
            formacao_enc = texto[inicio_linha: fim_linha if fim_linha != -1 else m.end()+60].strip()
            formacao_enc = formacao_enc.split('\n')[0].strip()
            break

    # ── 5. MONTAR HTML ───────────────────────────────────────────
    # Cidade + telefone
    meta = []
    if cidade_enc:
        meta.append(f"<span style='font-weight:600;color:#004D40;'>{cidade_enc[:60]}</span>")
    if tel_enc:
        meta.append(f"<span style='color:#4A5568;'>{tel_enc}</span>")
    if meta:
        html.append(
            "<div style='margin-bottom:14px;font-size:13px;padding:10px 14px;"
            "background:#F5F7FA;border-radius:8px;'>"
            + "&nbsp;&nbsp;·&nbsp;&nbsp;".join(meta) + "</div>"
        )

    # Experiências
    if experiencias:
        html.append(
            "<div style='font-size:9px;font-weight:800;color:#004D40;letter-spacing:2px;"
            "text-transform:uppercase;margin-bottom:8px;'>Experiências Profissionais</div>"
        )
        for ex in experiencias:
            empresa = ex.get("empresa", "").strip()
            cargo   = ex.get("cargo",   "").strip()
            periodo = ex.get("periodo", "").strip()
            if not (empresa or cargo):
                continue
            html.append(
                f"<div style='margin-bottom:8px;padding:10px 14px;"
                f"background:#F0FAF8;border-radius:8px;border-left:3px solid #004D40;'>"
                f"{'<div style=\"font-weight:700;font-size:13px;color:#0D1B2A;\">' + empresa + '</div>' if empresa else ''}"
                f"{'<div style=\"font-size:12px;color:#4A5568;margin-top:2px;\">' + cargo + '</div>' if cargo else ''}"
                f"{'<div style=\"font-size:11px;color:#9AA5B4;margin-top:2px;\">' + periodo + '</div>' if periodo else ''}"
                f"</div>"
            )
    else:
        # Sem experiência identificada — mostra trecho bruto da seção
        if ini_exp != -1:
            trecho = texto[ini_exp:ini_exp + 600].replace('\n', '<br>')
            html.append(
                f"<div style='font-size:12px;color:#4A5568;line-height:1.7;'>{trecho}</div>"
            )
        else:
            html.append(
                "<div style='color:#9AA5B4;font-size:12px;font-style:italic;'>"
                "Experiência não identificada no documento.</div>"
            )

    # Formação
    if formacao_enc:
        html.append(
            f"<div style='margin-top:10px;padding:8px 14px;background:#F5F7FA;"
            f"border-radius:8px;font-size:12px;'>"
            f"<span style='font-weight:700;font-size:9px;color:#004D40;"
            f"letter-spacing:2px;text-transform:uppercase;'>Formação</span>"
            f"&nbsp;&nbsp;<span style='color:#4A5568;'>{formacao_enc[:100]}</span></div>"
        )

    return "\n".join(html) if html else (
        "<i style='color:#9AA5B4;font-size:12px;'>Resumo não disponível — "
        "abra o documento original para ver o currículo completo.</i>"
    )




def setor_cv(assunto, texto):
    t = f"{assunto} {texto}".lower()
    if any(p in t for p in ["recep","atendiment","telefonista","secretaria","recepcionista"]): return "RECEPCAO E ATENDIMENTO"
    if any(p in t for p in ["enfermagem","tec. enf","tecnico em enf","enfermeir"]): return "TECNICO E ENFERMAGEM"
    if any(p in t for p in ["faturamento","fatura","analista de fat"]): return "FATURAMENTO"
    if any(p in t for p in ["adm","assistente adm","auxiliar adm","financeiro"]): return "ADMINISTRATIVO"
    if any(p in t for p in ["aprendiz","jovem","menor aprendiz","primeiro emprego"]): return "JOVEM APRENDIZ"
    return "TRIAGEM GERAL"

def novo_manual(nome, email_c, tel, setor):
    now = datetime.datetime.now()
    return {
        "id": str(int(time.time()*1000)),
        "nome": nome.upper().strip(), "email": email_c.lower().strip(),
        "telefone": ''.join(filter(str.isdigit, tel)),
        "data": now.strftime("%d/%m/%Y"), "data_iso": now.strftime("%Y-%m-%d"),
        "mes_num": now.month, "cidade": "", "tags": [],
        "preview": "Cadastro manual — sem currículo físico.",
        "setor": setor, "nome_arquivo": "",
        "arquivo_bytes": None, "foto": None, "manual": True,
    }

# E-mail de teste — troque por pessoal.expert@ntwdoctor.com.br quando confirmar
EMAIL_CONTABILIDADE = "esterteixeiradepaula@gmail.com"

def gerar_ficha_eptom_docx(nome_aprendiz: str, horario: str = "", salario: str = "",
                           cnpj: str = "", empresa: str = "", resp_setor: str = "",
                           tel_setor: str = "", resp_contrato: str = "",
                           email_contrato: str = "") -> bytes:
    """Gera o .docx preenchido preservando 100% da formatação original da EPTOM."""
    try:
        from docx import Document
        from docx.shared import Pt
        from copy import deepcopy
        import lxml.etree as etree

        import base64 as _b64, io as _io
        _docx_raw = _b64.b64decode("UEsDBBQABgAIAAAAIQBjLW1vrgEAAGMIAAATAAgCW0NvbnRlbnRfVHlwZXNdLnhtbCCiBAIooAACAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC0lstOwzAQRfdI/EPkLWpcWCCEmrLgsQQkQGLr2pPWwi/ZU6B/z7hpIwRtU1GyiRR77r3HY8XO6OrTmuIdYtLeVey0HLICnPRKu2nFXp7vBhesSCicEsY7qNgCErsaHx+NnhcBUkFqlyo2QwyXnCc5AytS6QM4mql9tALpNU55EPJNTIGfDYfnXHqH4HCA2YONRzdQi7nB4vaThhuS4KasuG7qclTFtM36PM43KiKY9EMiQjBaCqR5/u7UD67Biqkk5bImzXRIJ1SwJSHPbA9Y6R6omVErKB5FxHthqYp/+Ki48nJuSVnuttnA6etaS2j12S1ELyEl2iVrynbGCu3W/Js45Dyht6/WcI1gH6MP6fRgnNY0+0FEDW0Pt/bCze0EItH/fzNa606IhAsD6f8JGt/ueEAkQR8AK+dOhA+YPPVG8c28E6T2Hp3HPnajte6EAKd6Ylg7dyLMQCiIh3+Tvwga4732oZf8xniPfMoTEwN9EKysOyGQLjFonod3YmmzK5IqlwcxXYrxD8te32FZPQh7ncBtIlkfvD7I16MCtSGbL38Rxl8AAAD//wMAUEsDBBQABgAIAAAAIQAekRq37wAAAE4CAAALAAgCX3JlbHMvLnJlbHMgogQCKKAAAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArJLBasMwDEDvg/2D0b1R2sEYo04vY9DbGNkHCFtJTBPb2GrX/v082NgCXelhR8vS05PQenOcRnXglF3wGpZVDYq9Cdb5XsNb+7x4AJWFvKUxeNZw4gyb5vZm/cojSSnKg4tZFYrPGgaR+IiYzcAT5SpE9uWnC2kiKc/UYySzo55xVdf3mH4zoJkx1dZqSFt7B6o9Rb6GHbrOGX4KZj+xlzMtkI/C3rJdxFTqk7gyjWop9SwabDAvJZyRYqwKGvC80ep6o7+nxYmFLAmhCYkv+3xmXBJa/ueK5hk/Nu8hWbRf4W8bnF1B8wEAAP//AwBQSwMEFAAGAAgAAAAhAC3P/4VKAQAATgYAABwACAF3b3JkL19yZWxzL2RvY3VtZW50LnhtbC5yZWxzIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArJVLT4QwFIX3Jv4H0r0URh0fGZiNMZmtYuK2wOURaUvai8q/92bIMIxOGhdd3kN6zpfTm7LZfssu+ARjW60SFocRC0AVumxVnbC37PnqngUWhSpFpxUkbATLtunlxeYFOoF0yDZtbwNyUTZhDWL/yLktGpDChroHRV8qbaRAGk3Ne1F8iBr4KorW3Cw9WHriGezKhJldSfnZ2MN/vHVVtQU86WKQoPBMBG9AlGDIUZgakDz3cxySEePn86995lscOypwzp9mV/ydz3hQpdK4BDgoLoSVTwQ1yBwMbdeRYZZcELFPiGKwqOU7pc0QYXhUeYsgnUux9klTaY2/rmWWnJV47QTpLBwJ9uMkOpu49cnwBfkrINIqLLpYiM42Ir93ojATebdoZJZcFDdeX4s/XRwUF8KD791cPpjTPG8EP/kLpD8AAAD//wMAUEsDBBQABgAIAAAAIQA6jjM6aQsAAI5MAQARAAAAd29yZC9kb2N1bWVudC54bWzsmltz2jgUx993Zr+Dx+8Jl5I0ZUo7BEiTTptkSNrO7JuwBXgjSx7ZQNO3nX3Yr7H9KpsvtpJlGwgkY3KhBP/7EKHb0eWc39GR1bfvv/vMGlMZeoI37Mpu2bYod4Tr8UHD/nJ5tHNgW2FEuEuY4LRhX9PQfv/u99/eTuqucEY+5ZGlRPCwPgmchj2MoqBeKoXOkPok3PU9R4pQ9KNdR/gl0e97Di1NhHRL1XKlHP8KpHBoGKrxWoSPSWgn4vxFaSKgXFX2hfRJpLJyUPKJvBoFO0p6QCKv5zEvulayy/upGNGwR5LXExE72YR0l7qZUJKkPWSecU2XdrID8YglSZmag+Dh0Aumy3ioNFU5TIWM71vE2Gdpu0lQqT1OB21JJiqZCswzfdd08pmZ+f0SK+UcGtEish55pjA/ZjoTn3h8OvCDtmZmcyt7qwmo3hYQDB6nnA9SjIKpNO9x0k74VSZLk72CrETJs0sLHzeZiyEJFIG+Uz8ZcCFJj6kZKZVZatctbdb2O+VxesK91mlgTerKY7ndhl0utyq19n7FTovatE9GLFqsOZ8pioWcS51MPFdMWoJHUjDVckyYameX4haHbtwkEnrAuIZ7zNQx2o8WCnsiioS/UCy9wXBJYxpNKOW3ykvTYT3u6qnrzg1778D0+tNJOzjKVVCZjGDWIo/UOkLVgISOp8yjKT3C9PqHTR7O5p0wzZipxH8dwYTMtiD+l8woFq9nlqZQwi9Wgsqme30ub+3pOuah9PHu6Kz7+cunm7+6J2fWebPbtDqfz7udi+aF1e5YrbPTy27zsnnzz83fZ7qgqepO2yd/dC60kMiIymdNcyuEib0Yzg8Pq/utN1uhhLhiTgcr7mLUY0mS4NNj31RvFdxUaq9qZb1v0XWgDj33OzHjqBYn8aRUm73a8gYtythnEgtM9kg1fr1MWroty+s1kXPSVO6TEFfZ+mrNuFPfk2HUFXrWOstIkptWtgQb+fomkdanBXETLo4P1V0iy301ucp0DtnufJCeq38OVNqKrUWPum92+dcXa53ll1G9V7RZebrgKPPsOV3h5e0DIEpdtDymszQcvHmdDpe2cMzfNJeapJ7YEoNzDlXQpu6KcWbKqQ7gGNUdwh8NOzbVMCAOTdQck9KwySgSiziv2nee+lV7zzmH1TrrXZtdfzjUbBoPzSiRM10Sb2BpJJiqbx929o5iG1cdmswbZH5o1rcPPZcq+q7SsYxOnv1cfJ6DphXG6R1OUu94WlrZT0taeqCZsnXGPE85Xx3ddPxA0pAsCXVi3W4Ge3ccpYAKUG0kVK3T848g6v7OIApE5SfqVPjUcoWlzqpA8PDm55gyK6BMWESdX9z1flhcWCGNhAR493cGeAAvP3iXVFmv4IAP8AG+X3Lq6SdjpthaTqCjv1ESVcuJRXGXA4Ng8GkZ7Oz4xGPPCd8dX9XvYO/J3hpWwCAmJ4NgCuE8Ambsl4PAVypdwsmMUaUlT2/xG+eG8Tkb3noLvfVCxJTeUB4SFFURFAEzYLYEs2Mhb35KTxFGLRUB9QgbiocQhmsHCANhywj7SrQgdYCFhBnUfMrV77sp04npnK0yp8pz/7+M2l4tteS0xZOAvfROs40R6l0XKS6+SRJsJt+PJGPzr1Gr3KhhVg8zq/br2quDJswKZlUgbzV5vnjjyeKM//59/rgdJIAEkAASQAJIWDB6kAAStouEl/ueC9PNZ7rr84j4CAbrK04IUUUIARJAAoJpkFAwEnRimmXryKk2PF6CUTCK0wokJDUgASSABJAAEkACSAAJIAEkgASQABI2hwQ8EW+76a7PI+JTI6yvOCFEFSEESAAJCKZBQsFI0Ilplq0jp9rwRAxGwShOK5CQ1IAEkAASQAJIAAkgASSABJAAEkACSNgcEvBEvO2muz6PiE+NsL7ihBBVhBAgASQgmAYJBSNBJ6ZZto6casMTMRgFozitQEJSAxJAAkgACSABJIAEkAASQAJIAAkgYXNIwBPxtpvu+jwiPjXC+ooTQlQRQoAEkIBgGiQUjASdmGbZOnKqDU/EYBSM4rQCCUkNSAAJIAEkgASQABJAAkgACSABJICEzSEBT8Tbbrrr84j41AjrK04IUUUIARJAAoJpkFAwEnRimmXryKk2PBGDUTCK0wokJDUgASSABJAAEkACSAAJIAEkgASQABI2hwQ8EW+76a7PI+JTI6yvOCFEFSEESAAJCKZBQsFI0Ilplq0jp9rwRAxGwShOK5CQ1IAEkAASQAJIAAkgASSABJAAEkACSNgcEvBEvO2muz6PiE+NsL7ihBBVhBAgASQgmAYJBSNBJ6ZZto6casMTMRgFoy/ttPofAAD//+xZzW7jNhB+FUGXvSxgWbIlO6iN2o69TYFsjY2BvRQoaJG2uZFFgaLjJKe+Q1+gz9Fj+yZ9kvJHVCTHzU+3zq43o0NMDofk/H0zI6XR/257kpCFcLYnVyjpuSlN3IYizpkQbG3JOU2XCXHlNL/tuS09yFBMeq6nxjFLGO+5aCOY2c3pclWe+bzNje2JiIeMY8JzdVS+wvagOCGI726R0wVNkurtKfvIUaaHV4OELtPyAJIKwg3TimJyjvilvXLKFTWTrDyn+IMUzhs1W6dhU92gSVO+h3hKFmiTiPsr0wpJn2wu4MXPhKUil6wojyntuQNOUaK2rgZpXp3HuZ1oqbXmVh1PP2Yhv7XUZmgpI3VBhdYorm/cSbNXt69FTumW/p+/a/cYohG9cJheL/4W3PFHedBWHePp0BI3mYwyfI3MTbW4Ekz5uh70gARAAiABkABIACQAEgAJe3UDJAASAAnfBBKC1qGAUJJrEV5QIXRfLHRfLiM+N/VB4oTE+aXl/O8thA8tBCABkADNNCDhlSFB/Ri2Uo8num22q66wF/IfSDV2W+2WDUrLAX0bYBQwCtUKkABIACQAEgAJgARAAiABkABIACQAEgAJgISDIAH+Rfyth+7LZUT41AjR93paCB9aCEACIAGaaUDCK0OC+jE0MU/ueS8aen57ss9R9RXtqIJUcdSWYrYdSe05S0pFjMTZEBvpDtCQS2YitoSUMXvXqNtrFXgkpBzFumBc4igI/OJEI/v8IR/sRoEx6Y7xTqNW0BnsM159RRuvIB278WiqMo2SoecWpE+xPamaOg6DjUfRWTHzYSWQznGu14nN0hknOeFXxO2fZUhI46G3Tg2UFUnVcZNW2IzCnbT5gsIvEjxaIcVSjGa6IM7JkqaW88ual6a54DNy/S9mdmZn52Pn5+8dF8vnrYPfOJg4b87lY0Y38nG1D8qTjssXOckQR4LsdYeuRGGzE0X/h4gpm3LGFp8V905zJ97njF2uZfNwIRBXSYjioplJ0Vru/OUdG6L40uYvwzvW6cVwmjTyNSjcD1VEfUIpoZypoe/54RGim6T4Lpj2lrRx14860X7BLbFS55pBGLVKPY+uzkExewzU4+nsp3Pn719/e7CYPRQ1ny18FYfv//pDvt9oCA6EjGa6lk5ijpxJkeX0Fi3JWq9ilpBcuzCXyz+yK5LmNRX2hv9w6Iej7o4i3qh12taJazf86+w6/AedoDs+qvDXm2WH1+neg4BeqiFg/px4PpCJx3406NTeF4+pk/a7NYPWTDmZPGbKnMRiWsPeE+x5ITfpqjRoB2NPG25FECb8A1kQiZuYOOUHEWN51+G6CPMz3DHiLhgTT9tQ6KeWrV6xdA1NN2yTF05ZXqg3XP05phN0dAqQ42bX07WEcSqB23MTlOI8Rpn+LBIzLK8rDs+W57q6Sa+qfYH+XlIEsp0aR97NTZq3M2OAntsOdd4y6pXT5UboaRFE0kkqbRXZMfIKm2AWv+NUlxCakikVsdQiCK0Dja/0cM7wjR7ILRuVs/r/AAAA//8DAFBLAwQUAAYACAAAACEAYHWvXRUGAAA+GgAAEAAAAHdvcmQvaGVhZGVyMS54bWzcWOtuGjkU/r/SvoM1/+kMZFIIKqmAQBopbaI02/1tPIZxM2OPbANJVvs8+yD7Ynt8meEaQkm22mylhuPLOec7d8OHj/d5hmZUKiZ4J6i/iwJEOREJ45NO8NvtsNYKkNKYJzgTnHaCB6qCj6e//vJh3k4TiYCbq/a8IJ0g1bpoh6EiKc2xepczIoUSY/2OiDwU4zEjNJwLmYSNqB5ZqpCCUKVAVR/zGVaBF5dvShMF5XA4FjLHGpZyEuZY3k2LGkgvsGYjljH9ALKj96UY0Qmmkre9iFoFyLC0HSD/UXLIffQ6ljNBpjnl2moMJc0Ag+AqZcXCjEOlwWFaCpntMmKWZ+W9eVGPXxaDM4nn8LEQuA/8xDHlmUO+W2I92iMiRkTFsQ+EVZ0lkhwzvlB8kGuWnFs//jEBjXUBxeRlwTmXYlospLGXSbvgd5UsU9Q/IMsHedk09TIwX1NcQAXmpH0x4ULiUQaIIGQIvI5MWgen0Gz0KPMf19ITX/VDRtG8PcNZJ7jFI6hBUDeROKFB6O/8DueQRcdRBI0Ndh4KEF4QXV3oARbofnYlilIah15n7qvHTmAZVYEJ9TQRmYBOgadaODEZHevDOEdCa5GXvMYbWcUdP8FdH8YnzTNrDXib9t2upve64YRKNkkPxMO4Ygn99BLmb4cwhxuxGGWXQtyVsqK4a9nGTCp9IyCidbPMsF8tDsEd09xMsvK83LBXuPjUg1lWrb65VX2Bocquc8kSQ07gE2Qgm0eNRrPpjF3ZrjeOo5OFjJJVSziGyZrcgMKo2Yp6fRs3s3Url/bsZeL+egDEZ26zuZm4Rou/ZzL2CQ03Z3SMp5nePLleU114WSsF1YeCwlnqY2tiB9kJh3isqcFu/ctM92jE1eJmaop3kRLSSZZDwbUyzIow6F1dyXCGelCvd4Yz7ULmrO5abpskJZ5n0j70ysLKHPfH0VxcSyHGrlqxtYNCbnQVw+BVXevdLMuYt/1YA7KAtDaWoYQpfWvtNlSvoi4rysTA+LNoAygY5ojcQ/ygZZreQx4q2qCAO+MxJXrgbpoUD5AJVYDAqjiOjwM08oS7nwgIOWI2WxHHOfj5IscTmiNYJ1QR4Lu8Or/y18mX2bnERcrIUMJlYxaGpF3sXApyp3wPxweMejdguein4E/aVQVYU1XSbv0v1bok6gxrjKZycx4+L6pgRE8lBWlAtYsKFlAvlsZn18xWqFmAK3zkojJy1+72ltCVDI4dGzQuUpueXmxJKeYpxYkqA7AqxS5XII0yVgxZlhkNhkayTfMRBYjyIqlbB0MOXyrtKefiPxqtbhSdNHq1/nHUr8VRc1DrnsTNWjMaNOMobtX79f6fhrset6fKpBjOzgpWxnvfZ5h3rXssuDxzU8DaZgGVnxZi6IwwWJUkN+Ce0NJaUk1SQ47BVr8fLh1Yxyx8YVYKugcazT+LpGxkhv9+LHPzCQDRvY3jg4fj3PN0pYcL5gLG0zmFcW8I8DTgscLxDKxwV8srZpsLg8qqyPjKRuh2LHqD15Pw354tlcby2tWl62a211VNzvQ91zn9bNkxj+LG0baX1KybsQkvezWBnkblrlEVd+Nmv1GNqms7DeHfyfujbfNr9frPm1/fyaZJT061W5ZThb7QOboROeZLk239xEoZvcaM2+rA18a4PD234XUefuwbPXa/4YbcXqNWnw4UCMUIdIyZMr9D4Iw9YnA3uqXcuB1dqTnOEoE+Y5LiRNisWk3btfRayyTnnd5J1Bu09nweLbnyP5Je/8KjyUStilm8V4p5J75eir0OXmAwU6b8jlFIqqicwTi+FYlQiABEidFK4oCFJqsH0mh37QxGawZfK7H0Xe3N2H6K10z7YeSHZNdwWL1nd4K7HpxdDS8uL7p7+X/AEy9gs67XvuE8NzbeRF3/rJxxz5Hn6vvVR8iheOutveobpoam3zG8hZXCar3CK+OMaa3h0fvmyfbkeTP29qEi+d9/7TD1bdpF/0+hO+0mIoPvdab41ZbXitmx7/G2dnv+182VXtdqHTWj4ZrNT//Es3r9gCeyB7/UMbY1jMqANJGn/wAAAP//AwBQSwMEFAAGAAgAAAAhAEZAGhO+AQAA/wUAABEAAAB3b3JkL2VuZG5vdGVzLnhtbKyU32/bIBDH3yftf7B4T8DdukVWnKpq1Klv07r+ARTjGNXcIcDx8t/v7Nhx11ZR2uyFH8fd575wwPLqj62TrfbBIOQsnQuWaFBYGNjk7OH37WzBkhAlFLJG0Dnb6cCuVp8/LdtMQwEYdUgIASFrncpZFaPLOA+q0laGuTXKY8AyzhVajmVplOYt+oJfiFT0I+dR6RAo342ErQxswNnXNHQaaLFEb2Wkqd9wK/1T42ZEdzKaR1ObuCO2+DZiMGeNh2xAzA6CupBsL2joxgh/St59yBpVYzXEPiP3uiYNCKEybtrGR2m0WI2Q7bFNbG09+rUu/XpeDdZettRNwFPkF/sgW++VHyem4oSKdIhDxCkS/s05KrHSwJT4Q0fz7HDTy/cBLl4C3Oa84vzw2LiJZs6j3cHTgdW97HewhiI/31o4T8x9JR29QKuyuw2gl481KaKSJXTqSXet2Wr6cZI2iztHDkE76WVEz8hkipzN0t7P0ZR+tOJXzkjrtbgUovPoTWtdyqaOr1d+dqbF4st3cdtDfNccMvDVkvc2al3fDmLe0qUQooGm/w3uX2oU/1nim8mOyJ3GYfUXAAD//wMAUEsDBBQABgAIAAAAIQAJhWRUzQUAADIXAAAQAAAAd29yZC9mb290ZXIxLnhtbOxYbW/bNhD+PmD/gdCnDYgjyZZjx6jT2Y6dBkgbw8m6z7RE2VolkiNpK96w/747UvJLmiZp0wEtUBStjuTxuefuSN65r17fFTlZM6UzwfteeBx4hPFYJBlf9L3fbyeNrke0oTyhueCs722Y9l6f/fzTq7KXGkVgN9e9UsZ9b2mM7Pm+jpesoPq4yGIltEjNcSwKX6RpFjO/FCrxm0EYWEkqETOtwdSI8jXVXgVXfIwmJOOwmApVUANDtfALqj6sZAPQJTXZPMszswHs4KSGEX1vpXivgmhsCeGWniNUfeod6jl23ZZzEa8Kxo216CuWAwfB9TKTOze+FA0WlzXI+jEn1kVe65UyjF6Wg3NFS/jsAJ9DP3GbitwxfxwxDJ6REYTY7ngOhUObNZOCZnxn+ItCsxfcsP15AM37AHLxsuRcKLGSO7TsZWiX/MMWCy/1Z2BVSd53Tb+MzM2SSriBRdy7XHCh6DwHRpAyAlEneKy9M3hsJCl78Egls74XBONBK4zgnldTU4WTw9NgOO5uJ89ZSle5wZWTAfwZ1StTnIoGUWfUdMhTZT83ZpMz0FnTvO/NREKl59uFYWIVjEAOdhVp5wwB9d99L7KCpDHwDlGORS6AUjiJTjvnCOLvodC5rr41WpwzqnCfFBDKqNluOsOfVOm2g6iCreHQPJACDZoahvFA7TzD7Daj7WC2wuDSlRHORMYTWEkzpc2V1W21orZb+jPeGoeHiSk3q1yw1ERwo9GajjM4jLdZwTR5x0oyEwXlaG454PqBFYtiA1TD78KE4dzOntQzI7Rj55pB5bVlgUGt2Dx4EL4eXS6mSoj05dQhCdQmiVFtBjqjfU+axnC271bZq95UEGWP8ngJ9pJMm1ubVJSGW+kKrIVRK6iGs91QZ4XM2RSPC4xcfVqzNyxbLOFONNvhSbvTBKZkzpZwCqAAWcVcxB9Y4kS6EStzyUcsR6c8QvNclNfQK+RU4oQluLVD7uyuDf6Lrko8qxkWxTdb8xMloChCCFcFd9tB5zpNNTNnYafTPG0GEIf92XrogA5g39+DlVTRhaJyeR+5ddJtPYr73m5gdwYOOonBj6gTnXbBmXizlZ1LLE1ZbMZOM7cO4wsDAYaANDsgzCvB6ZfA5+avFVXwsIB4Czb63lyY5U2WQCPllBIRTxXJIOotj3BawDW8LOiCFQTGoBYD+NX1xXWlHr9bX6CbWTxRoIxnhvas49XMFeRQV68z/YIi7konF6MlHFY20BJcxnzbQ/q4/Zda3YM6p4aSlfq40j0NJbPYrBQDNJB6cksLpBej8fU0i9FnHEAoqswFdeamTvuB1NUb3HaKbFymPo70bkopUS4ZTXSdgEMUOzygNM8zOcnyHC2gTFSPFXO80uoywTtL8aBfaVNJLsT/NLuDIDhtDhujdjBqREFn3BicRp1GJxh3oiDqhqNw9C/uDqPeSuMRo/m5zOp8P7fBqkLr2gB3zuzz6B52S6j+Woq+cwK5ahXPIDy+lY1iJl6imIKv1by/t2ADs4sFjjRUCzIv34qkLoG4/y5VBX6B4L0XzIXn08+Bv9ssoX5eMFEQFCDSwMeC0zV44VRrFZzmAllZEzk/mPDdjGWPfCsR/tq1vauxP3b30pUKW0i2FQSLykGFfKp7+sZqPDQ4Z7MVJeNkQRUZig1YgKbxiISd4IiM4BVW4ohc4g9AvqBH5O3FMRmNp6TVPg5PgkYQtm2TVAcCWwYEPWgmg04wCVvPbiYPo2ebyfEwHLdGNnZPNZPfYvuXs9T86Pz+10OM72T9E0EqpplaQ0G5ZRB6wRlUChIDT2oEYeTgxIKn2HqOFbIwGwn7oUrk+Y2hylSJ+V4u8h9LqLADKZ/l4Jgn35d7D6eY/NIKfyWtbjNqBKfdztOv0Wf+tP3xGv14jb7WUS3L8phJIwrs2I7n6t493Xr3VNH8bhymSZFx+Nms8Fek+O2TvruL6tv/8T77DwAA//8DAFBLAwQUAAYACAAAACEAeKFQBr8BAAAFBgAAEgAAAHdvcmQvZm9vdG5vdGVzLnhtbKyUzU7jMBDH7yvtO0S+t3ZY2K2ipmi1FYgbguUBjOM0FvGMZTsNfXsmadOWD1WFcrEz45nf/O2JPb18tnWy1D4YhJylY8ESDQoLA4ucPfy/Gk1YEqKEQtYIOmcrHdjl7OePaZuViBEw6pAQA0LWOpWzKkaXcR5Upa0MY2uUx4BlHCu0HMvSKM1b9AU/E6nov5xHpUOggv8kLGVgG5x9T0OngRZL9FZGMv2CW+mfGjciupPRPJraxBWxxe8BgzlrPGQbxGgrqEvJ1oI205Dhj6m7TpmjaqyG2FfkXtekASFUxu228VUaLVYDZHloE0tbD3GtS89P68Hcy5amHfAY+cU6ydZr5YeJqTiiIx1im3GMhNc1ByVWGtgV/tLR7B1uevE5wNlbgFuc1pxrj43b0cxptBt42rK6q/0J1qbJ+1sLp4m5r6SjG2hVdrMA9PKxJkXUsoROPel+azbbe3KSNosrRxFBO+llRM/IZYqcjdI+0JFJb1pxlzMS+1dcCNFF9K65LmVTx/crt51rMvn1R1z1EN8N2wp8NuW9j0bXj4OaD5UphGig6R+E+7cqxTeL/LDYIcF7Rpi9AAAA//8DAFBLAwQUAAYACAAAACEAqiYOvrwAAAAhAQAAGwAAAHdvcmQvX3JlbHMvaGVhZGVyMS54bWwucmVsc4zPsYrDMAwG4P2g72C0N046lOOIk6UcZC3tAwhbcUxj2di+4/L2NXRpocONkvi/H/Xjn1/FL6XsAivomhYEsQ7GsVVwvXzvP0HkgmxwDUwKNsowDruP/kwrlhrKi4tZVIWzgqWU+CVl1gt5zE2IxPUyh+Sx1DFZGVHf0JI8tO1RpmcDhhdTTEZBmkwH4rJF+o8d5tlpOgX944nLmwrpfO2uICZLRYEn4/Cx7JrIFuTQy5fHhjsAAAD//wMAUEsDBBQABgAIAAAAIQCqJg6+vAAAACEBAAAbAAAAd29yZC9fcmVscy9mb290ZXIxLnhtbC5yZWxzjM+xisMwDAbg/aDvYLQ3TjqU44iTpRxkLe0DCFtxTGPZ2L7j8vY1dGmhw42S+L8f9eOfX8UvpewCK+iaFgSxDsaxVXC9fO8/QeSCbHANTAo2yjAOu4/+TCuWGsqLi1lUhbOCpZT4JWXWC3nMTYjE9TKH5LHUMVkZUd/Qkjy07VGmZwOGF1NMRkGaTAfiskX6jx3m2Wk6Bf3jicubCul87a4gJktFgSfj8LHsmsgW5NDLl8eGOwAAAP//AwBQSwMEFAAGAAgAAAAhADvOKlkPBgAAtRsAABUAAAB3b3JkL3RoZW1lL3RoZW1lMS54bWzsWUtvG0UcvyPxHUZ7b/2I7SZRnSp27BbStFHiFvU43h3vTj27s5oZJ/UNtUckJERBHKjEjQMCKrUSl3LiowSKoEj9Cvxndm3v2OPUbYKooD545/H7vx87Y1++ci9m6IgISXnS9CoXyx4iic8DmoRN71ave2HdQ1LhJMCMJ6TpjYn0rmy9/95lvKkiEhME9IncxE0vUirdLJWkD8tYXuQpSWBvwEWMFUxFWAoEPga+MStVy+VGKcY08VCCY2DbAxoUcHRzMKA+8bYm7DsMvhIl9YLPxKFmTnKajvQFVb88EZQbgmBY0Q85lm0m0BFmTQ/EBfy4R+4pDzEsFWw0vbL5eKWty6UpEVNLaAt0XfPJ6XKCYFg1dCLsTwkr3drGpZ0pfwNgahHX6XTancqUnwFg3wdzM12K2Fp3vdKa8CyAsuEi73a5Xq7Z+AL/tQX8RqvVqm9YeAPKhrUF/Hq5UduuWngDyob1Rf1b2+12w8IbUDZsLOC7lzYaNRtvQBGjyXABreM5jcwUMuDsmhO+DvD1SQLMUKVCimX0iTo14WJ8l4suoEyEsaIJUuOUDLAP4DaO+4JiLQVvElzYyZZ8ubCkBSItI1VN78MUQ3XMIC+fff/y2RN0cv/pyf2fTh48OLn/o4PqGk7CItWLbz/769HH6M8n37x4+IUbL4v433745NefP3cDVRH4/MvHvz99/PyrT//47qEDvi1wvwjv0ZhIdIMcowMeg2EOAaQvXo+iF2FapNhOQokTrGkc6I6KLPSNMWZ5dCxci9gevC2gD7iAV0d3LYUPIzFS1AHcjWILuMc5a3HhtGlXyyp6YZSEbuFiVMQdYHzkkt2ei29nlEJCT9LShkbEUnOfQchxSBKikN7jQ0IcZHcotfy6R33BJR8odIeiFqZOl/Ro38qmGdE1GkNcxi4FId6Wb/ZuoxZnLvY75MhGQlVg5mJJmOXGq3ikcOzUGMesiLyOVeRS8nAsfMvhUkGkQ8I46gREShfNTTG21N3F0JCcYd9j49hGCkWHLuR1zHkRucOH7QjHqVNnmkRF7AdyCCmK0T5XTiW4XSF6DnHAydJw36bECvera/sWDS2VZgmid0bCVRKE2/U4ZgNMDPPSXK+OaXJa42YUOncm4fwaN7TK518/cnfWt7Jlb8Pby1Uz8416GW6+Pbe5COjb35138CjZJ1AQDui75vyuOf/nm/Oyej7/ljzrwuYcPjltGzbx6UfvAWXsUI0ZuS5NE5dgY9CFRTMxlNPjfhrBMJdp4UKBzRgJrj6iKjqMcAqyKkZCKHPWoUQpl3DJMMtO3noDXiIqW6tPrpeAxmqPB9nyWvHaOWVjZqG5304ErWkGqwpbu3Q2YZUMuKK0ilFtUdrUZKc088i9CcWDsP5todKoZqIhWzAjgfZ7xmASlnMPkYxwQPIYabsXDakYv63gNn2FXF3ahmZ7BmmrBKkorrZE3CR6Z4nShMEsSrp458qRJfYMHYNW9WrdQz5Om94AzlwwjFPgJ3W/wixMmp6vclNeWczzBrvTslJearAlIhVS7WAZZVRmKydiyUz/ar2m/XA+Bji60WparK1X/kUtzKMYWjIYEF8tWZlN8z0+UkQcRsEx6rOROMCgt05VsCegEt4XJtf0RECFmh2Y2ZWfV8H8rz95dWCWRjjvSbpEJxZmcDOe6mBmBfWmsznd39AUU/LnZEoxjf9npujMhVPuWqCHPpwFBEY6R5seFyri0IXSiPpdAacHIwv0QlAWWiXE9E/aWldyNOtbGQ9TUHBsUQc0RIJCp1ORIGRf5Xa+glkl74p5ZeSM8j4zVVem2bNPjgjr6eptaPs9FE26Se4Ig5sPmj3PndEPdaG+rSefLG1e93gwE5TRryqs0PQLr4KNs6nwmq/arGMtiKvWV37VpnBXQfoLGjcVPsv+/NAv1B4/gOgjNjlRIkjEC9nBA+lSzEZ90DlbzKRpVpmEf+oYNQvBVO6cs4vFcY7Onh6X5px9urg3d3Y+snxdzCOHq0uLJVoq3GbMbOGPLd6/C7J34JI0Ykoa+8g9uJm2J/9GAJ9MoiHd+hsAAP//AwBQSwMECgAAAAAAAAAhAFNigk98SQEAfEkBABUAAAB3b3JkL21lZGlhL2ltYWdlMS5wbmeJUE5HDQoaCgAAAA1JSERSAAACPgAAAj4IAgAAADNQl/sAAABmelRYdFJhdyBwcm9maWxlIHR5cGUgaXB0YwAAeNo9xrENwCAMRNHeU2QEG5/PMA6CIKVLkf0VlCLvN1+u+xlyfLyKVxQ0TMXuZ82GFuY+hKVyxYI72N28cdCyh2JGsrHwjJQMVXkBogcUAtaWRVcAACAASURBVHhe7L35kx3Xdef5PefezHzv1at9w74vBHcSJEVSIimJraUl2ZJG8ozDbjs6YsYxPfPDREzE/AH+bWJmemJiHPa4x9Pd7rYkW63FtizRWriJBHcSBEAsxI4CUAtq396Wyz1nfshXAEiCLIBYC3U/8QKFAl7ezJcv837znHsWUlV4PB6Px7N04MXe4PF4PB7PrYWXLo/H4/EsMbx0eTwej2eJ4aXL4/F4PEsML10ej8fjWWJ46fJ4PB7PEsNLl8fj8XiWGF66PB6Px7PE8NLl8Xg8niWGly6Px+PxLDG8dHk8Ho9nieGly+PxeDxLDC9dHo/H41lieOnyeDwezxLDS5fH4/F4lhheujwej8ezxPDS5fF4PJ4lhpcuj8fj8Swx7GJv8Hg8AFBN08ND47VGXKmltTibr6bTVZmtuVocN+qNNI2zNEuzNHMauyBNU1UloksOFRUiYg4MlyKKjBaisBTalkLUWorKLaa9FLQWo6gQre/rWNXRcskRPJ5lDqnqYu/xeJYXopo5aSTZfK0+W6nNztdV9Z33B948dCoVShwyIafsnKoqFEogYoICBCKhkNmoqoh87D5UFRCQMkgy1tSoWAITmNRaCoxZ39/+jcfvAlG5VOhoLbW1FIuFsBBYa7yzxLPc8dLl8SBJs/l6Y3q2UmskA2Nzbx2faDQajSStpa6RSOZUUpmvpkmmjkmgQiRErJkVRwAxg8CAAqowFDCTiFzS8MrvOIIKmViNI2tISTNyCakCyIgUxEDA2tFiC5YNI7A2CoMo5GIUblvVc9eGPoX2dZT62kulQuTFzLPc8NLlWXao6nw9mZqtzFYbRPTW+wMTM7WZWjJfTyv1bHiqVokdACVWNiBDTIHaQJgMhKHq1IgCpMofvn1UFUY1/9vHOQybEDkYJWZVIiV1DCVogwJHhqFQgQqpqMtSJSiJClQs09re1mJkSiF1l0xbS/GRHetai2FobWd7S2drqRQFn7Rfj2fp46XLc/sjqtVGOlttEDA4MXv07MToZGVipjZZTycr9dOjswpLxpINjTGRZpGmSgZESoBCCEbICJQUBFVHJAAcsdAlzB2Wy7qnCMpQhkI1N9gIADSjQMic/zX/U4gBzs06hUJVxMVOY1E47WmzKztKrUXb3lrs72jbvKL7ns194ly5FLaVC6H1S9qe2w0vXZ7bE1GtNZK5WjxbjZ3IG0cHz03NnZuY331kOE3YcMBBqGEIw4ZyWRCQkqgRZ1UURBfkBEKipNT8jUgVIEcQvoR0kXyisbUAAQbCqgsLYqSAghhCECECLoyjTX/k+X8hBQmxY0MKhUiWOkmdc+IUjh7c2t/bUehuC1f1tD6wdW25GBlDbaVCMfQy5rkd8NLluX0Q1SyT+XqcZG5qrrr3xMjZierZyer4bPXY0DlrTGhCa0LLAcOIqjAphOFIxIhjclAVBA4BQSjXM4CAlEVICQQlNM0xJgjoEoEYqpe58kQEgkJBYFLJlZGsJgSnRArK9QzIPY8CcH675v8IaEAZ8nAPMAgKUkbG4jInmWrqXOK6y6X7tq1c019a3dO6ZXXvphUdzkkxCgPLl6WxHs+th5cuz5JHROM0qzWSOM0Gzk0fOj15Ynhiz9HhyWpqwsgEobXWSoY8FjCf9kEEUihBAZAKNW8EVTIOhppvXbg7CNL8mb+LACII4RK3j+LyFYGabyZo06TL9fL8gVLzcIHc6rpof0QQhst/URCIVSFQsCpAZEgNgZ1T51ySNpK4znDf+Oz2ret6163o2bamu2QQhUFo7Sevynk8txpeujxLFScqIqI6ODZzYmji4KlzE3O1F987K1QIwsia3GYR5JKli8/NuYpcUo1uCOf3u/ihLgpdGC7/yQSrUNHMSZq5LJX4rg19T2xfddfG1ZvX9JaLARMF1rAXMc9SwEuXZ4mhClFJUjc5X3v7yPBLe0+8fXgwFRMUS2QjUmf1EsYQLffLnNBcLROQAo6Niki9Ucwy3dhn/8XOddvWd29ft6KrXLJE5BXMc2vjpcuzlHAiY9Pze48Ojs9UXj86enCoysZkmTBIoGwMqbCXro9C+eKc5kaYwhFBRBKEAhOws0ggaXuRv/nEvVtWdN61aVVbMWK/Fua5VfHS5VkajExVjp0+Nz5bfebNE6PVrIEojVMjKVEzZD2fl5Wg4I+uNi136QIuKliqzRcBXAVl0AASwFmIaQkjFOjRO/s397fesbpzx4YVYeCDEj23HF66PLc0lUZydmz6yJmxX711crbSmGu4mbpQEDkTRJKF0ji/ppMnPQnIkfHS9THQRQtpBCjxPCEDAmigCAGrjupE1khHAZ2Be2zHykfv3tjV3rqqt93X7PDcOnjp8tyKpE4mZysjk3NvHR15/ci5MyMzcZwFYYBm0SUowAtB6JoLExGg+jEi5aULwAeli4EF9yFUOc9dg0KswoKcIs2cAqt7O/o6ytvXlH/3M5s7yqXWUvFjh/d4bhReujy3FjO1xuDE/Fyl9sbBM/uOjx4ZnoctRJYDEgVUJfcN0oLDCwtml54PLNDLiCZcvnzw3KhVIFd/JaekgFhRBgkxyDri1FGcJEinH9vasXF1/5ceubu7rdTXUTaXSsf2eG4MXro8twSqmK7Uz03ND07O/vPuk8OjU6dG5kqFkuUQyB2B/kK9Hpw/q4QFWcstVzr/SIC8FqNLNJ2tuf7O8me293/j4S0bV3S1lQqRL8/huRl46fLcZES0FifnZmrP7zm9/9TYXLUycG48DAuqDGVSBljzqG7PzUOJMrZiwjRT16h12MaT96z+wgNbt6/rK0ZBuRguUmvY47mmeOny3DRSJyIyPD7z3O4j//Ta8aG5ICp3RJg30gCxKhMRgQECxF+oNxlVJeNM6JQZxORco2ol/fx9G+7fvvLRu9eVI1sshD4fzHNj8NLluQk4EVU9ODj18r7jP9+1v55oIWwBrOYqldehVboorEAvcm15bgJGnYUTCtxCHUVWEbaZsEsqX7xv5T2b+p96cGu5EEWh9QLmud546fLcaDKR98+ce+vAsf/08kCCQmTISmacM3AEdQgc2QWhOq9b8NJ1czEqjGwhVJPyoBgH48iAAM1cUnvqrpUPbF39ufu3dZYLoTWLjOjxXAVeujw3jkaSjUzM/nbfyR++9H5KgaRiGUQsqqraNLHUNttkXVAsf4nefByxEjGkmfZNEBhWx0hSDhxCUSoEhLj20OaeLz648ZEd67taS4uN6vF8Srx0eW4E9TQ7Ozl/8NjZI6cnfntgtB5bNUFIDdKs2fKjqVHUfJ73inWLISAlznvBAJqbXwwldUIsMCAGVJ0YuHKEezb2fOOzd25b29ddLiw2tsdzxXjp8lx3xmbm954699v9A0dOjdVibSSGKVIlogYgF2ca6YfTjjxLCFr4AlVFM2h3e+GOdR1/8Pm71vd1tpW8gHmuJV66PNeRuVpjaGzmmdfff+f4uZHZmuGA2TSzikE+3v22hImIOFMTO5em1TtXlr775J2bV/dsW7vC+FJSnmuEly7PdaGepOMz1WfeOPLGoZFDZ2fLxTAwcOd72ef+QPUr+bchDCJwHsGhJJLW+1qwvq/1O0/cu2lNd19Xq/EF6T1XjZcuzzUmczIyMTMwOvm9X769+3TVFrsCG4WasCZ5L3tAvHTdxpACMGB2YAWJqiFN4qQUFB6+s+fbT2y6e31/FPgAes9V4aXLc81wIo3EnTo3+bOX9x46M3l6OuOgCCFSMJQvrojRrDTkr72ci3MAbhvygBsSBYGUTGKK9Ualjef/zdfuf/K+zd1tLYUwWGwQj+fSeOnyXBuSNBuenHvnyODfPrt3tCpBocURG8mMNBgiagG7MEfTQu+o7Labr68IOv/xRZQWagUudXQhXHQhpZy0KWMNUChSNJI8cWfvF3ZufPjO9aFhX8bX8ynw0uW5WkQ0E9l7fPjlPUd3Hx07VyE1VqBEZCQxmjDUIdC8o1Ze2Cmv7U5u+UlXPqcrQ5SMg4Wq0UYmIhRZNkZFCB/tN7aEUAKggDAWUvXAgFjURG2mZUGYuYa4yrefvOM7n7trZVdbwdfw9VwhXro8V8vA6MzL7w386KVDlUYWWKuSOwLPB7rflt6wK0able+Vm9LlHAUxFW1W7zZzK9esmqzR+GTNOCfM2e2YJKDIjUoB8lmHiOiz96x6eOuKLz+0JfK9mD1Xgpcuz6cnc/L2oVMHB8b+efeZ8UpqjWGFitwejq9rDDXvNCUiBUhTCmuJdpr67+7s/8Pf/cLrJ6b+vx/+dqaiQWgUt6V4fRhRRCFsVvudR7f8zpMP9He2+uh5z2Xin3Q8nwZVHZ2t7dpz7McvvTeXmPnUGFuAKtQx+c5al4YABSlI2ADsYNO02tcVfvGhbd2laEOHLUc6XjUEtrd7d7L84cYQJEtjDf/+rcFzFfeVh7ZsX9vb6ctHeS4DL12eK8aJ7j85/Oqh0z9/42QlJjIh5yvtSoAsoTn3vHV4Y3wP+T4IIBVicpIWjdvQ37lt/QporWg1YKgxugxCLy8+4cJhncLfHJysxO6Ld809sG3tip6OT9jW44GXLs+VUkuyV/Ye//HLB09P1GINrCWCkjpVkC41H1ez6O91RwlKBAWrBgTnkka1UalWezpbdmzsD4OiaiWMCmyYkRKWkdNMYVRApKBgz8mpkdHJQ6fHvv3UfSu620uRD533fCxeujyXS+Zkptr4xetH/vqf92YUFgqRRUqijNygYL0QBr00rAYRAbGC+LofMwmYoIyUssxqtmV1e6MWdHeV7ti0EiBFmFko20AqTIEum9VCATMklFihwjwwz4feOHt6qv5HT9+9cUV3T0frYgN4lileujyXRS1ORyZmf/bK/p+8eoqjNmuMampUqBlMmLOkzAWCIyNpHVlMUUmDFgMJpUFAioCuqZIR1KgKGUec1Oc295e/+/T9bS2FVLKN/R2ACIJEEiXJozmW1Hm8BhgVhROxREFYaHvrxLSm7/yLh7bdu2Xtur4OXzjK81G8dHkWJ8ncy++d2vXukTcPD5eKRUHmVAUgNfThOf5azvjXFVFybDrLQV+RRubTcaXQmAKERBXhNZYuVYs0pjBmm2bJ3Vv6PnvnmvZyGQCQABkhRJaKS2LDIRsjS2nJ8OpQBQQMsKoGmoSAWt4/0qi+deqdE+P/+isPrOtpD3zjSs8H8dLl+SREtdpI/u653f/5l+8WSy0cFMURk1JehHAJKdUHIUCckNHHd27/6gMbXnhtz7P7zs7EYS0sWAhf+5L2qqqCTJxEhjeuWxUVigoolMH5WlgSJ5lzi41ze7IQwEJArmVq2J4crR8+Mzk6OvFvvvnYHetXFiO7bNyonsVZbp4JzxWQiZybqfzFT1793rPHg1IPTAsoIgpYyIqwXPP5/UZCqmJco6cc3ru+97/75pNfuKMncrVqxpmJ+ForsgIOUBVJ672dretW9hSscVBVbS4TKpIkVlEoQa/17pcWRAxDjhi2WOg4Ppb++Y9ffm3/yelK/cYEgnqWBF66PJcmzrLjw2N/8ZOXfrH7HBfalaNMOROG5uEGcs3n9xuJAoaJNDt6bGCq0mgvF//V1z732I7VSGpOVK69KpMyqTrN4lW97b3lEgA4R2DAAAxovREniVMl8sYFCGqgJA4w0Zl5+7cvvv/ynuODY9OLbehZLnjp8lyCaiPed2LoP/7zm28cnQrCkCjLzQZABapEQksvEv6DKDEx08mh8VPjc5noqv6e33t65wPrOuqzUzDX2pHebK7pCNmq7rZiFADKqtC8OhKpolbPxCnyKM3Fxru9UZCSQB3BqapTGphq/OKtE7veO3FieGKxrT3LAi9dng8zMVvZffjMD184/PrxJKPQ0DxjnigmSkGZsMtYMiJHS/viIQBkx2v00qHRauyI4ns39X/3C/ev6uuemk/ZWoJSc9GLr7IskyqElAwZZCt62qMgAEDEtBDloqqV2YpLUwLTRRXllysKKCFjihkJIQXk2JT8zUtn/vynr54YGl9sc8/tz9KefTzXnIm56juHB/7x1SN7Ts4TRYYI6gyEIQQl1Tx6W5f45EpERBSn2cx85flXXh+YmAVsYPTJe9b93tN3tkeNRm3OcZBSJMREMSFbbMhPggBSIoVl7ulsD4MAUOSFeKnpnJytZQ0xwEdiNpcj0uxHqgpVgpA4iM6l9q3j0//XD1/cf2IwddfcqetZSnjp8lxgttr4u+f3fe+5wwfOzBsWK3VSp1pQjUiYhABiIVYQhK59GN6NQ1UV6GgpPbCtZ9OK4PjAwHQtBbgU8tceXvuNxzdKYzYRpFQQZqaErroeLgko0ygIujpbA8sAFnK3BRAHjFXS1JQIBJWr3NdtQe47ZYChBsrGJRHioNCyd6jxv/9w11uHBmpxutggntuWa+3T9yxNFJiYrf6/P3vp528MRoXW0BKQLcS6mYUmJrcLhDRxxUL41M7Nn7+vLwowNjI+MzVXDnuMdb0txW89+cCp4el3js3aslWCg6Gr1BICgSBabim0txYNAdD8wTEfN1Yam606WMNY0vEv15oLp50hkCTjkMOWgen4z3700p9887HP3LmxtRR9wvae2xUvXR6I6tj0/F/85KVn9w+WWlpZVfXiUkS320xKgFMllXXd5Ye3rAMw1daS1CVJXWiZVDb293z3K4+NjP1mqjGjQcFRxHS+acmngwAoaXdnW1sxXDC2CEogUvBsvTExPafiiA0R3aB6wEuK/HJkVYWSCcYb7v/5xzfjRB6/Z2N7ucA+KnOZ4R2Gyx0ncnZ85v/4wXMvHBqzLb2g4PZPK1JYa+MkeWf33nPz804bXZ3l7s7WwLAqMnWG3KPb137tyXtLqCGpi7DS1RleChFV5/q6yq2BBVxeYj/P63KC0YmZyZkKiJu+TD8RfwRSYgWrsqoqiSlOpsW//Pnbz71zdHqutozKj3gAeOla5mQip0Ym/7fvPfvWidkwarUKWrC3buPnfgUMswCHBydf3XciySxggkIUhsYSGWZAIsbvPHn/XRu6Aq3DOVK+KkUnAjNUO8thQACyXKCUFNAsy4ZGZ6pxBmLDpKq38cm/CijPKTRwVlVFAZqX0g+eP/DynmNT81WvXssKL13LFydybGj83/7w+X3DCRfaSRG6xCzl4IvLhyDEdorbn3n75FAtVlgoQZRVSOAyVZf0F4Nvf+WzK9oiTRoQ1atIZFPAKYiprVwm5vMRmqQAUZq508PjsRMHMMEXm70kQqREpMoiBpnVzGhGxFNp8fsvvv/CO0dHZ+a95C8fvHQtUxQYGJn8s5+8fPBcg4KiKljFakq6TKRL2djUtBw+O/3CW/urjRgMgQpUwEqWCIrsga1rH7p7a4vVLKkaAxC0uUKleYeXxfZzAVVl0nKpyJxnifFC3hjVnRsanwExQZ27LtOvAgrWZoIa5cF72uxT0xRSzQ/oCj7TjUZBefIAq2M41ozFMdFEI/jRK0d37Tt+bnp+sTE8twleupYjonpscOzPfvr6wdP1AEHoEiuxQhK2Sz3R+DJRECARp5boxZffOzw86QAwZ0qOGIaZLeCKzF/9/IPbVpc5nXJSy4jBgVUXSKpMGVk05WxRiIgKlLYVjCEAligw4hixwo3EcmayHhEKKsKRo2teJV0BAgxgF9Kr8wJUVomEtfkiOILjqzEvryMKFRJH6ogcsYAVbDSxUmUkY1X98avHXtl3fGTKq9eyYFnMU56LUcWJ4fF/90+v7jk5YW1AOF8JHldkRix1SDWAs8YMjlf+/rl3zkxXGbAkRoUAqBBp5tKtfW2fvX9rR7lYacRiAkcMMtQ8Y3KZtpeqQtUyF6OAiBZMH4A4Vj07Ojs2OWeJmRTXpdk0karRzGpikFl1FpnV1GpiIEZhVK04q1mgaSi3uuWtF2JeCaosjlTBZnQu/ftXDr2y54hXr+WAl67lhQID5yb/8h9fe/3YlAmKTHldieV4GRDUiAMMSl0vvXf2568cmq2nTIGBMJDPkCKpgX7hkR13bt+QOUocMlhHRkCkeW+UyzG5AAAKtjYMw4WsAwUxqFCpZgePna02YrABlK+Dx26hoK9TiABC5PLVvrx1cx7+QEwKFjZydbGUNxginH8UMIXhCn6069Bz7xwenfbrXrc5y3HOWrYoMDg2/Zf/sOu1o9McdSgZqOC2DiZcBFUlg0J7lVv/+ZX9L7xzOHMKtioKBVQtQ1T72loef3D7it72erVGZDLlvNY748rsIyYKApvHbwIAMcBTlfjIqRGykTJr0/14zb6OhW82L6ZkMhMmXIhNMeFiZqOG09lKtdKIE9EE1nExs6WESqLX3GN5/SDJC2oJqZJwYbRuf/Ty4WffOTo1X79m59Fz6+Gla7mgwPDE7J/9lxffODoZFVqNEkFAWM6h2HmdIRG0tLSOzdR++sLbu/afcmAiysCAYWPzO+SR7Wsf2b7SpjVyWR6jscjQl4KYmS++40iBoem5wdFpExZyFyI1y/ddGxYsPIJomtQb9VrsnLBxhGq9Llm8dkXrmr6ikVq1MlNtNBqZZiaQpWOFa3Pdjlk1X7BT8Gwafv/FQ8/vPjpXaSw2gGep4qtpLAtUcW567v/84XNvnJi1UZuFsEuUltfi1odQkKOAlAJNARe2lg8Mzvyn5w60dXTt3NAJmIzYgIgUQF9L9Pj2lbv3D0xUK8WCUajAKPgK/HsEACofeEyoNOKjAyPT8/VCW1ldygDrNclOIuQBg8SqiFOwxCtb0d5WGq/z2NSEgesumnu2rH3yka1ZWt9z8OTh0+NTldgRZ5opKV1VFtsNQxVGYQgwyPJavQp1ZOak+O9/tS8Kgqcf2lYuBIuN41l6mD/90z9d7D2epY1zcmZs5m9++fqLB0a40M4wVtJAs3zZ43Jn3tsNVbDjgEgDTYwk1hpny0Pj1crM1I61ne2tLZkSILyQpR2FdnB89ujASFQoKuXuPtDl1b1QUKYUInns/s1rOsuGc13gkyOT//TKvqHJelRsJZcycPWOEIJyHj9CRkFZlqRZsr6v5btP3f31p+4ll5w+cmjbqo7ffeq+3//qZ+5b27NtVff9d67duLp/RWepVIxGBkdArJeIcjyvZecXl3Qh3oTOx51c/IYLvzc9q5dzqq4IVbJKBiQGjqC8sFioNko02H/0zKrO0tr+Tmuu9qx6bjW8dN3mqOrw1Oyz7548fGZurpan9ghUlUiXjl/oOkAEMFyepAViVQpMAJUz49NDc7U7Nq/uiixnCSmJMjMKxWh+vvLekYFG0JaaUiSNQFMhs2jVJoIKTIowJPfYfZvWdbcYznKL4eXj4794430DCohJSWGAXDM+pc2jIKNpqPWEgpotxUSIpx9dX/jvv/X40w/v6G0vUVprpfSPv/OlJ+7f0hIa1YzIRcas6e7esXH1uUrjrQNnyIQfjNzJdceBEwIzmBUMZTgwlEMQNM8QI1KyDDIqBFViJSZSq44hSnyt1Ss3LptNUNHs0klEbDRjUCx08NiZjf3ta/o7lu9D2m3Kcp68lgWTc9Xdh0+/cWDw6EgNAENIJQ8v+5Sz420EqZLmAe6sYFLXEgVBWHz3/cH/8OPnR+caCAogQxBxaUi0bX3/+tVdcdJQcB4UfzlJXQpiaKAJZbXK9IS4LF/Nmqmnx8+Mz8wnhTBikuZSV77Fp4WgQhxzKAwjda1Ofu7+Tf/jH31j5x0bBS4RueeOLf/tH39n45q+VDWFAwzEQB2A0+Ozu97YnyAEf8jD1owcybtjC1iI89oWApuSTUCJagpkRCk4BTtiISsUCKyAla5X0CI1E6kv6JaCoUqakWbMPNWQv/rHVw4cH1psJM8Sw0vX7UylHh88NfbrNwbOTDTy6cnzCeSKUTSWhd88cObPf/LywEwDhpgTRgKgp61l85puJDWGUxiBwWWoV94yJjSJpo2zZ4egAIpAeObM0MGDR0iJ2KrkPrarDS8kwHFUt+2ZE1sd/c5nt/4P33p8Q19nqs4JmCgIrTHGiaqKCgAHAFSoJNnrB44fOjnMhYJc4jBIYQRWYIQ4I+soSGEaaaNemarNTTWq07X5iersZGN+Kk1qGSRj48gIsZBxsAJ7rU2uT0AX/iQbFs/OpH/+k1f2efW6vfBhGrctqZOjg2P/5bn3j57Llxp8A8PFUREGjC00XLDrwNB847k/+daj96zuATJAyqXC2hXdIR9jTfJKxc1lnk+ECErIKFBTfP3A4PpNZ+7evtoafuXg6bPnZsvFVhFo0zChTydd5ysmq4oo6okUksa3nrjvD778UH9nK6CZ5DYdMSAEJqgyASopGQPQ7mNDz7y01wVFtozsEk85SuyahaSUNK3XKpnE/e2FOzZ3r1rdb0MD1ThOJ8amTpwdG5qezTgqFMvGFsRBiKD86T7apyLvNK2qKsQIi8MVt/vocCmirWtXLbatZ2ngpev2RBXHh8a+95vd7w9nZIvQeLEJ1gMAIBJVYSJTStN495Gz2Q/nfv/Ljz1xz0YCgijs7+0sFW0sGakBGVyG25WgAk5N5Ip2sFr7m1/u6XrtMAijk/MaliMTQPRDcQ5XSp7eIKoMShvVMmdf+9wdf/jVnf3tJSBVkCEDgAgE11RIYlIoG8C8NzT5k+f3nJ2oh21dudB8VGVUFcwkKkk1rc5tWN358IP3PrR9zabOYrGYF2ZUEW00Gudm6ntOjrz23rFjA2MJRYVSe1Pgb1QOhgLIQ0VIVIlg6kLP7j1bsNLSUl7Z2eaXvW4DfJjG7cnw5Ny/+/nbrx+bMqElxJT7ha5iclw+CJOwMjQwpKoj47OHh+acYG1/ezm0NaJde440Eg40IEAvL79AATHKBFYzO5+cm6gNT1UbjmwQkeZO+4tNriua4i8sj5ExlUYWUfY7O9f88dcfXdFRBhygCgYxUW6LCEEApeY/8uDk9PefeeuVfYNc6mTLBg3TzLn+IAwwXL3GSlZWPQAAIABJREFUSfVzD2z6w68+/KWdW3as7GprKZaioBiiGKIUhW0tpVXdrVvX9m1b21ssRiOjk5W5WhhGN1YtSCkvsyEMRwCg83U5M90oGd2woiMK/CP7ksd/hbch0/O1v/6nt189OGsLbcQVkpg0UpgrnBNvMh+a6xT41LnAl08+tSvUaMIA2zAuFg4NJ/Vf75maGL1n66pT0/OZUyacP5m6+BMBESRwMYENBUFUUFVhXUieorw2x8Kbr+g7WqhhT0zEcb1Gqo/et/kPvnLvyo5ymmbWkpLF+fW4/LDVgZDf++fmaj95/u3X3j0mti0MC9C61QwXasxf2JEqSebINT7/mR2//6WHdqzuBBJoVbUEcqoZSACFBkyuNdL7N65c09fV1dH2D7/ZPVmp2KjM/CndoVeK5s7RPCgp9xwKQHZk3v705aN9bdGTD2wL7BKqGOK5BLRsKyncrozOVl946/B/+KeDErSRhfI8ISGJVJeGdCk4T9AhVdI83poAUgVBlFQBZRYAcu2jjDSf+EhZhQEnpBw62LQ6U6B6e1tRlCqpgiOjDIWQXkq6PnJcJKBMlUgN5apAuXSdD/L4NNLFcIQ04WLGkWaJzI1/ZsfqP/nO03etbYeqiBJzfnDadKIpSQpNwQZUGJme/d6Lh36za389lkJrpwNYM1YHypfxguYHUQFJCkqy+JFtHf/T7z29pacMqGiDQEABhIsqgLBqCs2YApCtAj99ce8Pfrl7LisVTZ5w9qGw+yv4vJfJ+SSyXLcW9kBEgWraU3T/yx8+/eC2VXxjLUHPteWa3/uem0mtkbz+/um3j40UWkohCbkMYlRDvT5zxHVEm4oFJSirGoUV2LyuenO+vw4zD+UWhxDUiBoiZs0irZVaiknYOTxnxiuWuERKAhVqGk0fQukjL7BqAYiUWChPUMCC2ZVXftKF1xVCKgQH1Oq1Ozb2/OFX7rtrbTsAELHh87EfRCBS5DpLASg6N1f9/nNv/uy3B6bTYlTuUJdazUhJESoWYk/UQBkgUaikK7uLv/f1Rzf3tAAKJVBBKCTKg1V44QUiCy6oWihagN996t7PfmZbYFzm0gsR7EBTXK4DzT0oVEmVFxqVIXDzBjpUMf/rD3YdH5ryT+1LGi9dtw+ZkxNDE8++fuDQwHgmIpRPQQsNOpYIDGGoEGVkUw5SY1PilDhlk7LNKHCwqpQXrFtssGuDgKFaNGgrhS3FZs3iT0Qv9XK4kDybv66k8PylcGQTKpISajOrWvlbX3zo4R2bLxkNQRBCSpSCWTg8PVv/3q/efvbV46TUXjSkQnS+T5c0twCax8nOSdxCyVcf3PTgun6niYpDLg2XOn4FpJkRoKppG9N3P//Alk5DLnOLnrnrhoISKqRkgzAanJz//i/fGJ6YFa9eSxYvXbcJqhibrf3VL948ODSbSt6PXvU6P95eD1jz8hZQkLBRIsABqVKqpHmekICouYxxg2iu+gOXdyYvKV0ffV0tAso4UsnCrPKtz9/z+fs3AVBxH1Iv1Xy1JwXUKZ+cqP7Vj195ZteJOneFhTa6RP1lghKgIAGJqFN1K7sKX9q5saCJkWShoVcezfHhE6LAgt2mhIxQvaO7/MRd60oFmzh3E710SkZgUlCxrevU6Ny7R87MVePFNvLconjpuk2YmK18/zd73j41z6agJnJQgctjDm7eXHHFNF2ESnlQHMiR1JHMIJmKtG4QK+UZUETXO2Djg+hFLPbe3DuXK92i6rX4aB+DAgAZECXV6cfuXv/0/RvbQwMR4gupZqoqktt2BEQZCu8NT/3Z3z334t4TLmrjsJR3yvz4IkkKuDRLjJFNq3s39HaQxkSyoOKfePD585IqJGVkD96xoa217ESbp+dGfnkLnH+Sy1QH57IfvrD/0Olz7pqUO/bccHyE4e3AfD3ef3LkzNh8EJUd4nxqIAgpKC/TQLhZpcBpYZYFsOhM3XxgR25QOclqXS1YubKtUa+eHhlLNaJSFwUFdfnyza1Gc2oUXTB7iDT/14vytoiEgAUJ/NQfggBSlTROelsLX3/ygfW9ncirzpM5P2hzL6rEJgNeOnz2b5955dDpGVvqYCK4mC6ZqK4LB0vNVtHlcnjH9g3MRgRKILASFoJMPhyqd95kAwjUbGq2ekVfW2vZTFbyt1zOE8A1h1UAUYGSCpnB2fhvfr27t6Nl6+rexTb13HJ46VryZCIDo1M/e/XQsZHEUCAaKITJQS9OdL0xM8V5nSIo0sxlzjkRQrPLorWGFAol+uTJiwAikqQ6u2nrun/9jc92FILn9594efeJfacmHWflYjFkQ5I3HCFVoYWSjNr8wAs/Pw10kT20MEgeDQgDVYIquaZVRURgJ8Y5cg5ORKCKDEwgY5jzT6o4/4cSMRMMkzXMrKxKzZRkVRUQf3jXwMdVQiFVSeKH7t+8fU2vYYKkYHvxxkTIC+k6l/38tfd++OK+05N1CTvYhsbFAZK8achHR75w2agA2tbWsmHDiuaIsPkP+hinDTXVK1dsk3+iYjkol8sAOZcyB7hwZd44GErIAHYKQMlGB8/O/PtfvP4//9dfWNHZutjWnlsLL11Lnqn5+o9eOrDn9Kw1odGMwAu5SQCaS+/XyeS6UHte1bCyxCmFNQRJo4a40la0veUojCIDyerV6fm5qWqIYlsxikLKrMR6YfbLzSwCIKSMzEATDVzYdXIsPXBy9BuPb/3243c+vGX9rn3HX3n3xInBmTm0RWWDYkPVQCNyxgoCEZBmnCuKAuf9e5f58QkQokSJErQ6WEIaqAucUTsLk3HaaVwpQyphnHJcR6GeGcriVq102kpnod5eSNuj+RWFE61FKdq4YJLQMFSrSTClPbNJcTaOJpP28Ub7VFKuxiWtlQMqUpBZmzGllgXSEGkFwFSDhCqtYjIXjBsx7EIon6/foaqErCeSL96/patsoDGaX4cKYCSDiymIAB6Yj3/24r7n33h/Yi4JozIZQy4zqudD7z6EEhjCQgKbwhpO+svR6o4SAKLCedVhEF1KfgjU/C8CYAklgIsG7aEzpBlsALaaAuQ4UKUbtmYp+SORqiVoHsBqSvsH6m8eGvzSQ1tKkW/rtZTw0rW0qcXpL1498MLeMzYqI0/K+eB3ep1EawFVsOZ1GihTaFKfE9V13W33bd2xY0N/d1uxGAaApPXG4ERl96nJvSfGp+enuKXIzKQKyAef3UmICI4VhDAotJ09N/rLF19fVc6efPDeO9eW1vW0PbBl9av7h946ePb06GTc0EKpJQitGBJkigQAEFLTclowHS4bAlhJ1VoYIgWnjJTIBHGXmmqjMFQxgUtWod7Faa1L6l1hdUXXuXUrTq7pONUfjbeHcTnQliAumVqJKwHVQYCSuKCGYiMNa1lQSbunkp6JuDA1Vx4cW3t8bsVZNRPSDi60gcpqAXEEQJkaahyBKSuBcJEtCCzIcndv19o1PdZYaJZ3YIEqRBSgIAKC3UfP/OT19199b6jRQEuh1RjO4wMJpB9z+ysRNM9YM5mSYdNZCNqjvNbwhU0uqVv5/1z0H7mVBgbKIRnSRBnEjLz0b4jmN3UjOL88umChq2Eba/jMm6fW9rY+sHX1zVmC83wqvHQtYZLM7Tt29pdvHoUpkBKL8jVsDn8ZsIojloVV90o9aQ/dl+5c8+TOHTs2r13RVrz4zTHw2Njc2wcHnnvz0PtDM3GxPYKj5nR8Ycpo9h+BWLgkzcoRNq3p6W/rgDKclovhI3es27ph5c4dXW8dGt53cGpofDppzFDBSsCxTUkNi5DSRU623Eq4nBOjClItsJqi1olSkUwIzggSQ9Jaj+qzIoGrreSJLd1jD7Yf3945trJ3vLNrpDUaDyhGBoCcbVepW0kXPpxykLRQUo7yWXoYLoKSODu0vv1Ipev4XN+h6a2npjdMV9ZWXK+NYuaqI0CFNTEuIIkci7BT+sC5ypRMqR1hCDAoUJAABsrGADRWSV/b//4zu/YeHDgHLraW2gDWhcj+T5ALkub6nJA650LSlugalHIyhqGQhUrDN0ivLsV5n6rAOUoOD89/75k3V/zRl1b1tC2ypeeWwUvXEmZwYu7Zd0/MJMYEEZxciIS/IRCUIACUSCFxvVpusV9+ZMcffOG+lV3tAIAEzfAtVrIR6Za+1i19925a0fHXv96z7+yMMJsPTIgKkBCTMlRYM03qbR3Rzvvu3r5lIwBFgjRRdp2F8pN3b3/wzm17tg2/+d7Rg6eHB6dmq5lBoWjIBs2ydedHvoIzoqCEiZEVpGZEISXHUcKIC6kkxs73bCpMbuvae0/vu/euGtrWM9hWqEKAmDQ2AqtOFMrpDACnTdHMV8YSZrAaFYayxCGDI6xsra5aOfS5WsvI1OnD45vfmrz/vekNk5XOLA0oQmYsw4QaWDHQ3KD8gD1DQKM6F1fqWioRhGEMMYgzkffPjP7qnZPPvXN0cq5RKrRG1qrKZUZl0vndkCqcJW0thFc/Uzh1ijyE4ybK1sUQIFDHlvaemv/VK0f/q6fvbm8tXNY58txsrv6C9Nwcakn2s1cP/3rPiAmLesFFcyPvO21WM2J1ScNktS898uAff/2x3oJVKMQREZAXxTAEhYuhCWzwmR3rUGz5v3/w7PBUQ9jwRQ15L6q1IKQSBrYWu/3HBu/a2LuqoxxaEJgcqUtJqWz5iftXP3zfit3vn3npneP7jo+OVjgVRZCRgYrR85Fyl3daCFBCxpmYWDU2aQu7TpFIsrk6py3R5D3Fg0+t2v/YhoFVrWdMkMFBKuQkVBepAbihVqFcaGQgqAGIAFZYEjKSV+kgMMDkWMUVZL4YuLhg5jd2HN/YNfSgHntnuGvXyQePnbt/rrqmHtmkNCpmPtA1UEs6h4usLiIKINWpkcrUtPaExCAyAkzPzu9+//Q/vPTe3oGprNAZtrUaxOpSAJdt6hBAClEShYsC6iiFi22yOKogoguLdTeb/FRaVSWmQtvuYxPbN44/fNeq0Jc3XAr4yvFLkjRzz7995HvPH3BkiQ0AhrIqgeSqHTuXCUFB6ohBWp+bfPTOtX/01YfWdZTz3oVMDDJKxhE7orxcUDPIC7Sms8yQAwNj1UZm7UKkBgggIWKo1YzYSlCoO3d2aGhgeCgotbSWO4wJA44IFkKqMVALmNb19u68d3NfZ+v81MzM3HSiLndhEudx2nna7KITZtNoZRA4cyZzVCJpC6QWyqn+8ODjm17/ziMvPr1xX4dMc1WQQLOI09CoM6ZuTWxEbBYGWQjKYEhBAgtYUsPKxoHFkBhVCw2cFiAIpUYQ4bYETjVro6ltncM71p0sBbV4Yl2tuiY1iZoJUM0gzxS7SLoABYw2Ht6xdm1/Dyica6SHhsZ++tybP/jl7jNj1UJbpwkihrA6VlxsTX0ynLfyImTMcZr0lcPP7Vi7ZV3fYtstwusHTx8bmnGwgWEjGQChALiZtV4IMKIk5MiOVZLRsZEHt69pLRVu1D3k+fR46Vp6iOrA2PSuvScGxisKYmIGEfLmGXQDM5BVCcqcZWnZuH/1tUcf3baCkEKZSAkuX21asKIIZJQCJatKUNm4qvedU5Mjk/NQNXlHjgvS5UJ1mXJiSmpDUR0erbz73vjI+GyxEHS0FQPDsCkB5AoAAy5Q2bSq+747+1MkJ0dn5mtxYEze1wMweUnfj/sYzaqIBCU1ylFaiJxxNhUNTYZCcmxL/7H/5p7Xv7N137rivGuIc+S4JIGKDcSQshMYQQGwihQcE6swMspDTsSIA7k0yjKbis3EJM4kyg1BkhE7yy5oCGeGlDIjtbZ2qm9fMdq76nSFxicnO7S2tsANokQQAMzNHGJVkAM3Ulnd17G6v3NgbOZHuw787a/eeOvwSBz1FsodDBdqZjTLT/7lJ8IxLKkqqTOm3ohXlO1T925Ys6Jrse0WIZeuDDYwdCtIlwKkFAqxUsrIrAyNj0fqtqxbUYy8O+pWx0vX0qORuZ++sO9X755JnSHKnXZK0MvsHXUNESYhkzWqd6wuf/3Ru/raWyEZsSXkNgwRiIguFAsi5OGIDLWBnazWjp8Znq0Thy0MELI8Q5ZVAQJbJZCqYQpMlDk+OTy2++Cpk+emEAVdXeWCicACl0IFpClca7Hlnq3renp7R85NjozOqW0xJrTaCCiWS5QsIlIWk6ZBBi0ESXvgIkZmuE6UOQ5NXG6pnnti22+/9eChh3tOlZMZiqEwYp3YhIjJODUxCIpQJYCSsoAVIAFAoLwEPhSkpMQgI2QVgZB1ZEBijMCETgsiLKRs0jDVlEIX93cNblp1ot3WRibtlNzhuGypAjWJ1STIVNoLWTnkqcy2nBkce2P/iV+/cWjfibFKbEyh1diIVVnz1lxA0wF4uVeHckYkUKvgLK1t6C9+YefWno7yYtstwosHh44NThKRNRYKpfyZ4aZByNPTWEDKAhJrDZysW9HT21Ey7CsN3dJ46VpipM69tv/kXz2zN0HIF258vQyH2HVBiLPG/FP3rn/87o0thcghX7vKc3sY542phRcjL/XBAMIQ+w6fHJkDh60MZ5Hmeb8AKRmFGnWsYlQZGZkktVFV7dnx2X3vHz1+csS5rLOrXCy0wARgAyFSiQw29HTdtXFlnOrA0ESSuagQODhVcwnpAjujjtWoDbPQQLOgVosqdQ7MfH+XG/3ywz//2t377gjPFrJ5JVUow1klCzKqLGqEWMEQppQpZeTBjUQgVuI8mC5X8Au7J807uRBYNUBKJPmqHEENNA2dKsIatUttffd0SzeGJjFT64BhmExhDNhwgyCCEHCNjKarrp4SUWCYLcg2+0nmJm9eafAKREI4AUAaipJzle0bOp9+5O6WqwvUEOC5A0OnhieMqCHrclv8Flj2kqZxrAwyZMbmk9Hp+Ud3rC4VrsHynuf64aVrKaHA5FzlH1/ee3QsVbIGN9A7eCmU4MDI4qd2brln46rIGlG9/DZIxVLh3aODZ8braIYFOoAUfNH2zSUahiMSMSFRIMRJI5mYmHrv1OibJyaHZuOwFHW0FgNmYoZWLc33dhTv2LCquxyeHhwZmqtrsS3QD58rBYQVGgYuCjWGnUqCai3EPJfSasc6TH3rM3//5bteWscTYbWkYtQkICDvN4yP8mniOwmgi5avchxBYa1jhisGyapSsr1QmZh2Q42VcWSL4LaGWjM6V6pUsbEk85RHaLAxnPdCpouWxOhKDwmAsCMwq82cCBr3bu7/4v3b7GV/rZckVTy399TZkSkDAlkAC2Uor2rYa8H5WH01bJRwdnSyJaBt6/p8vMatzFU9SXluMJlzL+4+8saRUQpaVW5KHbhLwMxBFDHzlRp/Lcas6e0qhzPVNOUQFx7CL4xBADSPTVPDAiUiE1DUFqe1yQodPzZzeGju7QMn7lrbee+G3ru3buhvbwEFBFnVGXz7iW193eW/2/X++0OzHym1RyDnbByktpAUmONGoVE3USxdthGsj4598/7nv7LlhZ5GBXHgTKo2ZaJmRUBckk85BX/UyWsdoAtWosvaMbizfwo75+ID4duVO+eDqGBckHSaMBEzgYyZiD84yFVfGJxP6JJlUWg621tDIojgU/nQVJWI0lTjekNUic2CuN5aEJETRwRr7a/fPf3gnRvv3dB7+c9hnhuMl64lg6geGhj9z7/aU6EWNWquUCeuEwQ41SxJVQQwxJd/qyuBVne3toQ8V48pjDQPj7xUZSDJexiCCCpCjo2GrQXVQNJ6tb5/ZubEqeE39rdsWT+yZUP/1rVd29a09RVtuRh+ZefWzp62f/sfnxmfVzQnzYXdA8KxIRu6QkqFjLqdGDtntwVnv3n3M1+8e1fXXB3zoZqo0l4PkBUTuvxIh6vBCoBMmDNjBGo1NeHszlUHG0G9fjB9e3qnmLZutLZW4yg6p+i65se0EFmj4lxbMeppbyUA6j6mZuEnkZeuN8Y0GnEjjlWVmYmuXlyvC6pKUGuDs3Ppz98+sbKz1H/VK3ye64SXriXDfC3+h5cOjjcKxZbQaGo+XELpZkFZllUq81mWAeGVTqOretvbSsHwZI0QANA8A+zD0kUACTEAVmHVvBid1bRAKSKbRUGaZicnsmMTp4qHRlb2td63sfOBDd2bVvT0dLR3llvaDE2ofiROgUgKhETtdMN0xWjR+uRG897vbH77X259uyWpox7AlKRQBzHEIt/t9YeQ+03FUSBqlMRRErnqZ9ftr2XFxp6e/dU7ZgrojrUlCeb5OokAESAiLaWW7rZ8+nbAp6vyRwBqtUocx7fCw9bHkaed5Z4DZfPq/oHPbu7qum9L4N2GtyReupYGTmTfscE3j4wGpU6VOKCUVBwi4Gb7XggApmfm4yQFoCLK5rJn+GzNyu7+7tZjZ2dU8y5Ql/44AqMgo47hCKT5ohhB1AoApiA0YWgydXHijp8aP3Vy8OU2s3FN39oVK+q1/5+992ySK8nONN9z3K+IiJRIJLRWBRQKpbsLpVprTjel7c4uaZzhfJk1W9tP+wf4A/YPzA53aVwzih3ONtlstmA3W7Crq6tQCigBDWQiIVIgdcgr3M/ZDzcSBVWNTIhEZiEeiwLKkPdG3vC411/348ff4ydrlplvVB4hNZQPQGd9MNcy3Q0Xr8HYy3t+/eUn361QNan2cdRi9cRSSQKoKGe3XtiDQJiVFJDAi2qgADG1WOOW+fLguWzbL//rkDktO22cdadLHissGgKR9763qzSwprsdtb0ripluvdFIs3ShBstdvtWyQFAtGYFzb5+6uHfr+q3r+u50SoeHQEe6VgGiGLk69/e/ODqXqI3BUFr+RPjbwQBDLUcT061amq8HSIVuX0TjVhTI1/dUdmzsfefU5ZYnGzKpu12+NBXRRVBbsAiAqiPjqbB2FVYhyUKSwFIlKAtVmi5999TVt09MWETdXV1sPN3QYxa+6RXjU2/nW5xRq/Hchgtf3HdqoDKtDSVjnYHhWpBUjBe1uZIuT1qMJ6skDDEqkKzY9pAHSOq9JWp9addvLknl4rmNVdMV2askhgrrRQSAUXiF53u7N4yAwClU2G0c6Nq4pguQW+ty/XaKuTMVqTwI5hpplml78yHkuinOioOIWcWp/uS9C2v7e/7ka5+J7EoIb3S4gc5XsgqoZ/k7Q2ND082uOIg1sSpeQ0GRvPvQHn6CEoQFRKXxOcymCiizX/QlEYQNsGd9pacnnBdkHASSWElvSXlQhrA6qArIgz2xJ0MKq86qYxEolIwgAFmokkhsov5K30B3f3d3TORuGukTCCTOZKzstSfzeMyO/t7mDw71n3dprqoWU7FLg5xhmhImysukWwCM5lY8Cyug7IVFQeUcNoA3Qaln/ss7Xj+85qOgHrCWCUKUEhyJhUQKo6wL5YDvklBATA24sEt3bSivjyLAgBedLK6AQBQeUHjVHMB4zaWZZRSmlUqkwPI16W9FP7ZmVoIyhCAECloo/d1rZ4+dH5UVPU18ROlI1ypg6PLVf379eLWZoxivFjzsp6mI/RCUDTVbSb2RAQDxoi+sXW1y47q16/u6kCdcxAzvLRVCASwkiCtEiynabQ4jUlhpteKZRphXWvT5A0cO7jgr3mRtR0hamNfSLVL6EBADpoxdF5Jo65rqt/a9uyE6OY9Bh96cKp5ZTAJuMoTFkt7Tc10kyohzAz2VjQN9AMS5u5DDIuWG2QCYmpnP8uyG7W0rGoIgNMaA3j89kmTuTsd3WG7u6RbvsAw0k+zyxGyjkUVhDBIi147EPOwugKCF+BhjGs3WTLWRA0Cw6OvSItl649q+zQNdRlqkImABL48sE7xFrWW46ssH+ide3fJOXzyR51ANlucCloQHSFuKLEMca/P5vg8Obz/Z0JqXkpeKJ+Nt4m2d4I3cXTLFxyiTkMKla7sra7u6ABSFwJbAwsGkClALuDwxk+VijLnrNbNlh6CU5PLPr5/84OxFL6vlsh8VOtK10hkZn/6n10/NNVSVRVTb8Y2HL10Aii6e2aS5m6s2kuz2deg/AVIYQNZ2l3ZvGuiNWH2q9zzrWjwENeoElXI68JW97+0qXaJUwGSYqL21d5muZBEo+S6GlXhebEsTXcfVV7ad29RzQnyLlEVDD1XKQEKgpU+QbkCgUKU839TXNdjbBeDa/uHFQ4RiAi7g2UY2MdsSBfM9XtpyoIAqVMEEZpM58+GZy9VmurKzSx45OtK1ommk+V//7Oip0boighTbYYrnh6C3mvItKx+LJ1Pu5cr4ZK3WWPjJnVFAyYjkTLRvy+DWwZ602QTbZewfmCSitPlk76XD645UuC5ZzGoDzZfvEhYNuVDIIMgNOXhmSg70n//qhg9CuhxoajQiLdJN5d7FQaFO85Bk98Y16/u7ANElddsfy5wWUeGJmfpMPRfF6tnhS0CxvZC9tT9+5/yRE5dyv5hqpR2WiY50rWjePzPy9vCUcBdpCBCKxe2PU92W0J88CIqHW8HGmNErV6fna3c64wYEUCigmzYMbNuyIc982513eVCGliJNX933T5t5WB0aETsjZiUuypPhZm5zp6FN49yYell67dQX+sd6umYjV4/yLtYuVQuC8r3uP1Nolqa9vd3bN66LrIGm+DiTYVEsRLQ9IKo0Nj7fajlg2WbU94wSKRciLBTMpfz6hxebreROp3VYPjrStUJRYLqeHD0/BnBgAPIAivg7AEAKd/KHiIIUhtVDxcbl0Zn62MwSpIuKGmNsAB3sKT+xbaAr8C1HHvZWN437xLWOkwjkyLQM7SlfObzhzcimObOajNVBlpYFvjwop4CDs1BSA2dgNd3TO/ps/6SjVobEuoDVKgmLwb2tF2aIs9zvWte9ZU0FgCoT0RK2PqtCVRXFznFVvTI+3fAG/NtKz6w0CEUZHIWCTHjs/NW3T13O8k6+xkqhI10rFdWjQ2NvjVTVI0ACpIBAi51UhZI9dOkyAssQRo4onmq5K1OpRIaRAAAgAElEQVRVt5SgihEAVoHY0JNb+3Zv6m5m3j3YvYaFjToBxsM4bj4zMLTVtBDDR1pyPsxJaSU+FApYT4FmMKlVX0lgRHorY58bvIRSlgSzRoTUCAkJfUJa5WLJEYvXA5t7NvRFAMARYJao5+oBDwvY3GfDo+MJBcR21SgXQUhQGPCLqKKWynd/c/Zqtbmk0GmHB8dKfEo7AGhm+T/++sOzY00Rbi8cf8w9dUz3lfYUkJhzz+NT8zP15h3OuBHV9sLI5nUDTx/YxS4h+DuddHdc32iqIBLpN7MHNoxbYQUzifEhaQh+QBdwjxR55V7Jk8J4IgVLunPd0KauWgAnEEEoYDHJvQ1rVH1eCXn3tg3dlfLCPy6tsha1/yRVvjpdHbkyjtW00FXEsW8gMObiZOP40ETamXitDDrStRLxou+dGnnv5KUQpBoI7MpcJbhWKpDAZILLk9Xx6fk7nHMj1P5De3sqT+7ZNhB59vlSk9kWx7W0EgXEq7Cmu0rje9ddNC5TFzEIKCvFnu+l33/QUKEiRWlRiPb3Dh3sG++SNFP1iAUmD5pgfw/jG0JW27a+Z9vm9WxMMWj6OD1oURBAhhAAmfMnLl6dmm8y3G1slVcPqjLfyH/w2kfT883V+yk+TXSkayXSyvJ/fftMHFWMd1JU+lh50tWutqTCqkqqQXThau3ixBwAXUpFFgUK76I9m/oObOtzSWPRpy4VBaSo8uihgc329lzZEF8mhWpQVIQU8IpweFgMSqRa5vrB/iu9VMtIFJaUhXOhJe1SuAmyvnVw5/rBnhKubYEnXcLaKhGIGAKgmflTo7OOYgNZzcoFBSJjJ+bSC+NzuVuZ8/JHi450rTgy5987P/6Dd68YDo06ZlFdWnLyMlK4CqoqgqgyVc/OX5kSVVrkchG1+zMVAbLB7tIrh3ZbeC8i9/kj07X5iqgovEJDm23tuRpoBkOeSYSgqVKq99DrLycKeJD12Fm52BPMOfIQCoXv0cBQFLHxj+/a0F2Ji/GJAlRkgy4OLXaGwQMy7+TM5TlnQiasaukCiMhfnmv9/Wsnr87W73RwhwfO4rqYDstIM8l/cfQyldZByCJVynSFrr6g6OQZSgzPpul47OrM9OQkAO/9IqVHAZBC0lJIT+7d1t/f6xQEGDbUTqlYWrjqE1i41QmAc5IrubX9TQASuJTY+wiUgJvQJWYkPCwIYIXjDZWJ7kpDTc6Sh86wjxca7SZuasbbbABThTg/2N+9d+u6OGBodt0IYrGzLg8U1hOCfGK2en5s1sFaswIDB0tCiXxYCUema+Oz9fs6rupwN3Ska2XhvUzPN8amq1FoCyNQXXygZtlRFLFCQNRAydix6frI+DQAQCAOCkBvXfS+xkJvxlAi1k2Dfa8++xj7zBPEENiReqhV3KM5kwIeaqCGoaJMFA7G2YbKFDy88apCasEgUroLt76HAimY1HX1BK3+yoRlr8oEhZpbn+tiJ4PCKhggUmV1Fh7ESqb9YvagHPLU/p0b+7oJRTbLNcFZrPIQoCoAmkl+4tzlyWpGNoLK6rGA+iQUoLH59AdHTs3UlpaO1OG+05GulUU9zX9ydGjo8misVYXmVFaN79FN9QFR7HoRsgoTiitJ1mXNlensvaFZDxgjpBmKyN8nqRcRARYgsuASEHSVgm9/Zvu2LsnVN0gUaaCONBaN70G6SMmBEpJQEFl1XmNFeQePbg8uKgGCsq8HaCkpKZnV0MO2faoUKsrs15Uvd0GcxrlxSrea9ytUSZnUAKwgEBlSQx7MDjYXaWVJtTZfrc7kLnnh4M7eUgQlkCWQAQi8+KInBDAE4Fo9//DchOOyEpG6VR4wBDRgz8bY9y5MnhubWQ23yaeZB7qHpsPSUMV0tfnG8RGFsGSeAoEtit7fQ8f94FAUiyHFXjOVyJp6NT81MjlebWzuiYGsWAzTj1MRb4au/U0GABN2b+x7/oktI++O5bklDgFVzgGz+FH/LRStpwQWOCKBgsn1hXOxbamCFUxesZDasFr6JC12qXmw6ytnMfm6eGGnpNAbnmsFGahFSlAH48nkQg1HmfPspqxIGPKactDVE/f0lDZtXPvEzvWBIVUluiZXS2h8ap/I49PVsxfGorCPACLWVZMA8wkoCAiYJ+abPzt27uD29V3xogvBdLjfdKRrBZHm7t3Tl0evzrMJRQkA4V6ynJePQhyIiIjHp6unh8c2P7VHKbiLcoJq+dWXDx05N3t5ChqWHWXgFoMgZqlvdR1EGoC8shOFqg+CelxpwoAciJfgFLHisApOKxUb2BRo3W4BCwBpe/gjPs/SPBfiqKu7vzKwveK2DZTXDfYP9FbWdFX6ukv9fV1ru8JrTa2qd7FIRcTNNPtoeHR6phb39RldHbfxnVBAACbYoycvXfzs1OM7N93plA4Pio50rSCma82fvXuu5bhEpGSg7Xjub1krWkEoCLBRaarpPjh/5fNP7WE2xYUvpdNSD/fE9k0v7d/yw9cv1BMTlS0oIcXiA1a3g1hY4MFOPIk3cdiqxNWVuvt4CagKSOJQIpswMgC3NhRBPNtEQ5cnZcp2ri/v2rJu29aNG9YObBuItgz29pei6w73UA81RHQvyQgXJ6tvnR7zHJXgjfp7LX+5Eii2HKhaG8416r86enrvtvXBUm1GOtwnOtK1UhDVUxcnhydqYRi1B7oKUiKsgh0xC32cmiiqpXJieOzSfG1Hb/e1sXsxK/st73DtOIBi0Ndf2Hvi7MVjV1qMrhCG9DYGU9fPBn57J0ttKwoHOCWrUmapRVTHPczjVgK6MKqxcJZyowKxgoAXHEkK+SHAe2kJ+rujl/dt+uKhHU8/tqXv5mCXQHJoBlJwqfiy7mK+heKXAufG5k9crJlSd6COVQSsd/VuK4niHrQEm/jgZ8eGv/zZJ/ZuGbzDSR0eDCtx/f/RZHym9v23zk01FGSKHpyUilQIuofB73JCUJDNKRqdmj/64Yk7HX5bCAic5I9t33D4mc3dvWjkmUjEGt6qMqIiKgW3fa/rIRCRV/LQkKSLNGDJ2z9ZnWgxtik0Bo7FGQVpqBrcZEeixD5tbi2n//7ze/73P/7SF57a1RcTtAFpQv1CHMx7QEwsVLppqewuqNabJ4bGJ5owYcWqM+pXph3MUii2MHqoQpltlCH8cHi0U4LyYdGRrhWBAqPT1dHpehjHqrhWi4sKPVg9KNQEQdPJB6eGEi0inku7fgKgAmQvPr9/+9bePM9UAggtzNtYwVrUZxZVD1G69tLCDgPttavrOksCBBBqH2NYArNqRatNcflKULAyF8+yEvQGl1sipE4Dzb717Jb/9OVDPSSQmiIThJ5iL5Q7baRudHJ+fGouV87J+ruWGW2PIc4OX/no9AiZkMhcq4yzekcJ10MAKYFM5nF+dHq20VqJRXIeAe51eNXhvtBK3fhMI018zMKqAKmqEgRYRc5vCg7RCKzUXPmty/TRpZlntnQbeKFgkf4aVPi6cwjovsH13z60Y+LS0SuNek93GOYNp72gblVJ0lnIPAOQSMAgBhkhY22o1oARkgs1YfGAcQgcBXmQVDJTafbP94612G9osZMYAHmrhhwDUKPOiC6EaVc0BACkpEQOAOV9ue9LrLM2sc6oWlB2TSpUxIYxoq4mbAw4Hzea2VzSmK42Zycnz025oxdrJqv94eEdX1u3XnxKHC56kiSq6skQYCSFTxCUWghfH2uenan2RUGomlG4EEhYLTfybVHRAEWlHrRU0Ur0Z0fHnj8w9fkntn4qRHmV0ZGuFcH49PyP3jwz10gDY296yFfZ4y5QZVU738iPvHviyU2HjYGqKH1ifvxNMAAQhIjxyvOPH7swdfn1k+IGVCMmlyaTofFffHb35sHHWDIvyJwmuW9kea3WmpqZmZqrNpuCsIQwzMmAGCoEgRYRH6Mg4lw0rKcBCCAvxAAIxWbkRV3kQ+fa4mFx6Y3M5D4AofCyuOH+UUTG5OJ++ObJ48NX1nYFkqYz89WZljQ8NG3OZ8Fcpl96asdzT+wDFHTXXs8CMkAwMtk6NXw1SfM1PTFUFuZbq+xGvh3XmkUJIGIvfO7ixCsHt9yj+VaHu6AjXSuCiZnq+GyDmFdHMuEnogqjakK2rMnR4+fGP3do60CFSABZZIogKdqm6KIDpfCLn9l39vL46eF527vBZVNWq1vXRN86vOHQ9m0o3AhVRciJZl4aSWO2Wj8xNP7rY0PnJxOJ+og40NRKCkBZFQ6ejFBGNO+7HQXW5OBiy7cqMSkv6MJKpx0eBUCYT4NEmIlJWUkWdG3hSBKy0dWUJkYakbQMxAkyisSEhkuN1vzmPvvK45sHKpF4l5ENFvlVLUALs0CYEKD3Tg5duDgaBzGr0bZoFYHEVdCqi4eA3Ll/fO2jLz+/e/u6fsMd9VpWOtL18Gmm+b+8dXaqnltjoffi+f3QUSULDQIAoItT9dc/Gvr2S090BUZ10R+LCulSUg/wc7s3fePw/osjR1qt/Pmntv/OFx7bUjb7Nq4N+ba7QXuwGS8e2EtBNPwvxxpeoiCwCkZuNBCocG68hTep4mo2UPOlPpPrQtEw0mJJZglWsw8XBkgFMJOtvpZaA5AaVxRFKSbvxTIMgYnZWKeci3ooW8NsyUYJQuXm4cc2f+HpvYCiKKfy23/rLbQX3dQQ24l66+jJkZnZZrmrR4kWFr8+hRARiGbq7r0zV7as7TW8JLnvcK90pOvhc/bSxHvDMzmMBQi6WmJWt6WIDlnNFdqQ6Kdvn33u4O59a3uW1h8WI3VSFV8Jgs89c+DU8PyPXj+trhQEcc9Af8PBI2WwV3KiWe4bzdbsfHN2vpG7xPt85PKkNUHQ7jupyNpUiKPcILKEDDLW6p9KervjatHq7b6egLuNly0/pACkqZWJ1mBLAwMhNcKwNysGkUogaQglS6QEUqO5epppNves63r1qT19XSWoAxu7xClXMUklENgA9MaJkTOXp0GBIaPtcdhqac4lY5htXPnwwvQ3PutDu8Rm63BvdKTrISOiQ1emE0QIFNKiVWNDdHsUAHsjKaAIez66WP/gzJXtveWomIktAgVARMRQ9iKssqWv8q0vHjp15cqJk2dq863NGwf7y3ElIGvYC6WZr6dJrdGcmm3N1ZpZnmUuFRi2pQp7FQFYNIpEMiPCjiQKxCesM653dj7eEUZKMJITRFfdwEEVhFqWTze6EkRdnKC91evmSS6RGBUSFUbh22VV87wVZq3PHTr0zGObAacgAVn4BcP+RUGAghRg4nor+fWxcxNVF8VdEFynW6v6pv5ECEAQv3l2+srk/N7NA8yLykXqcF/oSNfDRFXrrfSDoYnUYZE5eCscJSFSghcQmTBT929Hzj67Y8OuTWtu7L8+sWcUgNpdngG8Vx8QHdw28Ptfe/r/+bufHj87+dFQyxg1lDCLwjivuYiClQ1bA4qsKccWJXgrDsqeAsAociEFCfkin4Ba0jNe7cO6cXin0CKVk1a8dt3QiGLAMpnG02mXV7IwpADJrR+iiBsKF5sVCIAS1Zu1Q1t6vvLUjr5KGdISKgnUqgOZRfYMRZyVCp8n4t8cv3B6ZNIJl20AcQvTZ3xapQuAgsaq+ZGPLmxZ21O5wZSkw4NlUTdohweEFz15cepnx0Zh4sgCUCFezY85MRQKR4FCrWY9sXnn3MQb56c3re+PjahqW5aIbpkYtOGP+zljmAH1kF7D33x21/jFZ/75jQtN7rLGhr4B5B4mYK4QKUiJVYt5CBuIQLlYvVJRQs4A2AiIHFitssvKJ2r7v0SnymmWlZUQBikZODEr8Qto744qPikAKBTkAzLudOOJSVkTajP2sUcAniO5+blWsBABarTY36aZJ0PmOy88dnDrWgCqwUKRSgNdbMacBwFqpQUKpjL7w7eH5lpSDpk0W4mNeL9RwEraHdLJS1Opk8qdju9wH/k0jPRXL6r69ulRG1ZikkBTrLZ41a2QKikJrMJY+Ih8ovbnR4fOXZkBjOrC/lR8oo04oR2uImr7+RowkK8JzP/wjRef3Lsm0nn1NTLGUBiArYJFjUjgXSQu1jzShNWpwoOEWEkB50mhREqAE/IWUB+cmd452VxDEFKQt+RLUANeoakFSrqQgKqAMJTJpRKcurK75sohZ0ZI1QDu1kEBKRhCgMAIGCSuNf/845teenp3FEUiArYEMAhksZQAgEABcSq/+Wjk/QszmdiADbVNOq69PrVYSIx8ZKo+NlvvOGssJ0u4RzvcX1S11spm52tEogwFf1q+joUOi0iJu0vBqVNn3j1+JhHPbJR46WkQ5IQzseu74z/7nZf3D1CUNxI1KYXX658WhhkLGeM3vcN1BxJAxCqg6Wrf6al9eWwjjzBXIM5t2fEn2K8/VIRIAYIyvBGxAgOVUCbSgYnR7SyWjHWcKyckdiEZ/WNYfeidEWQUpwhT5wfK+R99+eDGwT4Aqqp3lQtoRKwXcKna0F//5lieZRxEnoPbfQWfXgyfHZ058tFwvZXc6dAO941H6Q5bYTgvF8anf/LukJNiM9OnbXSqChEJrMmdvvnR8MnhMQBSLLQsJRtFASUrYKvuya1r/+TrL2zq4Vaz7sHM7fjiUo1iCUKguh/4cHp71QZQQBLhzFnSFSldClIwKbGAgaJ0dh7Q8dmdM9XHVCwpeSPKOevN1WEERIAVz1Bh68DeNb7x6tNPbxuwgIi0C1cuHSJhS6nj14+PnBi+ElhrbOTbm5pXXiM+CFQgLgjikxcn03zVFyJYRXSk66Ehqh+cH4WNiYjgi8Wf1R4wvBElIqiUu3uOD8/94ujIfA5TpAgssV9jgmFmAlN2+Ml93/niZ9ZVJG/OEHGR3rIYB96bYNIM5sOr6y76QccMVm/roIYRyMp7LIQKCwdAC6dAFraz0v3e6IGq62EDsPFEIGXhm6SXiz3eRWO5hPLq/u2D33jlmd4KA56ZmfnupAtEQDDazH/45smGWoJCHaBFuPdOJ38KIIIwhMP4ai2rNdOl3dYd7oGV94w+MiSZGx6b4SDUBcs8BW6furC6ERNEdSofOTF67NQFAEUHLIueeNHCAhgUEB8G/K3DB/7g1b29Np2v1hUsXnjJPS8RKbOMVLe8d3lfghAhvMnZe5vZFTjtItXCZQuAkslsnFD3xdkN71/Z3wpztg0Fewqgxix4N99wNtk8iHLklE5tqvg//uarWwZ62yYnd40CMI3M/fzY+fdG5jQsB8iD9kLjYr/c1Q5BCQK2J0cb7w+NZbm70xkd7g8d6Xo4eJErk3NvnhpVmIUR76fzYSfAKwXl/pGJ+V+8cXRqrgaw16WtaHORQQADBOp9dzn6o889/dUXD6lqvdE01l4zzF20gimrhEoN3fDuuUPVpAvMIEuuG1nPte9iSRf5YNH2dBUMD6smbqD0wciGq8muLJoHmtLOjin8iLg9fScBCSH3Js8MNdK02/rf/fwzL+3fbKEi14yarr3uzMeHkgIYHp3+l9ePt0zZc8jwgebmrpbNVisKAEqUwLx7+nKt0brTCR3uDx3pejjkzn8wPDnXUKOOyTtiIcPFvqLVzXWdoBa2QkSCsiFl++b50e+9dyIBCGSyFlRu6DJVof6TJmMEgBhsQQbAQE/3//SVZ7/9yq5Q51tZw3PoYBlEoguRtZteN6AAA6EYX/LvTT99bOSzPiuFJIIgM2WvoSP1Rotk9LtyR7q/KMMYtc6iFRKTZ5ecnN/4oytf0QAMWBcFXlkzJe9JnAFgWMkAIM6NpFbrSRqCXn1617dferzEnpXAIbEHUmgGdaoqCv9JCqYCdQAU8AIvAtB0s/X9d4eGLl/tDWFEHWLViHQpGYqrGxVY0SDweSUww5ONmUZ6p1M63B8emXtsheFET1+aCYOSUQ91HiQwAB52F3lfuGH8rjAMCbRVLkdXW/yjt06/OzLOFLCk0OuXtRc59ida8CzYNlD5D9985iuH96hrttKWkJEi/x64o3QBUCLPVFafgH409MrQ/C7D3pi62BRKgEJBGkItrYDcT2ZhLcF1kwtN5Gaz8mtnDg+1DhrbYI2Mj60KIwe8Qlk9cVOIvHTnWZS3elszvZpUnj64+w++8pmBCgFV4iZRC6pQAzEQAyXAASnwyekGAhIlVQIE8tbpiz9+56wxXNLUqnqEgvDhN9YyosSKIBAXk16aSk5dnO2kyC8Pj9BNtnJQYGqucXJ4LLCm7UfQ/udPK17QYvIl7rp6JfnBj9+ZbLQQRERKcMULECV4MouvlaUq29f0/fG/+/wXn9tX0obPGo7Ym6L2h9yohTe3LQGOkAQ80LAbaeqt2pYfjr8wl/YHphXIrKU0UBhn2IekgX/4a18EdkIcNAfiJEgNvzH/1DvnD6+VRvHjhULaBJAVW3JEPJUbl/moO5440HXhcR46vKPxu1/ZsH/TWiAAYsBDE6/wSgKvcESekRtN6XYGxELsyAJCmhtyhnloYvanR07OTs2U47KoafuR3NPq2SqlPV6qtdJ3z1xspdmdju9wH+i4aTwE0twdO3/l3JXZnr7+T7NgtSEigDJVRGGfT8L3T4z+6PUPf/eLT/caVjhACaRgD+tBi7d/VXWksmdN93/+g89Hll97fzjxDBPQws7nO0EkZEw98FtagfnB5Sd2rjn1rc1HLPLigkgYamDgQQx5yClznjw3LJVg/cnpzf8w/NxF17sZWcPeEKFSALBQEpuLY+trz+5848tbjgaNtLx+/74109qYVL8HQRfZiGzFkyhSQsoAISYNyBMZc9PeOwUc4AELDyTgOMn9T45eOHL8SldcMRQQ8UKgV6/789OPglDcLaRRFI1N12Zrza6OI9SDpyNdD4FGkl2Zadgw5kciEauwFCYCiYKDMPHRP/3qvc27Nr24a1MEA+SGi2Hr0lqDiQEPuG298f/yR1+IQvvzt062skBsbNgQkep1861b3DuMqJV03jLlA2t44nKz9/vDX9wV1Q6u/ch7NRAiARwgRSbDQxGuj28QCbyq7R1v1Nb8+vTnj47vtn1T2uy56XgChFxqshwx3ECcZrvLtec3DHehqY3h5tF/a4Ub0X1Qup803c+E3TvYblTqUfSJMAFqHExiiRjhrZ+YAagHkWj41pmLrx0dSXzc31OGiHhcF5hd0tf4aaBYDQ1MMDHbGp2a37qu/05ndLhXzJ//+Z/f6ZgO95laK/3vr5+ZnM/MIxOvNWqgRknJKkgbaTZ6tf7E3q19lViVWIWICIbAt2Z2fxJKBDJFkKoc2AP7t8OnFy+PNjMyQUjqiXBdN3pLR0zC5Fo0oKZZQobATM6uN0J7N1ythCaX1JIHPKCeFQA/zA5ZATXaXy/hB+df+MeTX51DTyVwhrzcPEclkFfT8lSifKNtcbeb3DfY3QeVxlxICbtpSU9lzV/VGv+YNP+7NM6atBZ4tVIxGhkiaMJkiOz1LUYAK1g9SIlL443s73585NjZiajcu1CK7ZHVLQaxgQdI2MxU6xHLi0/s6FSefNB0pOshMDnf/N6Rc80MprDc/pSjUMsaAqScKmdqSKgyOtGy0ty3fbAnitQXW1iZaQlPvAM5EENJU5I0MuaJXVup0nPy0uz0XDOKzI1hw5uHCUJILQcuMmYyi+qVZC2lfaNOIp59Zs1YbsGaMQQEpSIIuUxfFgsIUAMU0z0oQIaUg+y9s8/81bmvncz3DJi4LNIMEuNvribDICuhopxgPrNXPve5l5986c9KfXvSbCrNpjw7a7QMxKmL51KeO+HnfpLUfpqkRzwuwpQZGzko3fpJSR1JSiYAzPd+/e4v3j7byIIgLDFwnW7hUZMuBYOY4QB4E+Qim/vjz+7fWooWW+Wnw93Rka7lJnP+tQ8v/PjI6ciGS4yQrVKuLRJpsaxFMFDDZM5dHO/p69m7eTC0Cp8K25QYS8kdIoDbJlCGiAIT7N7Qv22w++r42Mj4pA+7KIgIjuGUQKRMQvAMZSigQjZQB4iAjYuMyasSjdW6jak+vr5qEwY4DZ2jwPge4yMu6imLAcgHhcIsOqvkt1N435MqodalAoR1plyVOTOhGpd0bXh7fMc/vv/C+9UDvhJVVMI0MBIZJABpYV5BauE8OKFI0lrJj/3p7z3/rRefHege4Hh30P+cjwfTdJyyydDDOrb5Gosui4T9OGXnpP6WTv2MZ38lekWiijHri0tTNEUbKimpkin/Zmjiv/306Nh00hWVDUSoWFl8FO7k20PtYDQREUTjIHzusc0D3fGdzutwT3Ska7mZrSWvfTB8fGg8CsNHQrkK2hZXXNi3M2BY5zM/Ojq+qb+8a9MgiL2Kh2VabJoGt0f7hX4ZIgYotLxtsGf75rXzSX7+ylUvvhQGgBeyRGAUVhPERX+vBFKFIbFKHtYDNJ92jVUHAvgtW86HSLjZDWEJMw1ycE7wpEwIiATtGi73BbqWxkh5iYXA3lEZ2mWND0L/ztie777/5KnpJ53dYK1jNK231gcwLSVWYgco1DBUpdZq9UbZn3z9hd976akNXV1GVSjw4UZb2Rd17SPqT5J6ppNaapooZ3bGk81MkCU2G0d+BjPv+LnXWvUPhLMg2KDcnYA8U8jhZEv+r3/49fHheTZxxAyI3BCVfeQgoF2/R0GqrCDmZ/Zs2DrYfadTO9wTHelabuYayc/fPXd5qm7MoxlSaPf1RKAgnJmZnZmZXb++f/NAPyBGM0NMi5530cd/f7zWYpg3DfTs3LwuNJgYG5ur1iiMDUcAA0ZhFEwgQjFfIIBISUlBakGE8lTWdzZRMc09FY6lZjVVm/tACs8faMQakndU1Ku6X+pFVPwXNCPypSSKxSIOG2zM0eHnvvvBZ05MPu2wK+AyoQGuK7EQK3nPxpFRtkSU5HnSrG1dE/yP3zj8nZefXFOOF2KO6oiUu024z3Q9jZ4tWanVpBnTqnMmIGVjmQOwV/K2lpjaZZcdTRvv5clFNrGJB0EVpe+lHS8AACAASURBVOC//eTIL945m2pswjJU8cnFax412it+TI0kCwN+8fEtSzcn67AEOtK13MzUWz966+xcSxffQX8qISImwASXpuZn6snWTYMberpYm6SkMHfpBgsAcKJMura7tH/7+t6ueHq+NjnblFytCYhDgSkSPBakC0Axdyt+I1siH6RXXWl09JCmAwN9V3u754wLuFViUYIlCUBKyMCKtt3GXVP8/vavLsxHhEQ1Mow4mk9c8MsLn/nL4995b/4JZ9ZZrhhJDVWVc8eBYwNAYGCMMdxq1dNWY//2df/xm89+8/ChnrCwkFcQE4HgRbwj48PeoLLHdj9BdMjKQK7UpFpOLaWcVI0P1cQusMQZpVP53NvIZ0v9z5DZ+KsTl/72+6/N5YbCbiWzEC19dKdcN1EY8KeZc3n6led2xWEnf/sB0pGuZUWBo+fGfvTOeeGgHWt5hCEIBWFO0aWJ2Xqt8djWgd6ubgCqrKp3p17tflpzQlqOSgd2bFw7uMa7vD4312rUFYCxADMpIIqbpkxFzE5CSJePqq21H871z0ilt5wMBomVjDVnUiYRzr31emNJ67u6YG3nYYABVmERI3EWhIkRP1pb/73hV//q9Nc+au6TUshBxtQ0VGNOhQKvZSn2CDMTfNacD6Tx3P6tf/KtF7/69O6AISpFGAsAKViFyCk7gWeUQ7M5ip4zPU/7np0+HnCs8N7kOWVZHmQuhuEgUs+uYioHafBrJ6dKf/Hd1y5OVinuURNCi4XDR3qV63qKb5+IvEhPKTj8+Ob+7vJd3BAdFklHupaVJHPfe+PUG6cn4jiG93fV2X2KYFWFjULn3PjYaK3Z2rNja08c07XJyF21TxEBhIqKYzI7Bvse37ulK0C9OjPfaKReYdhYAlS1mDIp2jMnAQCNQhfH2vSViUnqGprYMFnticNqb99cWFI2OeBJJWUrIHNdlRqijy/8thd2K9RO1CFVJrWsIBIKTBOlUxP7/ubE7//1yBdGqHed9WW0iBpELSEniIASKRMBTN5nLqn2BvmXntn9p//upc/u3gDghgSSYqFRiWGsGgvDxPCqRlzcYyv7yn2fCUvPINyGoCszKg5G6xaOU3D8ZLDjfxvRZ//qpx++dex8VOmGiVSV4QzcQpysQ5viDogCc3D7wPb1nd1dD5COdC0rk7O1X7x/YWiyZa21neEqhKCsPgwoz7ORiemZDId2rC+Htuh2i4HsHd7jRggAFGCF9criHcH3xMGTe7Zu2TTYdPlUtZHkDlykapDqDfLIMCByJm+GXnx/j8ah0aH5NSdmBmZlTVyO+2zTirCqEkOJILTQfReuIYVktDP+rr20CFDe+mkUABGTEpMwsTCutHb8/MJLf/3R1/9t7DMuLvXbq2XnAxdYIWFyFDt0k9oQLUaaE2veWtcTfOdzT/3p77y0faAb2gCKYo/XClKrEoQMFZvZlIngOM+p5TWHRqQ9Nt4R9B2mtU/77qfZbHaeWRJ2qY3jLN71xtnk7378EcIea0JAod6o53aRuUc67n0rRGCi/dvW7N86eFdDrw6LoiNdy8rUfPNf3jw71SRjmdU/6jc26ULPrkEQZmIujE4nSb575/pyYFlzqFPAEyvAH5vjLcgTblUCLGQlkBCBmImZPCQH0caB/mef3NXXU5mbbszNz6eSqOkWeJLQszgm9t1Glc1MarXFg9b3VTJvkeXl0lW/5+z4ltGJsvq+UhSEkQuDFsFD2/74IIIaglUxpAJoO5amgHLb4hbKpEVmpBaOgQwlBApAczaTsvX4xK6///Dw98+/ejHfVurKBly91KyksSqJFSsU5BwLAlJvNBPftJof3Ln+P/7u57794v6e0MI3VJUoAhVCSVhYjxJiEJgEpCD2ZARRiDAkVpAAChZea+P9ds2r2v9Nz7tANZYP5+Z/dXJk5Mx4X8bbDDkVKRI1ua3ZHem6ASJi0OaB8lN7NhruNM6DoiNdy8rY9Pwv3r1UzZgCZZWH7ur6sOHrXtZyAOFTF6eqqT+wfV0lAqSlqjmsQm27HAd/LF10W+kqZASMIo+QQBYcAgCymHFg8+Azj28Og2j8ajo/aw1NWhdpiDxwpGGAjCkl3x37xFA9t84ZDkRK1sOWJ5uDJ6+uOz/dXXXWxuWeUGK0jFPWgAgMVVgx3DI+t5obFC/HlDNnBs56YSiMEimJQmHIRhDD440t71z+4t8ef/lvR1441TgA0xXbzKjzTD70LMwkws5zkRmSq2tlaau/HPz+K4/92XdeenLLmoA8VIQj4YjaE8CFBiEmYgNiEIiLvw0QFLvhiJhQvAzAQIJEgr6w56lo45fGqn2/fN9+cGX/ZGt7TjHDF7sLANL24R1uhCjN/MRE7Ssv7C6H9u6C3h3uSCcHZlk5cvzCxfGq6eqFeFK5bdf7qKIAjDHq3c9++bZp1f70Oy9v6uljTUNNgXaZro+tCG+vW5+EByWqkcA6SgbX0P/81adfeeLQO2+9dW7ke/X6luHpzXMlg7iVO7HeGnjAQ3UhniiaeQvjUW76jSeny5dm1v3yXPJE7/iT6y/u2Dg00H+52zYph0nYZBUt5XR9LU0RQEAwGrGEhAAsMHWoazR6T8zsPTqx54PpHRcaW2Z9WSgswYVoqC6kWCgrQcQIBQBT3kwbc90RHX72sW98/rmDW/rKUQhxUAEXDhxLaZjbQVqOs5aN7HS1//9+7wv/9Nb2btMbBZmxzXt+70cCIkpyNz1fX9MVm450PRg60rV85N6Pz9YyQUTtikcdbkEDY73EPz5ypu7pP/zO4b1ry4yG+kRtd3sXVqEJS8uHMJDYs3qISGQ5LkUTB/rf3LT3B3PbfnZ17uA//duaNybDaiSlcAPZEvG8KBeOG6RCZA0nhhrCNtcgQ7kh+642w3PZ5I8bO9aPPb+zku7pHt/WdXpdz/BgnK5xAYjBHqRoWwAT2EAoB2YcX57ddHV259Ts5uHa+nOtwVHtmwtiMc0+cn3NPLNIgva27MJnCETiOWvlSXO2O/SvPLH1yy8/8fTurWt7KgFBvYNmYAZsIeh3lztRZHWqaujJhCyY/O7r733/o8Z0d0WDsR5JkK8FPYI1TZaIKhF7pun5+u6NA4+OT+ky05Gu5aPeSmvNHMYStFjFuX9eDJ8iVBGUWhr+69ELrTT/s68/88T2dWTzVJVUWNWwIRCWsp+qbSEvxOCQoa2RxuR/cTN/v6Z5dSCqrhuU9X/4vx4Y3/nLMyfODTVm5nLblcRxxRhWMBMEbCRkzRQ5cWKRMwUEn0u52thar+ulKbxjdsXhoXKc9kXVbWaoFGsp1IDVAOS9eM5cdyPtnk56p3xl0kfTyOs+dukmSIhwMqTMSo/1cWLzxOapLaKoqgrv8laSadIcKAXPHNr28rO7PvvYpq2DfSEzkEOZSEDmuvjg3aOqCjVEoPjv3pj+7m9mMK87KnBeHEL7yAe4FwuRV56pNnkJlpwdlkZHupaP6Wp9vuXIBqxKVHi5moXpQ4drEJSiuNzKkt98MJw3a//+q8+/eHBHRGgpVJWLtSwVaOHotAhIBYnRLlK4+XfSsf+az3zXZlNICQmHpYld+6vrD+596lDvh8Mjx040zoxenZqeaglMVDZxmWwo6DViSJMALUZmxFvUvfZ5HRA0BHPVzM62NmOmT0HvlA4Z6w05a8Qowat4m2iPE/beK9rRTrJZYKtlljArkzeZzepRXg8o8tY68j7zacPnSUSyaaD/sS27Pntg26GdG3ZuXlMiAArk2k5nbMcJdSFaeBf9paoWBVaIGMb//MTpv/nl8OhMvM1m5SSr8sZWQOC5zuLWYlCCh601Wnf1VXRYFB3pWj6qjVY1ydgEgGd4hemo1q0QsQFI854oTDV+8/TEbPLWdL35tae2luLuvN0VyEJHvTiUQBWiRjrzk+Ty/0kzPy9rFlIMw4I8c3k+81rXtq88tX3rE9v7Xj6Qnb00fWro8pmLE6OzrcnafCMH2bkwiEIOmCxQ8px7ShWZog51hkJrLVRV50VNRn0qPvVpQp7IQAKnxjPYZlGQhXCBmCjroTxw2vJMXvodc0511bpJszy1LkFIfm2Ft27p3b9944E92x7bsm77ul4DAJlCFxytrk08F/6nCETfxYIXAYrCdeOdS0N/86Mj1ZFgIOpOy1kinnIKxEhHthYHgXOYN46f+73PPVmOOmUnHwgd6Vo+0lyaiRgmgqga1UXazD5aKIGgBqpeojCCCU5erP7FP/xmZnr/l154estAF8FBHBRgC5JbphmKdv4LA6KUA4bUcjafzv9zc/T/MPPHygrLJQFrmMEFoDxvHs+bk9asN1zeub575/qBLzy379zFsfOXp4YuTY5cHr9UnZpNZ7I0aElJEWsQeBsGyEI0mUpOuwSsnItpMVqVPAepGhEIsYUl0Zw0gYKkBIoFSKOmB3Jdk4Jz2/ReTUvKYjYQRxGt2WJ3bF732NYtj21fv3f75ootPmCm3kNBbITZgQE11312AhacrZaiWwoABCFWwA6PTvzF994/PeR7qZsMpg27gMpUjZyVzpRrcRAg8MfOXZ2ttUph2EkyfBB0pGuZUMXI1UYj0YBSAhSRAoTOovfNXHPTV6iKRky2VB6t5v/lR0PDk+kffuHAvs19kWXkUBUYEDGui5VBPTQHrCopPCgHCbJ5P/b3rdG/CPJjZQaj7JR82GBQkJciD1jWEGIAgaEMCEPC49s3Pr59I4DJ8dFjQxPHr1RHr85fnWvMNH0908QTXJiLNcyKRFWIwIYJUYvFe2kXbBQoPKkatepV4YRypwJSKAkANhXK41D6euxAd+/6gYHHBvue2GN37lnTxYUvRg5RByIm5oiEoVCowgPKsHR9XsYiI6gLCISVoQoSJR2enPvLH759/OQcR5XEeNUsTkkpIqijfGmK+ChDnkgyVKaq6fp+tabTbvefjnQtE17kN8cvV5tpKbQL49wOt2dBvKhYgjFE3V1l8e5f3/jw0uULf/j1lw8f2NlXjridgE43BMhIAStqvVFSY6SL8iutqb9Mr/y/JKeiGJwTlMCkCqOKoO7zIIg+Y4MBBW7OoFMF0eCGTV/dsOmrwHzqR69On70yM3RxfHJ2frqWzraklbbSPPOi3kOFQMZ5IyLKZNioSLEZzasCsIwAiI0GxpZC7omxpmL6+tZu3zCwe/umzRt6+/u6egAgAZoiCcGApz2XFBUAgIAFRTl54FrA8O7upUL/lEHCSvbK9Nz/97O3fvnuSFjuV1aBWxgXaOeGXRIKEJSjcHKu5rYM2E6W4QOgI13LRJp755yCiPj/Z+/NgyzJrvO+7zv3Zr6t9u6q3rfBYIAZzHBmsIMgiMUGSIkmaJA0ZdKyJNKOoBn0HwqHQxEKR1iyQnZYIVsKgaRtwJQoQRBpSoDEFdxEggRJABwSIIBZMEsvVV1VXfv26q2Zec/xH5mvumemp7urq+t11cz7RUXX0pkvM29m3u+ec885lxYMuosIuTcq0itGIMicACPHnllsXPrsH3zf+x/74Q8/fnZi2JF6PUsZAKCRIQQJGZwnGGaSa5/qzv3zEpfjGGKA5eWRjCCCqddgj1WqHxOZCL049h0MMFU1oyZOdLQUjZ6ZevjMFN771naWbbSTlUaysVVf29pc3WqsbtTXNurNVpqmUdDcMjJVdc75iCKhVimPDQ9NDNWOTlTHR4bHR4enhqvjw3G1HPfmQwK0kzIxiEtHXOoQWYgmMiohAhpTIgGCwAMl27MHz6AKOpF2N/mTrz/zm1/8GmpvVVFB0ttgwK4RMyWF2Gi0gw48K/vCQLr6RKubdNJgjsbcvTXoE25PEfMGEKpgxiprw53Q+rUvPT0zO/+DH37nex59oBwJLM/5dQaCllli5mJ6JivNhU+n135mWDvOASmgAiuBJr5LMajrWggj70TtUUhJLAPlRvHKl7Gg5eZHgGWgAaLGsvcnhv3xoSqOj5qdNkMKSwypapoiZMhClmdKOREfi4/gyRIkIiTXTgKQ/GCGQChMISZQqHdpQLIGTVx0jFq2fCRPp4wEESyCSf4BeyOXa3gf1YZHfVzrqncuu91eA24JIQSBta1W5AdT2vvCQLr6xOZ2q5MGo5jlFYMGFtedYkVAIZwZYGTsSmNfv7y+vPKF+sa7P/6R96qZQEEx0NiGmkdFkvnm4iebs5+pJQ3ngVCCeqM3GNExZGYw+jbP1M58zA1NgKC5V5pdAAASysgQ9X5DXiFQLFB7SQ6kp1QIwCHOBybXlxIt5Km4FIVlsAATZSmDJyBwhaLBFCUvodv5ra25n3OuUR37RGXsv0blKETNCMR2/TSwNwgITQBETp586Pz73/nwF74yX45qvf8dcJeQRnHN9mCCcL/Y65htwB2ytd1MAiAuXxNeBk/0bjBQiJJ2h7PtYe14ss7aahJ3AwTolUyiAMFK4mqSLbWv/Ww2+/8OpYsVF2Xmg8uCDyqZ+nYaJwZIF6YhHv04Rt/VFh/UAO2FLL4auV4WERCYmBrE6EGXK5kxlyWD5SV5e8WF8ysAYLkXURTeUDJEgDiY5FsaaaSJQhRekmbc/ma09ZXs6ifbL/7dZP23DNukAaqg0lRSMLnZqd45lhcvhBotOz5e+d4PPHb6CEO3DTPmVaiKFVMw8B3uAjMNAcAXv3ml1U1vt/WAu2EgXX2i1U3SNPOgwAymvF6Nb8BtMYqBeTFbowYwZHbhwtlHHj6dB+CpSYJOFyrmXLLduvbZzvy/rCQbFQoFKt7A4EIatwNVkhHp1ALQGn6ydPoHLZpSC0Ae4nHzDpo72pWrpOU6F1SQiQSBUc0YlAmRkWrUwKAa0FZrm9KMRlE6owMjlVglJuHyKuzXzXATdKB08mDJP1RqW6Uzq/Vf3rzyT1pzn0ZymXCKPD+gG5AWtlzvy6BWLKN1ewgKXFFLV9XTHjlz9D//7rd4dBKTgJiAt8ybIhfpO/vYAegVOV5Z324nA+naFwbS1SeyoJZZDHoFgAyDuPhdQQOV1vG+6csdumGv73hw8s0np4BMVBJzbWiA+LSRzn6+O/cvfbYYexIlg6OkLhMzqJgE+lZkHelWJu38T8jYYx5Sggjz2hy86Z25wX7q/UZHpGKpAYqMliBFUKYwoCjEbpAMDbU686r3IHtfUljezP944zEcUjFg+DzHnlTACwRJ3PhSmP5k49I/STef8gDhgsUBvhBRs8KpCgXCq8//tSAKJ4DRZSpDkf++9z3y2JuOhqzbDQbSa+YsM0gYTC7cMQQ8KBBTa3cG0rUvDKSrH6w323/+3JUkU4HRFANv4W7J/W1gBsnEd7qds0erj56ZKEsMjYwIhEO1ZI1k5d+2Fn/Gpc+Vo1wUvJk3BjIINEoqUebU1ZOycPwTtclPABVv4uFvuCO3HFQUAlMIGWECFSihNDhTb4tZ40+7jS9o50VnqaAClkGjXNe/G4Twhu87B9A4MGV1TEbfm5VG1eAkHhJX7i50Fz+zffV/Sdf/Q2Qtj5haMglgCklzxx5NaLvXGMIoAQLD0ZHRH/5P3n5qKLV0O4VkzhtBKBEGj+0dQqAo90hptDp6wyoCA+4VA+nqB9+eW/3zb8+2sixQDUqAEA48hruGMJOQotN424WpB08eAxDgMw8nqCm4/juda5+U7OvlWD0AFcCMwcjggzNEqTPTTjXtHnusdOy/j3BMzJnGN9Tx30UvY3BG5wwOYvAmFDRs/Uutl/5x86V/0J37WbRfijBEjqoPd/7JAVGQNuCj8vtk5CMtwrsOGUW+NOQ6XP9Ca/rvJyu/6MOW7ym6gioKGlSg7u40hrDcZHv/oxc++sSpkXLWCqHjyhmdIHEYWA+7QgEqfL3Zzgbx8fvAQLr6QbuTJEFUCMKYr5/Lu+tf3sgYSDIk3fGqe+dbT5wYH4Jpxm6gREY0vtGZ/lTc+laViE0YBCowFaakBUTISkCnyzSpPhid+HE39pgpyPjl0Z6885fC0IvGMCq9OtMw35n7zWjpC+XVP0vnPpOu/pFlqZkFdO/cj2ekwRlSVk7FUz/aLh/NVEU6sKykHDf4zW81Z/63ztq/oc7TvIVYNVLAEAqT8U5V8jok8hVPFCawD73v8akjI90sDRQlilWUB9wpRphBzEWbjW6W3emtH3Dn3OlbOmAvtLohgygJAWgsZhgG3DlmvWVO0rT95jOTD506SkCzxNihJa5zLZn9v9D8Y2/mtISshFA2OEhqDGIw9TSXWdouH/VTP1U98rcymEkGg0hvkqvQrTu9NUX0ncKUCmfMmMzE7W8Ml0O1hBhb6daXrTVP7q7jJ4O3MiHwkYx/t5v4WBJ8SMugz9/XGqOosbg984+6G/+HpUsMASmgBBQSdnyQu6cYUgXjeouBJU9xmgmsF6Qy4A7JI3kIibbbaRbu/OYPuFMGj2M/qDfToDCwiJ8GdtVFDsgxQg0hyx5589nJsRqQGUkMRbqVrf47Xf8Nhy68BGGgC2LBJ8GlNJMs8szA7USrpaN/vXTiR1IIrU2EPP6jlzkG7KbjN9CMMJhRCbOErVnXvWI0KAkoXgq4SgBa4h2/aw7qzKguA7LSaG3yh8EHDD4TBA91ICpVlsvtq9vT/6qx/PPAtKe6PC6kNyu4W8xgBqGI8vnFrV/4zadmlxo170qaeNVgUbghTW3AnWAAxDc7aRYGVte9505fpwF7od5IMlX2IsAMN3wfcHtIiJijSUjSU0eOvvvhB8ZrZYR19QnVWWOmde0zcbIcQ4Jzmcsy38qijrkAOlpMLQmzTBFNfKQ89TczfyqxIOqoDggo/GyhsDp2c1sMBivETrN2tn2ZSUsRQ4cBqC2pzhNw8Hf+rgkcYVALGhJaPPT+8rGPJXEI7GZhKGUlKTXpW0MWV7udjaWfa6/8e0Mdkmel3f4o1vvCdaFTIkTiPDmzuPDpX/6Pz17dDiiXRWINomBx/rtpmjc2+fBHnGy3OunAYbgP7D4YacDuWe8EWvAhE5YCKkYVCzboCe4Io9GrOGjXb8KnLkO33SLGjKNZiFVnk9VfLDWfdm2zbMiVHLDdgapDBXCZV7HUdbuZycRHqmf/R4w+SjCiQOJXHmqXRSUFCSiQWKkGgS51279XDWaWdqqZa0htu4FkK+QfbLsx50RAiWEeTn2pc/an6+GF8bUvljswJ92yNh3LQUrWHkvanPk5ZJN26hMdPyxIy2agf9XBLA8csLyOI5CfUTBz2hUk4kqA/8qVtU//ylenZ5uekRNAkbIEgqbOwm7b5w2LkqlArB0hbDYblfKrnrQBe2YgXf2g1WkDyBer6PUpu+jJ3uAYEeiULnUueG50O5/9rT+Ce/KDjzxcsVZ7+1cbjc9aLXNlWqgHg6pD9QlhCJ1no1JiwZL2MI5+pHTmJ2Xo3WreMQ9n3+vDX/gWHUjngZCtandB8skvpw4iWRJCAuS1Pm73cdcxACCoRXGNcvW8nPqxtHE16r7oKd6gSMyJBV/KPDvXWvP/MCqPR1M/AEbKrtzk0orwEwIosuENgGhqMLghBX7/ay989refurjc9r7q81XlaFqct/L6cjQDbk9uzjoGmrlB5fh9YK9v74A7od3u9Do6u+GHAXcCAVMJgZJixCzKED9/qfmZX7q4+Z8d/+53piMbL47NrUkF3ZIlDlHtydLwx6OR77FsrbP0M6ubvxs5Vzr6oejc33FDTwatqEIkOMtgbo9mBOEAMYIgLU3bcyFZh3cwwJRiahlCgru52UTu1jMnqAjgRz++fexKI/yzWmcrUlBMCcUITb1rafdSMvezZXfOH31C5eaTUgYCRmQ93SLMnAhY2g749T955lf/4GuzS1uVag2ST8ruuBN7O+z+Mt6gFB4VAd2z04v1ZrdWGswU3mMG0tUPut0OZKcKuQ0KaewWYYdiqVUyjVMT+vJTs42V3/x62rj2Q1PfGPZEa7wx8mj5yAdLYx+QoXebHxWDcwg2EZVPl45/HGPfpQoNoASBAi5f9mRPmEdvgS+GlrYvI3TAkkENCjGzFJoAsN0Nu6+flpCAh5lyonLix7eSp1uLvzKk5rIS4FNRjyAuKwdpbn2xu/xPWf4HbuicGQC7cXHewl0IEAYEKAO9owfw3PT85/746T98eqHe5ejEsRhpGsyMdoMD1V7mMBhwWwgTUki/uLY5KGO4Hwykqx9kWRAiED17azB8vXMMkIDIgpQtG3ez42MLZ8dWJyrdUhLGFqZX3DMonz/2lv8mHvsohh53Ls6IBKkQ8ZHvHi+/TVw1VI4qjKIiQfLKG4juQT+sAiJDJvQuraM97QygAwSaAQAzRbb7WNJC6PIVj/PqK6rwpfO1E/9tszvTXvnLWijTkVGdGhtiKEpMWhuf61TeVIv/rkXRLY5oimDixbUVX/rmtz/1S7/2xT9/UUdOTp08Y2I1z5KLnHMGM1Uz6xlcu7yINzoESEjYMV8H3FMG0tUP0jTJPTQKYz7YHzzMdw4tQ5RqNK7bp8sr7zn14uNTT58trzny+ZXSp//w5Fr0rv/iwn/59rEHHJCGRIzeZcqEEsfDFxRJ2zYcRiMEsYxwgFNQxfIFvm53+FuiCC4FvEtb7C55QOlumBUytVSgdx4ZD6AnXfkQp7CdHDQLEg1/rHxiMTT+cdi66F0UGcysi6MStWPL4nYrWf1cNvLd7siHXv2A9UwuUCKBm9+o//pXvvVzv/Sbzzz9POISuNK4uDlUrQ4PDU9Uh2vVauQj5x1JY6/K4uCh3QV5gW2DMoSBm+XeM5CufpBmwSjMl7wYzHXtBkIM6lzdRVmSVBbD5FfXH/j6QmZbm1uN7Jnl4W9cPQkpPzX3//29H/vB9z55fHQoIB2OQkWdZ4CFxCTQlT28QGhi5kgXRDOoUPYkXcVSmARgybYlm5FA0/zOkjRVmFq+ZO4tP+jlFB9gRQog83+CmsvgSxM/ylOLIflHIdRdVgrOEqvFlZpqEreaYfuF1sa/qE683yGy3lxVZLslIwAAIABJREFUHt9I5Itb+q5ienntU//2t//553+v3lJMnhJvzLoh7dY3W/XNjSWJhivV4eGRsfHRaq0W+0icE5B5JZjBk3tHEDRCzBAGNQz3gYF09YNuiNVSQPKEHfQWmHrDQ+BlDUFAURRmMgKkEoCY+STNttvtZ9c7i39haJ1GOAsovMd4Bbbxlaee+snLM//dj3//f/Wx958cj5ylkqWgg49ocUyQUDjAgRDAwRx0T7oFQIIRRBnILHxb2tM+wKSdFzU0wIzIF8RCERhxRxSbycuGOCqeUINKJTr2N9LOeuvap6quISGu6LpDNVRPttgNrfWou5YkK+XScUNq5qBGimgggrloq6u//9zVT/6rX/nSV55CFMvYqJHIFOop+dIrkpluNLc2Gpvza/HYyOjYxNjY6Fg1rkXO5VKMvDwEKb3bdz1mcUCBwqI8s2MgXfvBQLr6QZoCjIQAdobCA5B30lYsGFX8SchgwQhTiBcKO93Oyubq8vpyo76lMAQgJopFErvB2ibk0RMz7a3/6Z/+i289d+nv/K0ffOKBMwZRTZxTgQkkLxZRHLLY+87F5DVgyLtvQaK6ZOmGQMxr6EVlEBQ6UO72dvPGMxSDoxmI0un49E+oLncWPzeiBmwnrTXUvt8f/4CPz5ZrT8CPEV3tGW40oxMyXt5s/D+/8sWf/aXfWN1quZEJiFkIRF7NitZzVFJI5w0Ipmvrq+vrq+VKZXxscnL0SG2oVopiwlQ1NzpzhyZvU2//jUbu7IVB1BgG5Xf3gYF09QMz7ZU3xUC2bqB4pa03bEdRVsl8FGVBt9vbq6trq+ur7W5bQwLnCO5E610PHTAanSuPQsqf/+0/febFmZ/6Gz/41//Kdw37CmBAAjOSNAFcnl8H0uD2JlwABMVEkFqWImTwLKprGKAghHTFtns+GJB6amY+DS6OHiof/9vddiVd+/koVknRXr1WHns0HPuebZURZlRzYDCauBIJ4A+evvjJX/j8733lmy1X5dComlL1FukBhWwKYeh0OkvLC2vLK5VK5fjk1OjIaK1aExHNspA7TW/cc7Akwg55ReNBDcN9YCBd/SCvyW2DnM6X40g1zVeSN1MjxHsRSdJkY311aWWl3thOkiSYkoAr0rBe1S/mDkYqIomipBw9c2Xl733ys08988JP/7W/+s5zJ4AS0FajWIAZ4EEaqTtFd++eoiIwNDDtECnyMrX5fTYAJEn04jb2cjADjBSQpkSgd0OPl87/dEcutbZ/vwoXp8/a1hcwccZXxpEdBWKQ3sEDy+3Wz3/+d//1r/3xlbm1rqshrhhIofbmXdEznl6JCCAoDCxLNGSN7VanXY7jsdGxqcmp0ZERL5KlqWYhv9LcP8kbDM03+GO/V6f0gNdgIF0D7htm2Aleo/dGNjvtjY3Ven1re7veSbpqBhE6wY3W1vX9gXwOy4KBBh/gUYnMRytbzV/+jT9+/qVLP/LR7/yR7/3wqZGaAMECVUndmaPZu8cw76+pXdNGnorcc7uxd34euAdHApCXd3JQutRgCsfRx0rn/mF7Yaix+utDWcet/LbGj8UXfpgugpKCeqf7H7/ytX/zO3/0pa+/uLqlKFURR4Ag9xJSYLbjBriJeuWqU1wl4EwNIXSTVqfZaW5srY8Mjxw5MjE2OhrFkZiGNIiQhCkG4RwAYCBvYdkOuHsG0jXgvkFTEaF3qYV6a3ttc3Njc7PZ2Eyz1MzgBE4gkke1vfr1N4qBMCUywkAzBcxMAmrldqfz1W9cujy38qfPXP6BD7/vY+9+7FitDHHIJx4K2dw7uamRWmjSAJiCNClETMSc3BPZMpenFDsaBCkYADWJdfi95eh/ziTD/O9i6yIqfybHfxhDkgd5fPWbz/7vn/rsn79wDZVx1IYhhGUAWIgW7rzYsOVLzQEUgVkIYbtRb7S2N+rrIyOjR0bHxodHa6VyRKaZphZ27ti9aeZDzRu+AfaDgXQN6BdFF16YI07EO6Qha2w31rc2VzfX6o16lqYA6ITOgfk60q/ZtxoACAFCaZr3q6aZIVVxVq2hVF3eaH7+d5762rPTf/Lub3z8A+98/xMPj1XL+UnI9Sy7u+5aCo+jWkptFfUntHASGWDiTe7NK6bMABHzMEcD6AAY2inoam+PT/5916l20v/Q6H4hbr5/bOgn8uCXZppdWVqFH+HIlGVtWAfFSTP/fqfC9XJI0jkTmlqr1Wo1m1ubG2NDI1PjR8ZGRspxRZwLqqa6o1u8yY3M/+vuTmHAG517814NGPBqrjv4DMr8u4EgSXGZhnq7sb61vr6ytrldD2kC5+A9BQTzyRXexiqwPBYCcJYXdiqqH3kaYWZOODRsIUxPL356ZuZPvv7cX/nud3zv+x5//MHzkyPDAgEBM1UFSbmpaQcBoDdezHUUhMIEYOrQAZE5IFAdYoURSVxyUosAtTwbYjfZXS+HheLk+u9ABzMyROiaWjT8Trzpb2eVoW7zL5w9DxSNd+GBs5NnT65ebFm205R5GtrdqPV1oSvmbSlO4Jypttvtdqu9Ud8aGxk5MjE5NjQRxyXnxdRoUFVeT20jCm8tAQLZzQ926FHAem7pAfeegXQN2Beud43W+0ZCaECAdtqdzfrm0ubS1uYGkgTOoRQXmxXB8rylxQUUzq4AoyFC0Rsb8iJPAE2hqZmBxPg4gOcuLz/34ud/9Q+/8gMffs/3vO+dD587fXx82JMiBBR53T7LBcaBL6/adzOvnxEMuXSpUNWQOotVFHGUJqkhi4bFjYpBmRp7CV53xfVi8OydCUnEkSmYAIKhd5fOnjraeCEuDyksrwo8PFwdGa5AN5l2JaLR3b2l9bLdXuYFpAhLJVPtZMniysLK5tbU+NTUkamR0eHYRQRgatiJCLW8yUDC8omguzujg8yOaA2ka78YSNeAfcF6vb0VBXEsL0Dc7XbWt7dWV1bXt9bVAoSIS7f7sNvyit5hJ42m17+qwVVxdAhJ6+KVlf/zxX/3r3/7y9/3ke/6+HsefeLBcxO1SrVc8kJBCmSwfFnFfHlAAWGU2xgp5qAehUFBnzgIoXAcFRm1PA99D7r1GhjQhdQNNaCKzJX8OYyeM2AnBXakXBmtDUFoltm9P4EbMAMg4uBdSNKFa/MbGxtHjxyZmpwartacOFNT6Y0udmIbB9NgA+6WgXT1iTw+/nZbHXryrii/TitK5lHN4MQsbG3Vry0srG9sBM3gvThnVqQW7WvjEKAZ0szEY2IKmi1vtn7hF3/9l3/l997+4IWPfuhdH3jHY4+dmRwuRd5VREgY8qDFPPQiNwNv8sE7tpjABAaYgQmZAMEIuFHKKIDcLrzXECjDYpoDEAgYKEXN93yDahSN1moucv2poqdmVNB7E3aS7ty1ubWNtWOTU5NHJyvViiNNjWTRUC9/WgYM2BUD6eoHURxrq/sGCbXKHVoKGEyJPLq90dhauDa/ur6WZsFoiF0ei7czA7IfjcNeVDeBSBt5LlcwU1PEMUrlVhZ9+fmFrzz3ueMTv/WWc8c+9J7v+MB7nnzw3KnJsiu5yEKanxQLJ+ZNjoBCnr0xCgYjTTQ4k4AUsNKk+HEFgF5i8r2lsG0VIIUKzZAa4CxyZlDWRIZrFS8WsoCbX8I9Y6e1LT8lD0LaWWf22tXV9ZXJqaljx6YqcSUEhQVHRzOa7dT9en1iZoST/bR336gMpKsfiOyML1/Hr+l1DPkshtBLJ0kW5hdWVq51220lyDzgIA+v6FtrmDJfn57cmYYgTJxWhpB257ezpW9dfurpi//3537v9Kljbz09+b6HL/zoX/3wWCkCjCa3OFMChFBKkh9I0M1KkrkWW9XKWSlNBoPbQxTjrWBmbBIO8IISIAElwCy/SCOA4VolcuxaAKO7i87YPZZXlTLL+262uq2rc1fXN9dPnDgxOXG0FMWaKvMSJ69vb4QB4GCV5P1gIF39wLmo8JAIDcTr+l1l/rJ6182SpeXlpaWFRqMRLIEAoLFnl6F/riKjpFLKxdKCwXL7yES6CAahEZlGjZA1NtoL23PPPHPpy1/+6tmTkx975+ORd8VE1U26fS2i8+nB0k7FJ5bfblnZR3O+9gBkFAq6BIhfvf8eMTi1sgNgDjRh5oBs50TNAIxUq957mBpB68vg6caOmgSdZRmIerPenm5vbmyenDw+Pjzm6UOW4nVrkeSuZnA//AkDBtLVH8pRyMjAKAYiS2FQiN4u9PuAQ1Asd1lZHvVuAGjivTGsb61fW17a2NhoJx0ISIGwl+/UP9EqMAPzuSgUeVxFh6LMI+dI8R7eq9Eg7W52dbX1J09f+siTj0Z5WcKbLY8iBjCIOVA6ZaQlxN1KKjHOfLQ0/J0l51C5oFSHFOquC/a9gxBhCcj7SEPhl8ydpAYxAKOlcomEZX3SrWJ4cv03mFGEIjAkSbq8utxoNiYnJqcmp4aqNQRoprlXgr00cTv0a6sQFhsEprTgB1bXPjCQrn5QiWBEoDeomBL5Ahzoexd+L2HeSREgSKipOHFO2p3m0vLS0tpqvdlQU3hHJ3nMRm/i6T7AG6t3E72Wv7HUXv4/RpqWSwz+uenFjloFyItj8FUTXoQYuwJnrIZ4JIlQ7QZzQxh92I9+F1DLAEOHRmj0GnbbXinOSZB/OmG+F9ef/08tLnkCmvXzYXtVS9HMSIpzqtpobHc6nc3W1tSxY1PDk6W4BFPNNPc057v371z3BSkyvi0IbSBd+8FAuvpBJY7ETAgR0PIEIuxYIIcUQ14WFwqjwDmfhWxjfWNxcXFtYy3JMhPCexDIowj3oeO+95gZMoABNre4stlOxislmr62xZIaPFFzdjomXZTQ1xCfNvWqMIHsOMT6dPmvPEy1Uvbe5dLcp1N4DfLBgYiAzLJ0Y3W11Wq1x9vHJ48PDdXES0iDmUrxYtzfk907xSV4LwOX4X4wkK5+UClFYhktwMygBgfKYV/hKHcVQiyPbWt3WqtrK9cWFhrbdfE+KpUy7U9I9j2HIJVuZW1zeWPr/MQQdqpYvAKDMYBdWtXpmczGUm5I+bSLLsAiFLlLAr7WVFk/iEvOSc80PiBDJVK8N7Nuq3W1OdtoNo9PHRsfG4t9hMBgKrhdIt1BJ29oNdOpiaFqaT9SI97oDKSrH9RKXiyFpsWMC5C7pm691wFkZ/yYR4WZMwoyDdut7blr86tryxqCq5RJBDM4yVdGvG/d9l2QT8eJGLix3VxZ30jPH49vtbBXXPRT0dmuf6yVfGN04judO1pElRKAA/U1a0ntP5ava2x2cGx8M8vLG9J7U6yvL28368ePHT8+OVUtlwFQD87J3i1UgIbs8QfPDFf3nnQ/4JUMpKsfDNfKDiYWem4zUfAwevTzGYviZ4GKZVm2tLo0d22u02lRILHfma/ADVJ3WCAAGEmDa3WyhdWNNNPYuZt3pVSibFCIueoJOfJDSXbOT34/4EHrXTp7czf3yQl2MJ8x9pK9BYhcmnZmZ6cb2/Uzp0+PjY+aweULWx/iWFzL/zGzcEi9DwebgXT1g9GhyFGoRlPADvWDbGZmeWknNprNawvzy2vLIWQQsqgpUfTRPMTJpoQ4GK/MLnSStFbyML3ZxSQMTkUCU1ceHr/wY5n/aIYHQSdIiACLilfs0A1S+obQ1OidqW1srDbbzdOnTp05diZ2LkkScXKY1QtmYWJsqJ2kQ5V7nxrxBmcgXf2gVnYiQeGCeJfHX0OLwPID173f7HwIM92xEp13QcPi2vLs4tVWq0Wh+Dz4++VGFg/qkP+2aL58pV2aW+h0EwzXwJuNN+gglLxZvJgdUY4aotywvuHm8n54CwnrjR1uNIQPAC97TAzinKmKEzqXZOnM1ZlmvXH+9PnacC1LEkehkcbCUHsVB/ARywtDK+LUpFqKBxGG+8FAuvrBcDWmWKYAIwfzCM5CAA5Ob7KD3dzFF0hVM5A+ijqd7tzC/MLyQlcTilwPHzTgAF7SbjAQhWGcglispy29xVV5CAgyr1VPRvkPAOAAB+5IV19bhSBMcvOeJCDgweo9X6ZeWqQdGCAiwWxlc7XZbZ4+derYkSkanBpBtZvMDh9A3QKQDxaCSKo2NlSK/f6UAXtjM5CuflAqxZEjMgDamzG3gzrVdfOhrcGcdxRZ31y/Oju7ubkRBHSv68BfstlqdrLc3rrpZfIVP9yw0QFoFgJAmkDVDm4n/wpIIUyt1W5evny53WifPnaiVq6a0UK+BMHtPuEAkA92aLAQhiuxO+zxkgeSgzUWe73y2Nnj3/XohVqUD4StV5P8YD7Q9uovI1zszWF+6drFyxdX19cCje51WeMmr7eBPJ693W4muXl8qCi8gwSAbpoGMxymO0UKQSZpOr9w7aUrV9bqWypC54pZVhYr6RzcEF3mQVjqEIarZe8GVte9ZyBd/WCsVvmOB09VSuT1GGU56APIYqkPAhDnWkln+ur0lZkr9cYWPBjJYZ4+vxUkIQIhRNrtbjdJbrfHweK6hUUAaHXawcJhKxVIki72GXR5Y+3i9OW55YXEgo9iA8IhefJIFcvGhiqDyvH7wcBh2CdEGIsR1ufiCrtHi3PLZ2xEzLDZ2Jpfml9eXgohIHIg1QJ4q3SnQ0t+7QQIMlM9jJHNva6dALabzTRkhy/HlzBQvDfVrfpmJ03TNDs1eTKOozRLLb9FB1fCrDc+DUOV+DBZvIeHgXT1iWo59qRp3onkxtfBKW9wA0Rv4E6SIaRb9frMwtXVjRU4kZI3g1m+4u7N5s1fZ1heA+oQ02w2U1WIHMBn7bUpGl1NSWEp7nY6V69eDZ305ImT5WpFTVW1CEE5gNdlUCPMTh0dqZYGYfH7wkC6+sT4yFDkXehJ1wEcMebmhgEwNQHJJGRra6tX56422i14lyfZWB5S9wbAgCzL0jSf6zpkl2xUwgxSbydpZofNRM4HUAaK5RGeJZ8m2ezCXKvbOnf+fK1azV8i5tNKxevUm9+7/+8XCZrZBx5706AK1D4xkK4+MVKrxqXYmCpFLCOy+90b7rznRcRIIV2aGYLQdzVbWFmcmZ0JWYI4hpoGkCK5bt3fnmG/sbx5qGZF17nzx8OAwWgpaCn8egepukMTnFfAPEuDQDFOMrNYQsiW1691QvvC2fPjwyNiJuZoDoCRMMlXYQMUDPfxeg0Qo5iWokEHu18M5g/7RK1c8s5pyHiAusB8cGpGM5gWX+rjKMmy2dmrMzMzIcvgIxxyp9muKcbuFjkXxXzZHw8LJoC0MttuNLKgh2+u69WYQQRRvF3ffuHFF9fW1kDCUWnK/FE2w4EI3WURGB+OjA0lWXa7zQfcDYNBQZ+olHyt5PNpLqJ4w+7jXNfOoYtJBQEAEj6Kt9uNq1evLq2uKPJ1I2l6+EIV9g7NnHNeDt/Kaj1jxbW63e1G00zh5PUx/hCSUZRm6QuXL3XT5NSJUz5yWRaKiy6W4r7fk7AGA0MIo7V4EF64Twykq09EznnnJB8aHgxy9TIW0cZ0Il42G/Xpq1c2NzaUedpWHh3PV68R/EaAJA9YHYo7gQAggNtq1rfbKdSKVVcOyqN3VxAEYTAYySRLp+dmUw1nTp6J4zjppo4Cszxp4z4vspyPHUwnx4ajQRWo/WHQrH2CxHsfPjVSjc1g5I4n//5S1FEkVAAnKxsbF69cWl9fz0zztNB8ozekbhnMSqVSHEfFr4eQzUarvt2CGoGDM2a6O/I0/hxz4mKfhGxu9upL05ca3VZUiW/I9b/vTgIChOmR0aobSNf+MGjW/nH22Hi1FMPyiLX73PJ5J5BLk3MU4cr6yvT05c2tDaXB0WC39Loc7n7wdhQWaaU2HEcRcJj6fRb2iQCYW1httdo9a+XQXMKtMSJf7wvepZotLS9Mz043W82oHBME9AA8mwbi0TedHB+uyhsjHLf/3OcO9A3FcMWP1VxXzSBE6G+iTT4TABIkhPm0vRpUHJW2srEyPX2l3tiCkN69hplFA41iFKO3Q+hJuy0968RggqAjJVfNo8oPVf9DU6gCeP7idLudQjwOaMHMu8QMeeQn40gtLC8uXpm+XG9s+ciBAHRnKrfn3Ojr/QsKg3vHg8dqg8j4feN12PscWKZGKpPDvhMQxDkEos/F8QS9N5lQQo0mkSh1aXV5euZKs1ln5ESERerW9SCSfB2uvGhcLl997gv6iMHUTAFC7fhwVCsSog7R9Ro0wEIGvDg910kNEt10tbFDCg3FNKwZQfpITZeWly5NX95sbImniRJaLB53fYDYv77OQFImh8v9O+Qbj0GYRv8YqVWGy3Gxdsj1xZT6hxXxjWZQAPQM0JXVtStXLnfaLZZLwMtis67/ZPkSGrmdZbn8mZm+Hg0vIG8pBeyBM6fKcVQ0S79v192SrwXqXb2TzC9vJMa8gOFhOf07oXctLNZRjiKkYWV1lWZvOn+hWqki5NYzcV9c22aenBofUrNB5d19YiBd/SPy7uTRkbKbh4b+61ZOnr9lgHPM1JbXlqdnpjutXLfwWv2b0qnlofwqMCDQYCZ6fy5iHzGgyFmwILQHLpwrlUrAYSogYgBEFFza2F7e2AoAHGjh9SVePQpdMjohdX193WAPnH+gVqnBCDXc4DzoFzRTWjI2UpPXQTrdQeV1Omo+qLz7kXNnjw6Zqpn0M+r6eppm7vVzYt4tb6xdmZlpd9quUuYtIx4NTuEDo8AoQ5yhlLIU3OvWj58HXtZq0Ymjo7F3wOEqZEhSNGBuZX2t3gQFJBGuF1Z+PcHCuU2S4gKxsrb20uVL2+2WCc1JHpbIfpc6tNjp+HDl9bgq0EGhf73nAACVOKqWRDVApM+BUPk0vRHwjpFfWl+7OjfbSTpSKuVOF/a2uQmaor2NpGUajDBxCqfa1/PvDwRIwkALY0O1qSPj3omZ3fdo611hYKZh5tpCs9UFhFRBIPD6k65CtEiQSpgTeLe+sXHxyuV6uwnvILzVg71fWDnCILxwXxlIV18R4ZHRStAQDH3rDm+0p4wU5xZXlqenr7SaDUT+1kHweYWoIZ+999GTbzlZq4a6rV+z5TlsrlrSVYWaGO9i/fhbGXn3EQMMhMFZOHZ0fHJilLBikNHv7u8uMQJkloXnX7wc0gBxyH286NcDd7/InykhItlcX5+emW42m977Qtput/e9Qk2HSv7j739bJX7duiUOAoO5rr4yMTo0OTqcJPPVUrRvLxNv6GWLqIziF1pUipdWV2dnr7ZaTfh8UEiSpoV+GUmzvNCikRY0CtlH3vXQ3//JvxbEXZ5fnFlaffaFmb98+oWZpbXNZhfdjgWzuIRyBXEJXpCnhBkJIGju0wGt0AXL4/INgJrYzuCJxZkTuJ8q0YteE+DM8eOjcYx8un/X2nzfEBjArW725996MUFhdvTKabxuyZ8YFg8bzdvG2toVygNnLowNDWedhBDubFVszxt2vUeQqhZH7vzx8YHFta8MpKuvDFfiqeFhMZd685mKunv63gDIK26j0AnkkexUC6RJHK2tL8/MTDebdYjI9TWcrttdRiEyqhq8MUJo1ST9sY+9/4kHzgJ4/PypdpJsfqCxura+sLb+0ur23MLypbmlb1+Zvzi/lm1tIWRgirhkpQqimHTCCLTADBaY1zuAEQqYWVwcnzvaIGbXk3L6DwEHNcvE8NC502XngNR4uMLEGCy9uLzxzLXt4EqUQMAYH4KFufcA83/ygQ+BiJrp2uqyF0Tnzg9VhpJ21xWrzOHeqpfxxteHqbFS9ifHK7faZ8CeGUhXXxHyiYdOnJyorWoi/YpaM6p40Ln1zY3L05frjW2QvLE+zY1vblF625SwLDjVtz506oPvejyPqvfkcBwPT06cmZx4EkgR6o3u6tb27Mr6xcWNy9dWr80vLs0tzC0uza5uNjfXzHkVDx8jLiOOIB5kKMa/ZpbmtUUI5MXpDgJmqiGrDQ29+4m3RE6AlIfsNWEnCc9fmVuvtxFVHAIAowDkoYo2uXssH4OIWrayuhI7d+HcuTiKQxpIyf0Q97YhrCebBPKx1/hw7XY7DdgTh+udfD1wZmr4/PGh5elV7n+EHkkDKDDBVmP7yvT01tYWvbNb1VXrzYgI0enWIn7g3Y9PjdSYO86sSKQhQVqEcGSofGSo+pZTx/5ToGVYW9taXVq7trj80rWlZ2bmrm7Wl9e3llfrKxvtbLNpZvAx4qq5EmIPAZAVGW40GEA1kz5OTLwKYUhB0/MnJ55463nvCBPykMU4t7vZ089fRJoiKvc68sN1BXvGAIN4H9JsYWlZRM6dOed9FPRVCR12T9qGPV+kVb09+cDUyNDA6tpfBtLVb5zj8YmKvqTiHPazoAZJwISk42ajcXlmemNzg5EvJj/sFkm2VlQHVh2qld/9yJvz+OJXRPqaGS2YZSjmwV2VUfXo6Jmjo0++7QEAW910bnV9fnl9Zn7xxasLL0zPzy2sbNTb9XbaaNWT7QwQCOG9eW/OQQiBqFGD9Q7x6pPb54BjatCqd4+/6fjJ4TKhMH/opolW6s2nvvU8nAfB63XjsUfP2OEi9wmKc0mazi0uurh0+sRpODG13G+Y67nttVHy9DFRERqg2dGq/863nYnc4XIyHz4G0tVvKrE/e3xM0yuoSK8U7z5BM3XOtTrt2Wvz6+uriPyNfsJbDDctz2E1jI1UHjw5eeN2NyiHM1aKzaHI59SKwAwnkNFSNHrq2NtOHcOTDwNYbjYvz1y7OLswfW1pZnZhem5xerO7td1uJ6ET0ixJAIH3CqVYHrLYf+kyCIIO10rveeRCnqStdJJf4CEJxw1qz8wtP/PiVZZGLR+/7LFzPoQUZqaZAYxcqtnV+VkXx5NHJ8U5BHMAjQRtb4HzhBEhgAYHQtP2SLk8MVK93X4D9spAuvpNOY7On5iIxIUAt4+dMFRVnEuS9Nq1a0tLCxCH3FRi4ZS/Kb2ARJoaiKmxodPHRoAAvHIUuTNcJWhwuaAYRHdqxxmR5wDIX426AAAgAElEQVRQKdlUrTz1yJvf+8ibASiwsL797NX5Z168cvnq/PS1lSsLayur9U7SbRszeqgW2QMsUk7RF3Jv6OTR8fc/+QiQAVQIEe6nD3OXrDdbX3v2pXaj645Eail6d/yNQ25v5T8a1AiKJGn3yvSMOHd0/IiRqsXKXnu8sXmKPyhmJC3LsvHh+MjoYKJr3xlI133gLWemHn/w+DMzG/s62+WckJhfXJi/Ng8CXgDmw9BbvK153Q0joUGER8ZGJmpVQF8tXTRcn/bvfaQYJC/OKwARaCYGy5wGWmZwRmd0FJ6aGD418daPPfFWAI1MZ1Y2n7145etPX/qL2Y2r69srS0udra2QJEqqOBMWqWP72gvTrNuVSuU7vuOxt50/C60HqVmh97fb98Cwul7/xtMXAaf5msKFzcW9OsYOIwKDIA/L8FHSbc9MT8c+HhsaAfMyUQD2ql55u5owmBE2PlodqZVut8+AvTKQrvtAybuTx6rfuLxSikp23XrZs9cdAOAsGMzE6P388rWZldnEEolLRTXS2/XBCueMAqgFJ6iWIoEH0ptsmhtvr/i8nR6SACD5N0aUCLkFBWCnuIEFEIAMeXnk+Phbj4194jufaBlWGsni8tqzly597bmL33ppemZhaW2rkzYBKkplLVdAUkHT60EldwGpIqIoZSEwpE5BorV16viRDz1yBkBASZAJVOmKbLQDh8FCRgPMqdE8nHtpdfOrL82jXBbriikoYoAFhTM6QPNg7txRFmlKM4MYxQBjHl661878oFCkDebuQ7NIWq3t6atX3nz+wki5ZgggSb+Xy1U4Eg6pGVL4I0eGP/jEA+Vo0K/uO4Mmvg+IyIUTo2naRaX88qIa+St09wKWT22omY/8yvrq3Pxcp9uUKObdlMIRM7UbhejV3PTPfMWP1/955SY35EuRdKQDRmEjY+XzI6fecf7YD33gXetJOr++MTu3Mj238fTly89Mz12+tpKm6kmD2+taHgRQJGCDYiCCnp8a/+DbHwJAOloAaRDb89B83ygKpAeF926j3fqzp5+rL61idNyEvURzAMhXusH1yHAzILCI5zQQNBwm83K30FTpXH1jfaFUqpw5VymVsiTs0Rw15JmTEFga0uFq+ejY8O12GnAPGEjXfaAc+ScfPPngyZGVhvaKonLvugXAzIzmYl9v1Gfn5xrNJsXlf77drgXMy17kPRjZyxTuV4dmhpDRghd6z0pcOorqmyZGkvMnO52w1XzHv//6c//rz/3CerPLuGLkq92Yu8AMajAoJIgYgXa7MjTyvicefeD4UTMVVYgHZG+d235iAPJADIoTABevzv/pV75m3Q4IDQFOQMk7VlpgfsFGFeT3VBEBAhT2q8BINWP/7ngfyeukwMnS4mIpLp0/fdrHPkv3Oizp7W9Zt3vu2LGJkcFEVz8YSNd9wAlPjA+fOz668MJK5MUUL/Oy7Z78lcwFit51s2Rm9urm1qYRQoHxuhrdDprS1OAAqrGb5QbHHe17D2A+p9ULFlEDM0dUvFaG3PjQkc319bQTgDhjhWbgHhyGAExpUEYmgACt1rmH3vzB977Li1geKWkOuMOWuy8QRhCGotDrlasLFy9dgVNsLSNJIQ4KCFUIFyGK4DyiGN4XPa5KvnYbSCDkrW53+rAcMvKBmHmfZd35hWvlUunkiZMQWDCSJO98hLeDkYDAAqmWJpOj1bFBRldfGEjX/SFy8qbTR7707GwclXtl/nbsm929Pzu6BRicJNDZ+fnVjfVgRuFOBvEuMTinia1uNTa76Xgf1ylXccUwtvBkGaA0Gu1qvflbX/rLVlpCVFXEREIkt/u8W2BiRjglTADtwrrveuTs+x49h6LFHHpj6l23X7/IQzoJUGnAW9504X/4qb+53snazeZ2u7vZaNa3tuvbze1Gd3O7u15vtVotbDdgBp9rmBPvJXIGhKBm3KsNcoAhivsqcdztdq4uzEXl0vGJY2n7ZlO5d4oo6SwgZFMjlXc+dDq6Vb7/gHvGQLruD5F3D5ye7CTtWtlTBOpQqNfdmBFmBjPvfApbXF2ZX1zIzODyT9ulEO785DwsW17bXFjZGD89dYtd7jE03HDSzEe1jNOg37p0+fmLs4FliKdm5F6z4vIRg+UNX187fmrso+98aLwaATCRGzbbceoePEiARKCaKR594MyjD5wxILHQzbTeaNebrfp2c6vRWau3Vta3FlbWZuYWF9fWV9fr6/XttUYra6RKwsXwEXwZ4sE8M+HG6931iOpgko/kDMYoarQaV6/NVaPaSHkoDdldmFwADDAIyCzJ3vbmkw+dPnq7PQbcGwbSdX8oeffgyYljo5WuBimiFe6yZzQzNfVOIKxvbc7PzWZZSpHeLNquh9GWiygFZltb9UtX5x/po3TJ9UUR89UB8yr20uokf/pn32xvN1GNaF2vCcQC9/IAcyfSERrQanzko+/5yDvehp7g560mxTzQwS2p0TPYA82QQTyI1FOqUTQ+PoLxkVdsv7y1fW114+r80uzi0gvXVi7OL8/MLS+ubWy2M+00Yd7iyCIv+SNURLG8bDxxYJviTjGjE1O/ubU1M3f1LeffLE40hLteH8AonSS9cGqiVolvt+2Ae8Ne3vwBdw+Jk+O1T7zv4c/84QsSRRFVLOwuSNcgecQYDAL1Vm9uTc9Ntxvb8K4XUSi7NReMZiDoIIDTemv7y8++8LH3PVkiLI/nK97u/TNEdj6z5+YkAGw221/9xgvmYwiIDGK7a65XQYCal7RzaDRr4+MfftfbT00eyWXr5R99lz3avkMgn5aEwDkQTmgwZAAUkhb3iAKyd1mcGh2eGh1+4k1n889YbKff/Palp57+9jNX5l+cmV9aXNvoSiejWgbJc/CUpkZqrx12bnwxtjgs9MYjJFUVpKmtbK5UVkpnjp92pAB6fTxzR0aYg6oxYSSRPPngsWo86FH7xKCh7xuOfPT88Sx7wUUughGZ0dl1U+A25DlS5P/P3pnHyXFV9/73O7eq91k0i5bRvliWtdiWd2xsvAAmxmG1g9mDSeCRAIFAkpcEHg6EJLwQQlgCPAgmOEBYHDABm7AY2+AFvCNbErb2XRrNaDRbT3fXPef9UdUj2ZatGa09Un0/bcma6a6+XV11f/ecexYzGAMp10a27d7e09cNF4cUQuRQggNpanRGgRkyGKxVfnb/o2/7vd+d01KEWTI8GuppYgc73iHgDhg3uHNP/5qtuyxbgAjhPUODI6JnPnOsWCz7ZuoxOPCCC5/3vCWLAEC9yP73RaPqFoD6ZQAQDOLTRgDBge7rZ/mupubDqWctuvKsRWVgzdadDz762I8fXPPIE5t2bN81VKlWzamElEzslU7UKs6WirPXD3zURmQ0U2CfOBE1rW3eubVULExu6hA1M2hyQscEVYVS9nbmgq4ls6eM8VUph08qXccNEjOndYSIYLX9m+CNHY1D8JyA7O3ds337dsjhtlJxps6sJmJw4jLeZ1at3frTux98wwsvyIcZ08R1pnAY73APg4r367fu2N3ba7kWSDKFHmZbLyMiCQxApSzF8CWXX7BgdhcQmfmJ1uXkMKiNwBTizDEHWTJ9yrLpU1531WUbtm6/5+EV3/3Zfb/4zYb+4aqGGcuEsFriOTQaDQRjK31iw1o12rR5S2l+KZfJm8F0PKWzSIVGlfLszinZNBP5GNLQK8oTGyGntpVe/vyFOaekTxJjxywHyQKYDMKwr79v245tvlajCOsc7AAHhIQSkaBGeDPC5Uai4Kbv/XBH30CsFaZJcMSxnLL6+wcfXrWmMuIBx/qZOuz3J+hAYmDPi84/85LFC7KAReVDPHMTE3OBOWdx01FTWgVWyfihBdPaX/+SF37po39x0z//9fWvuXJGW4jd2zDQD++FJ0LeMvfdJzTVvQP9W3buqCU+8THfhAAAR2h15LzTZhWz6UbXscPdcMMNB3tOytGC5HDV37tyo/eedMrYlBnTnExSaUEYjtQqm7Zs6t29i4EDOapeBzvAgREowNGyQJBA4XZs32Zi55y+qJgJY5+ixr881HcZL9t7+r70/TvXbe5GviBUgRodyDGeqwNDmgiq5YKrfeCP3/CCpac4UdJDgsPKdJ5QRBQvTilGp3RGBwokJAMRVwjDOZ1tFy+d/6LnLenqbO3u7+/etN4q5SBwgQgQm/iH8RUcJ7g/gJFmNjRczhcKTcViIG7s0YYCMcXSOa2vu+LMUhqjcQxJLdzjzMKZk31UhRokGI18GMt9Y2YUMWJnd3dPb6+JUI5AFFzEEEgKAhmVJiZSdaWvfO+O9o4p7732yoITKI7aRteB6RmsPrGpG5kCJIBFjEcwjrP1bHgM773qxc+7eMmCjBOgYkzKZ5wkSOx5fQocjYs3sxDW3pRrb5q1YOrkl15+wU/u+vXN/3PHw6vWIcixtZNB5qmVzCYsIpFGm7ZtbcoVW4ulseeUqGmlFs2fPj3VrWNManUdT0jmc+H23Xu3de/1EIUD6MZw18QTdjab7dnTs3HzpnJ5mKGLt+wPU1FUQoMYCUJMk54O2WJ5KFq56rdaHV46f1Yhn3NJ5+D9hmr1fx4FU+zhJzbf+INfegkRCjUSqCFEEmJ5aNBMUau0T8p/+N1vXj6nS0TNVBkYRA7vHE4gaDWhF5iYCU0Q72ElWdowqNd4kVTIZqe3tZy+cO7ZyxY2d3Zs6u4d2NltZggzOHQj/3hj9UtYALpquQxlS7GUzYTeFAeOlbf9F4hGTptUeM0Ll03vaHYn0Zrn+JNK13EmEObD4M7HdwxELhAL8LRUUACAuTjRiiAt7mRhCFi2yrqN6/v29tGRIhwNNzs8WA9zHI2sNwBhZrBnz2NrNm/q7gvz2c6OSbnA1X9PQEFLgrBjomFD3DGSdYEZ29jiqYRJfRCSkfpfPPrk9+54mBnHuOwRpG4bjVW66slJTP6XsGolrA296dVXvvklFxdCB/MwAQJiwk7E44eIVymx95dxPajRC4kkKRABxbRGi3LZ3KwpHcuXnnrqnBnKkc3bNlf7qwhDBpk4ASzQKMngPtQcqWNNrLukiJj58shwkM8USwUhDIrYbR4HUyIQSy4PI4yIiAh29ry2qy9cVMyGJ89l0wik0nX8KeYyT27t39BTDsU783aA2JlYuhjfNrG1o0437Ny0c+cOM5XwiO3N1O/KWGfqyuSrMEWheWiw9vDKNQ8/uXZ3b0//8OBQZbjiEYnAiTGo1qx7cGjrnr7q8GBTMYQREmBUK8YoXTGxLpkJZag88tMHHr/j3hWSD2FGCuAAwwFl/llwAC1u7SEkrFYTrS4/ZdoN7/r9Ga0lBw8zwiW6NdajTnySkymJZiWyte8xuiFExhXSAUgpDJbO7jp78ZxCqbCzf6Rnx06LlEHGANEoOeyEmMiTDx3fVQahRpVhXy3m8835oppnkqtCI2giJgRBM0IJT9RqtSuXz1y+YHoYTBCpPlFI97qOP82l3GWndz26flu5pv6AGy1JqpcBZjQ1FWH/QP+u7Ts1ihiEY3AxHhYkYWoWsaXJKsHKVetXr1zdNblt9vzZs6dPn9I1o3NSa05cZbiyedfO3b3bLzxl+p+84dW0OGt47BsH+5HYXgQwNDzS3b0bURUo1UvD1k2zsWMO5gxxLWJvtXJnR+kPrnvZGV0dEcxMx3OskxYCgFktUhMumDbjQ2957emnLfvS179378Or+werKBQilyGQ9B+YUCShGS4o9/fv6t7Vki9lMqGvX7w0Eh70BhiFEJqFhtZC5uxFs91R7XeeciBS6Tr+CDmjozSvI7dq+yAYPHPfm2ZGAGZQmCmtElU2b9taqdUQBMfQvWWmKpmCtHehNrJl78iWX6+9u7YaYQgJoAoDRipgJXjlxajXY4xnu6ftEDwXiTAZYaYKJ4PD5V09fRBnh9qrngCMxjisW60yHGb00guWvfbKS2AaIKIlvk0DcES8ricstLicvocCzuGaC5aed8rsf7np5m/8zy+39w8jU0IQwg5pvXJcIWlmdM6Ant7eYqE4Y/oMhQklUS9EpCoCowBC1QB65fK586a1he6IuT1SxkgqXQ3BjMktV59/yq6f/rZnSAH/jFnexGBxiw5Hhe7q6e7p6zGBHJMdhdhppBQzqkHNIShKS1HpYYD3UEtcTBUNQ+aa25PXJQcYR+jjPomzJIGsXKn09A/EvqwxC+BTMIOqVxIitIgjA2edOe9d172sWWhWI5B0WjmEQ59M1DfAwECS0hSRB21We9PH3vP7ixct/PiN31qzpSeyvARB/NRxdIprAOKFEZ2r1Kq7dnc3t7a0FFs0SqpwCCxe26iJkAIVi06Z0ZH2RD4uHIuJL+WgtJQKkyeVMs5FeiCzggYqaBAodWBkcPOOrQrIsfv6BBCaEh7w8fYA1DtfcVp19M6piBeJ4hYc9VpWo9EZ49OE2MQEAREAIzU/WK4k8+ChzYQ0ZWSMIN76ds+c0vbWV7zk+fNmAp5wQAAGoBvvOE8yGFulPi70S02uSVP4KACuf8mFn/rAu56/eFYw3OdrVaWo6iF+X8cdkYHhoe27dkVIWqHRQBNRBwuUzgizWkYqp82dnF40x4VjNvelHIT21ua2Sa3l6oF8YnV/mwTi1e/q3jVSGTqWM62ScZkBZz6wmrOaM0+DUlSYPBzVEbTQWTFZh3qYWiy5Yx5rMtWZwTSu2FEz1A53AjSTGlyEkSGE7uoXXfz6Ky8BRkxV6TzFgx6xSaeHleZ8AmOIs+kUFkEjeqU3R3MhRKCRql62eO5n/vIdV11yAUaqWqkyPJbe7COHmThR0z19fT29PYELYs2mCSw0yyicqi9k7JorzpjR2SxpTPzxIJWuRmFaR8uVZ02bVhwm1UQM5sw7qIkzYawO5mt7+np3bt+OOGT5WMF62uq+6EOqUWEy+mCcF6RGWhjGYxu3vYW6TEMAKFQxutSP3wIwmonZs6ihxRvq8QjhzbzBG4CgCCW6t7/i4jPe+XtXFcIwqikorI+SQD0SMuVA1ONbBeIgrD8SR7GQQCCyeO6Mj7//LW++6iIM7Pb9feZCpaOpUz9hqh3GAZUiI5WRbTu2DVeG4BBQvUDpA5ZDi6rqJAjPPqUrzeU6Xhy76S/lucmEbunsjhmTpOKjCA4MHc3Be4qZmUXOSXmkvGPHjqha5Xhq1Rw5aMnDDBo79kZ/FDuQkIiHAEg2kMaDjUqXAUwyg8NAspkMIPXihYib+T7L5zcDDEmZc4Ma1ChgFru6zzln8btfe9Wirg41M5cz7JOueMSHJrcnCyRBBzqIgyMcku1MAk7inUjilKmTPvS2a95+zYtRGdaBASYmy7G/XA8dAhQxaP9g//bdO0xU6CPRyHmHkRC1bCAXL5szf3pneq0cL1LpahQIdLQUL1m+MMg4DTJeQqMjQKiZUqTmo56+PX17+xCGJg3t1jrkoXHfnwQkbv1XyOUmtTYBVcRuSwMNzpJOlM84ghFqgMEZAiQPonvL9Fmd777+movOXAyA5h33zbvpBDQuDnzGWPcOms6d1v7+P7j293/vd1AZsOEhcxnP8JmvaGgMoHjVHTt3Dgz2G0iKGQ3UqJZ3/txTu3KZtPjTcSOVrgYin8ucvmBWe3PRK5RO4QzioHREGPQN9O/o7vYaiWvwfM9DVi7gKVdk7PhDc6k0fVoHWGFs1kEIuEhFD6hdsXsRiOMvEMAc+vo6mvCet77q5c8/J+McfESomG/kkziBMYXWFkxp/4u3vOq6V12B2rCNVD2dyQSLICcIslwe3rZjR0U9QdKpywZOL13SuXh2Z0PfhSc6qXQ1EE6kc1LTwq5JtfIwNI7fAlRdICO+urNn90B/P4MQAK1xbxoCY22XeSCSFT05GlI/qbk4d+YUBIB5UgxiGgdxPEsVKFWoQkFQNMLQ3vaWzDvf9Mrrr768OROaT/RvwlQqmnCQsAhWWzS57Ya3XfeyK86XkQFEFTM100MPEz32JK5k7trdvWe4HzBxYQ2uqeBecPqs9qbCwV6fchRJ797GorWUf8lZc0osi9ZgqjAD1UlP357de3rMVJwDWC8q0YAwjmIHDt36Stx4lNgDVcwEC6ZPCQoFq9UoEm9k+TjL7RnEe1zJUaKqlfunt4TvfP3Vf/rGa9oyoaqCAGMv4jGM0TyZMIi5jIEwO7Wz9cN/9PqLz1no/KBVK2acGLplYLKXS4hEUW3rrp21alUYeJP5szq7OlvSAI3jSypdjUXo3PL5XS9YMpVRmaZqEOdGKtVdPbtHRoYZBvGdz9gx1ohQzaIodtkdyhBH54P6ixXAKV1TF8+dhXLFVCGEcypmfJaSDSKgwEdWHpzekvvj1770L9/y6qbQqVVomhisRAMbrhMbJaN4o1E9fHTGzM4PvfMNi+dMlihCYvI2+pLhqZt5hsD19HTv6eut1bStKf/icxZNbmt59lenHAtS6Wo4spnwwtMXVGrDNUIBoe3evWNv/16AxzIg/pAgqOqrI9Vq8s8k3m/cIhY/n2Sc2jVr8qRXX36eU28qkIBOKCamVKubWYlzBwggGZi5wT1zOgt/ev017/v9V4eAWoUwikD2Vc0Y76iONLZfLOToaPadsf0fgD71YQaz0X/sf8Dj/rH2x7yaXrZ4/p9ff92syW2sDEGUsEAhRmvIKr3xgAz1nBASICK/s7dnYKS/qzUzo72UVtA47jT4VHgyEjg557SZFyzuqrigFoZRdXBvX0+1VkMQNNi09EwI8z4ql2vVKuJ73uOQ6tntW/bSAWgr5a+9aPkZi+cwGgFoEoRAqLF0qQEqohKaOUgOVWJvzxkL2j7+ntf+ybUvyQDeg8yRuSTgngCftrI+TpjC1PbF8iusCvgDqZcB3qDxc33ySHRrVAPtaVp2PHBIeqnBBeay8UrhdZed93tXXVJqCVHZ68TCiIEGFoTWeNJlSPLS45QPRAYD6LoHB8Jg+DUXzp3dmZpcx59UuhqRUi5z7mkzfGWIwh19/XsHhqAmItLoxQkMQm8sj9QqwH7r1yPA3Jldf/H213bkqujbhupw5PIVV4okhASgwBS+atGIdW/B0M5rX3rRZ254/9UvuEiE3nvXkGEtBkZkTVgTRIKIpow8PcyLPv1hJopAESicwhESJ6UHUAcVKBEXvx9f7ZKjTXy9mpmQ73rdVZeftxS1KKqiGuQ9CY04UZosO0F1ZNeOrc2FTCYz0QL9T0RS6WpEMoG74qwFFy5o16i6fXdfperjOaDxlQsSmMnegaGhciWeQOtGzmFhZpkguOrcZZ/48z9cOrMZu7b5ngGtAUZ4RbWM/l70bEX/rrNOn/PZv/7Dj7/rzeedOi8bikYjhJo26PxoEA8RQGAOPg6/MdSLZwkgBjEkpUooMIEXeDEv5gURUIPVYuv2YO923CBp3s8oZd/7upcvX7bE+gZrzkWkaIQJEbVBQkIM7b1o0YzW5kKD34UnCanHthEh0V7KX3r67P+5/4eDA0MmcaUCNHx0lsE5GPsGBgeHy8hnD/b8sUKSZqVs5ppLn9fV1vqt2x/44X0rt2zehUoZzpiXqVObz1x05nlnLHnR8sXL588sZELAQ6PYQWgNuUIj4AzOwLgLm0FFKE74jLuSo38zabPJCN6bWZKcTi8M0JC175Ny7DTo8IWL5v/+VZdv3Njb298rLXkXqYc04qD3Q0QMMLO8i972iitmTe042CtSjgWpdDUozskFS+bma3u9ryJ0E8OvYqCIGQYGK3v6h9HeGv/0CE1NhCIXBJefteyUObOvuOTcVes29/f200lba37ujGlLZkydNb2rJYjfzcMUEodkkI3qXRAgKf1BYcA4ZXd7Obrv4cdWP7nOe21tbSkVcsVioTWfaSkW2lqbO9vbWjIOCOHCUXM2Mp80HNtnex2hs36kEGeqIfmay85ZsWbDl77xfau1CTKQuHFOo0JSlULfP3j5C85ZMm+ma/RQqZOFVLoaFAKdrU1XnHvaY1t+VYsCBPKsGbiNBQFXrlR37NqNuV1HfAI1NQNntjXPPG+pnre07M3AnKtfxxZZtQYX0MVVoEZpVB8PgSARoP7y8Lod3Y9u2Pbzh9be9+vHNm7ebEQxX8jlsvl8rpQPm4vZtkktU6d0zOicNLWjdWpH+5SWpunNhUltLfl9rQ4N8EAD1q1gxKyrRVNaS2986YV3r3x01aMbtHWOIao3x0GD+TzjTgIGqHm6CO99wzXNhTQNuVFIpatxcSK//8orb39gzUPrt1mQpcXB4qwXq2io+xxAPAMbhNWa37ZjNwBA61PAkSDeO/NmVqOMCPNFV98wt5p6L1FE50CJa/U+5XWHS3y2k2w6gxEKECYwY5ydquPbO1agv1LdMzDUu2dg4/Ydv1755ENPbnxw1Ybe7f1AgHwelMpgDf1V+L0wD/UwDwGyQXNzaUpnx8zJbXM7WufN6Vo6f86czkkdrc3trcWs2++mNrX6VmM9n23fn8cSAxRBnGh/9qnz3vaq3/nr3355pBZZ1gu8mREExfadweN6edMI7xRmoTIbVgeuvOqc8xbNTdOQG4dUuhqa02ZNe+1LLvjtl781FEV0ebEAJgaoALBGzEqmgVqp+m27+wEANZgzuiNQ/KM+8zIgkIEFMCQTHAEEEgRw2UTeABlVm2c53jjxgBkCQ1y03oiqqcC7eC9k9brNj6xYU2hpaW4uTppULIYu45wTFzgnZPycyEeqVvN+uOZ7B4e27Op9cseu327a/uSaravWbRzs7YcFyDehrV1c4vkzjBZ4JQgzg/fw2j9Y7e/Z9OTKDTBD4No6Ji2dO3PZgumnnzp72YJZsztaWppKhWxICs3MVGEK1gvk7iuBdYROzpgIYcwYrFp0uVddcP5PL3nih7ffy1ze1NOMcGBYD1A5zl3TTEwsCtSUuQiFme3B+19zeSHN5Wok0i+joRHylZedf+s99//80bUoBT4KQBJK8wd76XHCDHX6/Q0AACAASURBVCIjVd20o9sA0o3PEhk7T6tAeCBr4ojPy3FERP2wzjSiAmGwYuvOD3/lmzd/59ZMe2dLS/O0llJzU6mlUGgqForZfCEMM9lMJNo/OFSp1voGhnb3lbfv2L2zZ3e1UoESTpArsq3TJACEmjQqw9M+QpwhK4Q4ZAooFk1gVES+t3/vXb/acdcv70Y2M39W1yVnn37+WUuXLpgxv7OltZjPhXBQB2cQwJKUNkMsh8cGAqQYaFQi6upoeeurLv/5Qw8OV725IGn5hVGj5njqFgizQA2RU68+x/IVF5917qmnpCZXQ5FKV6Mzf/rkN179wvtWbytXFIGAIlCxeHP/6KjCYUAzc66qunHHrr2Rbw1CGnjsZsijRRyzzn2WroAhLUI2u3nv8Mf/89abf3w/ps6tZqx7uNLdp6j2xK5dqMEbqBAPABJAARcgzCDfySLEPAiNFVEVtDEZHHHNDW/IAI4olaRYEtBH0dodA2tvuffG2+6b3pG/+Iz5L7v0/AvPXDitqTVwTrifKBzjeI7Ea+wAhUaBBBefNvO1Lz7/K/91l89lTZwZE4e4RfG5Pm4YnBLMROKJaPGM5ndde0Xg0qmysXA33HDDwZ6TcpyZ1NL00Or169dsYbEJINSLRWJmcYmaRoIkKRr5QOzK55/T2VQQM+xzVU1UFEgMPYN5NW8GYRDWYH9/4zf/7ebbvZvEbIurIUTgc0Vkc8jlkc0hk0U+j3we+SzyeWTyKBSQLyIbQgR0JgHExZ2dSYtP10FGU0eIwEM8aSIIgZDMIJNBNoDD4PDw6rUbb7v3/h/c++j2vuHpHZM6W5pAgtR4q86Msp8ZeVQhoDAQNKqHWSabbWtt+cE9jw5VPBhAAoBQTXrWHJtRHRgLTMQC72v50L/q8rOvu/x5oWvAyJeTmlS6JgAtxXxnS/GWO35VVTBwZirqSdj42xAfbWgmoHq1anXpqbOWzZ7unENc7aexRjo+DDRQoDBTtUgtzDgP/N1NN3/hm7cNDBlzLVJDvqahRxR4xuUtqBSjKGj1OH0jjPC0uOF9fFIUAKGMnXljPlNiDLyIST10RwhSPTAMeJNAmakit3tv9MhvnvjJnXevWrPelZpmT+sMCBN6MzlmiwoDDEoajabw5oJwUlNp456hxx5f5Q0MsvG6IA7tOUajOjAUkCK+Ujlv4bRPvPtNk0rFg70k5ViTStcEgMCU9kkDQ4P33fsomltgnmoErfFSTAQkRQ20WldH6YpzloXOJSGGx3EuOmziTyCAwdTMZUIh/+nbP/nXm366s7ecKeUNEVA1Z95BJbZoLJmw4zKuFsTbfgSY/MwQl8aAcp+8j+c0kUZRoRJGM/FGNapSQAc6OhERqK966x6IVq7bdvfDK1au3WjUObO6sse4spjGok0jRWmKMJspdXTcedfde/qH4EKIxKeLfJqPebQq4zEaLUVrlUprLveu37vy8rNPk2N5llLGRipdE4NM4PLZ4M5HHu8bjOgyTiNCG7LwtgFQiJkR/qrLLmjNhrC4S1ajDfWA1PUG0H2NZUxBGIiIUHEZJT9/2x2f/LdbNu8ccYUinarU4LwG4h0B1H25BIjk/4PYWbfv50Bc1Qn7uVRtPGJiTCpEmZiJGsyooIAZMEh2GS0KzEsQaKZUNenp6X18zYYHV6xcvW6jZQvzp08WADBoDTClVA0kxAA+TS7G4ck8MMkKhgQlLrToOKW1aeO2bavWrK/WwEwOZoh7rMYjoMBIBERAiglF4vWaCZisAkZbBxCg1R0RYx3q6GekWfIiMyNtcOhlLzj7A9e/Ip9NKxY2Iql0TQxITu2YVGrK/eDWu9nS6bQsVKWrpxQ1CDR6mKc5Vekvl59/4fI5HU0OFYDgRNkt8DCvdB70QNzYSzXuFKaUUCE3/vjef/i372/c3od8hgHVFHRknFLGUdcgY7FK/sfqPcJGQxDiqTfOQqvbGeORrvr8nGhtfODYZOH+70JJNNgJwjCKbHfPwENrtz+yatOW7t0tTfnpHW2gA6qRcgQS2H4pzdz/73EM7ekwUSIBYi9l3Ek0SzR3tP7s/kd6e8sMCzCDeNCTBJ1ZQHMiWXrxURXVso1UtVI1VcBEKKRZnBXmIARN4Jmo0FhGS0AAJg1M44c4Lde6OlpvePurl8yddrAjpBwfUumaMATOTWtvu/fx9Zt37EXGKFA6NJYtQ0ANKuagUimX2ztLF502Lx8GE0q6DKDRjToJxdQ0IixwmbLyWz+/728+fdOGLX0oFOHUTAESjnCxJxDPNWs+LXSOSIr4xT8e37fJpzwSpQSwX30KgBIvcUwjaAQQmSwLzQbXvWXn3Y8/+cSGbUO1qL2tta3U5ITQsqv3holVNx7bkczMq/9FAPBT2jseWLPxt+u3+ZohFJiKBU4zNBGh1qo23K/VAUgtmw3yQRiIeDMdHtaRETVitHsACZLJZx/jNvA+x4UazcToTImB/j9981VveMmFgWs4n3xKTCpdE4lCLjtv5rRvfv+ntVyozrHhmh0x3pYQOEBMq317d1154dlTW1qTqXUCEKuVi0WgnoGspHdBZrim37rroY986qZ1G3dkmoqInYjxVNxYn+8pQ4ktsDjSEwDNXDxhNxVhbuOTW3664onuvXtLzcWpbc35UCWuUJ/syNXzwMYmBePFWy1kIEHugcdW9+zuQT4LiFieyCKq+ZG9zoantWcXL5518XlLrjj39Oefs+ScpaecNm9m+6RmF2JouByVawARBBC3n8qOdbCx1CmcwQGEBBgYvOz8BR9++++1N5cO9uqU40aarDCRCJxcfMYpb3z582/8yS81WwLrrvrGgQDFLN6eyK9Zv+Xux9eeMm1yIXCjde/H5RI7bsTpuz62oUiXHa74m+9+5MOf+tr6tbtk8jTTCqmkOzpT+hHGbL9NNLP4qxAzy4pN66gNVb72ndvveWTle1//0ldcdHZXR9E5qPkk6PFIWFzPgkWAwF54zrKzT5u7ZuNm9VUnRRPxFsEPTW4Kzl+y8HdecO4l5581f2pnrv6yCNg1VLl/xcr/vuvhO3/1+Kbtu6sjEXIFiDPUK5GMkcTVSpA0s1qlOYv3v/53Z0xuP9grU44nqdU1wRDhslNm33bvgz3D/ihNJ4eFxPsZjibmaEP9NbXnLTmls6WkoxUijtY8eMRIYtQNaohUg9CVq/5bdz7wkU9/dd367ZjcZQbQgwSTjZsG/1B8KvGPAAttUKzmswUUmvt29d1294Prewe7Zkxrbyo458xM4qSAuMTgwd5lvBhQQ0RFPnA7e/fc/+QTQ4PVIGyyWsW5oXnTWt7/5ld/4O1veMHShZNLxQBRHPQPmMCaMplFM6ZdfdHyBQvn7e3v3bplS7VSpXN0Ah2HO4IAICYCMGPe+ne++doXvfbFFzXlR4UypRFJpWuCQbKQy4Zh7o57f1NVNFaUBgCa0cQEEIMhdBu3bD9r2eIz53bFoWHSeAH9B8JIM1iNlECo9sUf/+qGf/3mpi07Mu0dBrV4dqQchfn8mCL18ugEmMsbw9WrN9x1/yOtnW1zuyaXwtCAJKDi6FiXVURhTSUIi62lh9euXbdum0oprA1dsGzaJ//q3ddceEYpFIsiRECcx6iEN3gFFIgIv3Da5BdeeHYml330sVUjwxWGGWAcViIBUIxCmK/0z5/W/Nm/ePvsKanJ1eik0jXxCJwsmz/znsfWr9vWDWEcZcXYmhnzHXu0oCGWLnPmwEC0r9+BZ506t72l9BS3VSPxVAeTwjypJAI6A/7127f+07/dsnXXkDQ3O1FYZDAwYMMtHMYLIyl4ZsUs0KrzNaEwzPf19N557wNRrbZs4ZxSLgcAlgQvHuR444YRNQQYSXtL6dGe3vseXMUh/9Lzz/zXv/7D02dOFR0hKupQcRJJQARgQAagMzplAK3QDxZymTNOO3XyrBkPrHhisG9AMrmx780l2390Qi2V+z7xv992wZKFaXRG45NK14QkcHLeaXPuefDJ7Xt6LRsE3kIPFfGBUMfj6D/S0EgjQNBoRsJcduO6TXO7Os9buoBCryqkwY7e/smYMcAD6sE4Dl7jlmgKpwrzkOyQ+r/5yi2f/c7Pt/QMWS5jhBogLjZEjvf4jwBxg2YARlGKSWwwu5FK9JvfrHmie+/8U+ZMay5AR+CrcBnUI/EPO8kLAAg4pUigIiLs37P31w/cv2BG2ydvePdp0zuFJB0kIAIHidPiBEkoZeL4pANzNORDd+rMKXNmT7/nsTUDfQMIA7jAIHHuo0CdeZBPHTUJIz1IQ2gDw1desvwdr3rhpObS4X+0lKNNKl0TEgKtTflqrXb3iseqXuny9KKEiTZAJ5T9bnwDgqDWP9C9t++U+TPnTu0gPFkjYDjuJktc5TVZnxOIu2/RKOIg4ZaePX/zmS9/9fs/27m3imwekqRPsfHqbx0Go5fLfiGSBIKg4m3tmnVPbtjYNbV9XlcXxMVhQXyq/XWYJ4ImIOMEAedcS2vzpeeffuXZS+q/ZlLd6hmyAyS+vuTrsFo2CGZMnzpr1owHHnyov6+PAnFBnGYsMNL0AOWqaRRjAF9rC/3X//5P5k3rTGtnTAhS6ZqoiMjpC2Zu3bnrwRVrkW1RBmRN4NFg5eRJIJPduWP7QG3k7NMWtDU1AWUgMLjjO0MYki0eMSPU0cRMQBEH8q7H197wuf+8+Wf39Q1Hks8R+yoOnhDm1nNCpYC5XFSrrVu/YfWmzaVJrcvmzAASj3Q9igXPlJNDgSCgZvlC/rT5sxfNnZUNgjHqBwFAQai5mtdsECyeNaWjKbdizbq9A4OgWGydxZVFnj5YA526HNXaOPi5v3zrRacvCibGXmxKKl0TmTBwi+bM+M3arRvWb7d8zlkUqDc2yr2XBMObIcio2qZtm8jq85YszIZ5r7C4aOyYJqijhY5Kl6mqkQJxFW/fuuO+j37p2z+5Z2XVFVkoOkTxNPeUCL0Tl3jjFKBksgZs3bJ9xbrNYal4xoLZDvSq8e8PqAbjxer7swbLOGnKZAthoKYy1stY41YxSlGjqM+ILT91wbbhymNrNpSHypLJjCruAQYrgWmAkaE3vvDMP7r2qkI2g5QJQipdE5tJpWJ7a+mhx1fu6e93LhQPPe5+uP0wq2+L5Au1wYG1mzY0FXJLFy7OBi7ycMe2+uszMUAMhBkhEkJkU0/fl2/52ce//J0VK9aidTJyBcCgNQohwcmgWwBQr6JhBmYyCMKe7d0PrXqiqZifP392PnCkxU48HLZ0ASAsDukkFIgAlXFkyyW5gvGfoXOAbO7uvXPVhhW/XV+pemZysXRx39PqkGqG8vDzTp3+L392/ZS2VqRMHFLpmtiQ7GhtilC798FHIhTFst41YgNlEszmB3tHVqzZ2jJp0mkzp+SzwXFXgXqVV1JcRfXXG7b88zdv/dS/f6975x43eTIDBx/B1CCkO/5hJccKQhJfqpmZInDI5Ab7Bn716ONNk1oXzZpaCDPJrt8ROCUGKJP6WfEmohBjrRlmyTaYCszRDXt/35qt//er3//mj+7qL1ddrkkZGCCmgtHyuvXXUjA8vGBa6RPvfcPyhXOf7S1SGpNUuiY8hWxmclvr2l3dq1ZulVxJJar/5ohEgR0+JNSxRoQWNA/0lu/9zWNtrYU5XdOKuRAAYr/ivjrlzxzzIXyQ/aIP9vt3YkoASIrhEgKQu/uHvv+rB//P52669cf3RpJne6cIRKtmkRFgCDo2QADMsYH1BmA0RfwQh2xhuH/k3gce6OxoWTBzRiEbgnokAkXjRPWkDHH98Vzf9/5frZmNRo70Dg3fdu+j//szX//pXY9UjS5XJIO4KIrAi+lTfem0yHfkg79480uvveLCMW6tpTQOqXSdCLS3NC2aM+u2O+/ZMzLCAHHIc7yPQHNJ6NgxvzeTPFZK4loyQhwyrjw8fOf9j5WzpSWzJhdyGYFBIyBSq/cWYaxjBgNGXY5jnFyS548+AMCICoh4Uoy8RTUgIg1wlSh6onvgH79+69998VsbN2xD8yQWihitxkBHEnKSzWzcd76NjNPXSCKbrQ5V73rosbapnafNnZmjWa1CFwJJUeFkZTCeq60eOhhX3I/fOf75MzDAoIAn/GhNsahKRMpwa9/QF/77zg988qubtnWzVJJMFnRmcYoeYvNM6gsYGsR8pjJ8/Ssue/urXtRUyD/z3VIanFS6TgQItJYKs2ZO+cntd4wEOQDxuhlx63iD8oCTwVGnnnwTr6mRNBoUVzX3qwceW79148w502d0tFGokaeC6hBPlaxPZRznrkqSV8b9H3ERXRdX+FBVQlyGEvaUa9+87c6//tRXb73zgbKGaGqGJPVAAFhc1eqIeMUmHrEpnMTnEUxWErlCbWD44ZWrZ0ztXL5gDoNkUyr5uoBxLDISuO8I+z0O/MT6G1m8soGKI11m1eadH/jcf37267eNeBeUCnTxE+PVG1i32sVU6ZQBKb5/z/LTZr7ntS9dPGfGAd8tpcFJpesEIXCus7XFOa5avWE43uWmI0QMhNpxkq5nQsRl/2DOrduw7vaHVwyoLZw7uzmXowT0ZN1jZMmzk22XMQ5fAQ8mOzX7pkVzFgkACp2Ic72m37vrVx/9wtduvOXHm3oHfKaETBgnJdPiVKGUp2GJRZXJlvf0/vbJtdO6pi6ePR1EvYtxbBwf7DCHSgQoIQYqnFfQAheQcst9Kz7wiRt/8suHfaGJ2XxitT8jRYSAQ6R05vJWrZUYffGDb7/0rMUn6cpk4pNK1wkCgUIuM2tqxz0rN27cuJmZDBmYUUwJMx7FOWVcECBpZuIkcvk9veVHHl756988rsCsmV35nIPEehURnhbHYZOjW1Q4mBVkFm/aC4xQmhcozVNCUAAMAN+9+6G//cI3brz5R49t2DGooWYLCDNGwkyS9ruNcbIaiXghYBSIkzC7e9uuDd175s2ZMW9KO6imPjlnR+3URTAP79QYASaScYM68vFv3faxL/7XY09ur+WbGeYsaX52AAgIIhNncNjZ/fG/eusrLz0ndGONB0lpNGijU0LKxMcMT2zddc17PvL49j3ItxiE3gvMEhdKQ0DEQfMUy1ikvjIEjsycXFi6eM7F5y773XPPXDprOvbFmEUwr0YgjL1YByngqzWgCjODozgwHP3gG3f2/uy+h37060fvX7Vu0869qoJCHmEGJGJ9t0jiqkgNltbdCBCAiY+T8UxtZERGyi+5+Ix/eOd1y+bNqmkU0gAB3FG60GqIIotCdYELAazcvOOTX7vpu3f8ZvceQbaJmSAx06GE54GmNcea0tmO7vf8wTXvf+PV0yY1PfM5KROFVLpONNTs3kdXX3/Dp57oLiPfDO8JTfrONxguMkC8OPU11AaASmtb6YxZ05fNn73s1FPPWjzvlDkzWsYtIgboqPKNAJu27161ZvODj697aNVvf7N+8+buvfBEtonZHASJs6sez4HYQ5neFAdAaOJjRyxUSB0YyaJ83UvP/+if/OH0prxZFDuCj5J0qdVgKpKtAf/9y4e+8O3/uePBFdXIodAMcTAPGEXiOJ0DDIGEc7ny3t89f/EN73zzwplTTrLYmxONVLpOQGqR/8aPf/FXn/vW1t4q8kWoF+pRmlAOA4NVjTRkIVkIUS1juB8Vj0x2SkfLovlTl5wyc9HcGbM7W+e0t86e3NHUVBpLvvVgpDt2927Z2bN+d9+azTtXrF73xNpt67f3VgcGkM+iWGKQoQoVtAiIAIAwUikWm2CmB3uTkw9zhPP0iM+dASo62F8q8p2v/90PvuXaghgAUx6+wzUJBNz3D6v/m9v39H/ltp9/5ZY7nli7E/lO5kJDGVpj7IsWwmSfd7MeoGFxXWhvZ0xv+vSfvvHCM05zacGnCU4qXScmvQNDX/ivH336mz/aMZS1TF5sOG5kziSPhoDUt9aPF6aMQNICIABEzJxpRNGqR3kAlWEEKLU2TeucNH9ax2mzZ0+ZMmVq56TWpmxTzuUDCYQiMCBSq6oN16LhcrRjz8Cm7r71G7et3bh1y47dO3v7Ua5AHIolhBk4giQIMzERRPH5MMTSFU+AdkB308mOCUGlgpbk4BnMOxsantyBj/7Z9W+57PnOEBmUCHjoLleL6/nH14SpeU8SLlDggTWb/v2WH//7LbcPDXpM6hRmaTVlFdCk/gdjq9ApROCd1cRMKREDq/mpRf7bDX981XlLDzaElAlAKl0nLOu37/zU17/3xe89MBw2U2qwKFnGEvaUZelxJB7RPgU1EKZmFodUwCsqVVRrqHkAyGakkGkthC3FoJgJsiGdg4IVj5HIBsvVoXK1f6iCoRoigxBhgGwGYVwII/EKkvsHehyedXDSYU87Y2IZq9Fb9/zZxa996C/Pmz/bCyJYCDnk+AcDIsADGRjVw3uG2YFy5c5HV/3jN75/1533I9eOphaoFy2TNAkAt/8izCBgQFNnNUHkgahmk5vyf/e2l73lZS9M/YQnBql0ncg8tm7TB7/47VvueBildiPiVl5EXCyuQb94M3vKNcnkP9DDDLUqoggaPwATQOAcECIMIQ7OMRiNBMEzL+8kVSnlSOCMDm7ED6Oy55WXXfTJ97xl1uRmNU+6sRdzOgCmUIXQKAB3Dwx9+ye//ORXvvPkuq3omCqZJlWFmdgISUsicZ7yRTtTZcaLA4TVoUlh9N7XXfXnb3p5Jg0pPFEIDvaElAnM0nmzPvK/XrN6y8bfru9HvtVoMBhNTMUMScpUY2MAjAKKqIKZInPOzIyAGpJ9FUmiBCUgImqVpoyLq6YcXdRMJcgrwu/e/tB5Sxf/r5dd2lrKmHqTQ25qY9Aa1EOyqvZk78A/33Tz1/7r1qHBCFPnkIFCQINpvYbH079lMRWLPEMwi8gXvP+DV7/g3dddlerWiURqdZ3gmNndq554+Ts/0jsSukLJq4BGrREKISZI7q2oOQNAFTE4gGZirEFqBGA01hOQLfaCPt21lXI0ECjNIskbMxjZO7PDffWD77x0+WLANDbxx2zj1q1tkgaNSELCu1Zv+PBnv/aLhx6vupDZJkOAZLllVO8QAXxmA0kxL9CIoSIne3uvf8XzP/S2V8/onHSgt02ZqEyMmSvlkCF5/qkLvv3x/91q/TYyIGJJydQxzSeNAwEhhAoxE1OHSKC02PGpYt5Z5KwqcW+tifbxJihKKl3gzUURctmtm3d+5js/Wbe1G8kVNu5vgaY00GUGJPzCrXe9/YOfuOvh1dWwJLkW0BE+TjOneVKVcqDGx/AUDycwdm987ZXL3/u6q7ra04YmJxqpdJ34hM5dsGTh5//2fcXagEYV1K0SJJPLoUwxxxgVqQZSDRgFFkkUSS2SqgqUoTJUCZWB0ildvDtysOOlHBmMNGHGRzlfEUZaLN12+6M/unuVeVBohvEskUiSzkFkQ8+e933qxhs++x+rdwzUgiYGecCJeWcVZxWnVWcRzBSicAeYxEy8y3Nv38svX/7u637nlJlTZCxJFSkTirQQ1ElBGLiZUzrmzpl+5y/uH/FOgpzQmUVALe4zK4SRYIPe4rZPZA37BUjW647DkiTUsTmnUo4QBEDQ4lhAsyCs7R3atmdwycLZs6e0EgMRrcqAsTmFpCJ/feVkgPnYv6tVokYJAfzgoVUf/Ndv/PBnD+wZrCKTlyCMsxloPrGnR53DEFDBGgmaEEIjIXBifT3LFkz9q+tfdd6yRdkw3dE/AUm/1JOFlmL+FZecVy5XP/j57/RUqi5bMK2OxmUZzWAAj3e4/IGh7V/hYn8zUYH69kfyq5RjiAGGmgMhqgojmouPrFj1nz++Z8ncyW1NGdNICQWcedBhX0yFGTxhUJqZkJBc38jI1370i8/d/OPHV25CpsBCycxgUXxVKoCnRi0SRnijhzEu9gGYwXSgb1p79rMfeMf5i+YHLnUsnZikVtdJRC4Tzp85ta0p99CqJwaGKshkACapn4yrtWs696ccGknAVxBgZGTbnj1zpneePn+uAKQTqDCRLgEIrT+M5iUIIW7l1t2f/tZtn/3q99at24qmVuRz8WLluR0BBACBOSKgEFbzlYGprbmvffRPLz791LRkxglMKl0nF/lsZv7MadmMrHhy/dBQVcI8IKSABNTBT4itr5QGJJEuA/LZgd3dNdjlZy8tFQrOR2CkNMS1TKBABFMDSUdxkfHnK9d97Ks/+Lebfz64d0Q62hmGpkktrucOUKSJoJ6SrFVfHZjenvvcB/74d85dNsbIxpQJSipdJx2FXHbejKnZbPDEkxv7hysuDEAXNyQWptKVcugkkfCZAKZ79g42FYoXLlsAXzX4iI4M4vw7mHqDMTAGewZGfnjP/R/41FfvvHsFci2Y1AYYTZODHSy2nnACMcC0atWBRdNb//l9b33ZhWenV/AJTypdJyNNhdzCWV2FXLjisdVlr+ZCI0ATREAaoZdyKIwqjZghCIf6hkci/+KLz24uZE3VkKFI3ETNjMqgCreht+/zt9x+w2e+vHHzbrR2IJ+DRYyDcsagWwCIuN6u11r/qdNb//k9b7n6onPTy/dkIJWuk5RSPjtv+uR8IfPAgytGGJpI3Ogola6Uw8FApz4EowgjPpo0KXfBovmiUHUiTGIF6SLDgxt2fPgzX/vif/18hFkpNTtnolUgqtd4eY6LcJ9jgCQEWh2c3lH8x/dc/7Lnn/vsr0o5oUil6+SlmM/Onzmtpbl07933V8McKCBpxiRlebQ4bypmKWNHaB6m6sLhkZH+gT1XXnJOczbnrEqrgUaGFfA79zzyl//w+XseXKmFSczkhKBFSccUyAHTmQ0CUKAOkYM30uBMRPv7ulpzN37oXVc/b3l6mZ48pNJ18kKymM/OnzFlytS2u372y1qYh2QAI5S0/QpupBNCylghTAmNg/6qUWW4Mmf2zLPmdpEVUin5rUPlj37tex/7f1/ftK3H8k0UIQxqBmcIe721LgAAIABJREFU4ipfB77kSJACT3iBGQmXsb7eyZMy3/zY+y9dflpaEv6kIq1heLJjQHdf/w9/+eD7/u+X91gBuSyhgCJO740zSYFUwFLGiEEtblRdi0KNrjx/0Vf+9r3tuQCQX6/Z+Mn/+K9b7/zV3rIh2yyBe2oVzfq1diAIBaiMtU0dvOztPXvJ3M/81TvOXDA71a2TjTQl+WSHwOTW5ldeel5LqfhHf//lnQMDVmwCXRztLIyrGEi6wEkZNy5T89GvH37swSc3XrrslJvvevhfv/G9Xz/626oFLDRBAlP/VMV5rqvMmQfMMwuXha+yd/O1Lzn/LS9/0enzZqW6dRKSWl0pCcOV6i8eXv3nn77pN09ulZZJ5kLzJojEIgMtLXeZMjaMCtDgYA61kbA28JqrXjB/Ztd3fnzP46vWIldiscUAmI8Tkw92vARnVUJrrmQ1CYf73vKyC972qiuWnzpf0rzjk5JUulL2EXm9f9WaD3z2ptvvX8WmTsvmxFdT6UoZFwQAejqAVO+sVgysUh0ZKUfIFhEWAAEhFoH63JbW/hAe5n0lKhLvuPZFb335ZYtmdx3sRSknLKl0pTydVRu3/OWnb7rljkfR3M5ARCPAUulKGSM0AlCKgSDFVCvDgEfWAY4+JAISRATxOsYUeNIMVh6e3OTee91L3nT1C7s60j4mJzWpdKUcgE3de/7P575+y50P7LUswyy1ZnGBg2SiMYx9tZxykhFLl4FKkqACYiTURkwhFog5UgBvos8iXRa3DCVgpFHMxCqVBZNL73/z1a9/8fNLuewz3zflpCKVrpQDMzhc/sgXvvqdu36zfigDqdKPiIa0kBDAG70yNcVSDgphNhaz6imYp3rGpeVdGEnoIiycXPynd1/3ogvOcg3amSflmJJKV8qzsqu377t3P/KRz/3ntoo352gh1dEIeNCDOv45KSXloBjM1ICkJqK2YmTxgtmfeN/15y6cfbDXppwspNKV8qyY2UC58tDqtX/wt59fu2uQ+ZIhhBmSjhX+YAdISTkUzMwgYAhFfmDXH77m8jddffkZp8xJm2+ljJJKV8pBGKnWtuze846Pffmnv3wYTZMQZmlG2jOSclJSDh0zi+ciAk5IcbURnxse+OifvfHVl5w1fXJb4J7SZzLlJCeVrpSDo2Z9g+Uv3Pyjf/7a97v3VqS1HQqYR6pdKUeIfdIlpI+0f+/0KW1f+tA7zz1tbltzMb3OUp5GKl0pY2VHz57b7nv4X75926OPrkNrh0gAYF/H9pSUQyEOWDVTNQroUB4pSvS7ly6/4Q+vnds1JROkxlbKAUilK2Uc9A0Or9i49f9989b/+O7taG5lUzO9UkmYwYwKwiCpnKU8jdFGJjTbl1hBwMxEAGfiLKqhr3ferM53vPrFVz7vjGXzZz7XEVNOblLpShk3W3v23HjLzz733R9v7xtirtn5EIAx8i4CAIujEFNS9mEgQAK0/SpokEBNmYHkLVL0777srNnve+PLz192akdL03MeL+VkJ5WulENhuFq78b9/+j+/vP9HD6z2LDCTV9Ko/5+9+46OKzsPBH/vC5UDUMg5AwQIEsy5GbrJZneL6qw8CpaTrLE147jrnbV3ZlfaGWuPPZ5dWfbYsoLldkutltSZTTbJbmaCASByziigco4v3Lt/FIhGkwCq3kNVoQq4v6PDQzVvFYmq9953w3e/CwGCS3rVBLFEbEPWwhNnIZhRCCEAwpFCg/rZo21/+MVntlSWknK6RFwkdBHy9Y/PfO/VN1670eEMY0AbAFBSWAQgSmIX8SgIMAAQARoDGoAHq6QipoTw3irjS0/s/8bnPq1Xq1Z/E4KIIaGLWJMwx//4/Yt//dM3Z+wcD4w0BJgKA4gBoPCDawuSTjQBAEQiBhSiGABoADCFeQYjBlC7m8r+9g+/1FpTrmDJGUxEokjoIpKge3L2O//02lsX7nGMGqoUgAIAYbQkdJHotdlhjJGIIQ0AAyFFoSiIereUmr5w5slvfua0UaMkk4SEJCR0EUkgiKI/HH3/xt1v/td/9PgAMOVSAGFRJKGLWIBj6e8MoBUgGoZB10snd/7Z1z6zpapcp1aRq4OQioQuImkEEU3YnN/+4Tuv/uIcr1EBrQ4gBDCAEECIwcd7wBYeVCQPcaNafKbEvmGIEQQQQYh4BNzeyrqi/+sbn33usZ1qliWThIQ8JHQRyYQwtrh8Y1Ozf/r91zq7Jji1HrIKCiKAIxgjiCkKUBBCiCEAUKARufiyWux0kwe/W4xYCACAAYUpBgIKIoHFAsICHxRBKPC/feO5Fx4/WF9RmqNTL/+mBJEAErqI5BNF8UbvSNeY+ds/fMM6bQcGPaVkYyd+UYACEEIMKQBFSsQkFzGbwQfjKvzxr5DCPA0EBGiBUgHIYlEEwQAI+k8c3fNnX3vmYEutUatZ+S0JIiEkdBGp4g1FxmbnXz179YfvXHe5/ECnAwoFwDBWVwECALFI0uiz3dJJXwgAABTEAoQigjTCNApF6Gh0x/bGr3z62MGWqt1NVSQdg0gKErqI1HJ4fJc6Bl+/1P7+nW6/2w+0esgqY9tRKURCV9bDS8ZeEGGAMWAUIlTgSAREfVsq87/45MG925pO7NmqJNUIieQhoYtIhzGr6xfnr1y93XlreMYdEjGjBDRLkbiV5TBYGEVDDCiMAcJAFAWoAIgtzlGe3Fn1uy8e391cr1Yq4r0TQUhDQheRJgjjaYv9n3757sU7/b1mbxghDBhIkcMDsxiGC2miEEMaY8yLSMC5Kmbvtqovnzn26UM7DRplvPcgCDlI6CLSKsJxAzO2v//VxbvdQ4NmVwQASFFg4eQLIotAAACGGFMAAABFQCGR5YSqiopvvvTYb545rFWRkk5ECpHQRaQbwpjjhYlp8/feuvaDX1/iEA+1GgAhRpjCEACAIAQAAYgBxACKAGCIWIhYEt7SA0EGAEDjKIN5CAACtABpjFmAGEyJgBYhECAWaFFElFJkNIALgYBrX2PZ15879flnHteplBRFsjGI1CKhi1gfgihiAMbnHd//xbv/8PO3OYEB+nzAqABGAIkAiiCWQQ8QABhgGmAytZgmCDJwocYgDwAQKUYELC0iSsQ8gzEFAQ0BgJAXcTAMoqFt2+v//UunXj6xT8UwGhUpjUGkAwldxHoSEYpwvMXt/fHZaz9867ptyi6oVUCpBBBCiClMAUABiAEQMRQ+mYlNpAqGFAQYYAQBxgAgqACQplGEEsMizSBIARECfxiEhMOHWv78a0/va21iaSpHR3ZrEelDQheREeacnpEZy7zN9Tevn+u6P8KJAOgMlEINEAMwAJAHMIoxgPDjK5bURUwZDADAGGNAA4oCkAYAQogBiuBwGPh8NBa+/OLp33jhdG1Zfp5BSxIIifQjoYvIIKFI9O7Q+JTD88Gtzjc+vO23+4E2B6oNFIUx4AAGGOOloYtEr5TAIsIAQAggTVE0BZDARUCYpzFoaix/9lDrC8f2lBfmFeflUhT5/In1QUIXkXGigjA2Zx2bmb/VM/zWtfu9g9OAYoFKDRUshACJKFYlD0II4dIFsKU1iYilVvpk8LJzsBiLGEAAGYAQiIZAxG/I1Zzau+PE3u0VhfkHW+sLcvSPvoog0omELiJzeYPh7tGp9t6R8zd7e2edlnkrZhhAs4ChAaQgjB28CwDGAC59PD/0ON7UIQ0+iE/LfTL4wa+L/x0CCJAoAE4AnKDUa+qLjPu2Vh7f3bKrsbaxslTBkELvREYgoYvIdBiAoWnLpY7+nsHRG10j8/6IKyiKiAIUoCgMAQIYA4ghEAHAGFIYsHhho9hCfSIAAcQozl+zAWEIMAUwAAgAiAGLFmMYogAFAIAACBBiACDAGCOMRQQwYCAq0KnbakqbGyqP7mo+vqs5R0s2aRGZhYQuImtwgtg9NnW+vefOwET/6Oy8OxTgRAgxpGhEYYgRAPDBE/lB0IoFsM254xkiCuPYwBQulGqiAIAIg4WlLAABDQDEUBSxILI00DK0Uat+cl/L/q11pw62lZgMLE0KDxKZiIQuIvt4A6EJs/mty3dfv9wzPDEnUBRiVZhWAMgCIFA4CoEIAIQAQ4ABxgBAEW6mmS6IAY4d7gkxgABQEGAKCbEwhgFAGIHYXCKtArxIc+FcneLIrqaju1uf2L+tpaKY7CgmMhwJXUT2iU1uYYw5EXeNTP7qYvsHd/q6BqYAYIFKDRQsoMCDAnuxYQfEAC9NrN/Q8EJ2O6AwpAGkwMIYCwOMAMYA8UAUAMeBaCQ/z7Cvtf7Msb3PHNxZmKOnKMjSDMkbJDLfJrmZiQ2LE4RwhKMZemLeca93+GcX2s9d7QUIA5oBaiVUqiiKQRhBHAUUhLGsjgevXUy1z9I8+2W3CmAMMI5FL0gxNIUBxgBTNBIxCIdBNAJQtLg0/8T+7U8d2XVyR5OCphQMo1OrSMQisggJXcQGISIUCIUxgDZfqHds5m7v8MXbfZ2dA7wAAUMDBQPUaopl4EL8evDg33ihC1IYAwBEgEQgCiAUBBwPotGi6rITu7cdaGtqa6qpKysQBcFk1OtVyiz8uQmChC5iI4pwvMPjc3oDURHdH57sn7S0D89294+FHC5A04BlgYIFLAso6kFaOIQL6fUY4Fi2/eITHT9I+nhI6nI/Ft82TlRZiFwUBBgsFG8SYnntIuAjkBKKqop3NVRtq6880FpfqNcZtZrSQpNeq2bIWTNEliOhi9jgQlHO6QuOzjkpgCfNlp7RmfF5x9C0dWhkSuSjgIKApoFSDRgWMAygaIgwxAhSECMEIH6QOg4wBiBWCxhiiBEEGAAaQWaZGJboQAYvREoAl/5Kxf4WsJBpASgYW6uLRTIIKQgpgABGIhIRQBgIPBB4gEQgCPoi0/baii1l+TVlRU21pRSAlcX5pQWmAqOepUm4IjYOErqITUREyOULztrdTm8gFOXuDI3PWWxWh9fqDXj8Easn4Hf7AaQAhIBhAAUBBQGkKIaOnb8CAAZABFCEGFEYYUAj8CB0LbmNYKJBAmOAPg5amFpICoztQwMAQQQWBoMYIBFgEQgiiOW28wgArNaqK0pMeTpNQa62JN+0raGyyGQ0alQNZYX5uUaNSkkSBYmNioQuYvMSROT2hyxOt8XumrE4BEG0OD33x2dtnoDXFwzwYpgTvIFwNCpghAGAIJbIQENAQUBTC1OLMUuCBJXYdBwEsd3UsW1WAGAAUGxwRwNEASyC2DZqjAENlQrKoFfoVQqDUpGr1+QZDSVFeTsaKtUULi4wleSbikw5OTotCVXEJkFCF0F8DGHs8YdmHC6z1WF1+W1uLwWp83f6J83WKBfheV7gkIApHkKBggJaiFw4tjb24FZKMHSBWOjC4EESPwAY0ZBiIGQBZCnAsLRCoVCxTK5W+9yJXSolYzJoK4ryKoryC3IMOrWSLFkRmxYJXQQRx4jD1z9l9vs8Fovd7nDbfJzNH7L5vS5/NMoDJGKMEEJIRCIAUo5iwRgBACHFUJBiKAVLMxTUarVFBk2JSV2QoyvKzS0rKirKMRi0ul11RXTi70wQGx0JXQRBEESWIRMOBEEQRJYhoYsgCILIMiR0EQRBEFmGhC6CIAgiy5DQRRAEQWQZEroIgiCILENCF0EQBJFlSOgiCIIgsgwJXQRBEESWIaGLIAiCyDIkdBEEQRBZhoQugiAIIsuQ0EUQBEFkGRK6CIIgiCxDQhdBEASRZUjoIgiCILIMCV0EQRBEliGhiyAIgsgyJHQRBEEQWYaELoIgCCLLkNBFEARBZBkSugiCIIgsQ0IXQRAEkWVI6CIIgiCyDAldBEEQRJYhoYsgCILIMiR0EQRBEFmGhC6CIAgiy5DQRRAEQWQZEroIgiCILENCF0EQBJFlSOgiCIIgsgwJXQRBEESWIaGLIAiCyDIkdBEEQRBZhoQugiAIIssw8RoQBJGhMMZhjg9GuAjHRznB6vH1js+NzztsHr8nEPIGIxwvQAiUClatYPUa1dHt9c/sby3Ny2Fo0mclshvEGMdrQxBEpuAFURCRIIoIY7PDc3tg8tbARP/kvNnpcftD3mB4ldeW5hlP7Gz67u+8WGwyrNKMIDIfGXURRNbwhyK3+ifuDE11j5vnnZ4wx4+a7b5QJMEO6LzL5/KHUGKNCSKTkVEXQWQubzA8bXOPmm1DM9Ypq8sTCM07vR0j07yAEEYYAxGheO/xCc2Vxdtqy8ryc8ryc47vaNxZXxHvFQSRiUjoIohEXekevT044faHVAq2pjhva01pc2WxSsHGe50EZodnfN4xZrabnZ5wlI/ywoedQ1a3L8IJUZ4XERYR4gUx3tusSMEwFAVZhlYwdF1pwTefO/b5E3tYho73OoLILGTCkCDimLG7B6bmR2ZtH94fvtk/Ho7yNEXpNcpcnaa5quT03hYFQ+fqNbl6rV6j1KtVRq1ap1au8oa8IAbC0UAkGoxEA+Goxx/yBiMRjveHI+0Dkx3D095gOBTlYiMqT2C15SupOEEAAEQ4HgDg9AUBAIU5+tN7W+K8jCAyDBl1EcQyBBHN2N3zTu+cw/NR1/C1nlGnL+gJhENR7qGW1cV5NEVplKxWpVQqWJWCUSsVKgXD0DQEgKYpCkIIIcYYYYwQFkTECUKUFyIcz/EixwvBSDTC8bH8C6vbH4su6aFVKb714uPf/vqz8RoSRGYhoy6CWBCMcG5/yB+O2D2Bs7d7h2asNrffEwhN21zByMMRa9GkxbnSH2U+lqGVLHkIENmHXLXEJiWIKMLxEY6PcII/HOmbmLs7PGV1+92BkMsbvN43Fu8NNoL60sIjrXXxWhFExiGhi9jIMMYiwoIoxn7lBTHCCXavf3jG6vKHPIGQOxDyBSMuf/Du0NSU1RXv/TaaPU1V+1tq4rUiiIxDQtcmxQtibH9PbLETQgghgABSFKSpbCq1gDAWRRRbRkIYiwiJIuJF0RMIW1y+YDgSjHCBcDQU5fzhiNsfcvlD4/OODzuH4r3xxldsMrTVlWmUingNCSLjkNC1uWCMBRHNu7wjszZ/KCIijDACAFCQUjC0RqXQqVVGnSpHq8nRqSkKUpCCEEAIIQAQwnhvn0IYYxz7FcfGUigc5e3egD8U8QbD/lAkEI76QhFPIOTwBrzBiN3rf+dmT7x33bxUCvZTB7Yd39EUryFBZCISujYXbzDyzs3u7791uXNkFgCAwcf5pRAsRCYIQZ5B+9XTB016TUmesTTPaDLoik2GfIN2vaKXPxSxewNuf8jtD7r8IYvLZ/f6x+cdv756HwAQS5KN/Sx4YSi56tsRABSbDE/tbWkoK4jXkCAyEUmO31xe++jeX716rmdiLm4VBp1aSUHI0DRNQYqi2urKj+9oqC7KqyrOKzEZKwpzqZSFMYTxvNPr8AZsbv+0zWX3Biwu33vtvYFwVEQIISwiJCIUy7OI92bEMqqKTP/pS09/6eS+5O6nJoi0IaFrE5m0OP/g//v5uTv9cePWoxiaMmrVKgWrVrIGjerMwe1bq0pqSwtK8gwlJmO8Vydq2ua6Mzg1Ome73js+MmvlBDG24YkXRF8oEu/VREL2NFb93rNHnzvclqvXxGtLEBmKTBhuIgNT88OzVhlxCwAgiChWfCFmyurK0WmMWlVbXcXBlpqtNaW1Jfn5Rp3soRjGuHN09ldXO9++0W3z+B3eAKkSm1yHttbm6DQN5YUvPrZzZ32FVkWyM4gsRkLXZhEIRwemLcHwiltrJXH6grFI1jc5/+H9oeJcw5aq4tN7WlprSrdUFssIYLwo/vJKxysXbs/Y3fHaEvGZ9NqiXH2+UVeSZ8w36krzjIda6/INumKTIc+gjfdqYj1FOH5oxurwBsbm7HeHpwEAFITH2xobyguLcg3lBTnx3mBTIKFrs/AEwhc6BpeOnJIiyguTFuekxXlrYOLe0FRFoam5qvjMge1tdWVGrTreqxcgjM0Ob+/k3KzDE68tsbyKglyNSmHUqo06tVGr3r+lur6ssDBXX5pnzDfqSAZ8tjA7PO+195693We2e8wOz7zLG/vvFzsGq4vzmqtKTuxo3N9cU5qXtFn6LEVCV3YIRjhvMBzheBEhCKCCpTVKhUGrUjCJfoNalQJhLIjyi47H1TMx1zMx9157743e8cOtdSd3bdnVWJlgH18QRZamlSxDMi8SkavXaFUKrUqpUyt1KmV1Sd6B5prCHH1Brr7EZCww6vQaVbz3IDKOwxv43hsf/eKjexOPVBcbn3eMzzuudI98dH/oyT0tv//88erivGXfZJNI9MFHpBnCOBCOhiLcnNM7ZXWOzznG5x0ufzDKCxSEBq2qKNdQU5zXWF5UU5Jv1KpzdHGGOLl6zZ997kl/MHJ3eEoQ5Sx3Je5m//jN/vFfX7v/1dMHnt7XWleav/oIjIKwvrSgobzw3J3+VZptWioFa9SqFQytVDAqllUp2Cf3NlcV5ZXmGUvyjBUFubl6TXZtJF8XnCC4fKFpm8sfitAUpWCZisJclqZVCtagVcmY5U66S51DPzl3y+r2rdRAEFHf5PzwrI0TxG9//VnDJu6gkNCVQThB4HiRE4RwlB+asXaMTI+a7e+195pXnkZTKdjTe1v2NlU9vb+1oiBXp1auUk31+I7Gbzx79K9ePTc4Y01DZunYnP2vXj33/u2+Zw+1ffGJvXkG7Sqp2Ajjw1vrfvHRvSmrazMnaEAINUqWpiiGphmaoihIQXisrfFwa11hjr6iMLe8IFevVmrVykx41GYRXhBvD069eb3rlQu3F2PDFx7fq1Mrq4pMR7bVN1UUGTSqddwtgDGesDgSyaLiBfH1yx1Ht9W/dHTnem21XHckdK2n2EEYsfIQ/nDk7tDU7cHJgWmL3RO40j2SyImCEY5/83rXOzd73rrRXVda8PLRnSd2NunUymX74BSEj+9s6pkwz7u8yT0FaiXBCHejb/z+6Oy13tGvnDrw7KHtDE0te7PRFPXEri1fOrn/79+6nPQFuUxGU4v1SiAAoLYk//S+luJcQ21Jfnlhbo5OnaPVGLVqtXLdHqkbw52hqX985+pbN7r9S3ZZvHrpTuw3VUWm7XXlnz++59nD29drXTAU5X3BCEIJ9duiPH97aPLp/a2bNlOUhK71NGl19Yybu8dnJy0ubzA8a3ffH52VcXC7iNC94emOkemhGcs7t3qfO7T9id1blr0DS/OMn9rf+tH94XvD04/+aYqEotxH94fdvtCtgYnfeuZwU0XRss3USraponAdu73r4utPH6otyS/Lz8k36gxaVZ5BV5ij16jYzCnBtQFgjM/f7T93pz8Qji7bYNrmNjs8IzM2bzD8+cf3JJ5hlCwY42mb641rXe5AKF5bAADwBiO/utL5Ry+f1CjZzXl5kNCVbvfHZkdmrbFYNe/y3ugds7r9sWK4gijKXoWKhbrucfPgtLVvcm5s3vH84bZHF3IhhAdaav/Tl57+yflbfZNzo2b7cm+WfIFw9PbQZM+EedLifP5w22eO71o2weTg1rondm159dKdREac2WV/c3VtSUFRrr4o15CjU+vUylydJkevqSnO16oUi9ODZMkq6TDG3ePmX3zU4fQFV5onxxgLIh6etf7yaudj2+vTH7oAAPNObyASSXAmH2Mc5nh/OFKYq9+MgYuErvQYnLGMzNisHr/N7b8/NnNncCoU5WKBaqVuoDyCiASRuzc8zQtiY3nhsjlISpY5vbelpiT/Z5fu/vyju1NWV4J3yxrFimK8195rc/vs3sBvPHXw0QdEZWHuN587RlHw5x/eDUezPtUw36h7/nBbbWm+Xq1qriyuKs6LLUayDM3QFEvTm7O/nGYQwvN3B+acnrjXOcLY6Q0kOO5JLgzA2Jyd4yX02DAGDm+gtiQfbMqriISuVJm2uWbsbpc3aHH73rnVMzJrC0aiEU4IhKOpzv8WEZq0OgemLMfaGpedClcp2O21ZZwgfHCvP82HVEU4/vbg1Pi8Y9Li/O1PHdlaXbL0T2mK2t1Y+b98/nRloem9W70D0/OrHE+c+UrzjH/82VNl+TkKhl5pkY9IKRGh8XlHx8g0l9g4HkKY4GpT0lndPl7K3hURIbsnEC8cb1gkdCXZjN1t9wT6Judev9Jh9/h9wUgoyqU5PAAAgmHuo67h4zsadzdWrtSmoiB3XRaWOEGYc3p/8N41py/4e88e3dtUxTL04p9SEDaWF/7BC8cPtdS+2977YeeQwxewuFZMF5atxGTUa5RKlmFoOnbOvVLBqhQMBWEswQ8AoFUrzA7PvaFpl19O5gjL0I3lhfFaESnEC+LZ9r72gQmOF+K1BQAAhLGQWJBLLgpCXygiSlkvEBGad3oRRgBsxklmErqSAGHsC0bsXv/FjsFrPWNWt29szp7+cLUUJwhdY7O9E3OrhC5BRBqVcqU/TbVwlP/F5XvuQOi3nzl8em/LQ0HUpNee2tO8par4yT3Ng9OWX1+7b3X7w1Fuzuld6Q2XFSsZzNAUy9AqltWqFVqVQq1UaJSKEzsaS/NzNEoFy9AMTasUjFqpUCkYmqJoiopFU61Kcb137C9++JaM0KVkmbpScqTIesIYW93+qz2jid+MHC+kf6CPMfaHo4FwVNK2EFFE7kBIFPHmfIpvyh86eThBcHiDnSPT7QOTQ7PW92/3JXftai3CUX7K6uQEYaWKGxQFtSoFBEDC7ZJUvCCebe/1BcKcIDy9r1WnfjiOVhTkVhTkPrFry2PbG8bn7N5g+FLn8PCsNRThorwQ5QWEUOz8SQpCiqIYmlIwtIJlYvtMlSxzfEdjYa5epWBVLKPXqHL1WpNBY9JrjVp1glnFBTl6pULObWLQqPZtqY7XikghhPH4nL1rbJaCMMGoEOUFfzjdZxSICPdNzg1OWyVlJ3GCeHtwMhiJbs6NE3LuSQIAEAhHx+cd94an7g5Ntw9MdI7OxHtFuiGMHL6g0xdc6VASmqL2+VBfAAAgAElEQVQSrNK0iGVok17LMtScw5vgsyCu631jnCCIIj5zcNuj0QsAoGSZvU1Ve5uqAABfOrl/cNoy7/Q6fUF3IMTxgogQxoCmoZJlNUrWoFHn6NQ6jSrfqM036PKNukffUJIoL4iinJ+Uoiiya3h9iQjZvYGxOQlptJwgpL/3KYhiz7j5eu+opATj2J7O//bbL6z9Is9GJHRJgzAOR/lpm+vCvYGff3jvZv94vFesGxHh2HBkpQYsTRcY9Sv96bJKTMZvvXiiqaLor39xweL0zdhdoSgfN3ErrrvD03/7y4ssQ5/e27Js9FqkVSlWmQLNKBACmt6MixCZQxRxMCItDgkiSnBVLIl4QRybc7AMLSl0AQCULBNN+782Q5DQlajYA3pw2vLzD++9eumOpK7cujBoVFurS4pyDSs1oCio1yghhCDh2MMydHlBzjP7W5/Z33prYOIn525duDcwbXNJveUegjG+Nzz9vTc+VDD0qT3N65I8spLYoczxWi2DgpAhoWtd8aLo8knLdEcIJ5iLmEQiwrMOt7zBPS+IGONNmLxKQlei3rrR/crF29NW14zdbXP74zVfT0qWqSjM/ezx3S8c2bFKM4amc3TSzslFCC3eYHubqrbVlHG8cKV75HtvfHSpc2j1164OYXyjbzxXf7O6JH9bTWm85ukDIaAoOc8FlqHVioSW04gU4QVRan6NIKL0j2MgBC5fCGHJPSSMAS+IGAA5F2iWI6ErjsvdIxfuDQzP2tr7J2yehbIX8V60bnJ06j1N1S8caXtsW0NpnnH1E9wpCKVWnhYRDnML+Vc0RWlVCq1K8fT+rbsbq+acnrO3+356/taj5zUkSBDR+bv9eo3q//7N58rys/48PZqiFOzHSf9E+nGCYFm5CvuyEEbhKM8L4tINGyklIjQ8awtHOXnPFQGJGONNuCuZhK7l9U3OX+ocHJtzjJhtt/rH/aGovFmjtHlq39b9W6rb6sobK4piJeTjvQLQNNSqFJKueYTRo0lQCoYpL8gpL8hpKCt8au/WztHp9oHJDzuHZJx3HI7yl+8P//SD9t946uAqU51ZgWVodSbNfG5CoQgndft/IBy92T/uCYYL0pX7EI7yHcPTXWOz8vrE67WBet2R0PUJ8y5v95h5yuo8f3fgRt+4JxDCGHBCuicQEqTXqB7f2VRdnFdZaDrW1lBfWiDpgEGGpnP1GknLVKKIV8nfzdVr9jdX72+ufnpf67OHtveMmy/cG7zeN7ZS+2XNu7yvfXSvsjD3+SM71quG90Pk5aGwNK1dv21zBABg0uK80SstkUoQkdnusbh8aQtdgXC0Y2TaF5KZkS9pF/NGQkLXgp6JObPd3T44+f7tvmmbKxXlG5Ki2GQoyjWU5ufkGbR7m6pO7W4uy89JZIz1KArCYpNxS0XxiNmW4JgSYZzISkBVkamqyPTsoe0ndzf/6P0bb9/otnsD8V60QBBR/9T8j8/dKi8wHd1eH695ymGceBbLJyhYWivre0kzfygSinIIYyXLmPTSNktkMozxjM09bUt0M/IiESG7J32L2WGOH5qxxmu1Ik4QEMKbMB+IhC4waXHeG57+57PXx+ccgXB03iWtXkN61BTnxXbUnty1ZVdDRV1ZQYnJKC9iLVWQo2urL0/wgDsAgCCKiYSuGJqiDm2trSoy1ZcV/NM71xJfA+MF8XLXcFWRqaG8YKVNaWmDAZY3k6Nk2bV/QSnCC6LF7fMGw0PT1ht9Y4FwVEBIrWCPtzXWlxeW5+dsgK1CvChKzYyPiXC8JY0PAV4Q7Z5EO3aPknVtbgSbOnTFSsG+eunO5a7htXR8UkGjVOjUSpWCVSvZpoqiF47sqCjILS/MLTDqc3RJO5FBpWAlzZUjjCVNMAIAyvJz/vDlJ4xa9V+/dmF83hGv+QJBRG/f6N7TWPU7Z46sc+IvxjInDBkqwYId6YQwdngDnSMz77X3Ds/a+ibnlp7Bfe5Of01J/tHt9c8d3tFUUbhSHZaswPFiQG7oMjvSFLpEhCYtjnWpVZ/tsvjSXKNwlH+vvfd7b3x4q38yQ1azNEoFQ1OxarBP7mlurSktMRnLC3IayotK81Iy+MAYsIyEozd4QUx86m+RgmF+46lDvCB+92fn512+BCOBNxi50DH45N6WmuWObsl8LMNk2lpXMML1Tc796mrnrf6J9oFlLvvxeUesRkz/lOXLJ/cd3lavVSmy9Agx2XUxOEG0e9M0Yejyhz64N2iVmAZJgE0bumJ52P/7D98cnrXFa5tasWKvEIKiXMNj2+srCnLzjbqKgtx9zdWVhaZ4r04CWsqmpUA4+tH9IUFEUjfbKlnm608fHpqx/uj9mwkmfXGCcKV75K3rXd968UTiwTXpMJaZpqFgmcyZMMQY+0KRN651vXb53gd3B1afH/YEwq99dG/MbP/s8d2n97U0lBWyDJ11Ra0inOD2yxnNCKIYSlcFXrvH3zU2G6/VarLue0mWTRq6bg1M/PPZG4mvvqTOtprSrTWlhTn65sriQ611TRVFAIB0nukudb8tBWGE42U8lJUs8+9O7b/ZN94zMZfg0lo4yvVOzoWi/DrOvGEsc61LwdAyPqUUsXsD33/z8i8+6kgwJQdjfG9kemze3jk6c6yt4ci2+qaKoux6SkY4Xt5EnIhwhEvTNIw3GJ5c21No054DtxlDlycQfvXinctdw5LqNCdRa3VpkclQlGtoriw+tae5oayQpiBD0wp2Hfq2FJT2d4oIByNRGQ9lhqZ21JWfObh9wuL0BsPxmgMAQCjKX+ocujM4ebStQdq/MnmQrKWufKPu0NbaeK3SZNbu+T9+8vbrlzvCUT6RuBWDMXb7Q7++dv/9O31Htzccb2t8bHv9jrrybHlQeoNhb1BmxnmUF1z+YK5Ok+of1h+KrDGbUWrXc8PYjKGre9x8tWdE3jz4WmiUiqNtDQ1lhSd2NO6or1CwtFqhMGhV6/VQjpH6l6M19EmVLLOjvlylYBMMXRhjTyA0OGM53FpHpau6wUMwllPDUMHQUqvyp8io2f4XP3rr3Vs98o6hinB8hOPP3em/0TtmeEN9rK3hxSM7djRUrHvmZ1ydIzPXekbjtVpGlBc6hqctLl+OTiPx5pAmwvGdozNrfBBt2vMJNl3oCka43glzOosQVhTkNpQXbq0ufWxbfWtNqUGrMum1SjYjPnkIAcvQUEoJNBGhtdxsdaUFCilBKBCO/s+3r37h8b1GJml5lZIgWaGLoWmWlvBjpsjQjPXb/3r2zetdiW9pWFYsgNm9AbPDc713rL6s8PSe5p0NlVVFpsws2YUwtrh8fln7fDHGTl9g2upqqiiS3LOTYsbuvjM4JW86ehEZdW0WNo//l1c7Zc8kJK6lqqQkz1hbkn+sraGporiiMDdt+/MlkZoAvcbQVZCjk5TiIYiod2LOEwgbtesTuigocUYVAABAbAY4XqvUGp61feeVs7+80rHGuLVUhOOHZqxDM9bOkemy/JzCXMMTO5t21JdXFeWV5BkzpEMGAIhygrxNXTEI49klewZSZNrq6puci9dqNRCCLM3/XLtMudTSxu0P3RueTl02vErBttWVNVeVfGp/a21pQVWhafUauOtOUiABCyeWyZl6AgBACLUqJSXxZlMpWIvLV16Qsy536Yzd7fBKqz4OAKBpiqYlB7wkmrQ4v/uz8699dC9FC7pzTu+c0wsAuDM4WV2cV5afs7+55uDW2privNI8Y6qXiOIKc9xaOliCiMwOT6qrVNg8/mmb5DqfD2HoTTpjuOlCF8a4riQ/8SQ3SbQqxZmD2//988daKksyPGLFQAClpj0jJD90gYWEqHiNPglhPO/0CGL5uoSusTmH1IMzAAA0Ra3jqMvuDfzz2es/+/BOiuLWUjaP3+bxAwAudgy21ZXvqK84vae5uaqkNN8odUCfROEoL29tL4YXRLPDs8apvNUJIhqfdyS4UWQV63iZra91u7bWS1Gu4cTOpvF5h+x6l6uoKy34rWcOH95aF69hBpG0JRkAgDCKyJ2AwhiHIpyk+h0xLn8oFV2NRMj7e1maXq/Zs2CEe+3Du7+80ik7m0YeXyhytWf0Zv94+8BES3XJC4d37G6qzDfo0nZ6yFLBCJdgKtCyBBHZ1pb4F9eM3X13aCpeq/jWpT+XCdbn7lpH5QU5B5pr3rrRnYrQZTJo12tJRjZJORoAAIzBWk5Ad/gCUktJAQBkn2a0RgmWG36Ugl2f0BWO8ufu9L1y8fbIrE1OUv+aCSK6PTh5e3Cyf3J+R33F547v3t1YKelAg6TwBEJO6WVfFmGA1zLfmIiecfPgtCVeq/jIvq5N5PiOxid2bbG6fUm/OvMM2gxJiU6c1AwlhGQ+zQEAGACnLyjjNNj1ujl5QZQ350ZTFJP20YaI0MC05Yfv37w7NJ3Sya5ExAJY74T5d848dubAthydOp1f4pzDI2OadxFCmOcFGdMDiRuctsg40O5RafxQM8tmDF15Bu1vPXPYFwq/ca1r7XPNS+UZdFlXclvq/Ykw5mQ9zUEs7dgb5AXJoYtl6HW5RWNJ4fFaLUOlYFVsus+Z9ATCPz538/ydfnmTnKlwq3/C4vLNOT2/8dShwhx9vObJgTG+3D3SNzkfr+GKIAQQwlmHu660QPLabAKivNA1PhuOyrm0HkLBTTphuBl/bAhhW135X3z5mT/57KnklrVVK9l1LFkkw8K+Lik3J8Z4LROGZodHECVHPrWCTcUTJC5eEGVMbwIAaJpSsGkddQUj3Pm7/Tf7xjInbgEAEMaTFuf337z8d29cjmUkpoEvFPEGw2sZd4oI90yYLS6f1I5dgiYtTmuStpauy32RCTZj6AIAMDS1paL4Wy+e+OGffuUPX36ipaok3ivi21JR/PiOJklhIBNIXUXHaxh1AQCmrE6pL4cQaFSKdflgeVHmhKGCYdK81jVldb5+pXMtQ40UQRjP2j0/PnfjR+/fkHHsgAz+UHQt6YUAAIxxOMrP2FwyulmJ6B6fnV1zWvwml9a7K9PkGbSn9jRvryt7Zn/rle4Rs8MzYXHOO7xTNqeMsXyRyVCakZUFVic1IiCM5T3NYysxb93ollGWO9+ol7r/LCk4XpS3BVDB0mpl+sbf3mD419fuX+oYlL0MmWqzds8rF24X5ui/fGq/SpHaqdRgJCrj/n2U3RMQUzDqEhF6+2aPjOOblyV1rXrD2NShK6Yo11CUazjcWuf0BS0uX//k3F/97Hz/lOTeq1rJSh3BrDuMAceLklLRMAbyuqK8IPaMm0dmbVJncqKcUJJnWJeJEdkThmkedfVNzr93qzcVSbNJNDRj/dH7N2qK807ubo7Xdk2CEW4tpTRiMAaeYDgVs69OX3B4xpqsTkbWPXOSZR16splJyTKlecZdDRWfO7GnWtbZhioFyzLZ93nKGFVIjT0xGIPhWZvUO42hqba68jTU8F4WJwhSpzdjFCytUqQpdDl9wTevd/VJ72yl392h6e+/dSXxw7LlCUW5SDJGXd5gOBVrXZMWZ7LiFiBrXcQilqHlTWgoWSYTKq5KJSMMybifMcYWt+/u0JTUbqxGqThzcNt6HTcc4QR5GYZKllGmK8PwZt/42ze75ZWaTTMRofdv9/3qaqeMDlPiukZn5l1rTQnBAHsCIRkbOVYnIjQwZfEE5G+Xfoi8fuQGQELXw0SE5D2tNEqFIo1zRMkidTc+BnKOXgxF+Usdg3cGJ6UuHjA0XVuSn+ZsvUVRnudk1aRg6TSNuuzewKX7Q2lL3ls7jMF77X3d42YZV1EiRITO3R2Ysq51JQlj7A1GpF6ucbn8oSs9I2vZc/YQSbP9G4m0x9Zm4AtG5M0RqRSMpOM8MoTU+QYsK03D6Qucu9tv9wYk3WkQQpNB07Z+xxtygsjLWthjGVpqn0AGjHF7/8S52/1pG3LRFKVgGAXDyP7pOEG42j3y+uWOFP2bnb5gsoJiKMJJulwTMWV1Xrg3kMRiCJt2X1c6OobZJcLz8tIQ1ErFOtYblS0NaRoYY6vbf390Nl7Dh+nUyk8d2Lattmy9JvQjUV5GrWEIYXrSC32hyIf3h8zOlB/PEaNkmaaKovKCXISx1eWbsDi8wYikiyeGomDHyMyU1bWtpjTpnRKLyxdIRlDEGPCitFsjERPzDhnz7asgGYbEAo4XRVlJZRqlQpmWOaLkkrr4hLHkCcNZh+cH716bkL44r2KZ6qK89YpbQG6GYX1pwcGWmnitkuD2wOTFjsEkduFXYdSqv3Ry3zefO1ZZaAIAOH2Bn37Q/oN3r8tI8hZEdLlr+GeX7pZ+5omkV5/xBsMh6b2NZSU9bokIDc1Yw7LWI1aS3ECYRbLvUZtqoSgn42kFANCo2HTmQycFBpgTBEl3KAZAamh3eANXe0alBjy1kj24tXZfc3W8hikU5eVkGLIMnYa8El8ocqN/fN7lS/oT9lENZYV/882X9zfXmPQLqZ5alen3nz/+0tGd3/ibf7vaMxrvDR4miOjO0GSyYsxSgVBE3lr1Q2KnHCTxoxUR6hoz//j9m97k5WgAAJKeSJItNuk86SrCUV7GWg4AQK1QyF4AWEcy4rSktetJi/OVC7cnLc54DR+mVih2NVTuqCuP1zCFBFGUOjsKAKAgpFM/jTM4bXn3Vo/bH4rXcK12N1b+y59/7fTeljyDdun8nlGr3lJR/D//6EvHtjes8vKV3B6c/PXV+2s5mmRZoSjP8ZK/smV5Asn8bAURDc1YrG6/1D4csazse9SmWpTnpc6hxWTdkCtGEBEGEu4lqROGUzbX2ze7ZWxk0ajY5spieRsVkgVhLGNCRsHSqT4AEGHcMTI9Me+Ud60mbldDxfe+9fl9W6pX6pY1VRT97qcfaywvXPZPVxEIR8/e6XMktTQUL4hv3+yeS8biHwaAF8QkpgICAMbnHes3+b3RkND1MI4XJD2aF2XjRYkxQNKffYm/ZMbufv1yh4whl0apOLW75eDW2ngNU05SXI9haJpK5fgbYTw2Z3/tw3upXuXa01j13d95cd+W6tWbPb2/9cTOJhk5ShPzjsFpi4xx/0ocvkDX2OwaCxjGYIydvuCY2Z6sf57TF5y1eyTNWCRC1rNqI0jhDZalOF7OHFFFQW55QW68VplIRpxO8CUY47PtvT//8K6Mm78oV3/mQGvZeteExFjOQhJNUTSdwo4MxrhzZOZm/0RKN/buaqj4L187c2JnU7yGwKBRPXtwe3NVcbyGD7O6/efuDlhcvngNEzVtdcsYJa/EF4qMmG1JWZCL8sK9oalb/ePyFiNWkephd8YioethnCCIouSrf19zdWN5UbxWmUjqoxknPIfWMzH3xvUup0/ylIuSZfY112ytKY3XMOXmnF4Z1fAoCFO66ukLRrrGZuVVHcvRqfc2VZ3c3Xxoa+1KJ/5ACHc3Vv7nr555at/WZRs86nBr3fEdjVIHXv5Q5FrP6NicLV7DRJkd7qREmkUTFqeMatGPCkW49sHJ7nFz0iNN0t8wW0i71DYDXhDlJe1k44QhAEBECQ6iPpbISCQQjv7ze9cvdQ7Fa7iM6uK8Lz2xt6FM8vJJ0t0fnZVRs4eiIASpuhp4QeyeMF/tHpWxfKjXqL7+9OFvvXiioiDXEwj/7MM7Pz3f3jc1v3R3sErB7ttS9Rdf/tTjCYy3Fuk1qoMtNW9d75qQODlscfl6JuYOt9Yn5WQApy8o42NZxaTFmZRcdk4Qkji4XCrpw7hsQULXw8IcL2OCS6dWSi0s+yiE8eKYZnGJBQIIIaQomKLtTamYdBJE9Prljjevd8m4r2iKaqkqqSsriNcw5ZDEhJRFLEMn5UG8rDDH3xmcvN43Fq/hw/Qa1Ref2PuXX35Gr1EBAHJ06m98+uiBltqfnr/162v37Z4Awphl6Gf2t/7p507trK+I934PqyrKqyjMlRq6rG7fxY7BTx3YViOr5vVSCGOXP7iWc1AfZXZ41r6giDG2ewMyNsAlIsoLGOOk7+zOfCR0PSwUkbOvi6XptVRkERHyBSM2j9/lD/pDkQgnxLJFaIpSK1mdWlVsMlQVmVKRxIhxQqOoxGGMb/SN/92bH83Y5Rym11RR9KWT+7ZUSF44STrZ2yQYOoWhK8oLU1YXhFDSt6ZkmVO7m//rbz0fi1uLdtSV1371zOHWup6JOacvuLWq5MWjOwtkbROuKy3Y2VB5o29c6u0zY3OPmW1rDF0YY4c3YHH55BXuWknX2Kzds9bjjH2hyNs3ui92DMZrKIeIEAYpG+NnsOQ/CrPdKxduz0p/5k5YnJc6Bw9ura0uylt9+OULRVy+oD8cdfmDgVDEH4oGItFQhLveN3apYwhhtBhLFq9IioKleTlfPX0gz6CtLS1orS7NM2hX+SskkTpXjvFqaRoYY5vH/4P3rg1OW1dqs7qW6pLWDFjlAgCEo5y8IalSwax9CL4sQURDM9bBaUu8hg+rLyv448+eNGrVj/6RQaM6c3DbU/u2IoQZmpLdPTLpNYe31r1x7b7U0rc9E+bXr3Su/RCvoRnru7d6vcEkVIFaSkYlsIe4/KEUnV6NMZCxML8xyLxMN6pghPOFIjKmiW72jfdPzmvVit2NVS8e2aFSsDq1MjaIxxjwghjmeF4Q2wcn3rvVK4hIRA/+J2IRIYRxIBxdpQqAJxD+7s/OMzStUjDPH9nxqf2tR7c3rP35GIxEaYqSOtuwyscDIfyX8+1n2/vkrZbXluQ/s6+1IjNyNUNRTt7mViXLSE1YSFAgHL3RO3arf0LSkKvEZPzdM4/tbqhcqYGCYdZewgxCWJxnKDEZpYYuQUQWt88TCOfolomsCcIAjMzavMGwpE8mEfMun4jQWvJuOF6YSc1sIQAPCi1KvIU3gDVfsBuLJxCSEbcAAKEoF3tYT1ldtwcmaJpaenYXwlgQUWxOQ/ZRtoupej9+/+aVrpGvPHngC0/slTe3s8gXjFztHpG6PEBRUBDRo3Nigoh+cu7m//jlJdkbOQ+01B7ZVre+25AXRThB3uyTgknVqCsQjg7PSk7XbqkuefZwW4r+SUvlaDV5si5Ilzc4YXHIWGBbhDE2OzxSpxASMWNz+0NR2WE1GOG6x+Uk+ySI56UVctswSOj6BF8oImNT11K8IEpdqZbKGwx3js74QpHq4rxnD22P13xFGGN/OCL1qKcIx9/qHxdE8aHQJSJ0tr33//3Vh/JO+YMQVhWZnt7XUlGYEUMuAECE4wVZa10sQ6UoTvCiKLVyUonJ+Oyh7eVp2SGn1yiXnZOMyxsMW9eWgEdTlDsQkrrMlohZh9sXkj8itLp95+8OTKVy1BWvycYkfxS8IU3MO5KyFT8NxubsP/2gfdRsj9dwRbFZSqkJBRQFl63WeL13/G9ev9A7Obfsq+LCGD+zv/Wx7Q0pmmqTIRTl5GVa0xSVinRQhHH32Oz90Zl4DT+hvqzg8Z1NUueE5TFq1bl6TbxWyzA7PRdl7aOIQRjPu7wuX3CN/c5lObyBtTwTnL5g78Rcig4nAwBwvMzNPNlO2mNrYxubs39wbyC5NTdT6uzt3tevdMieJMEYRKKS96woWWZPU9VDo4r+qfm/f+vylW7JFcRjIIQ1xXmfPrht3ctnLBWOcPIyDCkIU3GKUjAcvdk/Pi7l7BgKwsaKouq1Je8lTq9RGbWfyGBMkNsfutE7JvtK5gWxvX+iZ9ws7/ta3dC0dVLKZ/4Qm9sne/48EYIobs4JQxK6AADAF4r0Tc7/7S8vvdfe6w+tdRtH2oSj/Eddw/1TkvPNYhBGMrZbYvzwHMWs3fMPb19983rXSi+Jy6TXfPnJA/uba1IxWJEtwmVWcrzDF5yXOKtWbDIcaa3TpOXcSwAABWFZfo5WJeevYxla9g4qXhA7R2e7x82pmDAcMduGZ2XW+3D7Q3eGpmQUlEncpgxbAJC1LkFEZofnwr2BVy7eGZuzzdqTUHM6nXrH5272jW2TlU2OkJwTiR4qFOENhr/3xoc/ePeavLk1AABDU7saK58/3CZvmSR1orwg71GoYGk2BZXjPYGQ2SHt+myqLN5WWxavVTK11ZVvqSy+Nzwdr+HDeEH0hSLyrgFBRMOzVpqmkKyuRlyyh4ODM5YL9wZTejCN1MPzNoxNHbpEhLrGZv/uzY/uDE71T6Vk40VKQQjzjdrSPJmTbBgDefFmceEkygt/+8tLP3jvurz3iSk2GV56bGfbup7LtawIJwhIzqOQZWg6BaMujhelTvCW5hmLTYZ4rZJJwTDyBnm8KMooFxnDCYI7EEp6WvyiCCezYsWcw2uRlbWUIAxwmOMTrCm6wWzq0HV7cPK/v37xaveoPamHBqVNgVH3xSf2ndy9JV7D5YkIyXhYUBSleTAj9N2fnf/+m5fX0qlUMMzjO7e8+NjOeA3XQZQXZE4YUimZMIxwnC8kLb0wz6DN1cnJm5BNwTIaWROGvCDKmzDkBKFjeNrlC6YscgFfKByK8lInQhHGkxZn0rdIPyQQjsrbz5PtNnXoah+YfPdWb6wIWLy2GaeiIPdPPnfqK08ekL0LCmEcltiLBwDQFNSqFCJC//31i//4ztW1zONTEO5tqnrpsZ0mWWlpqRblZa51sUzyj5rEGF/uGpG66KJWKGRXx5AHQiBv9y5Ccq5GAIDbH3q3vTcVRdkXObxBTyAkNXT1TsydT3HaF8bYFwqn7gfPZGm9rDNKIBztGptdpYBFJivLz/nLr3zq5WO7DBo5CV0xCOEIL/nH94eiPzx7471bvbMOzxqPuFUr2af3b318V5pSt6WKcDJHXSoFk/R8E5c/NG1zS/33UBRM82eLEOYk/iNjZMe8YIS7Nzwl9ZORxObxewJhqemv/ZPzo2ZbqodEgXA0G3vea7d5Q5eSZTKzsx9Xscnwv37h9GeP79aplfHaxhHheKnXvYiQxeVLygkOnz2+56unD8pbGkkDXhDlnWmbiv3IvlBE6nxasZ3j9kEAACAASURBVMmwpSLdZ8iJCIVl7YKKVZqO12oZvlBkLVPWiQhI//AxxhMWR0qHXDEYg8251iWnm7MxsAzdUF6YiqdM6hQYdQeaa/7L1z79707tX3vcEhFKyjF6MjA0dWJn09efPrTSaYeZAMtNPJbYGUhI+EGlscTlGbTpP7nb4vQ6fHLG4ixDy7ikY0Udpe4ZkGpg2nJ7cCJeq08Yn3fc6BuXGvCkQgi/e6snS5fq12jzjroAAPuba47vaPzg7kC8huugxGQ0atUqBWPUafQapYJhNCrFzvqK3Y2VR1rrkjILhDCWsa9r7ViGbqkq+caZx3Y3rlgQNhNAAGVEr6JcQyrKB4einNScGgXDMOntmYkI9U7OS83gj2EZWq+WPPs97/Je6BhMXa2KGJvHPyCxWn/HyHRPavaZLYUw7h43p3TLc8ba1KFrZ33F7z933Gz3pDQzvjBHn2fQKlmGZWiKgrEzI+GDDHPqwf+LHSZJUxRFQZqiTu7aUlmYq9eoCnMN+UatkmWULKtRskkJWjEY47UktctWU5z39acPndzdnOYMAqlm7G4ZS6ENZQWpOCczHOWlnr7B0CmpR7UKtz80PGuVPdSQUSfQ7Q/1p+Y8kYdgDDhBSLBKmSCi83cH3KmfLYxJ6Tpfxkrom9jATu7e4vKf+uvXLgzOWNbeRWIZOs+gVSvYWI6ZgqFZhn752K5tNWUGrUqrUsbOz6UpKhaiIIA0DSkIKYpiKIpl6Nhu1iTGp1VgjNN/0Zv02peP7fra6YMyZofSrGfCLCPnDWOQioAhCKLULysV85ars7r98hZBlSxTUWiSmiuLMLa5/Ra3hL9RrWQrCnLVSsWs3S0pOTYYiQbCUZM+oQfmlNV5qXNIdgiXSurJDxtDQt/EBqZSsJ87sUfBMH/35kedozOxvYfxXrQMlYLVKBV7t1Q9f7itpjg/z6gzaFSFufq1ZACmwdqjtSQ5OvXvv3D8jz9zMvPjFsZYXgSS9aL4RISk5ozwopjO7xdjPGNzTUs8rCumtiT/c8d3x2v1sFiZJSQlNbypvOivf+/lbbVlf/mjt//h7Svxmn/MGwy7fCGTPv4RrxjjjpFpqf2MtSCha5NSsszLx3blG7X/41cfXuwYlDeHdnR7w+dP7Dna1lBVZJL3yEs/hLHUlX/ZIIS5Os23XjzxH156PPPjFgCAE0R5ac0QpiQfXRCR1LLooQgXkrg8thbBCHejb3zCIrlSLYQw36irkl4jeGzO/mHnkKSy7koFo1KyeQZtW10ZBWHiX7HbF7J6fPUJTAWHovydoSmpC5NrEeEFJLenlb02b4bhUgxNHdxa17iGhMO9TVXPHGitzp64BQBACEc5IdF7d22MWtW3XjzxrRdP6LMhbgEAwlGZh3UxDE2noGy8DGNz9q7x2XitkqZrfPaSxEASk2fQHmtrrJR+TpvZ4RmckZY9YdRqtColAGBHfUVzVXG85h+ze/1jZnvczb8I4xm7698u3El1EY2lopywCSsZktC1gKEphLG82UIAQGm+MUerTkV3O3Uwll9XVBKjVv2tFx//Dy89bsyejygU5aRO0MXQFEXJ2lq7utj6aLxWnyCvWopsQ9NWSQeyLMozaHc1VORIL1g1ZXVJPX242GSI7ebMN+qqiiSM8yYszndu9cT960QRdY+ZQ1FO9pNEhkAkms6Z4Qwh7WbY2MJRXt71dmhr7dbqUtkjtnWU+ISJbEat+g9ffuL3nz+e4ct+D4lwvLy4zjJ0KqIzQ8s5eTltRc56J+deu3xPRnUVJcscaKnZWl0qdbpifN5xq39C6neUq9cU5OgAALk6jaS9Gbwg2tz+uD9gIBztmTBLmtplaOqxbfXv/7c/+PMvPpXIWtqjIhyfhhs505DQtSD2qJKxjwcAsKO+orYkP16rjIMBRqnsrFEQVhfn/clnT33j2aN5Bjn35DqSHbro2O6HZGNZJsHM7KW8wbA/LXlul+8P3x2akvGJFZsMT+3dWlkkebawf3K+a0zagdHVxXm76itiH6NBqzqxo7GutCDxL8sfisw7VysDLyI0PGt79eKdCCdhvVyvUZ3cveXUnub9W6rl1RPhBVFSrsrGQELXArsn0D4wIS8vqCw/R59VQ4oYCCAvpXsoCcvQ9WUF//Glx7/x7GMFRl285hknHOXlrR/Etj3EayWZWsGqpD/X7J6ApBRwee4NT//ySqc33mTao9RK9uDW2l0NlTKi8s3+8WmbO16rT6gqNDU+qIxFU1RDeVF9WUHiX5Y7EBqbs6/yiOAFsXfCPGFxSgrhWpXi8NY6AICAkLwhsiAiMurapBDGNo+/b3Je3pRxWX5OKo4WTLUoL3OCNC6Vgm0oK/zGp49+9fRBeXMg6y7M8fIuBgo+dBhncujUShkZLhPzjlGJxealwhj/+NzNWwMTMp6eJSbj84fbpJa1BQB4g+HucbPUDeP5ObrcJWVLtSpFRUFu4iuaDm/gYufQKrVCPIFw39R84sO4mMpCU1t9OQBAxTKJx9GlQlFOUrDcGEjoAgAAUUS+YFhG1y+mMEevYLMvdNm9AbvXH6+VZEqWaa0u/c1nDv/2mSPZtb61VFRu6EoRk0ErI5FhxGy72jOa0ufahXuDl7uGZWwp0SgVO+or9m2pkTFL1jMxZ5WyEzmmstC09BRmlqF3NVQW5xoSDDbBCHezb3zUvHxXAGM8PGu70jWSaCQEAACgVrK7GipNei3GWKdRSc3EifGHIhl1raaHnE9q40EYS+3BxUAIC4y6IpNB3jW3vjz+0KjZLm+OYiU6tfLg1tovPLH3D144nrEl4RMhiDIXPlMkV6epLcmXWm8iEI7eHpqcd0p+yicCYzxjd3/nlbN9skoxNZQXvnCkTd7pDbf6x2VUSiw2GZb2pRQMs7+lZndTZeL5L4Iojs05lr1lQlG+Y2S6c3RG0g1VYNTvaljIFtEoFfIeI+GozHXZrCbnk9p4MMZhjpfxqFIpmG88e7QuC3M0QLJLn8U2lj69r/WPP3PyP770uLybMHNwvCDvLIkUHZHF0NTxHY0yDjEx2z2dI9PxWslh9wb+n59/cLN/PF7DZbAM3Vhe9MSuLTIWiaO88Ma1LqkVp5QsU5afszT2MzS1vbbMpNcmPk3nD0c7Rqatbv+j8cnlD8oohVpekLO/pTr2e61aIe/KmXN60lZ0KnNk9/MlWWTvgKEgLMvPSfqRuGkgiCiJpTQghOX5OX/0mSe+85vPPrO/NV7zLMAJYqb1ZJUso5F4UC8AYHDa8s6tHnmTCqsIR/nzd/r/9YP2xNeKFtEUtbO+4rPHd8mIWwCAGZtbUvZ5zMndza01ZQ/9R0EUTXpN4t0sfyjy1o3uK93LzAqaHZ4B6aFrV0PlloqFndEqlk1w6vIhb93o7hk3x2u10chc3dlgEMJSy3LH0BRVnGugaTkX3PoSRDFZPbWKgtznDredObDtUGudjHWLzBTlZU7CpCY3HgAAjFp1QY4+XquHIYxv9I1f7ho5vbclXlsJ7o/NfPtfz/pCEUnzYzEY47a68mNtjWqJ858xg9MWGUXZKwpyH92hQUGqpiQ/8dAFAIhwfOfozGc/WXExFOWudI3cHpxc4UXLy9VrqopMsd9DCDUqmROGAAB5Pe+sRkIXAABgjOWtc0IIc/SaVGSUpZqIklDA8GBL7Y768sOtdUe21RfnGhJfM8h8C9v8Mkm+UXdoa23H8PSMXVpS+PCs9Ufv39xWW5asgz2HZ23/+Sfvjs7JWShVMMyh1tqvP30oRyentIo/FLk7PCXjWOTakvxHi2dSFNzZULG/ueZK90iCA9MoL3SNzV7uHjnYUrOY2GVz+7vGZ6U+Q2qK8/ZuqV78v2oFm/jU5UMC4ehmK2NIQtcCeV3scJTPM2hpWmZfaR2JCMnIClMr2dqS/Ori/Byduiw/59lD26uL83K0mg0z2FqkVMjMVI4dvxavlRw5OvWxtsb32vukhi5BRJe7hl+5cPtPP3cqXtv4pqyu//Nf3r3cNSwjbgEAaBo+f7htR325vBHGtM39Xntv3IJMj6ouzlsmdEG4s76iuar4Vv94gqGLE8Q7g1Pv3uo50Fyz+B/H5x1jZvsqr1rWsbbGpRU9NCpFgv+GRwXCUV4QM/wAvOTaRD9q0qmV7OM7t+QbddnY2RFFxPGSFwwMGvULR3Z+7sRutVJh0KiyrkZG4t6+2RO36k+a0RRVXpBTU5J3uSte00fYPP4fnr3eVlf25J41TRtOWpzf/tezb1y/Ly/Hh2XoMwe2v3xsl+yNKGaHe9bukdTRpCDcUV9eU5K/7H2qUSryDbrEJwwwxu5A6P7obNfY7M76CpahQ1HuSvfI4Iw13ks/QadWbqksXpqFS1PUsbZGuzcg49DnKC9sttAlp+OzIWEMpPYi1QrFkdY6GUe7ZgIRIRnnMug1yj2NlS1VJTXFeRs4bgEAOoanZRRBTzWDRn14a11FgeSySQCA4Vnbd3/+gYwsuEWjZvt3Xjn7bxdvy1tZgRAeaK75veeOlphkzlsijMfmHFKLHlEU3LelZpXJ0mJTolu7FvWMm//l/K3YavGMzX1veFpqvKkvLWh45AiVbTWl8raUBCNRTpA8iZLVSOhaICMznqEpo1Ytu/+4vkSEwtJnJ5QskxWnba2djFy+NNCqFEe3Nzy+qylew+Vd6xn97s/Oz61aiG8l/VPz3/7X935y7paMeWYAAISwoazgC4/vPdjy8TybJAjjaavrrRvdMqoyluYbtSt/oQ3lhY3lRZIGXjaP/3LXSO/kHADA4vbNSpzCZRn6xM6mbbUPZzzyosxqhMEIJ28cnL1I6JKPZeiCHJ3UsVqGQLIyU2iKUsrKCss6Gfu1lhfkHGmtzzNoZeQ48IL40w/af/DuNUlrReEof3d46s9/8MZPP2iXNFO3lEGjevZQ2zMHWmV39XhB7BiZvtgxKHVBCCH8/7d3Z8FtXecBgM85917s+06AIEiQIMF9FUXJpFZKtjZbihKliWWrafbJ0rSTTjvTZOqH5qXTt7aTpTPtTBInqdPUdhzXtSwrlmXtolaKlLiKu7gvALHepQ9QZFsiCdxDkOIFzvfgB+uAEgHc+99zzv//x+92yJZfTPPnObZV+8R2fpmcC/zkrQ//649Xf3/uptjdR6NG1VxR+GSPNOwTc2Isi3E5Sxrm14gAANAUshl1yUZtUIKAU5LM0BReLQ6RLnKG3l7tO9xc84uTl/DWiP7xl++YdOqX9zYlvVnHWHZiNnjyasc//eZk9zINkFJBU6ihxHNsRz3eUmcCzwv3H0wrZLSohVwIoVmnLnbbZctPqqx6jcduEtupZHI++JvTV39z+mqygUuwGjTGpdbb5QyNd05pNM5iP1VIFAldDyHx57JTCEl0owsAwPMCRuiiKYRXi0OkC4Qw32F+YWvV/15sfzC7gDE75Hj+7/799XA0/o1DLWql/MltHkEQ4hwXCEUvd/a/+v6VX5++suTPSRGCsLLA9Z3DO0Sdj/WkGMuOTs+LbfKrlDGHm2vsRu3KV7dJq8bbZMJAIVTpdbnMSzQdpikk9i6UEIrEMLKuJI2ELgAAgBBidMSgEFIwUr2PczyP8cyOIKIlWAmQYSiEWutLv3qw+V9e/yNGhRMAIByN/+jVd7qGx//qs7vLPDmf/CNBEAYnZi929L114fZ7bZ2rPzMlx6z/6oHm3fX+ZANXwnL83cHxruEJsctiNIWcZv0nu+4uyWrQrtsmbp7NeOSZGq9zie5x4p+fH5oLhiNxceuoUkdCFwAAQAgx8kohBHIZjflde9p4QcB4TEMQpr6bLWkYs5n1JGfoE3ub7g2Nv372BsYjCAAgEIr86v0rlzvvP/9MVXGuXadShKKx4cm5ruHxC3f6RqfnI7E4XkbGJ9kM2q8fannhmepVzmnC0djFjr5TbZ1ilwoghP48R9IaMn+efXedf2B8Ziaw2lCdlEYp9zjMYnMaVxaJxVb/kCEtou/XGQlBiFFUiyDCCHgbBI8360Jwhe1uYj157KavH2pZCEXevXxH7DJaQiQWb78/OjQ5q5QzFEI8L0TjbCQWX32blQS31Xh0W+0XdzfajaL7Vz0mEmfvj0+LDaVyhm4qK6gocCatLrfoNfXFeb8/f3MdQld+jsW4zPk1FEJ49extXYOXO+9vqyyS6JM0BnIbeghvHSzp01yGQQhJ8VDNjAQh3FJW8JV9W0OR2EerOJRrfjE8vygi4TBFHrvpSHPNN5/fXuAwJxubHARgeFL0KScyhva5bIVOa9IpDoLQaTGsw5qhN8eyb1O5bZlYztAUXuxZjMSm5oOxbKpKzpbfc2UUBdUKOcbzzgZfVloBLwgxkWsvCUkfYLOcIODUCOKR0XRrfWliynXhTh/GNHqNFDqtn2mp+cr+5qInqm4xhKPx691DD2ZEl6PJaMpp0ad4N8+1GGqK3F3DE+lqS72kfIe5papouQOvaQphX1+haCwUiaX4y2aAbPk9V4YgSpoovCS8hZqNgOX4mMjlFwAAEARO5D45saY0Svm+xoenzJy/0yt2K2gt+N2OF1sbj+2oT0vcAgDMBBbfuXKna1h0dj5DU0VOW7JRD+WY9btr/R/d7sH4i1LntBhshmXTHWU0nXSCuJyFUCQYiRqxju6UIhK6AAAAIahTK+PiDwGS7qwLDy8IYvcbss3atd9djlLO7N9cwfMCheC59t6n+AFBCKu8ri89t+VIc22udYnkbzyLkVjHwJjYlpIMTfnzHLU+d7KBD8kZurzA6bIa1yh0IQjNOnW117VCopNSzuDtXAAAFsPRkJiKN6kjoQsAABCEVoPGZdGPTokoHOEFPs5KdQqCUccGcFPqibWmkDFHWmrkDC1j6LO3ukUV7aYLTaFNJflfP9Ty/NaqpMnoqRMEYTYQmhNfA6BTKXbX+r1iTjC3G7UV+c7zaxP+aYpqqfJtry5eIdlSrZBhZ/CG05ERKiGYET7zWPSa5zaVi6qoFwQQjsWkO/ESH7kAx/NiC2uI9UEhdHBL5fc/1/rCMzXrvGoEIVQrZNuri3/40v6X9mxOY9wCAATD0av37k/MBZINfJxSzojNbDTr1K11/scK3dKFomBBjrnS61whOKmVcowC04RwJCa2RZakkVnXQxCInoXglUZtEALA6WEoCCBLZl1ivwwbAYRwZ21JjkXvMOl+cfLi1MLiOjxX0RRyWQwHmiq/fXhHiduebLhow1NzJ9s6B8Znkg38FAihXq10mMT1aaMQqip0VRfl3ugdTvtbp5AxPpdt5fNL1Qo5/oJhJIbRUFu6MN+mzAMhEDtVj7PcxOxCslEbFMvxvPjQxfESXiPNEiW59u8fa33lzw+m60zkle1pKPvB8f2vnDhYnJtqQoQoM4HFB9OirzKdStFcWVQsPpQaNKqawty0H+ijlDO76/x1vryVbzIquQy7+CQSiwcWxR28ImkkdH2MQuKyeyKxePfIhBQfz7FxHJ9VW8EY1jM5fkkQQptB+/Leph9/7wsv7dmc9rtwAoRwW1XRD47vf+XEgWM76vE62adiIRheCIkuO1PKmcoCV4FDxEZXglYpb60vrfOtqt3ik5QyWZ3PXeV1rVwJqlcrsb8894bHT127i13eJzlkwfBjYqfqoWjsVt9IJBYXtUO2QVAIQfEVJCzHR2IkdG10EEKVXLanobQgx1JR4Hzn8p1rXYMLIs9CXI5KLmsqK6jz5e2p99f68gwa5doV5o/PLvzhYrvY1UIAgEouM2lVYq9okOhubDcf3VZ75d59vP6QS4IQlHtW2uVKUCmYPfVl/WPTGHt7cZabmAsEQlHp9gQXhYSuj2lVClEXYSgSO3m18/LdgW1VRcnGbjgIq+MwL+D0myeeChlNl3lyckz6LeXet87fOt/e1/9gCu+cyYQck768wLmpxPNcY7nPZTXrNBixQZTukcm2rgGMxDmdSuHAXS9Vypk99aWn2u7+9kxbsrEpUStku2r9Tx4s+SSGoso8Dr1aiRG6AAChaGw2GCKhK7sgBM06tahSdl4Q+semXvvgqoyhqryuVTYYXX8YCzwcz2dJ6EpXH7+nzqhVPVNemG83P7+1+uzt7ncu3ZmYC9wbGk/2uofsRp1SzuhUijy76dj2+vICZ45JZ1+vY+om5wLjs6Jv4lqVoqXKt5pEQbfN+O0jO673DPaMTCYbm5zVoD3YVJFKoRuEUKdSyBjRz5QJC8HwbCCUls5bGx8JXQ8hiAocZoykuzfP3bw7NH6wqXJPfWmh0yKhxUOMylmeF+Li3yIp2lxa8GBmYU17Aq0nl8XgshhK3PZtVb7p+cW3Lty63T8yNb+4GInGWS7OcgIAgVBEKWfkDI0gUshonUrpthu3VfqKc216tdLjMK9RIsZyguHopc5+sZXIAACHUbe92mfVa5INXBaCsLmi8LtHdv7za6cGJ0QvVz5GKWNqfO4UV3RUSjl2M6dAODIXTNsi5waH+R5lHgrB6kJ3Rb6zrXsw9YkFLwij0/Oj0/PtfSNnb/c0VxQ2VxaVuO14baVEmZwPTs4FFiMxlVzmMOnE7sZDCEVNMRN4QcBpHyVBz20q++h2T8aErgSzTr2lzAsA2F7tu90/OjQxMxcMR2LxSDweisSvdg1UeV0mrZqmkFalsOo1bpupwGFe5yqxR9rvj55qu4sx/TXr1GLT4pd04tkt98dnfvX+5QczolMcH9Eo5Y2l+fn2VGdCWqVCKcNcv4nE2MXM+saugISuhyCEuVbD53c2dI9MYJx8MzkffOOjG2+eu7mtyre3obTK68qzm2wGrUYpT/tCYiAUaesafPdqx537YwuLYY1KUVuUu7ehrL44L/W/C0GIsVHB8XyWlD3G2dTbqnwKx2O+cD1pVYqt5V5Q7k028Glq7x/F2PKRM3Shy2rViytGXpJGKX95b9PCYviNczcxJn8Jfrfjm89vT70nvc2gVSpSvYofE42z2VPaRULXp+TZjGadejYQwrv9CIJw5mbXmZtdAIDtVb5Kr6umKLfOl2c36hiaYmiKQoimEE2hFFcPHsNyfDAcffXUpZ/+4eyd+2OP/v87l9pPXun84Uv7d9WVpBi9IMQ5sYXlOIw9cylicPugcjyf9mrWLMRyfO/oJMZpLN4cy+5af1pmXQCAygLnn+1sCISjp7AOjKYpVOiyFi51IPJycsw6t9XI0FTqaz+PBMPRoYlZXhAwv7uSQkLXpzSWFhxurvmPd85jP2Q9cuZW95lb3QAAi15jM2g9DnOR0+q2Gj12k9dpdZr1Jp3q0W5TKjUxcZa71Tfy49+fee2Dtifjx7XuwZ/+4axZr27056cSkyCEqQx7DM9jHpUiOXKGwXh/AAC8IOA99xCf1H5/9L2rnRgLtiatuqLAmfosJ6mdtSUURcVZ7vT1u/OLEVHPJX6349j2OpNWxGI+hdCxHXU3eoau9wwlG/u48dmFd692nHi2ybKKfT6pIKHrU3JMuoZiz3+fubb60PXI9MLizMJi1/D4+xAhBBGEMoYqdFohgDSFVAqZRilXyJjEtEzO0IIAEIQsx9EUFee4aJzleSEaj88vRqbngx0DY0vOe3hBOHm1o6HE43c7UtmcoBDCSGTKnr0uGU1hhi5eEHV3I5Z0rWtwcl70aiEAwGHWadO90/xMhdegUZp16v85ez31uReF0K66ki3iV2Uteq0hhUt4SeFobHw2QEJX1oEQPrup7Hb/yL++8UG6ChITdzKeEwB4mJsXioLr3Q8fqRCCCD48Xw5CgCAUBAAhePRfAQiCAARB4AWB54UVquVZjr/WPdj/YDqV0IUQxEiGnJwLvv7Rjb88uivZQMljaAojjQUkPm4SuVaH4/mfn7w4Jr7/EwDAYdKn8v0XhUKoPD/nB8f3V3ldP33rbPv90WSvADKa3lbt++z2OqtB9K6bRa/GzvMKhqMTswvl+fiFAVJBQtfjNEr553c29IxM/vZMG0aufIoerSl9MqSt3oc3u09e7fDmWJKWJVIIKRjRoYvj+an5YCgaS3FHTboMWpUsWe8DYo1c6x4aGJ/B6GlUke881FRpSnfoAgBQCOVaDS/t2ezPc/z2zLVz7T2f3Gx+Unl+zt8ca20o9mBsO9mNuh3VxW1dA8OTc8nGPi4QioyJP05aikjoWkKJ2/4X+7Y+mF344/V7ycZuLAuhSHv/aCoV9RSCchnOp48QDEfjGR+6ChxmGVZ5DdnoWqUYy17s6AuLz4kHAFQX5fo9DryV3lRoVYpdtSW+XNuFO30XOvruDo6/19b52CzbqtfU+vK+cahle3Vx0s5PS5IzdFNZQXGuHSN0ReNsujp+bXA4F2fGQxBuLfd+64XtkWj8WvegtHLqxmcXUrns8WZd4GGmhpTeEDxWg7bKm9vePyr202clkR2/UfGCMDw597sPr89jNUH35li0SsylttS5rUZ7c/Wu2pKRqblGf/7t/pG5QCgciyMIFXJmW5XvQFNFvS8vldyr5RTkWPKxmmLEWS4SE/eNlSgSupamkDH7Gst5XvjZ2x9d7OjDyHR6WgKhyOR8MGmCLE1RaiXOzEkQQDQLrg0E4ebS/Lcv3hZbWsTzPC9+pYtIEAThVu/wpc77Yh+PIIQFDvOWMm8acwtXIKNpi15j0WtKPY6RqfnphWAoEkcIKmR0ns1kE7+/9RijRtVcUfT2xfZxkccqBSPR6z1Di5GYGrc4TCpI6FqWQsYcbq6WM7RFr373Ske6sjbW2sRc4E7/aGNJvlK+0qSKQhBv0Q/vjEop8uc5Stz2yfmgqLwLjsy6ViEQil7tGqQpJPbpCAJwpKWm0Z+PUWi/GjKaLnCY0942kKbQ7jr/81ur/vP/zou63MLR+Pn23q7h8doid7Kx0rauH7PkJA5N/86RnV96bqvbalw5GGwQA+MzPaOTSbe4aYrCfj7FPlJIWjaVeJ7dVKYV+S6R3HhsvCD0jU2dv9MX50TXDvKC4Hc7kcSdzwAACFlJREFUMJJmNyy7UXuwqdKbI6KcOYHj+Z6RiWSjJI/MupJrKi3wux0+l/XUtbsf3e6dCSyy3MbtmEBTyKRVJa3ZkjGUUZP+RKxMopAxm0o8hU7rzd5hMpFaB9EY294/kmhGIwqEUKOU+3JtSb/2EsLQVFNZwa5af9/YlNiJ1+D4bLJRkkdCV0oMGuXXDrbU+vJ8rhtX7g10DoyNzSxszOhV5XU1VxTJ6CSfLIUQdvkLRr6vRFV6Xc9vrRqenJ1MX4k6sZxwLHZvaBxCKPbKUsqYL+/bWl2Yu3a5hU+FWaf+yoFnbveNXLrbn3r0isTio9OiUxMlh4QuEep87vL8nP6x6TfP3bx0t39wfOZW30iyF603i16bYhWk3aRrKPa0dQ+KulMgCPFSfqXIZtDuqvO/fbE99dCF4Goyy7La+GygZxTnfCyGpp6pKFRgFXtsZBDC4lz78T2NMZZt6xpMcerPcnw2NOHNtA97TVEIqeSyMo/DZTGceLZpJhA6397bNzY1MD6TrgNVV8mq1+yuK8kxp9R7VKtUYJwMFI2zePkdUgQh3FTi+evP7f7bn70+NJnSIkyWbASmXZzlbvUOv3ulQ9SDFACAplB1ocuft4blXE+RWiH7/M6G2WAoEIreHXqQbDgAAPCCMDUf7BubwtgnkxDRdy4CQmjQKA0apctiKHJag5FoMBw9sXfz+GxgZHpuYHxmYi7QOzLZMbBSsf1a0CjlB7dUHdpSZUhtE4tCiBY5f9KpFK31pWqFuMwFSZMz9J6G0u6RiX9740wqifIURWXPgmoacTw/FwwHxJfT6tXK1vrSQqc1I0MXAECvVn7pua0KhvnlqUvX/tRAbgUsxy2EInNB0U33pYWErlVRyhmlnLHqNYns2FA0NrMQCoQjo1NzN3qGo3F2fjEcCEcSHXjjLNc9MrGaY+tWUOAwb68p/vaRHb5cW4q3TgiBQkZDIGKaoFbIy/Nz8NpwSJdJqz7eujkUif/69JWkZ+YyFEVT4h4ICJDo7CyylitBo5QfaKrEWD+QEJtB+2Jro92o+9XpK6faOlcuk2c53qxT5ztMK4zJAJn8ea8/lVymssoAAKV5jt11fgBANM4Gw9H5xfD8YjgSi/eOTF7rGYrGWI7n4xwXZ7lonI3E4tE4G4uxHM9zvMDxfJzleEGQ0dT1nqElt2ddFoPDpEuc6CNnaJfVuLu2ZHt1sai2mxRCYnts0xRSyWUphsZMku8wf+1gs0WvfvXU5Ru9w8sN0yjllV4ndv5LNqMpZDNoc60GUd2PaApVel3lHhFfe4my6DWHm6u9Tktxru309XvdwxPLnR/dUll0sKlS1EkrUkRC19qSM7Scoc26h1+jLWXe43s2gz+dHhJjuXA0Fo7GQ9FYJMayHMdyfCJ0cRzP0NSHt7p7x6YAAI9iRWKGtLO62GM3xTlOEIBCzrjMBqdFL3YLiqbQo39YihgavxpM6vId5i89t9Wf5/jJW2c7B8b6xqYeG6CQMc2VRYe2VFmz4MiJtGMoqqLAtakkf3jyRrKxH7MZtEdbarMkb0ghYxr9+d4cy86a4tPX7125OzATWIzE4rwgQAgZitKpFRUFzhd3NzaUeJL9MMkjoevpQBAqZIxCxqx8usGOmuIV/nSVGJqyG1NK6HhERtPYxzFkAKNWtX9zRakn58Kd3t+dvdE9PB6OxmMsy/MCQ1P1xZ4v7GrYVOLJkjtpekEIPXZToz//g5tdqXeuqS5y720oSzYqo1j0mgNNlc2VRSNTc/1jU3PBcJzlKAqp5DKnRV+al5O09XZmIKEre9EU5TDqYOJYsNTIZbTFkO1TikTjn/2bK7uGx8em52cDoWicVStkxW57oz8/2auJZakUsiMtNb2jk6990JZK+3OdStFa53eYxD1+ZQa9WqlXK8uyYKV0OSR0ZS+GRk6zPtmojyEIrQZtZmfcps6gUZJAlV4IwkKn9YutjTOB0Km2zqTRq8rraiorWHkMkalI6MpeDEXZxTyxeuymL+7alPHbv8RThCDcVlmkUci1KvnvPry+wokNDE0d3Va3mTw9ZKvMrIQgUgEhLHJZd9aWpLg3Y9KpvU4y5SLWFoSwujD3H14++KMvv9BSWbRcJ/hci8FjN5HOJVmLzLqympxhdCpFKv0L5Aztz3OITesgCAw0hTx204m9TZtLC87c7LrQ0X+po/+xc+s/t6N+V23Jcj+ByHgkdGU1BGFpnuNSZ//o9PzKAcxtMz7bUJZnM64whiDSSKtSbCrxFOfa9jVWtHUN3OobeTCzMBNYRAj53fbjrZvFViUSmYSErqwmY6jjezYPTsz84r1LK49020z1JXmZdB4SIQmJVDpfrvVAU+VcMLQYiSEEbQathRTPZTcSurIahVBxru1bh3dE4+y59t6RqaUbGRi1qn2byvNsGd5ahtiwZDRt1tFiK+iJDEa98sorycYQGc5lMdQWuU1adTAcXYzEYiz3ycVDBOHRbXXf/cwucuMgCGKDEH2qG5GpQtHYxY7+09fvvXe1s2NgjOX4GMsiCMvzc/75G0db60uT/QCCIIh1QkIX8SmBUOTNczc/uNndOTB2o3fYYdJ97+iul/c26dVZ0V2GIAhJIKGLWMJcMPzzkxdP37jnNOv//sV9Losh2SsIgiDWDwldBEEQhMSQbhoEQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExJDQRRAEQUgMCV0EQRCExPw/4EocCzVNM40AAAAASUVORK5CYIJQSwMEFAAGAAgAAAAhAPmPhAP3AAAAcAEAABwAAAB3b3JkL19yZWxzL3NldHRpbmdzLnhtbC5yZWxzjNCxasMwEAbgvdB3EIJCO9SyM5QSLGexCxlCQ3E2LxfrbItKOiGprfP21VJooEPHu+P/7rh6t1rDPjFETU7yqig5QzeS0m6W/NS/PD5zFhM4BYYcSn7ByHfN7U39hgZSDsVF+8iy4qLkS0p+K0QcF7QQC/Lo8mSiYCHlMszCw/gOM4pNWT6J8NvgzZXJ9krysFcVZ/3F439smiY9Ykvjh0WX/lghICXIedWj9bmP2YYwY5J80gazLtrtcIr5HUPnE9mhpS9nCFQcjuDR9NqeAyjqjv3r4W5T3lcPhaK0/jgHUvnSbk0YHBgumlpc/an5BgAA//8DAFBLAwQUAAYACAAAACEANkldYtIGAABxFwAAEQAAAHdvcmQvc2V0dGluZ3MueG1snFhLk9u4Eb6nKv9BpXNkEW9S8XiLz12n7CS1ci65QSQ0YpkkWCA1snYr/z3NlzUj96QmOzoM2B+60f11owni/U/f6mr1ZFxX2uZhTd5565VpcluUzePD+l9fso2/XnW9bgpd2cY8rK+mW//04c9/en/ZdabvYVq3AhNNt6vzh/Wp79vddtvlJ1Pr7p1tTQPg0bpa9/DoHre1dl/P7Sa3dav78lBWZX/dUs+T69mMfVifXbObTWzqMne2s8d+UNnZ47HMzfxv0XBvWXdSSWx+rk3TjytunanAB9t0p7LtFmv1H7UG4Gkx8vS/gniqq2XehXhvCPdiXfFd4y3uDQqts7npOkhQXS0Ols1tYf6Doe9rv4O15xBHU6BOvHH03HPx/xmgdwa66i2RTNCn8uC0m+pkDqPOdx8fG+v0oYKqhHBW4NH6A5Tlb9bWq8uuNS6H3EBNe956OwDAiD3ue90bgLvWVNVY5HllNBi87B6drqE8F8moo/tegwvFF1O31aDpdmXxsHYfCzJNKMxRn6v+iz7se9uClScNgSkvmODTtT2ZZqyyf8P+WXBOxazu9AW8+NmVxS/Wlb/ZptfVvtU5CJfJhM4BFGUHTlxvE5Obdgp7+Lpo0Gl+ftJO571xs8EYlJytllmF/bvtY9iIDupkdrdw+5NuTTJF1X14b3fdIJjD7FZPO/MNSDVF2UNjaMui1t9gQY+PAW8xE5fd0dq+sb35p3v+BH4MXG5mJu/EY8zbe13TFD883Nl5KV3MvFCcus9ttJ86Gag0uoZyetGdPtvCDOVxduXbK35QmJM3O3BbE3ptMdIyDH6F8Ja5nhcFXpT6UxwDekPgT2XZK0ggGYoQFmQcRRj3YhwRLEteQUQUxygiue9LFFFeRnDffJ5yhSKxF6RzLu8Rnoh5L7xEiEe4RH0jHhPL/rlDiEw5yjVhlGW4DpOKo5ESQahE80OU58chisSChGikJBFJFGEIVZJnc/O4Q0IuCMooTVgm54b0EmFMCIrmhzFJGcoB4zzD88O4JCHKDpNe5CcoEkFAuAex9ASaU5YQSnAPUuisqA7nIlAob1yyCM8ch8wFaDxciTDBrfle4qH7h/vANr5OyFU89+w7JFa+h9Yoz1iWob4JQvwQXUcwIhhaiUIKL0R5E1JBLaCIopTg6wQ8weMRoZQUjUfENJJoTkUieYTWNVRowFIUYTxUqNdScopnTgbSZygHMvTEcoi4R+CH60BvCVAOZCYyH9VRhMQByo6CMwDeYRVhPEVzqigkFeVAMU4VmjnFVIB3PtilnOC++V4Uo3tbRR7FK15FPMb7gYoJx98LKlYZQ3eWSmRAcZ1ERhzNgkqUH+ORpoQKtPeqVIoA5dqH9u+j7PjQYAmu4zN4O6JIIMMYzZyfwQsI7eQBpRne/QPKFUF5CzhLFMpBoGSqUA6CRDAfZTRIpIrQHRxkjOAdCTZWKFDfQkJSicYTEiETlFF4ATKK7qxQkChGfQuFiAO0h4Q+C/CdFYbUZ2i2AQlf8S0SWYjHE1OV4l7HLPVxdhIpBFoHgAQKzU9EvNRDOYigpiia7dfPo1Ew9EsUCb0UPzlEoUwTfJ2IyhiNB9pBgr8XYkl8hTIaK5Hg3T/2iZfh6/hERmjmYp/yCO29sc8y/LwTQ2HjHSkOuReg8SQMtjDKTiKHgwCOiJCi1Zsoznw0HkA4QfOTxB4N8HVi5uP5SVIuQ3SdFE6kEVq9KVWhj3KQciYl6lsaEBGgHqTQ+33U6zSkEcV1gDX8KycNGcFP+GnkZa/EE5EUPzmkEfcEmp80EQQ/i2XwHSFxBM7XCvUtCwVLUd4y+JLgqNdZyrLpu2Q7QfBNWu+GW6zhI3kaZbbpV/WkEev64Eq9+jzcc22HGQf3NSqbBT+Yo3XmObI/HxZws5mArtZVlTmdj0/DhUZijuO4+qzd483aGEy9c6i0MMe/5YtsuOUx7mdnz+2EXpxuPzaFubkOR4pZs2z6T2W9yLvzYb9oNdpdn0HnpvjHkxvZuZFy2fXw3W8GVj7p2wVN22+iXyeK88rth7sB81m37XSlcHgkD+uqfDz1ZLgV6OGp0O7r+HB4pDNGR4xO2Pig8yEymD0PbjK6yJ7NY4uM3WR8kfGbTCwycZPJRSYH2enaGleVzdeH9ffhID/aqrIXU/xyw38QTSSMN0Z/9Appnl3pqz33L+YO2DC5fWmh0L2GFI+peqE8FvadL8OFXV5CEe6v9eF2FfaXyfGq7Pq9abXTvXUL9tcRI2JX2PxjMVwzTvLfkwiOpCojm0BkfMMlJZtICLWhvpfFfhikgQr+M++v5a78w38BAAD//wMAUEsDBBQABgAIAAAAIQDBqZFxPA8AABWMAAAPAAAAd29yZC9zdHlsZXMueG1s7F3bctvIEX1PVf4BxafkwRYp6mbXard0jV2xvVpLXj8PgaGINYBhANCy9nfykMpDvsI/lrkBHLAxAHow1m628mKLAPoA6NOnp3tw++6HL2kSfKZ5EbPsdDJ7Pp0ENAtZFGf3p5MPd9fPTiZBUZIsIgnL6OnkkRaTH77/85++e3hZlI8JLQIOkBUv0/B0sirL9cu9vSJc0ZQUz9maZnzlkuUpKfnP/H4vJfmnzfpZyNI1KeNFnMTl497+dHo00TD5EBS2XMYhvWThJqVZKe33cppwRJYVq3hdVGgPQ9AeWB6tcxbSouAnnSYKLyVxVsPMDgBQGoc5K9iyfM5PRh+RhOLms6n8K022AIc4gP0aIA1fvr7PWE4WCfc+P5KAg02+5+6PWHhJl2STlIX4md/k+qf+Jf+7ZllZBA8vSRHG8enkgiTxIo8nfAklRXlWxKSxcHWWFc3NwuJ0chennOd39CF4z1KSTfYEdEKye77+M0lOJ+vy2fn7Jmi9aBFHHJHkz27PhOGePjb1v3HE6/qX2mrn9Di7nOtbFXJ8LV2+YeEnGt2WfMXpZCp2xRd+eH2TxyznYXU6efFCL7ylafwqjiKaGRtmqziiH1c0+1DQaLv8p2sZGnpByDYZ/3t+PJMuT4ro6ktI1yLQ+NqMpHzX74RBIrbexNudS/N/VGAz7bM2+xUlQm3BbBdCHj4KYl9YFMbZtmNuds5dboXa0fypdnTwVDs6fKodHT3Vjo6fakcnT7UjCfMtdxRnEf2ihAh3A1D7cCxqRONYxIbGsWgJjWORChrHogQ0jiXQ0TiWOEbjWMIUgVOy0BaFRrDPLdHejds/Rrjh9g8Jbrj9I4Abbn/Cd8Ptz+9uuP3p3A23P3u74fYnazyuKrWC11xmWTlaZUvGyoyVNCjpl/FoJONYsgXxgycGPZp7OUkPMCqz6YF4NFpI5O/+CJEidR/PS9E1BWwZLOP7Tc4717EHTrPPNOE9ZECiiON5BMxpucktHnGJ6Zwuac47eeozsP2BJnFGg2yTLjzE5prce8OiWeTZfRWil6RQBzTZlCshkthDUKckzNn4Q2PEW354ExfjfSVAgvNNklBPWO/8hJjEGt8bSJjxrYGEGd8ZSJjxjYHBmS8XaTRPntJonhym0Tz5TcWnL79pNE9+02ie/KbRxvvtLi4TmeLNqmM2fO7uImFi0nj0cdzG9xnhBcD44UbPmQY3JCf3OVmvAjEF3A5rnjN2P+csegzurGNaP/RQF9c78lX2ywi64E6Js814fzfQfGmvxvOkvhrPk/5qvPEKfMuraFG/vfLT7txuFmWrpocH3C1JNqreHS9GUo6PsK0AruO88CaDdlgPEfxOVLuCTh+JcXuU4w9sizVeVrtZyevhaUgPR5mw8FNHlkZAvXpc05x3bZ9GI12zJGEPNPKHeFvmTMWaKfl9SckgyV+l6xUpYtlKNSCGVwLV1ejgLVmPPqGbhMSZH96unqUkTgJ/Bcaru7dvgju2Fl2ocIwfwHNWliz1hqknCv/ykS7+6ucAz3iPnD16OtszT7NHEuwi9jDIKCQWeULiVWicxV7GUIn3d/q4YCSP/KDd5FTdAFJST4i3JF2rosODtnhefOD5x0M1JPF+Jnkspo18ierOC5gxq1hsFr/QcHyqe8cCLxNHP25KOT0pS11p7Q9ufJnQgBtfIkg2+fAg4tfDyTbgxp9sA87XyV4kpChi6xVWZzxfp1vh+T7f8c2fxmMJy5ebxJ8DK0BvHqwAvbmQJZs0K3yescTzeMISz/f5egwZiedhxk7i/S2PI29kSDBfTEgwXzRIMF8cSDCvBIy/gccAG38XjwE2/lYeBeapBDDAfMWZ1+Hf00UgA8xXnEkwX3EmwXzFmQTzFWfzy4Aul7wI9jfEGJC+Ys6A9DfQZCVN1ywn+aMnyKuE3hMPE6QK7SZnS/FkAMvUPd4eIMUcdeKx2FZwvkj+SBfeDk1g+TwuDzOiJEkY8zS3th1wpKUxcXj4otfsbkXT8W30TUJCumJJRHPLOdlteb98uyahnqYHl+wGTXu+ie9XZXC7qmf7TZijaa9l1bA3zPp32Obzo/0Os7c0ijdpdaDwWYuj+XBjGdEN44N+420l0bA8HGgJ93nUb7mtkhuWxwMt4T5PBlpKnTYsu/RwSfJPrYFw3BU/dY9nCb7jriiqjVt32xVItWVbCB53RVFDKsFZGIqrBZCdYZqx2w8Tj90eoyI7CkZOdpTBurJDdAnsPf0ci5EdkzTl/uqbK3Z3N5dF9KDM+dOGqXn7xgWn4c98veaFU1bQoBVnPvzCVSPL2P04ON3YIQbnHTvE4ARkhxiUiazmqJRkRxmcm+wQg5OUHQKdreCIgMtW0B6XraC9S7aCKC7ZakQVYIcYXA7YIdBChRBooY6oFOwQKKECcyehQhS0UCEEWqgQAi1UWIDhhArtcUKF9i5ChSguQoUoaKFCCLRQIQRaqBACLVQIgRaqY21vNXcSKkRBCxVCoIUKIdBClfXiCKFCe5xQob2LUCGKi1AhClqoEAItVAiBFiqEQAsVQqCFCiFQQgXmTkKFKGihQgi0UCEEWqjqSUR3oUJ7nFChvYtQIYqLUCEKWqgQAi1UCIEWKoRACxVCoIUKIVBCBeZOQoUoaKFCCLRQIQRaqPJi4QihQnucUKG9i1AhiotQIQpaqBACLVQIgRYqhEALFUKghQohUEIF5k5ChShooUIItFAhRFd86kuUttvsZ/hZT+sd+8MvXemDem8+6W1CzYdDVUdlxxr+LMI5Y5+C1ucS57LfGAYSL5KYySlqy2V1E1feEoG68PnjRfcTPib6yHcy6Wch5DVTAH4w1BLMqRx0hbxpCZq8g65INy1B1XnQlX1NSzAMHnQlXanL6qYUPhwB4640YxjPLOZd2dowhy7uytGGIfRwV2Y2DKGDu/KxYXgYiOS8a3040E9H9f2lAKErHA2EYztCV1hCrqp0DIUxlDQ7wlD27AhDabQjoPi0wuCJtUOhGbZDuVENZYal2l2odgQs1RDBiWoA4041hHKmGkK5UQ0TI5ZqiICl2j052xGcqAYw7lRDKGeqIZQb1XAow1INEbBUQwQs1SMHZCuMO9UQyplqCOVGNSzusFRDBCzVEAFLNURwohrAuFMNoZyphlBuVIMuGU01RMBSDRGwVEMEJ6oBjDvVEMqZagjVRbWcRWlQjWLYMMcVYYYhbkA2DHHJ2TB06JYMa8duyUBw7JYgVxXnuG7JJM2OMJQ9O8JQGu0IKD6tMHhi7VBohu1QblTjuqU2qt2FakfAUo3rlqxU47qlTqpx3VIn1bhuyU41rltqoxrXLbVR7Z6c7QhOVOO6pU6qcd1SJ9W4bslONa5baqMa1y21UY3rltqoHjkgW2HcqcZ1S51U47olO9W4bqmNaly31EY1rltqoxrXLVmpxnVLnVTjuqVOqnHdkp1qXLfURjWuW2qjGtcttVGN65asVOO6pU6qcd1SJ9WWbmnvofF9JoEtPw7GNy4f11S8ott4YCZSryjVFwHlhq+j+jtKwlgcSaC/LaUXywPWFwzl33nBuzq9zXR6djJ/cXWmtrJ8+2r3K1aB+bmqlpX6W1itH78qfq12va+v6RW/Xoh9GcuML13Jc+zxSu2Hu3KTMPW9KdMP289Eyf0tSEGjHwVtwEvinXrVcgV2sSK5WrflrtpCR6fdtVfnB9PDK7XVWrm2UI+N8m0W4k1a/FRmU/nEkfp5timZ3kSTTJYlzeut5K+djSQ8U+82evM5qXevPal3rKldKBdcFPL/TzSv/TDXKWNLUCVgkyC1rJ+gkPuNhPpNXJawFVFGOZP3axLlDNBmeR+vhQzth+0FdbVd4+K5OnLLEZdCuB1Hy4VNE5J1SU2J3xouOl76DpEf0CJRbPE/XmcioB70Z9HUoUZfiILi6y9okrwlamu2tm+a0GWp1s6m8t0LO+sX6jWCVvtcjjdWgL3mwaif3SGivjugb4SwOP2WpgnPqqTF4fK+nLG+Hphhwk3BnSNTdHvuFQ+Qs8ISFnplX/IhXNbvaSTe/0d7EsvFyezoXOdskFhEHhByliyKH+834rAFvLKIZUgtxbtr38htD6t7wn4Jq33wcFi1JpDdseEsj9XrHXTSV793Msm+LmsaqV7eDOKa6i+4GkmygllDf/ylz9lmpq+wOnO9DqgWtTapOTmZH0+vG9TwzCIHdv5/tZ0oCVRuXLOCZ9X9qvAztpGKqzc5OZxWY2OFt6VcDxL7apBQAbB/LEsOGAC9lCI/dTmE66qsVV++3KLS7NmHW4cRpUOQDTZ3o0Os/PovsTrYst0WC31xYCddO/XbO6RdGO9ZRNbgvPW3jDCqkED/l8QfQRJbKnfjQq75+m+LGFQs/f6V0PBDfdbbt2TvnvV2zdjQvjqbz6peW59uKBq/7RbT6bV2x6ZaKL57ocpE50ZHVBQRXZAEDoCN9/RgBG+ADpF9fxnbdNWUx8VM3wBq6zPJijeKfSLbbiW/sKx+7ShuVjcz20irlj2F4nZ9uUuSXB9ENOBbfP2nORjtkLXTI3WRhiRsBDu/AyLqbm3rctmfhSzlcoko9LiclxEXtFv93OjuLI6sXk3V9NzxyfT84lKtwTRu5yznhaoaBWXjZuYGcWK/8kFQ/sHHNFp/VZsnF7OW122dk23d8jlZVw2hk3HMAyiir8aZ/+xmrnpT0/3/y110+wBxwzNGTpZ8/BZdMmyf5VTl9l08bYowhwmohbmevLP3qLzYOrjWN9/r4ko1nsKH4sEZPSsbigz3pdyQRL9STflBmnhIxNtZPJARvv5HrAxmI7KvmnC0Jd62dGHOCT7BdOtC/fvN5/vqfNzFhdjGMlet0nN3zFXVY2/hYQ/K+eVsf6YJ00H5EEfsQbxMM2fNuVM1IyOCkl5eWde8a6zpaSTkmwtPJ2nMB5pXgjWT29aVktvWNWFRGovP48il7ZDL1LDbiDIwEuPGNnw6bVuLSKZwtbdUesHyNa+wRI5qKbbrb7m1pQ/rhJMB2VVq60D93YazqM55H1XKT6CaJXu9sO6L6yWiVKx+Osbrunx2c9cTrx7mkXZJ2uVebiBqeLmJ+yDSCLDuSLCz/8fhZ+jEuxwxtiXMLjtqQBlT4vzm6msWTCf71VDclORvyk/NhnLlR7oARDQ+KIVnYfBUzMHx4dlls978NhdbYXFa/VV8/18AAAD//wMAUEsDBBQABgAIAAAAIQD2v9Zn4gAAAFUBAAAYACgAY3VzdG9tWG1sL2l0ZW1Qcm9wczEueG1sIKIkACigIAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAJyQTWvDMAyG74P+B6O7azdJ263EKWk+oNexwa6u4ySG2A62MzbG/vscduqOO4lHQnpelJ8/9ITepfPKGga7LQUkjbCdMgOD15cWPwLygZuOT9ZIBsbCudg85J0/dTxwH6yT1yA1ig0V67Vm8LWr9mn6dGxwVSZHnGVtgy80S/Chrdo6bei+rMtvQFFt4hnPYAxhPhHixSg191s7SxOHvXWah4huILbvlZC1FYuWJpCE0gMRS9TrNz1Bseb53X6Wvb/HNdri1H8tN3WblB0cn8dPIEVO/qhWvntF8QMAAP//AwBQSwMEFAAGAAgAAAAhAHQ/OXrCAAAAKAEAAB4ACAFjdXN0b21YbWwvX3JlbHMvaXRlbTEueG1sLnJlbHMgogQBKKAAAQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACMz7GKwzAMBuD94N7BaG+c3FDKEadLKXQ7Sg66GkdJTGPLWGpp377mpit06CiJ//tRu72FRV0xs6dooKlqUBgdDT5OBn77/WoDisXGwS4U0cAdGbbd50d7xMVKCfHsE6uiRDYwi6RvrdnNGCxXlDCWy0g5WCljnnSy7mwn1F91vdb5vwHdk6kOg4F8GBpQ/T3hOzaNo3e4I3cJGOVFhXYXFgqnsPxkKo2qt3lCMeAFw9+qqYoJumv103/dAwAA//8DAFBLAwQUAAYACAAAACEAqchcqowAAADaAAAAEwAoAGN1c3RvbVhtbC9pdGVtMS54bWwgoiQAKKAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAskmyCs4vLUpOLVYITs1JTS5JTQkuqcxJtVWKcQxw1IsI9lFSAAv4JeYCBYFiSgoVuTl5xVZJtkoZJSUFVvr6xckZqbmJxXr5Bal5QLm0/KLcxBIgtyhdPz8tLTM51SU/uTQ3Na9E38jAwEw/KTMpJzM/vSixIKMSahhVjLKz0Yd7xo6XCwAAAP//AwBQSwMEFAAGAAgAAAAhAC7TA/nyAQAA8AMAABAACAFkb2NQcm9wcy9hcHAueG1sIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAnFPBbtswDL0P2D8YPm2HRk6QtEWgqBjSDT2sSIC47ZmV6USYLQkSEyT7p33Ffmy03XjOttN8eo+knp9ISt4d6yo5YIjG2UU6HmVpgla7wtjtIn3Kv1zdpkkksAVUzuIiPWFM79T7d3IdnMdABmPCEjYu0h2RnwsR9Q5riCNOW86ULtRATMNWuLI0Gu+d3tdoSUyy7FrgkdAWWFz5XjDtFOcH+l/RwunGX3zOT571lMyx9hUQqjV4rHJTvwYo3Od1vnpMPow/jgpHRyn6Kpk7gqYM1XiacaKncg1bjGosRQfkiwsF82wiRQflcgcBNHFP1Wx2I8WAy0/eV0YDcbfVo9HBRVdSsmqvkDTnpRiWSL7WBvU+GDop9jGk8quxbGAqRQfYWYBtAL97s9czudFQ4ZIbokqoIkrxOyAfEJphr8E0/g40P6AmF5JovvO4J2nyChGbNi7SAwQDltKurCMtrnykoPKfP2hfOSn6SAuHhUNspo3NDlwWtqR1wfjSX26owrgq+Xb0D7vjod3WQ2d2YGfo7PyPP1SXrvZgucOiR9zib/HJ5+6+WY+3Ll4GB3N/MbTbeNA8levZ7XADBhm54SgWPNJ+Kn1APvANQtXo81m7xeJc83ei2ann7gGr8WyU8dcu0TnGq9C/LPULAAD//wMAUEsDBBQABgAIAAAAIQB3C1VhYAEAAK4CAAARAAgBZG9jUHJvcHMvY29yZS54bWwgogQBKKAAAQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB8kl1vwiAUhu+X7D803FdKXRtHak324dVMTNZly+4IHJWsUALM6r9fW23VzewS3ocnhxey2U6VwRask5WeIjKKUACaV0Lq9RS9FfNwggLnmRasrDRM0R4cmuW3Nxk3lFcWlrYyYL0EFzQm7Sg3U7Tx3lCMHd+AYm7UELoJV5VVzDdLu8aG8S+2BhxHUYoVeCaYZ7gVhmYwoqNS8EFpvm3ZCQTHUIIC7R0mI4JPrAer3NUDXXJGKun3Bq6ifTjQOycHsK7rUT3u0GZ+gj8WL6/dVUOp2644oDwTnHILzFc2fza+Uhk+22nbK5nzi6bolQTxsM/nUEoDGf6btLCFrWyfKCd3HTKse9PSSu1B5HEUj8NoEsZJQe5pMqFR9DlIeyg7tnSYB0TQ3I4euuiT9/HjUzFHR18cRqQghCbk4Pt1/iRUx7H/N6aNLiRpQVKapJfGXpB3Q1/+sPwHAAD//wMAUEsDBBQABgAIAAAAIQCJEvoqpwIAALcMAAASAAAAd29yZC9mb250VGFibGUueG1s1JZdT9swFIbvJ+0/RL6HOCb9oCKg0lFpN7vY2Hbtuk5r4Y/ITin99zu22/LRFBoEaCSKlBw7b+zH7znx2cWdksktt04YXaDsGKOEa2amQs8K9Pt6fNRHiaupnlJpNC/Qijt0cf71y9lyUBpduwTe126gWIHmdV0N0tSxOVfUHZuKa2gsjVW0hkc7SxW1N4vqiBlV0VpMhBT1KiUYd9Faxh6iYspSMP7NsIXiug7vp5ZLUDTazUXlNmrLQ9SWxk4raxh3DuasZNRTVOitTJbvCCnBrHGmrI9hMusRBSl4PcPhTsl7gU47AbIVUGzwfaaNpRMJ8GEkCYih8zX9ZDnQVEHDtVDcJT/4MvlpFNWhQ0W1cTyDPrdUFggTOLv4BHdwDheBuxylviObU+u4F4sdcQyXVAm52kRt0A0NlajZfBO/pVb4wcUmJ2bQsHATXKArjDG5Go9RjGQFGkGk1+9criPEfyscp+vIyTaCfYQFnfCYRR0WdLZ94JtpJLFD5NdKTYxsBNGBM8N+nXoAhMBTrxkEeRsQD6cUQWSPI/cgNpFGEP3HkQNBjMzCCm69OfbYogdWOA328LbIW9lCmSm3TThKccenh5siP/kIU/yFDPeVze3xxc7Rwhd0UZtPZIshDKs5PQi+hLTIgyHi2cYQbimc+0R1YkSlmFixh8Q4pERgAEny7iTyJhIk730IiWs6h9L2rCXiz8Nb451BZE0gunjXEuQlENlrLGFhY2ICCSrrHxDcDDlmzdtOfd8w/nA7pfr/WJChZ9Z9uCC5h9+wINmztSp2OW25IIF6cikpu9lLYxgo+N/6+1esQIMMn9Do9ce7efoyDdyWxogqKFj7fOH3dtEVfq/XjsTr9nhPfEF8xdpGPAlyIImXfbG+cef/AAAA//8DAFBLAwQUAAYACAAAACEAGmsxhU4CAAAnHgAAFAAAAHdvcmQvd2ViU2V0dGluZ3MueG1s7JnPb9sgFMfvk/Y/RL635ocBO2paqao6Teq2aut2d2ySoPHDAhIv/etH7DRN1x2aUzlw4vHgff3Qx5gnc3H1R8nJhlsnjJ5l8BxkE64b0wq9nGU/H27PymzifK3bWhrNZ9mWu+zq8uOHi37a8/kP7n2Y6SZBRbupambZyvtumueuWXFVu3PTcR0GF8aq2oeuXeaqtr/X3VljVFd7MRdS+G2OAKDZXsa+RcUsFqLhN6ZZK679EJ9bLoOi0W4lOvek1r9FrTe27axpuHNhPUqOeqoW+iADi1dCSjTWOLPw52Ex+4wGqRAOwWAp+SxAThNABwHVTD8vtbH1XAYEIZNJEMsuA4NWbNy+nfRT0c4yzBjFGJBqGJ+bdnszjG1qGfhm+c4bCNzxhX/ygoP3u1iu/uN+MN1r57Xx3qh//CGP69buLP8co8Obk4WOe9zN2xld3fC93RhpAvB67c0oIY8yOy1y/iKj02Lt8cpPCc2fFz2aL3EUADHECkQTjihwIMowLmnCEQUOAghkkEKYcESBAwUcALMy4YgCR0kooRCPR33C8c44aNgdpCoxSDhiwMFKyDCCFCUcMeCoMKEMYpCO8ihwQAQLXMGCpUo3Dh7hKGcEYZx4xMGDElwVFSCp1o2DBwu7AxQEs8QjDh4VpYyxKv1HjINHWYDwyQoFb+IRBY+qRFVVEZZ4RMEDAcIohbhM50ccPCCAACEMUn31bjzGdrgfNJ0XSjzyW2Ovrekdt+PTuNx+07++3A29WkrT33/9NKodXe9e/gUAAP//AwBQSwMEFAAGAAgAAAAhAI601zNVEAAATxwBABIAAAB3b3JkL251bWJlcmluZy54bWzsndtu48YZgO8L9B0MAwHSC6/nfFjECYZDMk2RBkF3g1zLsnYtRCdIsp3tVZFH6U0foe/TF8grlEd5eDAlUiY1ys7NyktyyPkPM/PNP6evvvl1Prt4nKw30+Xi5hK+AZcXk8V4eTddfLy5/Ol9eCUuLzbb0eJuNFsuJjeXnyaby2++/vOfvnp6u3iY307W0YMX0TsWm7dPq/HN5f12u3p7fb0Z30/mo82b+XS8Xm6WH7Zvxsv59fLDh+l4cv20XN9dIwBB8tdqvRxPNpvoPXq0eBxtLrPXzatvW64mi+jmh+V6PtpG/11/vJ6P1r88rK6it69G2+ntdDbdforeDVj+muXN5cN68TZ7xdUuQ3GSt2mGsp88xfqQ76ZJ/OX4YT5ZbJMvXq8nsygPy8Xmfrp6FqPr26Kb9/lLHpuEeJzP8ueeVpAcZwN/PXqKfp5feEj279JE81ma8+Y3QnCAReJX7FIckoXiN/OczEfTxfOHO6nGUC6k7V6Ayi9YfTzOON+ulw+r57dNj3vbd4tfdu+Ki3aLd2VGNkXbHJeZd/ejVVQC5+O3331cLNej21mUo8hkF5HWL2K3vvw6qnJGt5vtejTe/vAwvyj877u7m0uQPLLYTO+ie4+jWXQFhphDKC6v4zvzh9l2+v3kcTJ7/2k1yZ+5/3S7nt79Pb43i++lz27nq1n+BAwEwQSz9M7sMb4xjX7iL0Z/blezqKYCBDIAQJjkIaow19td8jRdVFuG893Fu8l4Oh9lH4ve9X7y6+7eF/DN7vrfxvnV2eTDNr28+nEd/0wXsZzx5ZtLjpKs3I8WH5OKGzMQP3u9e3id/YTLxXYTPzldbONcfBhFgmePJs9cJ58tCwrLgkJ5oKCz5dNk/f1ku52s64VFrYWFhDRKWy8CqojgJVei+jaqtB8ncf5biPSP5Xy0qJcI10m0nn68f1kkFGWoIBIUB4iEa9yvm0iN7khaWwgJ0cFCpMbJjrBQk9PR1iJFEnQQiQ7mdKy90xFcqjUOcjo2jNPx1haioEu1wIdzOtFeJF6qFg4SSQzmdLK90zFSqhpecLrrQvMev6Wx7U8lKrT9SESUIVCak85tv/IE1yQjCFPLn3Xb31MJcSzgWMCxgGMBxwKOBbqzAKqygA6DqB46kgWEJ5mUPMliUcsGCwSUCohx8EqO8uUX8C9tPUWgxOqGoySffhkIRpvxdHpz+X46n2wufpg8XaRmju5MRput2kxHtTfvVaTj2jvjTX2CMng8vb1N/9Wb5Pcpl0/K9Pbmn/kVRPIrOs5y4doskjS/ttpe/fi+mPXJ4uqnd/Gl2+ldlN/R+uqdyhRyGPYoHwAeBDr1q4L9bh9ms0kmS8l0//vXv9taDiKctOOHm65Gpb1owyAmSUDAscpKWZ/aYCRpYOzThgFbHCAWKL9/30CgUv9bog2D1aIaMqCa9u8biHBLS4qBeRhx5GE9gG8IkbCLfdowCJF4nmCI9O8bGFfw3BJtmICJcEAQzsYs+tQGh5aWFINNNcCBDFD/vkEgsqaktMRanCrHwFoIAsQBPBJrGWVaANqMtVn37tDOQoPRfv/tv22N1jrClQHtu0/z2+UsSZqiqnHhyBgYSBGgfTeqQTPLtnrZF/x6UTF6+bCeTtYxnRvaKV2NIb78YDutVcJsIG0cX1Vrv//2n7Z6q0TYDtXbz9HT8cyYjaG14rV2CqoG7VJHe2UFtS5w+2J2A5S4SgDQihK3L/J36hJXiTFaUuIq4cVTlbhqxNKOErcvYDlAiatEP60ocfvCnqcucZUAqyUlrhJbHa7EteTadGZXgWuRrzGWfipc53CtiPBIMb3T3M5mBtciJAmgOO14FGyUDfm2GKZoP3QLQbEzAsXZRWohyDoORqg265AUQrXZtb5DtcILted56sAytz/8jlqH3yFM25SdVQmiZ2LVzWo0nj4b6Cozmk1heaNPQZUPaNR7TeuPPgPRtDSisteiJwhEKw5ICOUAYXlJi2F5a7Rh9Bmiah0RxkXv2kBYJJ+1TxtGXyAIkGJMD6ANAZMCap82DPBnUjIY4ow7+gy9ItKyJThBIJpiQBDiAwSiGUumU9inDYPhWaAFhSAdvu81EA1kkf1OqI2WwE5T5ZjALqXEjBwbiEYe0h7N6aEe2LUmDGOcVmUFwMu+PiiwI5msCjlHtMtQztIpFwRowFFdv2w4jqeo3LLjczW2fbY2OJ6AEAGPDMDxstIz22PRE3B8qFXAiOyfThASxYFha7RhcHxAGfFQOAC5MtaytJ+A44UWBNNwAHIFpalX1mjD4HgoOABwgGl5mJQmDVijDTNmDzVBQvdfi2IByhxviTYMjpeaEs2yabi9cnzUqNiijZYcn7mKyfEeDwEUycu7c7xUcfvlHTShxK2ZcmumyhK5NVNuzdRL8xncmqkDnc6tmYpFcmumDmQBXmEBBKAi7NhBeEygFszPiMLUsmMBxwKOBfZbyLGAYwHHAmWRHAv0xwJZsNVkAYQ8CLSX5qQrC/iB0Ijp5gl52KfCJ37LwE2HGbcQsKRGObzdr8y4PWRU5/npeDBHraej11p/cpjbvep6Ey66qexzX29C04y01tvP089jvQlG6Niy+AddbyKJ1SXO2vUmrGPl/tolztr1JrhjVf5HX2/CYMe6+nNfbyI6VuHHl7iWeJvuXFvAW+ZTgOSR2wRLwAJCkdsmuNgBPExQF9pKr7rQ1l5xXGjrCKdzoS0X2vqMQ1uwekYAUgzCXbPdeZ9gJDWO1xvlgu7U7Ma53DiXg4H9FnIw4GDAwUBZJAcDPcJAKpIJAxjggMPwSBgADLKQSLUTdKdmc6BLBkL72ZrcglqzrzsYOKyIOBhwMOBgwMGAgwEHA0fAQPXUAAy44lyRNCudV7VzDQQHzTDgIgOvVEQcDDgYcDDgYMDBgIOBI2Cgutc6RiGgXPlpVtxyGAcDDgaOd08HAw4GHAw4GKhKZBEMVDeoxhxCCnCYZqUzDACovBAFO0F3anYw4GDAwcB+CzkYcDDgYKAskoOBHmGguvkt9jwMqHfknAFPES10/hZTzQ4GHAw4GNhvIQcDDgYcDJRFcjDQIwxUd9AkAIZQiiN3ynAw4GDgmCLiYMDBgIMBBwNlkRwM9AgD1S00CSKc+DoTrCsMhL5mhGA3TOBgwMGAgwEHAwWRHAw4GMglsggGqntoEsEgY4CkWem+h2YgkEeyt5hqNmAA0pBIhtPoxPGe8uUXsPXhaaJ0Yi5myadfJgL7Tra260A8LH3AmBjgkDSEi8eC7TXdUMf7GMjkhVKFvur/ICzISgdhWaMNg7aIwirQYoBD0kClAbBEGwas+cxTiHj9H4Qd9e4sLSkG52GuAwrpEEc/i+IhadZow0RED2oh6QBHxuEKn1uiDYMwsQKMcNr/IWmYQ0tLinn0swJE+eEQR8aVjhM8oTbacm1180wSCIwDH6XZ6r55ptIiBM17w5MQIh1qleRh8CAXlsXDY5FI0O+ckBaCpHNTYNrMcAWmza71zbQIU6XUAIc0Qkxb2u4ETCtCwgKEU+/utaXG6U7L9mnDYFqGqYZiiKNuES0e/GuNNgym1b7HaOD1rw0SIb6d2jCYlmukUTDAUbcROBdbamu0YTCt1hj4XraRVJ/aYLB08K812jCjph5kOiD9a4OX+8LWaMNgWsg5AR7uXxsCwKQps0EbLZkWVfeEjZoezSJMT7PVlWmV5ITH55PnutxZyGBahGWgRWahV4jqt+daCJpHx14A2/f3k/nk5nI+XSzXf42RNXlH/EftnfFmW5ug6gXj5Wy5znM+ethmRzg8IytMdtk9mDItHDEWXQbv7B4xpl0G76wdMcaoy+Cd3SPGssvgnd0jxuywwbuiSPaOGJMu1YLVI8YMNs8M2gswFo4Yi1LV8DojxigVqUAhgUCMU5xmpSuFxEfSILhn/7kAR8Ugj8QU1Jp9vdfIGgSlXhbNW/gmALE+spYA8GlGi0mAdUBxOs7xGjMAUOsZAIwXwypY0jMx6mY1Gk+f7XOV9wNsmhlgjoVTBVDI+o8bQliKlO016QnihjhERBE5wDwJhorREGu0YcYNo5qf4nydUp8xZQgs9Q2DYDUDSCsxwMwAykvTn2zRhgG/jGihQjJAhD1yQTu1YbAz86UIOc5iQX1qg5anxtmiDQO7hfB0GIb9ayPq8hTjhifURltir+4YTYN4MDztDxyx4ENy7QHcfJZUyKhSGNUQ+xBj4RVi5wkSnSHcoayDYxPcGSAviQahAodO5X1YrRr7u+1jlWWMRzBZlXxOfTO7cJ2GQgmRVx5D4vpe050A12HIhUe8/huaCq5bow0D1zHEgis0wGTNMq5bow0D1z0qCafkBLhujTYMXA8pVBzIASZrlnHdGm0YuK4lI54H+veNCq5bow0D1z1fBDIcYJi/gusn1EZbXK/u6c400QwfuyRLe17gSZLjQz2u+z7QEpBDp64OsyTrkGF+R3JFmxps7guFaKBOsiSr7fr6nrRhcK30EVJSDTCZt7okyxJtGFwbUKRgMED7hAAthaFt0YbBtVxRQvwhtEF4KZxkizYMrpU0IB4LBwjKC2FpSTG4VsaHrws2QOAVSUtLisG1QhEkggEWIWFemrBpjTbM6avaAxDpQZZkJQXUBm205drq8QQsCDEWKJvy0XlJVigFBH7zvkMcKQBANmxS4Fqephs0DI3huSBtOQyNszEDa8PQnkIsSBvwU4ShIQTFFVxRNXkmprY0Dh1iAAIhX2+KEG7dI4VYliYNuC7pcd0OJjHCVAzQCROkpelO0O3wpQd8MMR8j4i07dSG0e1QfiBQMMCYPuLYUrQ0p35joSAB/YdMMeSWdsKMbocfSOZ5A+wSgtnRW0X2pA2j26GlUpgFQ4TTmTW1aNtuR/UgFBYKnxCS4YM7IrXBageDt9vutFakhr4Fdtuduu1OK/Djtjtt5XRuu9NYJLfd6aEwUD0IhQPqQcSPPCKVMSC01M1L6ANKBcR1i9c6RjLcdqfJlVMGHF99C9v2C9jab4TqjFpjVHOKAAEBx2qAqa/nsGsrB4gFyu9/L8qz2LXVkyygmrpdWxNtYMSRh/UAvnEOu7YSzxMMDTAR+Cx2bQUIBwQNsVLtHHZt1QAHMhhgOcUZ79qKqkcTcU8BTr0j95ZQHmAqTAfdixZysToXq3OxuoPaXxerc7E6F6tzsbqySH3F6qpHE3FfhREMZIJ1hQFBIg4Tu1iCgwEHAw4GWlnIwYCDAQcDZZEcDBwFA4sEAhbZ+S2JBAUiyL+fT6de1CRLoyu1yfLlv3XJ0hW3tclQnvu6dOmKhvp0+e5AdenSKUm16WAS6XghXTrKVJsuWcXwQrI0qlL/uSZtpgBWL15DsvTondpkCXO9kAymgcL6zzUmbHAWlIe+ahM2uEuT+WCTvzTZDzY5TJNKYZPH5Dsa1SZscBnUVCSyU8Lrv9io1AavgU1emh3ZVJuwSanZtvj1H2wyI2rwm91uAbUJG/ymqbrIVvbXG6NJNdnSqfqcNvlNNvmxPmFTuga3aXK3LIBb/72mMpx19uoTNuq0wWt2K2+ShOnv7WQdNUtf/18AAAAA//8DAFBLAQItABQABgAIAAAAIQBjLW1vrgEAAGMIAAATAAAAAAAAAAAAAAAAAAAAAABbQ29udGVudF9UeXBlc10ueG1sUEsBAi0AFAAGAAgAAAAhAB6RGrfvAAAATgIAAAsAAAAAAAAAAAAAAAAA5wMAAF9yZWxzLy5yZWxzUEsBAi0AFAAGAAgAAAAhAC3P/4VKAQAATgYAABwAAAAAAAAAAAAAAAAABwcAAHdvcmQvX3JlbHMvZG9jdW1lbnQueG1sLnJlbHNQSwECLQAUAAYACAAAACEAOo4zOmkLAACOTAEAEQAAAAAAAAAAAAAAAACTCQAAd29yZC9kb2N1bWVudC54bWxQSwECLQAUAAYACAAAACEAYHWvXRUGAAA+GgAAEAAAAAAAAAAAAAAAAAArFQAAd29yZC9oZWFkZXIxLnhtbFBLAQItABQABgAIAAAAIQBGQBoTvgEAAP8FAAARAAAAAAAAAAAAAAAAAG4bAAB3b3JkL2VuZG5vdGVzLnhtbFBLAQItABQABgAIAAAAIQAJhWRUzQUAADIXAAAQAAAAAAAAAAAAAAAAAFsdAAB3b3JkL2Zvb3RlcjEueG1sUEsBAi0AFAAGAAgAAAAhAHihUAa/AQAABQYAABIAAAAAAAAAAAAAAAAAViMAAHdvcmQvZm9vdG5vdGVzLnhtbFBLAQItABQABgAIAAAAIQCqJg6+vAAAACEBAAAbAAAAAAAAAAAAAAAAAEUlAAB3b3JkL19yZWxzL2hlYWRlcjEueG1sLnJlbHNQSwECLQAUAAYACAAAACEAqiYOvrwAAAAhAQAAGwAAAAAAAAAAAAAAAAA6JgAAd29yZC9fcmVscy9mb290ZXIxLnhtbC5yZWxzUEsBAi0AFAAGAAgAAAAhADvOKlkPBgAAtRsAABUAAAAAAAAAAAAAAAAALycAAHdvcmQvdGhlbWUvdGhlbWUxLnhtbFBLAQItAAoAAAAAAAAAIQBTYoJPfEkBAHxJAQAVAAAAAAAAAAAAAAAAAHEtAAB3b3JkL21lZGlhL2ltYWdlMS5wbmdQSwECLQAUAAYACAAAACEA+Y+EA/cAAABwAQAAHAAAAAAAAAAAAAAAAAAgdwEAd29yZC9fcmVscy9zZXR0aW5ncy54bWwucmVsc1BLAQItABQABgAIAAAAIQA2SV1i0gYAAHEXAAARAAAAAAAAAAAAAAAAAFF4AQB3b3JkL3NldHRpbmdzLnhtbFBLAQItABQABgAIAAAAIQDBqZFxPA8AABWMAAAPAAAAAAAAAAAAAAAAAFJ/AQB3b3JkL3N0eWxlcy54bWxQSwECLQAUAAYACAAAACEA9r/WZ+IAAABVAQAAGAAAAAAAAAAAAAAAAAC7jgEAY3VzdG9tWG1sL2l0ZW1Qcm9wczEueG1sUEsBAi0AFAAGAAgAAAAhAHQ/OXrCAAAAKAEAAB4AAAAAAAAAAAAAAAAA+48BAGN1c3RvbVhtbC9fcmVscy9pdGVtMS54bWwucmVsc1BLAQItABQABgAIAAAAIQCpyFyqjAAAANoAAAATAAAAAAAAAAAAAAAAAAGSAQBjdXN0b21YbWwvaXRlbTEueG1sUEsBAi0AFAAGAAgAAAAhAC7TA/nyAQAA8AMAABAAAAAAAAAAAAAAAAAA5pIBAGRvY1Byb3BzL2FwcC54bWxQSwECLQAUAAYACAAAACEAdwtVYWABAACuAgAAEQAAAAAAAAAAAAAAAAAOlgEAZG9jUHJvcHMvY29yZS54bWxQSwECLQAUAAYACAAAACEAiRL6KqcCAAC3DAAAEgAAAAAAAAAAAAAAAAClmAEAd29yZC9mb250VGFibGUueG1sUEsBAi0AFAAGAAgAAAAhABprMYVOAgAAJx4AABQAAAAAAAAAAAAAAAAAfJsBAHdvcmQvd2ViU2V0dGluZ3MueG1sUEsBAi0AFAAGAAgAAAAhAI601zNVEAAATxwBABIAAAAAAAAAAAAAAAAA/J0BAHdvcmQvbnVtYmVyaW5nLnhtbFBLBQYAAAAAFwAXAO4FAACBrgEAAAA=")
        doc = Document(_io.BytesIO(_docx_raw))
        d   = FICHA_EMPRESA_DADOS

        # Atualizar data no parágrafo preservando o run original
        hoje = datetime.date.today()
        meses_pt = ['janeiro','fevereiro','março','abril','maio','junho',
                    'julho','agosto','setembro','outubro','novembro','dezembro']
        data_txt = f"Ipatinga,  {hoje.day} de {meses_pt[hoje.month-1]} de {hoje.year}"
        for p in doc.paragraphs:
            if 'Ipatinga' in p.text:
                # Preservar formatação do primeiro run, só trocar o texto
                if p.runs:
                    p.runs[0].text = data_txt
                    for r in p.runs[1:]:
                        r.text = ""
                break

        def _preencher_celula(cell, valor: str):
            """
            Preenche célula preservando toda a formatação original.
            Mantém o primeiro parágrafo/run e apenas troca o texto.
            """
            if not cell.paragraphs:
                return
            para = cell.paragraphs[0]
            if para.runs:
                # Preservar fonte, tamanho, negrito do run original
                run0 = para.runs[0]
                fonte_original = run0.font.name
                tamanho_original = run0.font.size
                negrito_original = run0.bold
                run0.text = valor
                # Restaurar formatação se foi perdida
                if fonte_original:
                    run0.font.name = fonte_original
                if tamanho_original:
                    run0.font.size = tamanho_original
                # Limpar runs adicionais
                for r in para.runs[1:]:
                    r.text = ""
            else:
                # Sem runs — criar um com a formatação do parágrafo
                run = para.add_run(valor)
                run.font.size = Pt(9)
            # Limpar parágrafos extras dentro da célula
            for extra_para in cell.paragraphs[1:]:
                for r in extra_para.runs:
                    r.text = ""

        # Preencher linha 1 (primeira linha de dados — índice 1)
        tab = doc.tables[0]
        if len(tab.rows) > 1:
            row = tab.rows[1]
            # Mapeamento exato das colunas conforme o template original:
            # [0] Empresa  [1] CNPJ  [2] Resp.Setor  [3] Tel.Setor
            # [4] Resp.Contrato  [5] Email.Contrato  [6] (vazia)
            # [7] Nome Aprendiz  [8] Horário  [9] Salário
            valores = {
                0: empresa     or d["empresa"],
                1: cnpj        or d.get("cnpj",""),
                2: resp_setor  or d["resp_setor"],
                3: tel_setor   or d["tel_resp_setor"],
                4: resp_contrato  or d["resp_contrato"],
                5: email_contrato or d["email_contrato"],
                6: "",
                7: nome_aprendiz.title(),
                8: horario or d["horario"],
                9: salario or d["salario"],
            }
            for idx, cell in enumerate(row.cells):
                if idx in valores:
                    _preencher_celula(cell, valores[idx])

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        print(f"Erro gerar ficha EPTOM: {e}")
        return b""

def enviar_ficha_eptom(nome_aprendiz: str, horario: str = "", salario: str = "", cnpj: str = "", empresa: str = "", resp_setor: str = "", tel_setor: str = "", resp_contrato: str = "", email_contrato: str = "") -> tuple[bool, str]:
    """Envia a ficha preenchida para a EPTOM por e-mail."""
    try:
        from email.mime.multipart import MIMEMultipart
        from email.mime.base      import MIMEBase
        from email                import encoders

        docx_bytes = gerar_ficha_eptom_docx(nome_aprendiz, horario, salario, cnpj, empresa, resp_setor, tel_setor, resp_contrato, email_contrato)
        if not docx_bytes:
            return False, "Erro ao gerar o arquivo .docx"

        nome_limpo = re.sub(r'[^A-Za-z0-9 ]', '', nome_aprendiz).replace(' ', '_')
        nome_arq   = f"Ficha_Empresa_HOVA_{nome_limpo}.docx"

        msg = MIMEMultipart()
        msg['Subject'] = f"Formulário Empresa — {nome_aprendiz.title()} — HOVA"
        msg['From']    = EMAIL_CONTA
        msg['To']      = EMAIL_EPTOM_RESP
        msg['Cc']      = EMAIL_CONTA  # cópia visível para o RH
        msg.attach(MIMEText(
            f"Bom dia!\n\nSegue em anexo o formulário da empresa preenchido para a aprendiz "
            f"{nome_aprendiz.title()}.\n\nAtenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço",
            'plain', 'utf-8'
        ))

        part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part.set_payload(docx_bytes)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{nome_arq}"')
        msg.attach(part)

        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as s:
            s.login(EMAIL_CONTA, SENHA_CONTA)
            s.send_message(msg, to_addrs=[EMAIL_EPTOM_RESP, EMAIL_CONTA])
        return True, f"Ficha enviada para {EMAIL_EPTOM_RESP} (cópia para {EMAIL_CONTA})"
    except Exception as e:
        return False, f"Erro: {e}"

# Assunto padrão para o candidato responder com documentos
# O sistema vai buscar e-mails com esse prefixo para salvar automaticamente
ASSUNTO_DOCS_PREFIX = "HOVA-DOCS"

def _assunto_docs(nome: str, cand_id: str) -> str:
    """Gera o assunto padronizado para o candidato responder com os documentos."""
    nome_limpo = re.sub(r'[^A-Za-z0-9]', '', nome.replace(' ', '_'))
    return f"{ASSUNTO_DOCS_PREFIX}-{nome_limpo}-{cand_id[:8]}"

def email_admissao(nome, dl, di=None, hi=None, cand_id=""):
    return f"""Prezada(o) {nome.title()}, bom dia!

Aqui é a equipe de RH do Hospital de Olhos Vale do Aço.

Temos o prazer de informar que você foi selecionada(o) para integrar nossa equipe.
Seja muito bem-vinda(o)!

Para darmos continuidade ao processo, precisamos que você nos envie os documentos
listados abaixo até o dia {dl.strftime('%d/%m/%Y')}.

COMO ENVIAR:
Responda este mesmo e-mail com os documentos em PDF anexados.
Nomeie cada arquivo com o nome do documento. Ex: RG.pdf, CPF.pdf
O assunto do e-mail ja esta correto — nao altere.

Documentos necessários:
  - RG
  - CPF
  - Comprovante de residência
  - Cartão do PIS
  - Diploma (se houver)
  - Cartão de vacinação (Hepatite B e Tétano atualizados)
  - Certidão de casamento (se houver)
  - Certidão de nascimento dos filhos + CPF (se houver)
  - Declaração escolar dos filhos (se houver)

A foto 3x4 deverá ser entregue presencialmente.

A data de início será informada em breve!
Qualquer dúvida, estamos à disposição.

Atenciosamente,
Equipe de RH — Hospital de Olhos Vale do Aço"""

def email_admissao_aprendiz(nome, dl, cand_id=""):
    """E-mail de admissão para Jovem Aprendiz — sem data/horário de início (definidos pela EPTOM)."""
    return f"""Prezada(o) {nome.title()}, bom dia!

Aqui é a equipe de RH do Hospital de Olhos Vale do Aço.

Temos o prazer de informar que você foi selecionada(o) para o Programa de Jovem Aprendiz!
Seja muito bem-vinda(o)!

Para darmos continuidade ao processo, precisamos que você nos envie os documentos
listados abaixo até o dia {dl.strftime('%d/%m/%Y')}.

COMO ENVIAR:
Responda este mesmo e-mail com os documentos em PDF anexados.
Nomeie cada arquivo com o nome do documento. Ex: RG.pdf, CPF.pdf
O assunto do e-mail ja esta correto — nao altere.

Documentos necessários:
  - RG
  - CPF
  - Comprovante de residência
  - Cartão do PIS
  - Certidão de nascimento
  - Comprovante de matrícula escolar atualizado
  - Cartão de vacinação atualizado
  - Certidão de nascimento dos filhos + CPF (se houver)
  - Declaração escolar dos filhos (se houver)

A foto 3x4 deverá ser entregue presencialmente.

A data de início e o horário de trabalho serão informados em breve pela EPTOM.
Qualquer dúvida, estamos à disposição!

Atenciosamente,
Equipe de RH — Hospital de Olhos Vale do Aço"""

def send_email_admissao(dest: str, nome: str, dl, di, hi, cand_id: str,
                        aprendiz: bool = False) -> bool:
    """Envia o e-mail de admissão. Para aprendiz, usa mensagem sem data/hora de início."""
    try:
        if aprendiz:
            corpo = email_admissao_aprendiz(nome, dl, cand_id)
        else:
            corpo = email_admissao(nome, dl, di, hi, cand_id)
        m = MIMEText(corpo, 'plain', 'utf-8')
        m['Subject'] = _assunto_docs(nome, cand_id)
        m['From']    = EMAIL_CONTA
        m['To']      = dest
        m['Bcc']     = EMAIL_CONTA
        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as s:
            s.login(EMAIL_CONTA, SENHA_CONTA)
            s.send_message(m)
        return True
    except:
        return False

def varrer_documentos_recebidos() -> list[dict]:
    """
    Varre a caixa de entrada buscando respostas dos contratados com PDFs.
    Reconhece pelo assunto padronizado HOVA-DOCS-NOME-ID.
    Retorna lista de {cand_id, nome_doc, bytes_pdf, email_remetente}.
    """
    encontrados = []
    try:
        conn = imaplib.IMAP4_SSL(IMAP_SERVER, 993, timeout=15)
        conn.login(EMAIL_CONTA, SENHA_CONTA)
        conn.select("INBOX")
        _, ids = conn.search(None, f'(SUBJECT "{ASSUNTO_DOCS_PREFIX}")')
        for mid in (ids[0].split() or [])[-100:]:
            try:
                _, md = conn.fetch(mid, '(RFC822)')
                msg   = email.message_from_bytes(md[0][1])
                subj_raw = msg.get('Subject','')
                try:
                    dec, enc = decode_header(subj_raw)[0]
                    subj = dec.decode(enc or 'utf-8', errors='replace') \
                           if isinstance(dec, bytes) else str(dec)
                except:
                    subj = subj_raw

                if ASSUNTO_DOCS_PREFIX not in subj:
                    continue

                # Extrair cand_id do assunto: HOVA-DOCS-NOME-XXXXXXXX
                partes    = subj.split('-')
                cand_id_8 = partes[-1].strip() if partes else ''

                for part in msg.walk():
                    fn = part.get_filename() or ''
                    if fn.lower().endswith('.pdf'):
                        payload = part.get_payload(decode=True)
                        if payload:
                            # Nome do doc = nome do arquivo sem extensão
                            nome_doc = os.path.splitext(fn)[0].strip()
                            encontrados.append({
                                'cand_id_8':   cand_id_8,
                                'nome_doc':    nome_doc,
                                'bytes_pdf':   payload,
                                'remetente':   email.utils.parseaddr(
                                               msg.get('From',''))[1].lower(),
                            })
            except:
                continue
        conn.logout()
    except:
        pass
    return encontrados

# ──────────────────────────────────────────
# ALERTAS — respostas de candidatos agendados
# ──────────────────────────────────────────
TEMPLATES_MSG = {
    "Confirmação de entrevista (WhatsApp)": {
        "canal": "whatsapp",
        "texto": "Olá! 😊\n\nAqui é a equipe de RH do *Hospital de Olhos Vale do Aço*.\n\nEstamos te aguardando para a entrevista:\n📅 *{data}* às *{hora}*\n📍 Rua Ponte Nova, 185 - Centro, Ipatinga/MG\n\nPergunte por *Josi* ou *Paula*. Até lá! 🤝"
    },
    "Lembrete de documentos (WhatsApp)": {
        "canal": "whatsapp",
        "texto": "Olá! 😊\n\nAqui é o RH do *Hospital de Olhos Vale do Aço*.\n\nLembrando que precisamos dos seus documentos até *{prazo}*.\n\nEnvie por e-mail: *rh@holhosvaledoaco.com.br*\n\nDúvidas? Estamos à disposição! 🤝"
    },
    "Reagendamento necessário (WhatsApp)": {
        "canal": "whatsapp",
        "texto": "Olá! 😊\n\nAqui é o RH do *Hospital de Olhos Vale do Aço*.\n\nPrecisamos reagendar sua entrevista. Poderia nos informar sua disponibilidade?\n\nAguardamos seu retorno. Obrigada! 🤝"
    },
    "Convite para entrevista (E-mail)": {
        "canal": "email",
        "texto": "Olá,\n\nO Hospital de Olhos Vale do Aço analisou seu perfil e você foi selecionado(a) para a próxima fase.\n\nTemos disponibilidade para {data}. Responda com o número da sua escolha:\n1 - {h1}\n2 - {h2}\n3 - {h3}\n\nEndereço: Rua Ponte Nova, 185 - Centro, Ipatinga/MG\n\nAtenciosamente,\nEquipe de RH — HOVA"
    },
    "Não selecionado (E-mail)": {
        "canal": "email",
        "texto": "Olá,\n\nAgradecemos seu interesse no Hospital de Olhos Vale do Aço.\n\nApós análise, não temos vaga compatível com seu perfil no momento. Seu currículo ficará em nossa base.\n\nAtenciosamente,\nEquipe de RH — HOVA"
    },
}

def varrer_alertas() -> tuple[int, list[str]]:
    """
    Varre caixa de entrada buscando e-mails de candidatos AGENDADOS.
    Regra simples: qualquer e-mail de um agendado → gera alerta.
    Não interpreta conteúdo. Usa apenas o remetente para vincular.
    """
    novos = 0
    logs  = []

    # Mapear e-mails dos agendados para acesso rápido
    mapa_agendados = {
        c.get('email','').lower().strip(): c
        for c in st.session_state.agendados
        if c.get('email','').strip()
    }
    if not mapa_agendados:
        return 0, ["Nenhum candidato agendado com e-mail cadastrado."]

    # IDs já processados para alertas (evita duplicatas)
    if 'alertas_processados' not in st.session_state:
        st.session_state.alertas_processados = set()

    try:
        conn = imaplib.IMAP4_SSL(IMAP_SERVER, 993, timeout=15)
        conn.login(EMAIL_CONTA, SENHA_CONTA)
        conn.select("INBOX")

        # Buscar e-mails dos últimos 14 dias
        data_corte = (datetime.date.today() - datetime.timedelta(days=14)).strftime("%d-%b-%Y")
        _, ids = conn.search(None, f'SINCE {data_corte}')
        ids_lista = ids[0].split() if ids[0] else []

        for mid in ids_lista[-200:]:  # máximo 200 e-mails
            mid_str = mid.decode() if isinstance(mid, bytes) else str(mid)
            chave_alerta = f"alerta_{mid_str}"

            if chave_alerta in st.session_state.alertas_processados:
                continue

            try:
                _, md = conn.fetch(mid, '(RFC822)')
                msg   = email.message_from_bytes(md[0][1])
                rem   = email.utils.parseaddr(msg.get('From',''))[1].lower().strip()

                if rem not in mapa_agendados:
                    continue

                # Candidato agendado respondeu — gerar alerta
                cand = mapa_agendados[rem]

                # Extrair trecho do corpo (primeiras 300 chars, sem quoted)
                corpo = ''
                for pt in msg.walk():
                    if pt.get_content_type() == 'text/plain':
                        try:
                            raw = pt.get_payload(decode=True).decode('utf-8', errors='ignore')
                            linhas = [l for l in raw.splitlines()
                                      if l.strip() and not l.strip().startswith('>')][:8]
                            corpo = ' '.join(linhas)[:300]
                            break
                        except: pass

                assunto_raw = msg.get('Subject','')
                try:
                    dec, enc = decode_header(assunto_raw)[0]
                    assunto = dec.decode(enc or 'utf-8', errors='replace') if isinstance(dec,bytes) else str(dec)
                except:
                    assunto = assunto_raw

                try:
                    data_msg = parsedate_to_datetime(msg.get('Date','')).strftime('%d/%m/%Y %H:%M')
                except:
                    data_msg = datetime.datetime.now().strftime('%d/%m/%Y %H:%M')

                alerta = {
                    "tipo":     "resposta",
                    "assunto":  assunto[:100],
                    "mensagem": corpo or "(sem conteúdo de texto)",
                    "data":     data_msg,
                    "lido":     False,
                    "mid":      mid_str,
                }

                # Adicionar alerta ao candidato (evita duplicata por mid)
                if 'alertas' not in cand:
                    cand['alertas'] = []
                ja_existe = any(a.get('mid') == mid_str for a in cand['alertas'])
                if not ja_existe:
                    cand['alertas'].append(alerta)
                    novos += 1
                    logs.append(f"Alerta: {cand['nome']} respondeu — {assunto[:50]}")

                st.session_state.alertas_processados.add(chave_alerta)

            except Exception as e:
                logs.append(f"Erro ao processar {mid_str}: {e}")
                continue

        conn.logout()
        if novos > 0:
            salvar_json()
    except Exception as e:
        logs.append(f"Erro de conexão IMAP: {e}")

    logs.append(f"Varredura concluída. {novos} novo(s) alerta(s).")
    return novos, logs

# ──────────────────────────────────────────
# BUSCA DE CURRICULOS
# ──────────────────────────────────────────
def buscar_curriculos(limite):
    logs = []
    capturados = 0

    # 1. Conectar
    try:
        conn = imaplib.IMAP4_SSL(IMAP_SERVER, 993, timeout=15)
        conn.login(EMAIL_CONTA, SENHA_CONTA)
        conn.select("INBOX")
        logs.append("Conexao IMAP estabelecida.")
    except Exception as e:
        return 0, [f"ERRO de conexão: {e}"]

    # 2. Listar e-mails
    try:
        status, dados = conn.search(None, 'ALL')
        if status != 'OK' or not dados[0]:
            conn.logout()
            return 0, ["Nenhum e-mail encontrado na caixa."]
        ids = dados[0].split()
        # Mais recentes primeiro
        ids_varrer = ids[-limite:][::-1]
        logs.append(f"{len(ids)} e-mail(s) na caixa. Varrendo os {len(ids_varrer)} mais recentes.")
    except Exception as e:
        conn.logout()
        return 0, [f"Erro ao listar e-mails: {e}"]

    # 3. Processar — prioridade para e-mails da EPTOM
    emails_em_triagem  = {c['email'] for c in st.session_state.cvs}
    emails_processados = st.session_state._processados

    # ── Processar EPTOM primeiro (prioridade máxima) ──────────
    try:
        _, dados_eptom = conn.search(None, f'FROM "{EMAIL_EPTOM}"')
        ids_eptom = dados_eptom[0].split() if dados_eptom[0] else []
        for mid in ids_eptom[-20:]:
            try:
                _, md = conn.fetch(mid, '(RFC822)')
                if not md or not isinstance(md[0], tuple): continue
                msg    = email.message_from_bytes(md[0][1])
                msg_id = msg.get('Message-ID') or mid.decode()
                for part in msg.walk():
                    if part.get_content_maintype() == 'multipart': continue
                    fn = part.get_filename() or ''
                    if not fn.lower().endswith(('.pdf','.doc','.docx')): continue
                    chave = f"eptom::{msg_id}::{fn}"
                    if chave in st.session_state.historico_emails: continue
                    payload = part.get_payload(decode=True)
                    if not payload: continue
                    txt = ''
                    if fn.lower().endswith('.pdf'):
                        if pdfplumber:
                            try:
                                buf = io.BytesIO(payload)
                                with pdfplumber.open(buf) as pdf:
                                    txt = "\n".join(pg.extract_text() for pg in pdf.pages if pg.extract_text())
                            except: pass
                        if not txt and fitz:
                            try:
                                doc_f = fitz.open(stream=payload, filetype="pdf")
                                txt   = "\n".join(doc_f[i].get_text() for i in range(len(doc_f)))
                            except: pass
                    elif fn.lower().endswith(('.doc','.docx')) and docx2txt:
                        try:
                            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
                                tf.write(payload); tfn = tf.name
                            txt = docx2txt.process(tfn); os.remove(tfn)
                        except: pass
                    nome_arq   = os.path.splitext(fn)[0].replace('_',' ').replace('-',' ').upper().strip()
                    emails_pdf = re.findall(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', txt)
                    email_cand = emails_pdf[0].lower() if emails_pdf else f"eptom.{re.sub(r'[^a-z0-9]','',nome_arq.lower())}@aprendiz"
                    if email_cand in emails_processados or email_cand in emails_em_triagem:
                        st.session_state.historico_emails.add(chave); continue
                    try:
                        dt = parsedate_to_datetime(msg.get('Date',''))
                        ds = dt.strftime("%d/%m/%Y"); diso = dt.strftime("%Y-%m-%d"); mes_n = dt.month
                    except:
                        now = datetime.datetime.now()
                        ds = now.strftime("%d/%m/%Y"); diso = now.strftime("%Y-%m-%d"); mes_n = now.month
                    st.session_state.cvs.append({
                        "id": str(int(time.time()*1000))+str(capturados)+"E",
                        "nome": nome_arq, "email": email_cand, "telefone": "",
                        "data": ds, "data_iso": diso, "mes_num": mes_n,
                        "cidade": "Ipatinga", "tags": ["EPTOM","Jovem Aprendiz"],
                        "preview": resumo(txt), "setor": "JOVEM APRENDIZ",
                        "nome_arquivo": fn, "arquivo_bytes": payload, "foto": None,
                        "manual": False, "eptom": True, "primeiro_emprego": True,
                        "cidade_longe": False, "cidade_longe_nome": "",
                        "motivo_mudanca": "", "motivo_rejeicao": "",
                        "obs_triagem": "📋 Enviado pela EPTOM — preencher ficha após contratação.",
                        "duvida_enviada": False,
                    })
                    st.session_state.historico_emails.add(chave)
                    emails_em_triagem.add(email_cand)
                    capturados += 1
                    logs.append(f"[EPTOM] Capturado: {nome_arq}")
            except Exception as e:
                logs.append(f"[EPTOM] Erro: {e}"); continue
    except Exception as e:
        logs.append(f"[EPTOM] Erro busca: {e}")

    for mid in ids_varrer:
        try:
            status, md = conn.fetch(mid, '(RFC822)')
            if status != 'OK' or not md or not isinstance(md[0], tuple):
                continue

            msg        = email.message_from_bytes(md[0][1])
            msg_id     = msg.get('Message-ID') or mid.decode()
            remetente  = email.utils.parseaddr(msg.get('From',''))[1].lower().strip()

            # Decodificar assunto
            assunto_raw = msg.get('Subject','')
            assunto = ''
            if assunto_raw:
                try:
                    dec, enc = decode_header(assunto_raw)[0]
                    assunto = dec.decode(enc or 'utf-8', errors='replace') if isinstance(dec, bytes) else str(dec)
                except:
                    assunto = assunto_raw
            al = assunto.lower()

            # Verificar se é CV
            tem_palavra = any(p in al for p in PALAVRAS_CV)
            tem_anexo   = any(
                (p.get_filename() or '').lower().endswith(('.pdf','.doc','.docx'))
                for p in msg.walk()
            )
            if not tem_palavra and not tem_anexo:
                continue

            for part in msg.walk():
                if part.get_content_maintype() == 'multipart': continue
                fn = part.get_filename() or ''
                if not fn.lower().endswith(('.pdf','.doc','.docx')): continue

                # Dedup por sessão
                chave = f"{msg_id}::{fn}"
                if chave in st.session_state.historico_emails:
                    logs.append(f"Sessão: já visto — {fn}")
                    continue

                payload = part.get_payload(decode=True)
                if not payload:
                    logs.append(f"Anexo vazio: {fn}")
                    continue

                # ── Extrair texto direto do payload (sem arquivo temp)
                # Usa io.BytesIO para compatibilidade com Streamlit Cloud
                txt  = ''
                foto = None

                if fn.lower().endswith('.pdf'):
                    # Tentativa 1: pdfplumber via BytesIO
                    if pdfplumber:
                        try:
                            buf = io.BytesIO(payload)
                            with pdfplumber.open(buf) as pdf:
                                paginas = []
                                for pg in pdf.pages:
                                    t = pg.extract_text()
                                    if t: paginas.append(t)
                                txt = "\n".join(paginas)
                        except Exception as e:
                            logs.append(f"pdfplumber erro: {e}")
                            txt = ''

                    # Tentativa 2: fitz (PyMuPDF) via stream
                    if not txt and fitz:
                        try:
                            doc = fitz.open(stream=payload, filetype="pdf")
                            txt = "\n".join(doc[i].get_text() for i in range(len(doc)))
                        except Exception as e:
                            logs.append(f"fitz erro: {e}")

                    # Tentativa 3: arquivo temp em /tmp (sempre gravável)
                    if not txt:
                        try:
                            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                                tf.write(payload)
                                tfname = tf.name
                            if pdfplumber:
                                with pdfplumber.open(tfname) as pdf:
                                    txt = "\n".join(pg.extract_text() for pg in pdf.pages if pg.extract_text())
                            os.remove(tfname)
                        except Exception as e:
                            logs.append(f"temp fallback erro: {e}")

                    # Foto via fitz stream
                    if fitz:
                        try:
                            doc  = fitz.open(stream=payload, filetype="pdf")
                            imgs = doc.get_page_images(0) if len(doc) > 0 else []
                            if imgs:
                                foto = doc.extract_image(imgs[0][0])["image"]
                        except: pass

                elif fn.lower().endswith(('.doc', '.docx')) and docx2txt:
                    try:
                        ext = '.docx' if fn.lower().endswith('.docx') else '.doc'
                        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tf:
                            tf.write(payload)
                            tfname = tf.name
                        txt = docx2txt.process(tfname)
                        os.remove(tfname)
                    except Exception as e:
                        txt = f"Erro Word: {e}"

                logs.append(f"Texto extraído: {len(txt)} chars — {fn}")

                # E-mail do candidato
                emails_pdf  = re.findall(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', txt)
                email_cand  = emails_pdf[0].lower() if emails_pdf else remetente

                # Dedup por e-mail (outras etapas)
                if email_cand in emails_processados:
                    logs.append(f"Já em outra etapa: {email_cand}")
                    st.session_state.historico_emails.add(chave)
                    continue
                if email_cand in emails_em_triagem:
                    logs.append(f"Já na triagem: {email_cand}")
                    st.session_state.historico_emails.add(chave)
                    continue

                # Extrair outros dados
                cels    = re.findall(r'\(?\d{2}\)?\s?(?:9\s?)?\d{4}[\s\-]?\d{4}', txt)
                tel     = ''.join(filter(str.isdigit, cels[0])) if cels else ''
                cidades = ["ipatinga","coronel fabriciano","timoteo","santana do paraiso"]
                cidade  = next((c.title() for c in cidades if c in txt.lower() or c in al),'')
                sk_list = ["Atendimento","PABX","Pacote Office","Excel","Agendamento",
                           "Faturamento","Recepcao","Financeiro","Enfermagem","Triagem",
                           "Vendas","Administrativo"]
                tags    = [s for s in sk_list if s.lower() in txt.lower()]

                # Nome
                nome = assunto.upper()
                for r in ["CURRICULO","CURRÍCULO","CURRICULUM","CV","VAGA","PARA","DE",
                          "EMPREGO","CANDIDATURA","SELEÇÃO","SELECAO","-",":","/"]:
                    nome = nome.replace(r, " ")
                nome = re.sub(r'\s+',' ', nome).strip()
                if len(nome) < 3:
                    nome = (fn.upper()
                            .replace(".PDF","").replace(".DOCX","").replace(".DOC","")
                            .replace("_"," ").replace("-"," ").strip())
                if len(nome) < 3:
                    nome = remetente.split("@")[0].upper()

                setor = setor_cv(assunto, txt)

                try:
                    dt      = parsedate_to_datetime(msg.get('Date',''))
                    ds      = dt.strftime("%d/%m/%Y")
                    diso    = dt.strftime("%Y-%m-%d")
                    mes_n   = dt.month
                except:
                    now  = datetime.datetime.now()
                    ds   = now.strftime("%d/%m/%Y")
                    diso = now.strftime("%Y-%m-%d")
                    mes_n= now.month

                candidato = {
                    "id":str(int(time.time()*1000))+str(capturados),
                    "nome":nome, "email":email_cand, "telefone":tel,
                    "data":ds, "data_iso":diso, "mes_num":mes_n,
                    "cidade":cidade, "tags":tags,
                    "preview":resumo(txt),
                    "setor":setor, "nome_arquivo":fn,
                    "arquivo_bytes":payload, "foto":foto, "manual":False,
                    # ── Triagem inteligente ──
                    "primeiro_emprego": detectar_primeiro_emprego(txt),
                    "cidade_longe":     detectar_cidade_longe(txt, cidade)[0],
                    "cidade_longe_nome":detectar_cidade_longe(txt, cidade)[1],
                    "motivo_mudanca":   "",   # preenchido ao redirecionar setor
                    "motivo_rejeicao":  "",   # preenchido ao rejeitar
                    "obs_triagem":      "",   # observações livres
                    "duvida_enviada":   False,
                }

                st.session_state.cvs.append(candidato)
                st.session_state.historico_emails.add(chave)
                emails_em_triagem.add(email_cand)
                capturados += 1
                logs.append(f"Capturado: {nome} | {setor} | {email_cand}")

        except Exception as e:
            logs.append(f"Erro msg {mid}: {e}")
            continue

    try:
        salvar_json()
        conn.logout()
    except: pass

    logs.append(f"Concluido. {capturados} novo(s) currículo(s) capturado(s).")
    return capturados, logs

# ──────────────────────────────────────────
# SIDEBAR
# ──────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="background:linear-gradient(160deg,#003329,#004D40);
                margin:-1rem -1rem 0; padding:28px 24px 22px;
                border-bottom:1px solid rgba(255,255,255,0.08);">
        <div style="font-size:26px;font-weight:900;color:#FFFFFF;letter-spacing:-0.5px;">HOVA</div>
        <div style="font-size:9px;color:rgba(255,255,255,0.45);letter-spacing:3px;
                    text-transform:uppercase;margin-top:3px;font-weight:600;">
            Gestão de Talentos
        </div>
    </div>
    <div style="height:20px;"></div>
    """, unsafe_allow_html=True)

    st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)
    st.markdown("<span class='sb-label'>Filtro de setor</span>", unsafe_allow_html=True)
    filtro_setor = st.radio("", [
        "TODOS","TRIAGEM GERAL","RECEPCAO E ATENDIMENTO",
        "TECNICO E ENFERMAGEM","ADMINISTRATIVO","FATURAMENTO","JOVEM APRENDIZ"
    ], label_visibility="collapsed")

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)
    st.markdown("<span class='sb-label'>Filtro de período</span>", unsafe_allow_html=True)
    filtro_mes = st.selectbox("", ["Todos os meses"]+list(MESES_NOMES.values()),
                               label_visibility="collapsed")

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)
    st.markdown("<span class='sb-label'>Quantidade de e-mails a varrer</span>", unsafe_allow_html=True)
    limite_busca = st.select_slider("", options=[10,30,50,100,200,300],
                                     value=30, label_visibility="collapsed")
    st.caption(f"Últimos {limite_busca} e-mails serão analisados")

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)
    st.markdown("<span class='sb-label'>Cadastro manual</span>", unsafe_allow_html=True)
    with st.form("form_manual", clear_on_submit=True):
        mn  = st.text_input("Nome completo *", placeholder="Ex: Maria da Silva")
        me  = st.text_input("E-mail", placeholder="candidato@email.com (opcional)")
        mt  = st.text_input("Telefone / WhatsApp", placeholder="31999990000")
        ms  = st.selectbox("Setor", ["TRIAGEM GERAL","RECEPCAO E ATENDIMENTO",
                                      "TECNICO E ENFERMAGEM","ADMINISTRATIVO",
                                      "FATURAMENTO","JOVEM APRENDIZ"])
        ok_manual = st.form_submit_button("CADASTRAR", use_container_width=True, type="primary")
    if ok_manual:
        if mn:
            if not me and not mt:
                st.error("Informe pelo menos e-mail ou telefone.")
            else:
                st.session_state.cvs.append(novo_manual(mn, me, mt, ms))
                salvar_json()
                canal = "WhatsApp" if not me and mt else "e-mail"
                st.success(f"{mn.upper()} cadastrado. Agendamento via {canal}.")
        else:
            st.error("Nome é obrigatório.")

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)

    # ── Campo de Observações / Instruções ──
    st.markdown("""
    <div style='margin-bottom:10px;'>
        <span style='font-size:9px;font-weight:700;color:rgba(255,255,255,0.45);
        letter-spacing:2px;text-transform:uppercase;'>📋 Observações para o Sistema</span>
        <div style='font-size:10px;color:rgba(255,255,255,0.35);margin-top:5px;line-height:1.5;'>
        Anote ajustes ou instruções aqui.<br>Copie e envie ao desenvolvedor.
        </div>
    </div>
    """, unsafe_allow_html=True)

    if 'obs_sistema' not in st.session_state:
        st.session_state.obs_sistema = ""

    obs_nova = st.text_area(
        "",
        value=st.session_state.obs_sistema,
        height=100,
        placeholder="Ex: Mudar cor do botão, adicionar campo...",
        key="obs_textarea",
        label_visibility="collapsed"
    )

    if st.button("SALVAR OBSERVAÇÃO", use_container_width=True, key="salvar_obs"):
        st.session_state.obs_sistema = obs_nova
        st.success("Observação salva!")

    if st.session_state.obs_sistema:
        st.markdown(
            f"<div style='background:rgba(38,166,154,0.12);border-left:3px solid #26A69A;"
            f"border-radius:0 8px 8px 0;padding:10px 12px;margin-top:8px;"
            f"font-size:11px;color:rgba(255,255,255,0.75);line-height:1.6;'>"
            f"{st.session_state.obs_sistema.replace(chr(10), '<br>')}"
            f"</div>",
            unsafe_allow_html=True)
        if st.button("Limpar observação", use_container_width=True, key="limpar_obs"):
            st.session_state.obs_sistema = ""
            st.rerun()

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)

    # ── Backups automáticos ──
    if st.button("📦 VER BACKUPS", use_container_width=True, key="btn_ver_backup"):
        try:
            sb = _get_supabase_client()
            res = sb.table("hova_dados").select("id,updated_at").like("id","backup_%").order("id", desc=True).limit(10).execute()
            if res.data:
                for row in res.data:
                    data_bkp = row['id'].replace('backup_','')
                    hora_bkp = row['updated_at'][:16].replace('T',' ')
                    st.markdown(
                        f"<div style='font-size:11px;color:rgba(255,255,255,0.7);"
                        f"padding:4px 0;border-bottom:1px solid rgba(255,255,255,0.06);'>"
                        f"📅 {data_bkp} &nbsp;·&nbsp; {hora_bkp}</div>",
                        unsafe_allow_html=True)
            else:
                st.caption("Nenhum backup ainda.")
        except Exception as e:
            st.caption(f"Erro: {e}")

    st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.08);margin:18px 0;'>", unsafe_allow_html=True)
    if st.button("LIMPAR MEMORIA", use_container_width=True):
        st.session_state['confirmar_limpar'] = True

    if st.session_state.get('confirmar_limpar'):
        st.markdown(
            "<div style='background:#FFF5F5;border:1.5px solid #E5BCBC;border-radius:10px;"
            "padding:12px 14px;margin-top:6px;font-size:12px;color:#9B2C2C;font-weight:600;'>"
            "⚠️ Isso apagará TODOS os dados (candidatos, contratados, histórico). Tem certeza?"
            "</div>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Cancelar", key="canc_limpar", use_container_width=True):
                st.session_state['confirmar_limpar'] = False
                st.rerun()
        with c2:
            if st.button("SIM, LIMPAR TUDO", key="ok_limpar", use_container_width=True):
                for k in ['cvs','agendados','contratados','aguardando_retorno','cvs_antigos','ex_funcionarios','favoritos']:
                    st.session_state[k] = []
                st.session_state.historico_emails     = set()
                st.session_state._processados         = set()
                st.session_state.candidato_foco       = None
                st.session_state.contratar_foco       = None
                st.session_state.perfil_foco          = None
                st.session_state.pular_idx            = {}
                st.session_state['confirmar_limpar']  = False
                if os.path.exists(ARQUIVO_MEMORIA): os.remove(ARQUIVO_MEMORIA)
                _sb_set({})
                st.success("Memória zerada.")
                time.sleep(1)
                st.rerun()

# ──────────────────────────────────────────
# HELPERS DE FILTRO
# ──────────────────────────────────────────
def _por_mes(lista):
    if filtro_mes == "Todos os meses": return lista
    alvo = next((k for k,v in MESES_NOMES.items() if v==filtro_mes), None)
    return [c for c in lista if c.get('mes_num')==alvo] if alvo else lista

def _busca(lista, termo):
    if not termo.strip(): return lista
    t = termo.lower().strip()
    return [c for c in lista if
            t in c.get('nome','').lower() or
            t in c.get('email','').lower() or
            t in c.get('telefone','').lower()]

# ──────────────────────────────────────────
# CABEÇALHO
# ──────────────────────────────────────────
n_tri  = len(st.session_state.cvs)
n_ag   = len(st.session_state.aguardando_retorno)
n_agd  = len(st.session_state.agendados)
n_con  = len(st.session_state.contratados)

# ── Header mobile (só aparece em telas <= 768px via CSS) ──
st.markdown(f"""
<div class="mobile-header">
  <div style="display:flex;align-items:center;gap:10px;">
    <div style="width:34px;height:34px;border-radius:50%;background:#26A69A;
                display:flex;align-items:center;justify-content:center;">
      <svg width="16" height="16" viewBox="0 0 24 24" fill="none">
        <path d="M4 6h16M4 12h16M4 18h16" stroke="#fff" stroke-width="2.5" stroke-linecap="round"/>
      </svg>
    </div>
    <div>
      <div style="color:#fff;font-size:15px;font-weight:700;letter-spacing:0.5px;">HOVA</div>
      <div style="color:rgba(255,255,255,0.5);font-size:9px;letter-spacing:2px;">GESTÃO DE TALENTOS</div>
    </div>
  </div>
  <div style="display:flex;gap:6px;">
    <div class="mobile-stat"><span class="ms-n">{n_tri}</span><span class="ms-l">TRIAGEM</span></div>
    <div class="mobile-stat"><span class="ms-n">{n_agd}</span><span class="ms-l">AGEND.</span></div>
    <div class="mobile-stat"><span class="ms-n">{n_con}</span><span class="ms-l">CONTRAT.</span></div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Header desktop (escondido no mobile) ──
st.markdown(f"""
<div class="hero-card desktop-only">
  <div class="hero-left">
    <div class="hero-marca">Hospital de Olhos Vale do Aço</div>
    <div class="hero-titulo">Seletor de<span>Talentos</span></div>
    <div class="hero-sub">Gestão de Processos Seletivos</div>
  </div>
  <div class="hero-stats">
    <div class="stat-box"><span class="n">{n_tri}</span><span class="l">Triagem</span></div>
    <div class="stat-box"><span class="n">{n_ag}</span><span class="l">Aguardando</span></div>
    <div class="stat-box"><span class="n">{n_agd}</span><span class="l">Agendados</span></div>
    <div class="stat-box"><span class="n">{n_con}</span><span class="l">Contratados</span></div>
  </div>
</div>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────
# BARRA DE SINCRONIZAÇÃO
# ──────────────────────────────────────────
col_btn, col_alerta, col_msg = st.columns([1, 1, 2])
with col_btn:
    if st.button("SINCRONIZAR CURRICULOS", type="primary", use_container_width=True):
        st.session_state.executar_sync = True
        st.session_state.limite_sync   = limite_busca
        st.rerun()
with col_alerta:
    # Contar alertas não lidos
    n_alertas = sum(
        1 for c in st.session_state.agendados
        for a in c.get('alertas', [])
        if not a.get('lido')
    )
    label_alerta = f"📬 VERIFICAR RESPOSTAS" if n_alertas == 0 else f"📬 RESPOSTAS ({n_alertas})"
    if st.button(label_alerta, use_container_width=True,
                 type="primary" if n_alertas > 0 else "secondary"):
        with st.spinner("Verificando respostas de candidatos..."):
            novos_al, logs_al = varrer_alertas()
        if novos_al > 0:
            st.success(f"{novos_al} nova(s) resposta(s) de candidatos!")
            st.markdown("""<audio autoplay>
<source src="data:audio/wav;base64,UklGRnoGAABXQVZFZm10IBAAAA
EAAQARAAIAIgAAABAAEABkYXRhSAYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA
AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA
AAAAAAAA" type="audio/wav"></audio>""", unsafe_allow_html=True)
        else:
            st.info("Nenhuma resposta nova encontrada.")
        st.rerun()
with col_msg:
    if st.session_state.sync_msg:
        m   = st.session_state.sync_msg
        css = "notif-ok" if m['tipo']=='ok' else ("notif-warn" if m['tipo']=='err' else "notif-info")
        st.markdown(f"<div class='notif {css}'>{m['texto']}</div>", unsafe_allow_html=True)

if st.session_state.get('sync_logs'):
    with st.expander("Ver log detalhado da última sincronização"):
        for ln in st.session_state.sync_logs:
            st.write(ln)

# Execução real (após rerun, fora de colunas)
if st.session_state.executar_sync:
    st.session_state.executar_sync = False
    with st.spinner(f"Conectando e varrendo {st.session_state.limite_sync} e-mails..."):
        qtd, logs = buscar_curriculos(st.session_state.limite_sync)
    st.session_state.sync_logs = logs
    st.session_state.sync_msg  = {
        'tipo': 'ok' if qtd>0 else ('err' if any('ERRO' in l for l in logs) else 'info'),
        'texto': f"{qtd} novo(s) currículo(s) capturado(s)." if qtd>0
                 else (next((l for l in logs if 'ERRO' in l), None) or
                       f"Nenhum novo currículo encontrado nos últimos {st.session_state.limite_sync} e-mails.")
    }
    st.rerun()

# ── Navegação mobile funcional (só aparece em mobile via CSS) ──
st.markdown("""
<div class="mobile-chips">
  <span class="chip active">Todos os meses</span>
  <span class="chip">Triagem Geral</span>
  <span class="chip">Recepção</span>
  <span class="chip">Tec. Enfermagem</span>
  <span class="chip">Administrativo</span>
</div>
<style>
.mobile-nav-bar {
    display: none;
    position: fixed;
    bottom: 0; left: 0; right: 0;
    background: #fff;
    border-top: 0.5px solid #E2E6EA;
    z-index: 200;
    box-shadow: 0 -2px 16px rgba(0,0,0,0.08);
    padding: 0;
}
@media (max-width: 768px) {
    .mobile-nav-bar { display: block !important; }
    .stApp { padding-bottom: 76px !important; }
}
</style>
<div class="mobile-nav-bar" id="mobileNav">
  <div style="display:flex;justify-content:space-around;padding:8px 0 14px;">
    <button onclick="navClick(0,'Triagem')" id="nb0"
      style="flex:1;background:none;border:none;padding:6px 4px;cursor:pointer;display:flex;flex-direction:column;align-items:center;gap:3px;color:#004D40;">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none">
        <path d="M17 21v-2a4 4 0 0 0-4-4H5a4 4 0 0 0-4 4v2" stroke="currentColor" stroke-width="2" stroke-linecap="round"/>
        <circle cx="9" cy="7" r="4" stroke="currentColor" stroke-width="2"/>
        <path d="M23 21v-2a4 4 0 0 0-3-3.87M16 3.13a4 4 0 0 1 0 7.75" stroke="currentColor" stroke-width="2"/>
      </svg>
      <span style="font-size:9px;font-weight:600;letter-spacing:0.5px;text-transform:uppercase;">Triagem</span>
    </button>
    <button onclick="navClick(1,'Agendados')" id="nb1"
      style="flex:1;background:none;border:none;padding:6px 4px;cursor:pointer;display:flex;flex-direction:column;align-items:center;gap:3px;color:#9AA5B4;">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none">
        <rect x="3" y="4" width="18" height="18" rx="2" stroke="currentColor" stroke-width="2"/>
        <path d="M16 2v4M8 2v4M3 10h18" stroke="currentColor" stroke-width="2" stroke-linecap="round"/>
      </svg>
      <span style="font-size:9px;font-weight:600;letter-spacing:0.5px;text-transform:uppercase;">Agendados</span>
    </button>
    <button onclick="navClick(2,'Favoritos')" id="nb2"
      style="flex:1;background:none;border:none;padding:6px 4px;cursor:pointer;display:flex;flex-direction:column;align-items:center;gap:3px;color:#9AA5B4;">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none">
        <path d="M20.84 4.61a5.5 5.5 0 0 0-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 0 0-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 0 0 0-7.78z" stroke="currentColor" stroke-width="2"/>
      </svg>
      <span style="font-size:9px;font-weight:600;letter-spacing:0.5px;text-transform:uppercase;">Favoritos</span>
    </button>
    <button onclick="navClick(3,'Equipe')" id="nb3"
      style="flex:1;background:none;border:none;padding:6px 4px;cursor:pointer;display:flex;flex-direction:column;align-items:center;gap:3px;color:#9AA5B4;">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none">
        <path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2" stroke="currentColor" stroke-width="2" stroke-linecap="round"/>
        <circle cx="12" cy="7" r="4" stroke="currentColor" stroke-width="2"/>
      </svg>
      <span style="font-size:9px;font-weight:600;letter-spacing:0.5px;text-transform:uppercase;">Equipe</span>
    </button>
    <button onclick="abrirMenu()" id="nb4"
      style="flex:1;background:none;border:none;padding:6px 4px;cursor:pointer;display:flex;flex-direction:column;align-items:center;gap:3px;color:#9AA5B4;">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none">
        <path d="M4 6h16M4 12h16M4 18h16" stroke="currentColor" stroke-width="2" stroke-linecap="round"/>
      </svg>
      <span style="font-size:9px;font-weight:600;letter-spacing:0.5px;text-transform:uppercase;">Menu</span>
    </button>
  </div>
</div>

<script>
// Mapa: índice do botão → texto da aba do Streamlit
var tabMap = {
  0: 'TRIAGEM GERAL',
  1: 'AGENDADOS',
  2: 'FAVORITOS',
  3: 'CONTRATADOS'
};

function navClick(idx, label) {
  // Atualizar visual dos botões
  for (var i = 0; i < 4; i++) {
    var btn = document.getElementById('nb' + i);
    if (btn) btn.style.color = i === idx ? '#004D40' : '#9AA5B4';
  }
  // Encontrar e clicar na aba correspondente do Streamlit
  var tabs = document.querySelectorAll('[data-baseweb="tab"]');
  var alvo = tabMap[idx];
  tabs.forEach(function(tab) {
    if (tab.textContent.trim().toUpperCase().indexOf(alvo) >= 0) {
      tab.click();
    }
  });
  // Scroll para o topo
  window.scrollTo({top: 0, behavior: 'smooth'});
}

function abrirMenu() {
  // Clicar no botão de abrir a sidebar
  var btn = document.querySelector(
    'button[data-testid="collapsedControl"], button[aria-label="Open sidebar"]'
  );
  if (btn) btn.click();
  document.getElementById('nb4').style.color = '#004D40';
}
</script>
""", unsafe_allow_html=True)

st.write("")
# ──────────────────────────────────────────
sc1, _ = st.columns([2, 1])
with sc1:
    # Não usar key de busca persistente — campo limpo a cada sessão
    # para não interferir com a triagem
    termo = st.text_input(
        "", placeholder="Pesquisar por nome, e-mail ou telefone...",
        label_visibility="collapsed", key="busca_global"
    )

st.write("")

# ──────────────────────────────────────────
# ABAS
# ──────────────────────────────────────────
abas = st.tabs([
    "TRIAGEM GERAL","RECEPCAO","TEC. ENFERMAGEM",
    "ADMINISTRATIVO","FATURAMENTO","JOVEM APRENDIZ",
    "AGENDADOS","AGUARDANDO RETORNO","CONTRATADOS","FAVORITOS","BANCO ANTIGOS"
])

SETORES = [
    "TRIAGEM GERAL","RECEPCAO E ATENDIMENTO","TECNICO E ENFERMAGEM",
    "ADMINISTRATIVO","FATURAMENTO","JOVEM APRENDIZ"
]

# ── ABAS 0-5: TRIAGEM ─────────────────────
for i, setor in enumerate(SETORES):
    with abas[i]:
        base     = [c for c in st.session_state.cvs if c['setor']==setor]
        if filtro_setor != "TODOS":
            base = [c for c in base if c['setor']==filtro_setor]
        base     = _por_mes(base)
        fila     = _busca(base, termo)
        fila     = sorted(fila, key=lambda x: x.get('data_iso',''), reverse=True)

        if fila:
            idx_ex = st.session_state.pular_idx.get(setor, 0) % len(fila)
            per    = f" — {filtro_mes}" if filtro_mes!="Todos os meses" else ""
            st.markdown(
                f"<div style='font-size:12px;color:#8A94A6;margin-bottom:14px;'>"
                f"<b style='color:#004D40;font-size:17px;'>{len(fila)}</b> candidato(s){per}"
                f" &nbsp;·&nbsp; Exibindo <b style='color:#004D40;'>{idx_ex+1}</b> de {len(fila)}"
                f"</div>", unsafe_allow_html=True
            )

        if not fila:
            st.session_state.pular_idx.pop(setor, None)
            st.markdown('<div class="empty"><div class="e-title">FILA VAZIA</div>'
                        '<div class="e-sub">Nenhum candidato com os filtros atuais.</div></div>',
                        unsafe_allow_html=True)
            continue

        # Candidato atual
        if st.session_state.candidato_foco:
            c = next((x for x in fila if x['id']==st.session_state.candidato_foco), None)
            if not c:
                idx = st.session_state.pular_idx.get(setor,0) % len(fila)
                c   = fila[idx]
        else:
            idx = st.session_state.pular_idx.get(setor,0) % len(fila)
            c   = fila[idx]

        # ── FORMULÁRIO AGENDAMENTO ──
        if st.session_state.candidato_foco == c['id']:
            tem_email = bool(c.get('email','').strip())
            tem_tel   = bool(c.get('telefone','').strip())
            via_wa    = not tem_email and tem_tel  # só WhatsApp quando não tem e-mail

            st.markdown("<div class='form-sched'>", unsafe_allow_html=True)
            st.markdown(
                f"<div style='font-size:16px;font-weight:800;color:#004D40;margin-bottom:4px;'>AGENDAR ENTREVISTA</div>"
                f"<div style='font-size:12px;color:#8A94A6;margin-bottom:20px;'>"
                f"{'📱 Sem e-mail — convite será enviado via WhatsApp' if via_wa else '📧 Convite será enviado por e-mail'}"
                f"</div>",
                unsafe_allow_html=True)

            c1, c2 = st.columns(2)
            with c1:
                st.caption("NOME DO CANDIDATO")
                ne = st.text_input("", value=c['nome'], key=f"ne_{c['id']}", label_visibility="collapsed")
            with c2:
                if via_wa:
                    st.caption("TELEFONE / WHATSAPP")
                    et = st.text_input("", value=c.get('telefone',''), key=f"et_{c['id']}", label_visibility="collapsed")
                    ee = ""
                else:
                    st.caption("E-MAIL")
                    ee = st.text_input("", value=c['email'], key=f"ee_{c['id']}", label_visibility="collapsed")
                    et = c.get('telefone','')

            st.caption("DATA E HORARIOS (3 OPCOES)")
            cd, ch1, ch2, ch3 = st.columns(4)
            da  = cd.date_input("",  key=f"da_{c['id']}",  label_visibility="collapsed")
            h1  = ch1.time_input("", datetime.time(9, 0),  key=f"h1_{c['id']}", label_visibility="collapsed")
            h2  = ch2.time_input("", datetime.time(14,0),  key=f"h2_{c['id']}", label_visibility="collapsed")
            h3  = ch3.time_input("", datetime.time(16,0),  key=f"h3_{c['id']}", label_visibility="collapsed")

            # ── Indicador visual de conflitos em tempo real ──
            ocupados = [h for h in [h1,h2,h3] if not horario_disponivel(da, h)]
            livres_agora = horarios_livres(da, [h1,h2,h3])
            if ocupados and livres_agora:
                hocs = ", ".join(h.strftime('%H:%M') for h in ocupados)
                st.markdown(
                    f"<div class='notif notif-warn' style='text-align:left;font-size:12px;'>"
                    f"Atencao: {hocs} ja esta(o) ocupado(s) nessa data. "
                    f"O candidato sera avisado automaticamente se escolher esses horarios.</div>",
                    unsafe_allow_html=True)
            elif not livres_agora and da:
                st.markdown(
                    "<div class='notif notif-warn' style='text-align:left;font-size:12px;'>"
                    "Todos os 3 horarios escolhidos ja estao ocupados nessa data. "
                    "Por favor, altere os horarios acima antes de enviar.</div>",
                    unsafe_allow_html=True)

            msg_conv = (
                f"Olá {ne},\n\n"
                f"O Hospital de Olhos Vale do Aço analisou seu perfil e você foi selecionada(o) para a próxima fase do Processo Seletivo.\n\n"
                f"Temos disponibilidade para o dia {da.strftime('%d/%m/%Y')}. "
                f"Por gentileza, responda com o NUMERO da sua escolha de horário:\n\n"
                f"1 - {h1.strftime('%H:%M')}\n2 - {h2.strftime('%H:%M')}\n3 - {h3.strftime('%H:%M')}\n\n"
                f"Endereço: {ENDERECO_HOVA}\n"
                f"Ao chegar, informe na recepção que é referente à entrevista e pergunte por Josi ou Paula.\n\n"
                f"Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
            )

            # Mensagem WhatsApp formatada (mais curta e com negrito)
            msg_wa_conv = (
                f"Olá! 😊\n\n"
                f"Aqui é o RH do *Hospital de Olhos Vale do Aço*.\n\n"
                f"Seu perfil foi selecionado para entrevista!\n\n"
                f"📅 *{da.strftime('%d/%m/%Y')}*\n\n"
                f"Escolha um horário respondendo com o número:\n"
                f"*1* - {h1.strftime('%H:%M')}\n"
                f"*2* - {h2.strftime('%H:%M')}\n"
                f"*3* - {h3.strftime('%H:%M')}\n\n"
                f"📍 {ENDERECO_HOVA}\n"
                f"Pergunte por *Josi* ou *Paula*.\n\n"
                f"Até lá! 🤝"
            )

            if via_wa:
                with st.expander("Visualizar mensagem WhatsApp que será enviada"):
                    st.code(msg_wa_conv, language=None)
            else:
                with st.expander("Visualizar e-mail que será enviado"):
                    st.code(msg_conv, language=None)

            bc, benv = st.columns(2)
            with bc:
                if st.button("CANCELAR", key=f"canc_{c['id']}", type="secondary", use_container_width=True):
                    st.session_state.candidato_foco = None
                    st.rerun()
            with benv:
                if via_wa:
                    # Botão que abre WhatsApp com a mensagem pré-preenchida
                    tel_limpo = ''.join(filter(str.isdigit, et))
                    if tel_limpo:
                        url_wa = f"https://wa.me/55{tel_limpo}?text={urllib.parse.quote(msg_wa_conv)}"
                        st.markdown(
                            f'<a href="{url_wa}" target="_blank" class="wa-btn" '
                            f'style="display:block;text-align:center;height:48px;line-height:48px;'
                            f'border-radius:9px;font-size:11px;font-weight:700;letter-spacing:1px;">'
                            f'ENVIAR VIA WHATSAPP</a>',
                            unsafe_allow_html=True)
                        # Botão para confirmar que enviou e mover para aguardando
                        if st.button("CONFIRMAR ENVIO ✓", key=f"conf_wa_{c['id']}",
                                     use_container_width=True):
                            c.update({'nome':ne,'telefone':tel_limpo,
                                      'data_entrevista':da,
                                      'opcao_1':h1,'opcao_2':h2,'opcao_3':h3,
                                      'canal':'whatsapp'})
                            st.session_state.aguardando_retorno.append(c)
                            st.session_state.cvs.remove(c)
                            st.session_state.candidato_foco = None
                            salvar_json()
                            st.rerun()
                    else:
                        st.error("Informe o telefone para enviar via WhatsApp.")
                else:
                    if st.button("ENVIAR CONVITE", type="primary", key=f"conf_{c['id']}", use_container_width=True):
                        livres = horarios_livres(da, [h1, h2, h3])
                        if not livres:
                            st.error("Todos os 3 horários escolhidos já estão ocupados. Escolha novas opções antes de enviar.")
                        else:
                            with st.spinner("Enviando convite..."):
                                ok = send_email(ee, "HOVA — Convite para Entrevista", msg_conv)
                            if ok:
                                c.update({'nome':ne,'email':ee,'data_entrevista':da,
                                          'opcao_1':h1,'opcao_2':h2,'opcao_3':h3})
                                st.session_state.aguardando_retorno.append(c)
                                st.session_state.cvs.remove(c)
                                st.session_state.candidato_foco = None
                                salvar_json()
                                st.rerun()
                            else:
                                st.error("Falha no envio. Verifique as configurações de e-mail.")
            st.markdown("</div>", unsafe_allow_html=True)

        # ── CARD CANDIDATO ──
        else:
            if c.get('foto'):
                b64 = base64.b64encode(c['foto']).decode()
                av  = f"<img src='data:image/jpeg;base64,{b64}' class='avatar-img'>"
            else:
                av = f"<div class='avatar'>{iniciais(c['nome'])}</div>"

            manual_b = "<span class='tag tag-manual'>Manual</span>" if c.get('manual') else ""
            cid_b    = f"<span class='tag tag-cinza'>{c['cidade']}</span>" if c.get('cidade') else ""
            dat_b    = f"<span class='tag tag-azul'>{c['data']}</span>"
            tags_b   = "".join(f"<span class='tag tag-verde'>{t}</span>" for t in c.get('tags',[]))

            # Tags de triagem inteligente
            primeiro_emp = c.get('primeiro_emprego', detectar_primeiro_emprego(c.get('preview','')))
            cidade_longe, cidade_longe_nome = detectar_cidade_longe(c.get('preview',''), c.get('cidade',''))
            if primeiro_emp:
                tags_b += "<span class='tag' style='background:#FFF3CD;color:#856404;border:1px solid #FFEAA7;'>👶 Primeiro Emprego</span>"
            if cidade_longe:
                tags_b += f"<span class='tag' style='background:#FFE4E4;color:#9B2C2C;border:1px solid #FEB2B2;'>📍 VT Alto — {cidade_longe_nome}</span>"
            if c.get('obs_triagem'):
                tags_b += f"<span class='tag' style='background:#E8F5F2;color:#004D40;border:1px solid #B2DFDB;'>📝 Com obs.</span>"

            # Estrela de favorito
            ja_fav  = any(f['id'] == c['id'] for f in st.session_state.favoritos)
            estrela = "★" if ja_fav else "☆"
            cor_est = "#F59E0B" if ja_fav else "#CBD5E0"

            st.markdown(
                f"<div class='card-cand' style='position:relative;'>"
                f"<div title='Favoritar' style='position:absolute;top:16px;right:20px;"
                f"font-size:26px;color:{cor_est};line-height:1;user-select:none;'>{estrela}</div>"
                f"{av}"
                f"<div class='cand-nome'>{c['nome']} {manual_b}</div>"
                f"<div style='margin:8px 0;'>{cid_b} {dat_b}</div>"
                f"<div class='cand-info'>{c['email']}"
                f"{'  |  '+c['telefone'] if c.get('telefone') else ''}</div>"
                f"<div style='margin:10px 0;'>{tags_b}</div>"
                f"<div class='cv-resumo'>{c['preview']}</div></div>",
                unsafe_allow_html=True
            )

            # ── Alertas de triagem ─────────────────────────────
            if primeiro_emp and setor != "JOVEM APRENDIZ":
                st.markdown(
                    "<div class='notif notif-warn' style='text-align:left;'>"
                    "⚠️ <b>Sem experiência detectada.</b> Este candidato pode ser de Primeiro Emprego. "
                    "Só avance se houver vaga para Jovem Aprendiz aberta no momento.</div>",
                    unsafe_allow_html=True)

            if cidade_longe:
                st.markdown(
                    f"<div class='notif' style='background:#FFF5F5;border:1px solid #FEB2B2;"
                    f"color:#9B2C2C;text-align:left;'>"
                    f"🚌 <b>Atenção ao VT:</b> Candidato parece morar em "
                    f"<b>{cidade_longe_nome}</b>, que pode ter custo alto de vale-transporte. "
                    f"Verifique antes de avançar.</div>",
                    unsafe_allow_html=True)

            # ── Observações de triagem ─────────────────────────
            with st.expander("📝 Observações / Anotações"):
                obs_tri = st.text_area(
                    "", value=c.get('obs_triagem',''),
                    height=80, key=f"obs_tri_{c['id']}",
                    placeholder="Anote qualquer observação: dúvidas, pontos de atenção, histórico de contato...",
                    label_visibility="collapsed")
                if st.button("Salvar observação", key=f"salv_obs_tri_{c['id']}"):
                    c['obs_triagem'] = obs_tri
                    salvar_json()
                    st.success("Observação salva.")
                    st.rerun()

            # ── Redirecionar para outro setor ──────────────────
            with st.expander("🔄 Redirecionar para outro setor"):
                st.markdown(
                    "<div style='font-size:12px;color:#4A5568;margin-bottom:8px;'>"
                    "Use quando o candidato se candidatou para uma vaga mas se encaixa melhor em outra.</div>",
                    unsafe_allow_html=True)
                novo_setor_op = st.selectbox(
                    "Novo setor:", SETORES,
                    index=SETORES.index(c['setor']) if c['setor'] in SETORES else 0,
                    key=f"novo_setor_{c['id']}")
                motivo_mud = st.text_input(
                    "Motivo da mudança:",
                    placeholder="Ex: Candidato tem perfil mais adequado para Recepção",
                    key=f"motivo_mud_{c['id']}")
                if st.button("CONFIRMAR REDIRECIONAMENTO", key=f"redir_{c['id']}",
                             use_container_width=True):
                    if motivo_mud.strip():
                        setor_antigo = c['setor']
                        c['setor'] = novo_setor_op
                        c['motivo_mudanca'] = f"{setor_antigo} → {novo_setor_op}: {motivo_mud}"
                        salvar_json()
                        st.success(f"Redirecionado para {novo_setor_op}.")
                        st.rerun()
                    else:
                        st.warning("Informe o motivo da mudança antes de confirmar.")

            # ── Mensagem de dúvida ─────────────────────────────
            with st.expander("❓ Tirar dúvida com o candidato"):
                st.markdown(
                    "<div style='font-size:12px;color:#4A5568;margin-bottom:8px;'>"
                    "Use quando precisar de informação antes de avançar (ex: mora longe, vaga pretendida).</div>",
                    unsafe_allow_html=True)

                duvidas_prontas = {
                    "Cidade/localização": f"Olá! 😊 Aqui é o RH do *Hospital de Olhos Vale do Aço*. Estamos analisando seu currículo e gostaríamos de saber: você reside atualmente em qual cidade?",
                    "Vaga pretendida":    f"Olá! 😊 Aqui é o RH do *Hospital de Olhos Vale do Aço*. Vimos seu currículo com interesse! Para qual vaga você está se candidatando?",
                    "Disponibilidade":    f"Olá! 😊 Aqui é o RH do *Hospital de Olhos Vale do Aço*. Qual é a sua disponibilidade de horário? Você teria disponibilidade para trabalhar em escala?",
                    "Personalizada":      "",
                }
                tipo_duvida = st.selectbox("Dúvida:", list(duvidas_prontas.keys()),
                                            key=f"tipo_duvida_{c['id']}")
                msg_duvida = st.text_area(
                    "Mensagem:",
                    value=duvidas_prontas[tipo_duvida],
                    height=100, key=f"msg_duvida_{c['id']}")

                tel_d = c.get('telefone','')
                email_d = c.get('email','')
                dd1, dd2 = st.columns(2)
                with dd1:
                    if tel_d:
                        url_duvida_wa = f"https://wa.me/55{tel_d}?text={urllib.parse.quote(msg_duvida)}"
                        st.markdown(
                            f'<a href="{url_duvida_wa}" target="_blank" class="wa-btn"'
                            f' style="display:block;text-align:center;">WhatsApp</a>',
                            unsafe_allow_html=True)
                with dd2:
                    if email_d and st.button("Enviar por E-mail", key=f"env_duvida_{c['id']}",
                                              use_container_width=True):
                        ok = send_email(email_d, "Hospital de Olhos Vale do Aço — Dúvida", msg_duvida)
                        if ok:
                            c['duvida_enviada'] = True
                            salvar_json()
                            st.success("Mensagem enviada.")
                        else:
                            st.error("Falha no envio.")

            if c.get('arquivo_bytes') and c.get('nome_arquivo'):
                with st.expander("Ver documento original"):
                    if c['nome_arquivo'].lower().endswith('.pdf'):
                        b64p = base64.b64encode(c['arquivo_bytes']).decode()

                        # ── Botão de download sempre visível ──
                        st.download_button(
                            label    = f"Baixar PDF — {c['nome_arquivo']}",
                            data     = c['arquivo_bytes'],
                            file_name= c['nome_arquivo'],
                            mime     = "application/pdf",
                            use_container_width=True,
                            key      = f"dl_{c['id']}"
                        )

                        # ── Viewer via PDF.js (Mozilla CDN) — funciona no Chrome/Streamlit Cloud ──
                        import streamlit.components.v1 as components
                        pdf_viewer_html = f"""
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    body {{ background:#F8FAFB; font-family:Inter,sans-serif; }}
    #pdf-container {{
      width:100%; height:720px;
      border:1px solid #E2E6EA; border-radius:8px;
      overflow:hidden; background:#525659;
    }}
    canvas {{ display:block; margin:0 auto; }}
    #toolbar {{
      background:#323639; padding:8px 16px;
      display:flex; align-items:center; gap:12px;
      border-radius:8px 8px 0 0;
    }}
    #toolbar button {{
      background:#004D40; color:#fff; border:none;
      padding:6px 14px; border-radius:6px; cursor:pointer;
      font-size:12px; font-weight:700; letter-spacing:0.5px;
    }}
    #toolbar button:hover {{ background:#00382E; }}
    #toolbar span {{ color:#ccc; font-size:12px; }}
    #canvas-wrapper {{
      height:672px; overflow-y:auto; background:#525659;
      display:flex; flex-direction:column; align-items:center;
      padding:16px 0; gap:12px;
    }}
    .page-canvas {{ box-shadow:0 2px 8px rgba(0,0,0,0.4); }}
  </style>
</head>
<body>
<div id="pdf-container">
  <div id="toolbar">
    <button onclick="anteriorPagina()">&#8592; Anterior</button>
    <span id="info-pagina">Carregando...</span>
    <button onclick="proximaPagina()">Próxima &#8594;</button>
    <button onclick="zoomMais()">+ Zoom</button>
    <button onclick="zoomMenos()">- Zoom</button>
  </div>
  <div id="canvas-wrapper"></div>
</div>

<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<script>
  pdfjsLib.GlobalWorkerOptions.workerSrc =
    'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

  const base64 = "{b64p}";
  const raw    = atob(base64);
  const arr    = new Uint8Array(raw.length);
  for (let i = 0; i < raw.length; i++) arr[i] = raw.charCodeAt(i);

  let pdfDoc    = null;
  let paginaAtual = 1;
  let escala    = 1.4;
  const wrapper = document.getElementById('canvas-wrapper');
  const info    = document.getElementById('info-pagina');

  pdfjsLib.getDocument({{ data: arr }}).promise.then(pdf => {{
    pdfDoc = pdf;
    renderPagina(paginaAtual);
  }}).catch(err => {{
    info.textContent = 'Erro ao carregar PDF: ' + err.message;
  }});

  function renderPagina(num) {{
    wrapper.innerHTML = '';
    pdfDoc.getPage(num).then(page => {{
      const vp     = page.getViewport({{ scale: escala }});
      const canvas = document.createElement('canvas');
      canvas.className = 'page-canvas';
      canvas.width  = vp.width;
      canvas.height = vp.height;
      wrapper.appendChild(canvas);
      page.render({{ canvasContext: canvas.getContext('2d'), viewport: vp }});
      info.textContent = 'Página ' + num + ' de ' + pdfDoc.numPages;
    }});
  }}

  function anteriorPagina() {{
    if (paginaAtual > 1) {{ paginaAtual--; renderPagina(paginaAtual); }}
  }}
  function proximaPagina() {{
    if (paginaAtual < pdfDoc.numPages) {{ paginaAtual++; renderPagina(paginaAtual); }}
  }}
  function zoomMais()  {{ escala = Math.min(escala + 0.2, 3.0); renderPagina(paginaAtual); }}
  function zoomMenos() {{ escala = Math.max(escala - 0.2, 0.6); renderPagina(paginaAtual); }}
</script>
</body>
</html>
"""
                        components.html(pdf_viewer_html, height=740, scrolling=False)

                    else:
                        st.info("Visualização direta disponível apenas para PDF.")
                        st.download_button(
                            f"Baixar arquivo — {c['nome_arquivo']}",
                            c['arquivo_bytes'], file_name=c['nome_arquivo'],
                            mime="application/octet-stream",
                            use_container_width=True, key=f"dl_{c['id']}"
                        )
            elif c.get('manual'):
                st.markdown("<div class='notif notif-info'>Candidato cadastrado manualmente — sem arquivo físico.</div>", unsafe_allow_html=True)

            st.write("")

            # ── Confirmação de rejeição com mensagem editável ──
            if st.session_state.get('rejeitar_foco') == c['id']:
                # ── Detectar cidade automaticamente ──
                _class_cidade = classificar_cidade(
                    c.get('preview',''), c.get('cidade',''))
                _cidade_nome  = detectar_cidade_longe(
                    c.get('preview',''), c.get('cidade',''))[1]
                _fora_regiao  = _class_cidade == 'longe'

                st.markdown(
                    "<div style='background:#FFFAF8;border:1.5px solid #E5BCBC;"
                    "border-radius:14px;padding:22px 26px;margin-top:8px;'>",
                    unsafe_allow_html=True)

                if _fora_regiao:
                    st.markdown(
                        f"<div style='background:#FFF3CD;border:1px solid #FFEAA7;"
                        f"border-radius:8px;padding:10px 14px;margin-bottom:14px;"
                        f"font-size:12px;color:#856404;font-weight:600;'>"
                        f"📍 Cidade detectada fora do raio de 12km: "
                        f"<b>{_cidade_nome or 'outra região'}</b>. "
                        f"Mensagem de localização pré-preenchida automaticamente.</div>",
                        unsafe_allow_html=True)

                st.markdown(
                    "<div style='font-size:13px;font-weight:700;color:#9B2C2C;"
                    "margin-bottom:14px;'>Confirmar Rejeição — revisar mensagem antes de enviar</div>",
                    unsafe_allow_html=True)

                motivos_rej = [
                    "Selecione o motivo principal...",
                    "Perfil não compatível com a vaga",
                    "Mora muito longe / VT inviável",
                    "Fora da região — vaga presencial em Ipatinga/MG",
                    "Sem experiência para a vaga solicitada",
                    "Candidou-se para vaga não disponível no momento",
                    "Currículo incompleto ou sem informações suficientes",
                    "Já possui processo em andamento conosco",
                    "Outro (descrever nas observações)",
                ]

                # Pré-selecionar motivo se cidade fora
                idx_motivo = 3 if _fora_regiao else 0
                motivo_sel = st.selectbox(
                    "Motivo da rejeição:",
                    motivos_rej,
                    index=idx_motivo,
                    key=f"motivo_rej_{c['id']}")

                # Pré-preencher mensagem conforme o motivo
                if _fora_regiao:
                    msg_rej_padrao = MSG_REJEICAO_CIDADE.replace(
                        "Olá,", f"Olá {c['nome'].title()},")
                else:
                    msg_rej_padrao = (
                        f"Olá {c['nome'].title()}, tudo bem?\n\n"
                        f"Agradecemos muito seu interesse em fazer parte da equipe "
                        f"do Hospital de Olhos Vale do Aço.\n\n"
                        f"Após análise do seu currículo, informamos que no momento "
                        f"não temos uma vaga disponível compatível com o seu perfil. "
                        f"Seu currículo ficará em nossa base de dados e entraremos "
                        f"em contato caso surja uma oportunidade.\n\n"
                        f"Agradecemos sua compreensão e desejamos sucesso!\n\n"
                        f"Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
                    )

                msg_rej_edit = st.text_area(
                    "Mensagem que será enviada:",
                    value=msg_rej_padrao,
                    height=200,
                    key=f"msg_rej_{c['id']}")

                rc1, rc2 = st.columns(2)
                with rc1:
                    if st.button("Cancelar", key=f"rej_canc_{c['id']}",
                                 use_container_width=True):
                        st.session_state['rejeitar_foco'] = None
                        st.rerun()
                with rc2:
                    if st.button("CONFIRMAR E ENVIAR", key=f"rej_env_{c['id']}",
                                 type="primary", use_container_width=True):
                        if motivo_sel == motivos_rej[0]:
                            st.warning("Selecione o motivo da rejeição antes de confirmar.")
                        else:
                            if c.get('email','').strip():
                                with st.spinner("Notificando candidato..."):
                                    send_email(c['email'],
                                               "Hospital de Olhos Vale do Aço — Processo Seletivo",
                                               msg_rej_edit)
                            else:
                                st.info("Candidato sem e-mail — rejeitado sem notificação.")
                            c['motivo_rejeicao'] = motivo_sel
                            st.session_state.cvs.remove(c)
                            nt  = len([x for x in st.session_state.cvs if x['setor']==setor])
                            idx = st.session_state.pular_idx.get(setor, 0)
                            st.session_state.pular_idx[setor] = (
                                max(0, min(idx, nt-1)) if nt > 0 else 0)
                            st.session_state['rejeitar_foco'] = None
                            salvar_json()
                            st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

            # ── Estrela funcional via form (sem botão visível extra) ──
            ja_fav = any(f['id'] == c['id'] for f in st.session_state.favoritos)
            with st.form(key=f"form_star_{c['id']}"):
                st.markdown(
                    f"<div style='text-align:right;margin-top:-8px;margin-bottom:4px;'>"
                    f"<span style='font-size:11px;color:#9AA5B4;'>"
                    f"{'Favoritado — clique para remover' if ja_fav else 'Clique na estrela para favoritar'}"
                    f"</span></div>",
                    unsafe_allow_html=True)
                submitted = st.form_submit_button(
                    "★ Favoritado" if ja_fav else "☆ Favoritar",
                    use_container_width=False)
                if submitted:
                    if ja_fav:
                        st.session_state.favoritos = [
                            f for f in st.session_state.favoritos if f['id'] != c['id']
                        ]
                    else:
                        st.session_state.favoritos.append(c)
                    salvar_json()
                    st.rerun()

            # ── 4 botões na mesma linha ──────────────────────
            b_vol, bp, br2, bac = st.columns(4)
            with b_vol:
                if st.button("← VOLTAR", key=f"vol_{c['id']}", use_container_width=True):
                    cur = st.session_state.pular_idx.get(setor, 0)
                    st.session_state.pular_idx[setor] = (cur - 1) % len(fila)
                    st.rerun()
            with bp:
                if st.button("PULAR →", key=f"pul_{c['id']}", use_container_width=True):
                    cur = st.session_state.pular_idx.get(setor, 0)
                    st.session_state.pular_idx[setor] = (cur + 1) % len(fila)
                    st.rerun()
            with br2:
                if st.button("REJEITAR", key=f"rej2_{c['id']}", type="secondary",
                             use_container_width=True):
                    st.session_state['rejeitar_foco'] = c['id']
                    st.rerun()
            with bac:
                if st.button("ACEITAR", key=f"acc_{c['id']}", type="primary",
                             use_container_width=True):
                    st.session_state.candidato_foco = c['id']
                    st.rerun()

# ── ABA 6: AGENDADOS ──────────────────────
with abas[6]:
    ag_list = _busca(_por_mes(st.session_state.agendados), termo)

    # ══════════════════════════════════════
    # PAINEL DE ALERTAS — candidatos que responderam
    # ══════════════════════════════════════
    alertas_ativos = [
        c for c in st.session_state.agendados
        if any(not a.get('lido') for a in c.get('alertas', []))
    ]

    if alertas_ativos:
        st.markdown(
            f"<div style='background:#FFF8EC;border:1.5px solid #D69E2E;border-radius:14px;"
            f"padding:16px 20px;margin-bottom:20px;'>"
            f"<div style='font-size:13px;font-weight:800;color:#92540A;margin-bottom:12px;'>"
            f"📬 {len(alertas_ativos)} candidato(s) responderam — aguardando atenção</div>",
            unsafe_allow_html=True)

        for cand in alertas_ativos:
            alertas_nao_lidos = [a for a in cand.get('alertas',[]) if not a.get('lido')]
            hf = cand['hora_entrevista'].strftime('%H:%M') if cand.get('hora_entrevista') else '—'
            df = cand['data_entrevista'].strftime('%d/%m/%Y') if cand.get('data_entrevista') else '—'

            with st.expander(f"📩 {cand['nome']} — {len(alertas_nao_lidos)} mensagem(ns) | Entrevista: {df} às {hf}"):
                for al in alertas_nao_lidos:
                    st.markdown(
                        f"<div style='background:#FFFDF7;border-left:3px solid #D69E2E;"
                        f"border-radius:0 8px 8px 0;padding:12px 16px;margin-bottom:10px;'>"
                        f"<div style='font-size:10px;color:#92540A;font-weight:700;"
                        f"letter-spacing:1px;text-transform:uppercase;margin-bottom:4px;'>"
                        f"📧 {al.get('data','—')} — {al.get('assunto','(sem assunto)')}</div>"
                        f"<div style='font-size:13px;color:#4A5568;line-height:1.6;'>"
                        f"{al.get('mensagem','')}</div>"
                        f"</div>",
                        unsafe_allow_html=True)

                # Observações persistentes do candidato
                st.markdown(
                    "<div style='font-size:10px;font-weight:800;color:#004D40;"
                    "letter-spacing:2px;text-transform:uppercase;margin:12px 0 6px;'>"
                    "Observações sobre este candidato</div>",
                    unsafe_allow_html=True)
                obs_key = f"obs_{cand['id']}"
                obs_val = st.text_area(
                    "", value=cand.get('observacoes_rh',''),
                    height=80, key=obs_key,
                    placeholder="Ex: Candidato confirmou presença, pediu para trocar horário...",
                    label_visibility="collapsed")

                al1, al2 = st.columns(2)
                with al1:
                    if st.button("SALVAR OBSERVAÇÃO", key=f"salv_obs_{cand['id']}",
                                 use_container_width=True):
                        cand['observacoes_rh'] = obs_val
                        salvar_json()
                        st.success("Observação salva.")
                        st.rerun()
                with al2:
                    if st.button("✓ MARCAR COMO RESOLVIDO", key=f"res_{cand['id']}",
                                 type="primary", use_container_width=True):
                        for a in cand.get('alertas',[]):
                            a['lido'] = True
                        salvar_json()
                        st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<hr style='border:none;border-top:1px solid #E2E6EA;margin:8px 0 20px;'>",
                    unsafe_allow_html=True)

    # ══════════════════════════════════════
    # TEMPLATES DE MENSAGEM
    # ══════════════════════════════════════
    with st.expander("📋 Templates de Mensagem"):
        tpl_nome = st.selectbox("Escolher template:", list(TEMPLATES_MSG.keys()),
                                key="tpl_sel")
        tpl      = TEMPLATES_MSG[tpl_nome]
        canal_tpl = tpl["canal"]

        st.markdown(
            f"<div style='font-size:10px;font-weight:700;color:#004D40;"
            f"letter-spacing:1.5px;text-transform:uppercase;margin:8px 0 4px;'>"
            f"Canal: {'📱 WhatsApp' if canal_tpl=='whatsapp' else '📧 E-mail'}</div>",
            unsafe_allow_html=True)

        tpl_edit = st.text_area(
            "Editar antes de usar (não altera o original):",
            value=tpl["texto"], height=160, key="tpl_edit")

        if canal_tpl == "whatsapp":
            tpl_tel = st.text_input("Telefone destino (só números):", key="tpl_tel",
                                     placeholder="31999990000")
            if tpl_tel:
                url_tpl = f"https://wa.me/55{tpl_tel}?text={urllib.parse.quote(tpl_edit)}"
                st.markdown(
                    f'<a href="{url_tpl}" target="_blank" class="wa-btn">'
                    f'Abrir no WhatsApp</a>',
                    unsafe_allow_html=True)
        else:
            tpl_dest = st.text_input("E-mail destino:", key="tpl_dest")
            tpl_subj = st.text_input("Assunto:", value="HOVA — Processo Seletivo", key="tpl_subj")
            if st.button("ENVIAR E-MAIL", key="tpl_enviar", type="primary"):
                if tpl_dest:
                    ok = send_email(tpl_dest, tpl_subj, tpl_edit)
                    if ok: st.success("E-mail enviado!")
                    else:  st.error("Falha ao enviar.")
                else:
                    st.warning("Informe o e-mail destino.")

    st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

    # ══════════════════════════════════════
    # LISTA DE AGENDADOS
    # ══════════════════════════════════════
    if not ag_list:
        st.markdown('<div class="empty"><div class="e-title">SEM ENTREVISTAS</div>'
                    '<div class="e-sub">Nenhuma entrevista confirmada ainda.</div></div>',
                    unsafe_allow_html=True)
    else:
        st.markdown(f"<div style='font-size:12px;color:#8A94A6;margin-bottom:16px;'>"
                    f"<b style='color:#004D40;font-size:17px;'>{len(ag_list)}</b> entrevista(s) agendada(s)</div>",
                    unsafe_allow_html=True)
        for c in sorted(ag_list, key=lambda x: (x.get('data_entrevista') or datetime.date.min,
                                                  x.get('hora_entrevista') or datetime.time(0))):
            hf = c['hora_entrevista'].strftime('%H:%M') if c.get('hora_entrevista') else '—'
            df = c['data_entrevista'].strftime('%d/%m/%Y') if c.get('data_entrevista') else '—'

            # Badge de alerta não lido
            tem_alerta = any(not a.get('lido') for a in c.get('alertas',[]))
            badge_alerta = (
                "<span style='background:#FEF3C7;color:#92540A;font-size:9px;"
                "font-weight:700;padding:2px 8px;border-radius:20px;"
                "margin-left:8px;letter-spacing:1px;'>📬 RESPONDEU</span>"
                if tem_alerta else ""
            )

            st.markdown(f"<div class='card-agendado'>", unsafe_allow_html=True)
            ci, cd, ca = st.columns([1,3,3])
            with ci:
                st.markdown(f'<div class="avatar" style="width:64px;height:64px;font-size:18px;">{iniciais(c["nome"])}</div>', unsafe_allow_html=True)
            with cd:
                st.markdown(f"**{c['nome']}**{badge_alerta}", unsafe_allow_html=True)
                st.markdown(f"Data: **{df}** às **{hf}**")
                st.markdown(f"<span style='color:#8A94A6;font-size:12px;'>{c.get('setor','—')} | {c.get('email','—')}</span>", unsafe_allow_html=True)
                # Observações salvas
                if c.get('observacoes_rh'):
                    st.markdown(
                        f"<div style='font-size:11px;color:#004D40;font-style:italic;"
                        f"margin-top:4px;background:#F0FAF8;padding:4px 8px;"
                        f"border-radius:6px;border-left:2px solid #004D40;'>"
                        f"📝 {c['observacoes_rh'][:120]}"
                        f"{'...' if len(c.get('observacoes_rh',''))>120 else ''}</div>",
                        unsafe_allow_html=True)
            with ca:
                tel = c.get('telefone','')
                if tel:
                    mwa = (
                        f"Olá! 😊\n\n"
                        f"Aqui é a equipe de RH do *Hospital de Olhos Vale do Aço*.\n\n"
                        f"Estamos te aguardando para a entrevista:\n"
                        f"📅 *{df}* às *{hf}*\n"
                        f"📍 {ENDERECO_HOVA}\n\n"
                        f"Ao chegar, informe na recepção que é para a entrevista e pergunte por *Josi* ou *Paula*.\n\n"
                        f"Qualquer dúvida, estamos à disposição. Até lá! 🤝"
                    )
                    st.markdown(f'<a href="https://wa.me/{tel}?text={urllib.parse.quote(mwa)}" target="_blank" class="wa-btn">Confirmar via WhatsApp</a>', unsafe_allow_html=True)

                if st.session_state.contratar_foco == c['id']:
                    tem_email_c = bool(c.get('email','').strip())
                    eh_aprendiz = c.get('setor','') == 'JOVEM APRENDIZ'

                    st.markdown(
                        "<div class='notif notif-info' style='text-align:left;font-size:12px;'>"
                        "📋 <b>Passo 1 de 2</b> — Envio do pedido de documentos. "
                        "Data de início e horário serão definidos após o recebimento dos docs, "
                        "no dossiê do colaborador.</div>",
                        unsafe_allow_html=True)

                    st.caption("PRAZO PARA ENTREGA DOS DOCUMENTOS")
                    dl = st.date_input(
                        "", value=datetime.date.today() + datetime.timedelta(days=5),
                        key=f"dl_{c['id']}", label_visibility="collapsed")

                    if not tem_email_c:
                        tn = st.text_input("WhatsApp (só números):", value=tel, key=f"wa_{c['id']}")
                        msg_adm_wa = (
                            f"Olá! 😊\n\n"
                            f"Aqui é o RH do *Hospital de Olhos Vale do Aço*.\n\n"
                            f"Temos o prazer de informar que você foi *selecionado(a)* "
                            f"para integrar nossa equipe! 🎉\n\n"
                            f"Para darmos continuidade, precisamos que você envie os "
                            f"documentos abaixo até *{dl.strftime('%d/%m/%Y')}* "
                            f"pelo e-mail: *rh@holhosvaledoaco.com.br*\n\n"
                            f"📋 *Documentos necessários:*\n"
                            + ("• RG\n• CPF\n• Comprovante de residência\n• Cartão do PIS\n"
                               "• Certidão de nascimento\n• Comprovante de matrícula escolar\n"
                               "• Cartão de vacinação\n• Certidões (filhos, se houver)\n"
                               if eh_aprendiz else
                               "• RG\n• CPF\n• Comprovante de residência\n• Cartão do PIS\n"
                               "• Diploma (se houver)\n• Cartão de vacinação\n"
                               "• Certidões (casamento/filhos, se houver)\n")
                            + f"\nA foto 3x4 entregar presencialmente.\n\n"
                            f"A data de início será informada em breve!\n\n"
                            f"Qualquer dúvida, estamos à disposição! 🤝\n"
                            f"— Equipe de RH — HOVA"
                        )
                        st.markdown(
                            "<div class='notif notif-info' style='font-size:12px;text-align:left;'>"
                            "📱 Sem e-mail — pedido será enviado via WhatsApp.</div>",
                            unsafe_allow_html=True)

                    cx, cok = st.columns(2)
                    with cx:
                        if st.button("CANCELAR", key=f"cx_{c['id']}",
                                     type="secondary", use_container_width=True):
                            st.session_state.contratar_foco = None
                            st.rerun()
                    with cok:
                        if tem_email_c:
                            if st.button("CONFIRMAR E ENVIAR DOCS", key=f"cok_{c['id']}",
                                         type="primary", use_container_width=True):
                                with st.spinner("Enviando pedido de documentos..."):
                                    ok = send_email_admissao(
                                        c['email'], c['nome'], dl,
                                        None, None,
                                        c.get('id',''),
                                        aprendiz=eh_aprendiz)
                                c.update({
                                    'data_inicio_contrato':   None,
                                    'hora_inicio_contrato':   None,
                                    'email_admissao_enviado': ok,
                                    'eptom': eh_aprendiz,
                                })
                                st.session_state.contratados.append(c)
                                st.session_state.agendados.remove(c)
                                st.session_state.contratar_foco = None
                                salvar_json()
                                st.session_state.sync_msg = {
                                    'tipo': 'ok' if ok else 'warn',
                                    'texto': f"{c['nome']} contratado(a). Pedido de docs enviado." if ok
                                             else f"{c['nome']} movido para Contratados. E-mail pendente."
                                }
                                time.sleep(1)
                                st.rerun()
                        else:
                            tel_limpo = ''.join(filter(str.isdigit, tn))
                            if tel_limpo:
                                url_adm_wa = f"https://wa.me/55{tel_limpo}?text={urllib.parse.quote(msg_adm_wa)}"
                                st.markdown(
                                    f'<a href="{url_adm_wa}" target="_blank" class="wa-btn" '
                                    f'style="display:block;text-align:center;height:48px;'
                                    f'line-height:48px;border-radius:9px;font-size:11px;'
                                    f'font-weight:700;letter-spacing:1px;">'
                                    f'ENVIAR VIA WHATSAPP</a>',
                                    unsafe_allow_html=True)
                            if st.button("CONFIRMAR CONTRATAÇÃO ✓", key=f"cok_{c['id']}",
                                         type="primary", use_container_width=True):
                                c.update({
                                    'data_inicio_contrato':   None,
                                    'hora_inicio_contrato':   None,
                                    'telefone': ''.join(filter(str.isdigit, tn)) if not tem_email_c else tel,
                                    'email_admissao_enviado': False,
                                    'eptom': eh_aprendiz,
                                })
                                st.session_state.contratados.append(c)
                                st.session_state.agendados.remove(c)
                                st.session_state.contratar_foco = None
                                salvar_json()
                                st.session_state.sync_msg = {
                                    'tipo': 'ok',
                                    'texto': f"{c['nome']} contratado(a). Instruções enviadas via WhatsApp."
                                }
                                time.sleep(1)
                                st.rerun()

                elif st.session_state.get('editar_agendado') == c['id']:
                    # ── Formulário de edição rápida ──
                    st.caption("EDITAR DATA E HORÁRIO")
                    nova_data = st.date_input(
                        "Nova data:",
                        value=c.get('data_entrevista') or datetime.date.today(),
                        key=f"ed_data_{c['id']}")
                    novo_hora = st.time_input(
                        "Novo horário:",
                        value=c.get('hora_entrevista') or datetime.time(9,0),
                        key=f"ed_hora_{c['id']}")
                    ea1, ea2 = st.columns(2)
                    with ea1:
                        if st.button("CANCELAR", key=f"ed_canc_{c['id']}",
                                     type="secondary", use_container_width=True):
                            st.session_state['editar_agendado'] = None
                            st.rerun()
                    with ea2:
                        if st.button("SALVAR", key=f"ed_salv_{c['id']}",
                                     type="primary", use_container_width=True):
                            c['data_entrevista'] = nova_data
                            c['hora_entrevista'] = novo_hora
                            salvar_json()
                            st.session_state['editar_agendado'] = None
                            st.rerun()

                else:
                    be1, be2, be3 = st.columns(3)
                    with be1:
                        if st.button("EDITAR", key=f"ed_{c['id']}",
                                     use_container_width=True):
                            st.session_state['editar_agendado'] = c['id']
                            st.rerun()
                    with be2:
                        if st.button("NÃO CONTRATAR", key=f"nc_ag_{c['id']}",
                                     type="secondary", use_container_width=True):
                            st.session_state['nc_agendado_foco'] = c['id']
                            st.rerun()
                    with be3:
                        if st.button("CONTRATAR", key=f"ct_{c['id']}",
                                     type="primary", use_container_width=True):
                            st.session_state.contratar_foco = c['id']
                            st.rerun()

                    # ── Modal Não Contratar ──
                    if st.session_state.get('nc_agendado_foco') == c['id']:
                        st.markdown(
                            "<div style='background:#FFFAF8;border:1.5px solid #E5BCBC;"
                            "border-radius:14px;padding:22px 26px;margin-top:8px;'>",
                            unsafe_allow_html=True)
                        st.markdown(
                            "<div style='font-size:13px;font-weight:700;color:#9B2C2C;"
                            "margin-bottom:14px;'>Confirmar — Não Contratar</div>",
                            unsafe_allow_html=True)
                        msg_nc_ag = st.text_area(
                            "Mensagem que será enviada:",
                            value=(
                                f"Olá {c['nome'].title()}, tudo bem?\n\n"
                                f"Gostaríamos de agradecer imensamente pela sua presença em nossa "
                                f"entrevista e por todo o tempo e dedicação que você investiu no "
                                f"nosso processo seletivo. Foi muito bom te conhecer!\n\n"
                                f"No momento não temos uma vaga disponível na sua área, mas seu "
                                f"currículo ficará guardado em nossa base de dados e entraremos "
                                f"em contato assim que surgir uma nova oportunidade.\n\n"
                                f"Desejamos muito sucesso na sua jornada profissional e que novas "
                                f"portas se abram para você em breve!\n\n"
                                f"Muito obrigada e até breve!\n\n"
                                f"Equipe de RH — Hospital de Olhos Vale do Aço"
                            ),
                            height=260, key=f"msg_nc_ag_{c['id']}")
                        nc1, nc2 = st.columns(2)
                        with nc1:
                            if st.button("Cancelar", key=f"nc_ag_canc_{c['id']}",
                                         use_container_width=True):
                                st.session_state['nc_agendado_foco'] = None
                                st.rerun()
                        with nc2:
                            if st.button("CONFIRMAR E ENVIAR", key=f"nc_ag_env_{c['id']}",
                                         type="primary", use_container_width=True):
                                if c.get('email','').strip():
                                    with st.spinner("Enviando mensagem..."):
                                        send_email(c['email'],
                                                   "Hospital de Olhos Vale do Aço — Processo Seletivo",
                                                   msg_nc_ag)
                                st.session_state.agendados.remove(c)
                                st.session_state['nc_agendado_foco'] = None
                                salvar_json()
                                st.rerun()
                        st.markdown("</div>", unsafe_allow_html=True)

# ── ABA 7: AGUARDANDO RETORNO ─────────────
with abas[7]:
    pend = _busca(st.session_state.aguardando_retorno, termo)

    if st.button("LER RESPOSTAS E AGENDAR AUTOMATICO", type="primary", use_container_width=True):
        res_auto = []
        # IDs de mensagens já processadas nesta sessão (evita loop infinito)
        if 'respostas_processadas' not in st.session_state:
            st.session_state.respostas_processadas = set()

        with st.spinner("Verificando respostas de e-mail..."):
            try:
                conn = imaplib.IMAP4_SSL(IMAP_SERVER, 993, timeout=15)
                conn.login(EMAIL_CONTA, SENHA_CONTA)
                conn.select("INBOX")
                _, ids = conn.search(None, '(SUBJECT "Re: HOVA")')
                for mid in (ids[0].split() or [])[-50:]:
                    mid_str = mid.decode() if isinstance(mid, bytes) else str(mid)

                    # ── Pular e-mails já processados nesta sessão ──
                    if mid_str in st.session_state.respostas_processadas:
                        continue

                    _, md = conn.fetch(mid,'(RFC822)')
                    msg   = email.message_from_bytes(md[0][1])
                    rem   = email.utils.parseaddr(msg.get('From',''))[1].lower()
                    corpo = ''
                    for pt in msg.walk():
                        if pt.get_content_type() == "text/plain":
                            try: corpo += pt.get_payload(decode=True).decode('utf-8',errors='ignore')
                            except: pass

                    # ── BUG FIX: extrair opção APENAS das primeiras 3 linhas não-vazias
                    # (ignora o histórico quotado do e-mail anterior)
                    linhas_novas = []
                    for linha in corpo.splitlines():
                        linha = linha.strip()
                        # Linhas quotadas começam com ">" — ignorar
                        if linha.startswith('>'):
                            break
                        if linha:
                            linhas_novas.append(linha)
                        if len(linhas_novas) >= 5:
                            break

                    texto_resposta = ' '.join(linhas_novas).lower()
                    op = None
                    # Buscar "1", "2" ou "3" como número isolado (não dentro de outra palavra)
                    for opcao in ["1", "2", "3"]:
                        import re as _re
                        if _re.search(rf'\b{opcao}\b', texto_resposta):
                            op = opcao
                            break

                    if not op:
                        continue

                    cand = next((c for c in st.session_state.aguardando_retorno
                                 if c['email'] == rem), None)
                    if not cand:
                        continue

                    # Marcar como processado ANTES de qualquer ação
                    st.session_state.respostas_processadas.add(mid_str)

                    hmap = {
                        "1": cand.get('opcao_1'),
                        "2": cand.get('opcao_2'),
                        "3": cand.get('opcao_3'),
                    }
                    hd = hmap.get(op)
                    dd = cand.get('data_entrevista')

                    if not hd:
                        res_auto.append(('warn', f"Opcao {op} invalida para {cand['nome']}."))
                        continue

                    # Re-ler disco antes de checar conflito
                    carregar_json()

                    conf = not horario_disponivel(dd, hd)
                    if conf:
                        livres_h = horarios_livres(dd, [h for h in hmap.values() if h])
                        if livres_h:
                            livres_txt = "\n".join(
                                f"[ {n} ] - {h.strftime('%H:%M')}"
                                for n, h in hmap.items()
                                if h and h in livres_h
                            )
                            send_email(
                                cand['email'],
                                "Re: HOVA — Horário Preenchido, Escolha Outra Opção",
                                f"Olá {cand['nome']},\n\n"
                                f"Infelizmente o horário de {hd.strftime('%H:%M')} "
                                f"acabou de ser confirmado por outro candidato.\n\n"
                                f"Ainda temos disponibilidade para o dia "
                                f"{dd.strftime('%d/%m/%Y')}. "
                                f"Por favor, responda com o número da sua nova escolha:\n\n"
                                f"{livres_txt}\n\n"
                                f"Atenciosamente,\nEquipe de RH — HOVA"
                            )
                            res_auto.append(('warn',
                                f"Conflito para {cand['nome']} — "
                                f"{hd.strftime('%H:%M')} ocupado. "
                                f"Novas opcoes enviadas."))
                        else:
                            send_email(
                                cand['email'],
                                "HOVA — Precisamos Reagendar",
                                f"Olá {cand['nome']},\n\n"
                                f"Todos os horários disponíveis para o dia "
                                f"{dd.strftime('%d/%m/%Y')} foram preenchidos.\n\n"
                                f"Nossa equipe entrará em contato para oferecer "
                                f"novas opcões de data e horário.\n\n"
                                f"Pedimos desculpas pelo transtorno!\n\n"
                                f"Atenciosamente,\nEquipe de RH — HOVA"
                            )
                            cand['alerta_lota'] = True
                            res_auto.append(('err',
                                f"ATENCAO: Todos horários de "
                                f"{dd.strftime('%d/%m/%Y')} esgotaram para "
                                f"{cand['nome']}. Reagende manualmente."))
                    else:
                        cand['hora_entrevista'] = hd
                        st.session_state.agendados.append(cand)
                        st.session_state.aguardando_retorno.remove(cand)
                        salvar_json()
                        res_auto.append(('ok',
                            f"{cand['nome']} agendado(a) para "
                            f"{hd.strftime('%H:%M')}."))
                    time.sleep(0.3)
                conn.logout()
                res_auto.append(('info', 'Varredura concluída.'))
            except Exception as e:
                res_auto.append(('err', f"Erro: {e}"))

        for tp, tx in res_auto:
            css = "notif-ok" if tp=='ok' else ("notif-warn" if tp in ('warn','err') else "notif-info")
            st.markdown(f"<div class='notif {css}'>{tx}</div>", unsafe_allow_html=True)
        time.sleep(1.5)
        st.rerun()

    st.write("")
    if not pend:
        st.markdown('<div class="empty"><div class="e-title">SEM PENDENCIAS</div>'
                    '<div class="e-sub">Nenhum candidato aguardando resposta.</div></div>',
                    unsafe_allow_html=True)
    else:
        st.markdown(f"<div style='font-size:12px;color:#8A94A6;margin-bottom:16px;'>"
                    f"<b style='color:#B7791F;font-size:17px;'>{len(pend)}</b> aguardando resposta</div>",
                    unsafe_allow_html=True)
        for c in pend:
            alerta = c.get('alerta_lota',False)
            df     = c['data_entrevista'].strftime('%d/%m/%Y') if c.get('data_entrevista') else '—'
            cls    = "card-alerta" if alerta else "card-pendente"
            st.markdown(f"<div class='{cls}'>", unsafe_allow_html=True)

            cc1, cc2, cc3, cc4 = st.columns([3, 1, 1, 0.5])
            with cc1:
                st.markdown(f"**{c['nome']}**")
                st.markdown(
                    f"<span style='font-size:12px;color:#8A94A6;'>"
                    f"{c.get('email','—')} &nbsp;·&nbsp; Entrevista: <b>{df}</b>"
                    f"</span>", unsafe_allow_html=True)
                if alerta:
                    st.markdown(
                        "<span style='color:#9B2C2C;font-size:12px;font-weight:700;'>"
                        "TODOS OS HORARIOS ESGOTARAM — REAGENDE MANUALMENTE.</span>",
                        unsafe_allow_html=True)
            with cc2:
                if st.button("Agendar", key=f"mv_{c['id']}", use_container_width=True,
                             type="primary"):
                    c['hora_entrevista'] = c.get('opcao_1', datetime.time(9,0))
                    st.session_state.agendados.append(c)
                    st.session_state.aguardando_retorno.remove(c)
                    salvar_json()
                    st.rerun()
            with cc3:
                if st.button("Nao Contratar", key=f"nc_{c['id']}", use_container_width=True):
                    st.session_state['nao_contratar_foco'] = c['id']
                    st.rerun()
            with cc4:
                if st.button("🗑", key=f"del_ag_{c['id']}", use_container_width=True,
                             help="Remover desta lista sem enviar e-mail"):
                    st.session_state.aguardando_retorno.remove(c)
                    salvar_json()
                    st.rerun()

            st.markdown("</div>", unsafe_allow_html=True)

            # ── Modal de confirmação de não contratar ─────────
            if st.session_state.get('nao_contratar_foco') == c['id']:
                with st.container():
                    st.markdown(
                        "<div style='background:#FFFAF8;border:1.5px solid #E5BCBC;"
                        "border-radius:14px;padding:22px 26px;margin-top:8px;'>",
                        unsafe_allow_html=True)
                    st.markdown(
                        "<div style='font-size:13px;font-weight:700;color:#9B2C2C;"
                        "margin-bottom:14px;'>Confirmar — Nao Contratar</div>",
                        unsafe_allow_html=True)

                    msg_padrao = (
                        f"Olá {c['nome'].title()}, tudo bem?\n\n"
                        f"Aqui é a equipe de RH do Hospital de Olhos Vale do Aço.\n\n"
                        f"Gostaríamos de agradecer seu interesse e sua disponibilidade "
                        f"durante o nosso processo seletivo.\n\n"
                        f"Após análise cuidadosa, optamos por seguir com outro perfil "
                        f"para esta oportunidade no momento. Seu currículo permanecerá "
                        f"em nossa base de dados e entraremos em contato em novas "
                        f"oportunidades que surgirem.\n\n"
                        f"Muito obrigada pela sua compreensão e boa sorte!\n\n"
                        f"Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
                    )

                    msg_edit = st.text_area(
                        "Mensagem que será enviada por e-mail:",
                        value=msg_padrao,
                        height=220,
                        key=f"msg_nc_{c['id']}")

                    col_canc, col_env = st.columns(2)
                    with col_canc:
                        if st.button("Cancelar", key=f"nc_canc_{c['id']}",
                                     use_container_width=True):
                            st.session_state['nao_contratar_foco'] = None
                            st.rerun()
                    with col_env:
                        if st.button("CONFIRMAR E ENVIAR", key=f"nc_env_{c['id']}",
                                     type="primary", use_container_width=True):
                            with st.spinner("Enviando mensagem..."):
                                ok = send_email(
                                    c['email'],
                                    "Hospital de Olhos Vale do Aço — Processo Seletivo",
                                    msg_edit)
                            st.session_state.aguardando_retorno.remove(c)
                            st.session_state['nao_contratar_foco'] = None
                            salvar_json()
                            if ok:
                                st.success(f"Mensagem enviada para {c['nome']}.")
                            else:
                                st.warning("Candidato removido, mas o e-mail falhou. Avise manualmente.")
                            time.sleep(1)
                            st.rerun()
                    st.markdown("</div>", unsafe_allow_html=True)

# ── ABA 8: CENTRO DE GESTÃO DE PESSOAS ────
with abas[8]:

    # ── Função interna: e-mail para NTW com anexos ──────────────
    def enviar_ntw(func: dict) -> tuple[bool, str]:
        """Envia e-mail de admissão para a NTW Doctor com todos os PDFs do dossiê."""
        try:
            from email.mime.multipart import MIMEMultipart
            from email.mime.base      import MIMEBase
            from email                import encoders

            nome      = func.get('nome','').title()
            cargo     = func.get('cargo_atual','—')
            carga_h   = func.get('carga_horaria','—')
            vt        = "Sim" if func.get('vale_transporte') else "Não"
            linhas    = func.get('linhas_onibus','—')
            data_ini  = func['data_inicio_contrato'].strftime('%d/%m/%Y') \
                        if func.get('data_inicio_contrato') else 'Data não informada'
            data_exp  = func['data_inicio_experiencia'].strftime('%d/%m/%Y') \
                        if func.get('data_inicio_experiencia') else data_ini

            corpo = (
                f"Bom dia,\n\n"
                f"Seguem documentos para admissão a partir de {data_ini}.\n\n"
                f"Colaborador: {nome}\n"
                f"Cargo: {cargo}\n"
                f"Carga Horária: {carga_h}\n"
                f"Início da Experiência: {data_exp}\n"
                f"Vale Transporte: {vt}"
                + (f"\nLinhas: {linhas}" if vt == "Sim" else "")
                + f"\n\nAtenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
            )

            msg = MIMEMultipart()
            msg['Subject'] = f"Admissão - {nome} - HOVA"
            msg['From']    = EMAIL_CONTA
            msg['To']      = EMAIL_CONTABILIDADE
            msg['Bcc']     = EMAIL_CONTA
            msg.attach(MIMEText(corpo, 'plain', 'utf-8'))

            # Anexar todos os PDFs do dossiê
            docs = func.get('documentos', {})
            for nome_doc, bytes_doc in docs.items():
                if bytes_doc:
                    part = MIMEBase('application','octet-stream')
                    part.set_payload(bytes_doc)
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition',
                                    f'attachment; filename="{nome_doc}.pdf"')
                    msg.attach(part)

            with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as s:
                s.login(EMAIL_CONTA, SENHA_CONTA)
                s.send_message(msg)
            return True, f"E-mail enviado para {EMAIL_CONTABILIDADE} (TESTE)"
        except Exception as e:
            return False, f"Erro ao enviar: {e}"

    # ── CSS ciano para este módulo ───────────────────────────────
    st.markdown("""
    <style>
    .ciano { color: #004D40 !important; }
    .btn-ciano div[data-testid="stButton"] button[kind="primary"] {
        background: linear-gradient(135deg,#004D40,#26A69A) !important;
        box-shadow: 0 2px 10px rgba(0,77,64,0.25) !important;
    }
    .btn-ciano div[data-testid="stButton"] button[kind="primary"]:hover {
        background: #003329 !important;
    }
    .dossie-header {
        background: linear-gradient(160deg,#F2FAF8,#E6F4F1);
        border: 1px solid #B2DFDB;
        border-radius: 20px;
        padding: 36px 40px 28px;
        text-align: center;
        margin-bottom: 20px;
        position: relative;
        overflow: hidden;
    }
    .dossie-header::before {
        content:''; position:absolute; left:0; top:0; bottom:0;
        width:5px; background:linear-gradient(180deg,#004D40,#26A69A);
        border-radius:20px 0 0 20px;
    }
    .func-avatar-xl {
        width:110px; height:110px; border-radius:50%;
        background:linear-gradient(145deg,#004D40,#26A69A);
        color:#FFF; display:flex; justify-content:center; align-items:center;
        font-size:34px; font-weight:900;
        margin:0 auto 16px auto;
        box-shadow:0 6px 24px rgba(0,77,64,0.3);
        letter-spacing:1px;
    }
    .func-avatar-xl-foto {
        width:110px; height:110px; border-radius:50%;
        background-size: cover;
        background-position: center center;
        background-repeat: no-repeat;
        border:4px solid #B2DFDB;
        margin:0 auto 16px auto;
        box-shadow:0 6px 24px rgba(0,0,0,0.15);
    }
    .ntw-box {
        background:#F2FAF8; border:1.5px solid #004D40;
        border-radius:14px; padding:22px 26px; margin-top:16px;
    }
    .func-card-v2 {
        background:#FFFFFF; border:1px solid #E2E6EA; border-radius:18px;
        padding:28px 16px 20px; text-align:center;
        box-shadow:0 2px 12px rgba(0,0,0,0.05);
        transition:transform 0.18s,box-shadow 0.18s;
        animation:fadeUp 0.3s ease;
    }
    .func-card-v2:hover { transform:translateY(-4px); box-shadow:0 8px 24px rgba(8,145,178,0.12); }
    </style>
    """, unsafe_allow_html=True)

    # ── Sub-navegação ────────────────────────────────────────────
    sub = st.radio("", ["Equipe Ativa", "Ex-Colaboradores"],
                   horizontal=True, label_visibility="collapsed", key="sub_func")

    # ── Expander: cadastrar colaborador já ativo ─────────────────
    with st.expander("Cadastrar Colaborador Existente"):
        with st.form("form_colab_antigo", clear_on_submit=True):
            st.markdown(
                "<div style='font-size:12px;font-weight:800;color:#004D40;"
                "letter-spacing:2px;text-transform:uppercase;margin-bottom:16px;'>"
                "Novo Colaborador</div>", unsafe_allow_html=True)

            ca1, ca2 = st.columns(2)
            ca_nome  = ca1.text_input("Nome completo *")
            ca_cargo = ca2.text_input("Cargo *", placeholder="Ex: Recepcionista")
            ca3, ca4 = st.columns(2)
            ca_setor = ca3.selectbox("Setor",["TRIAGEM GERAL","RECEPCAO E ATENDIMENTO",
                                               "TECNICO E ENFERMAGEM","ADMINISTRATIVO",
                                               "FATURAMENTO","JOVEM APRENDIZ"])
            ca_adm   = ca4.date_input("Data de admissão", value=datetime.date.today(),
                                       key="ca_adm2")
            ca5, ca6 = st.columns(2)
            ca_tel   = ca5.text_input("Telefone / WhatsApp", placeholder="31999990000")
            ca_email = ca6.text_input("E-mail")
            ca7, ca8 = st.columns(2)
            ca_unif  = ca7.text_input("Nº do Uniforme", placeholder="Ex: P / M / G")
            ca_carga = ca8.text_input("Carga Horária", placeholder="Ex: 6h semanais")
            ca_vt    = st.checkbox("Utiliza Vale Transporte")
            ca_linhas= st.text_input("Linhas de Ônibus", placeholder="Ex: 201, 405",
                                      disabled=not ca_vt)
            ca_foto  = st.file_uploader("Foto do colaborador (opcional)",
                                         type=["jpg","jpeg","png"],
                                         key="ca_foto_upload")
            ca_obs   = st.text_area("Observações", height=70,
                                     placeholder="Certificações, CNH, anotações...")

            st.markdown("<div class='btn-ciano'>", unsafe_allow_html=True)
            ok_ca = st.form_submit_button("CADASTRAR", type="primary",
                                           use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        if ok_ca:
            if ca_nome and ca_cargo:
                foto_bytes = ca_foto.read() if ca_foto else None
                novo_func = {
                    "id":   str(int(time.time()*1000)),
                    "nome": ca_nome.upper().strip(),
                    "setor": ca_setor,
                    "email": ca_email.lower().strip(),
                    "telefone": ''.join(filter(str.isdigit, ca_tel)),
                    "data_inicio_contrato":    ca_adm,
                    "data_inicio_experiencia": ca_adm,
                    "hora_inicio_contrato":    datetime.time(8,0),
                    "cargo_atual":  ca_cargo.strip(),
                    "carga_horaria": ca_carga.strip(),
                    "num_uniforme":  ca_unif.strip(),
                    "vale_transporte": ca_vt,
                    "linhas_onibus": ca_linhas.strip(),
                    "observacoes":   ca_obs.strip(),
                    "foto":   foto_bytes,
                    "manual": True,
                    "email_admissao_enviado": False,
                    "documentos": {},   # {nome: bytes_pdf}
                    "docs_check": {     # checklist
                        "RG": False, "CPF": False, "PIS": False,
                        "Comprovante de Residência": False,
                        "Diploma": False, "Cartão de Vacina": False,
                        "Certidão de Casamento": False,
                        "Certidão de Nascimento dos Filhos": False,
                        "Foto 3x4": False,
                    },
                }
                st.session_state.contratados.append(novo_func)
                salvar_json()
                st.success(f"{ca_nome.upper()} cadastrado.")
                st.rerun()
            else:
                st.error("Nome e cargo são obrigatórios.")

    st.write("")

    # ════════════════════════════════════════════════════════════
    # EQUIPE ATIVA
    # ════════════════════════════════════════════════════════════
    if sub == "Equipe Ativa":
        ativos = _busca(st.session_state.contratados, termo)
        ativos = sorted(ativos, key=lambda x: x.get('data_inicio_contrato') or datetime.date.min)

        if not ativos:
            st.markdown('<div class="empty"><div class="e-title">NENHUM COLABORADOR ATIVO</div>'
                        '<div class="e-sub">Use o painel acima para cadastrar a equipe.</div></div>',
                        unsafe_allow_html=True)

        # ── DOSSIÊ ABERTO ────────────────────────────────────────
        elif st.session_state.perfil_foco:
            func = next((f for f in st.session_state.contratados
                         if f['id'] == st.session_state.perfil_foco), None)

            if not func:
                st.session_state.perfil_foco = None
                st.rerun()
            else:
                # Garantir campos novos em registros antigos
                func.setdefault('documentos', {})
                func.setdefault('docs_check', {
                    "RG":False,"CPF":False,"PIS":False,
                    "Comprovante de Residência":False,"Diploma":False,
                    "Cartão de Vacina":False,"Certidão de Casamento":False,
                    "Certidão de Nascimento dos Filhos":False,"Foto 3x4":False,
                })
                func.setdefault('carga_horaria','')
                func.setdefault('num_uniforme','')
                func.setdefault('vale_transporte',False)
                func.setdefault('linhas_onibus','')
                func.setdefault('data_inicio_experiencia',
                                func.get('data_inicio_contrato'))
                func.setdefault('ntw_enviado', False)

                ini_f = func['data_inicio_contrato'].strftime('%d/%m/%Y') \
                        if func.get('data_inicio_contrato') else 'Data não informada'

                # ── Header do perfil ─────────────────────────────
                if func.get('foto'):
                    b64f    = base64.b64encode(func['foto']).decode()
                    av_html = (f"<div class='func-avatar-xl-foto'"
                               f" style='background-image:url(\"data:image/jpeg;"
                               f"base64,{b64f}\");'></div>")
                else:
                    av_html = (f"<div class='func-avatar-xl'>"
                               f"{iniciais(func['nome'])}</div>")

                ntw_badge = (
                    "<span style='background:#DCFCE7;color:#166534;font-size:10px;"
                    "font-weight:700;padding:3px 10px;border-radius:20px;"
                    "letter-spacing:1px;'>NTW ENVIADO</span>"
                    if func.get('ntw_enviado') else ""
                )

                st.markdown(
                    f"<div class='dossie-header'>"
                    f"{av_html}"
                    f"<div style='font-size:24px;font-weight:900;color:#0D1B2A;"
                    f"margin-bottom:4px;'>{func['nome']}</div>"
                    f"<div style='font-size:12px;color:#004D40;font-weight:700;"
                    f"letter-spacing:1.5px;text-transform:uppercase;'>"
                    f"{func.get('cargo_atual','—')}</div>"
                    f"<div style='font-size:11px;color:#9AA5B4;margin-top:6px;'>"
                    f"Admitido em {ini_f} &nbsp;·&nbsp; {func.get('setor','—')}"
                    f"&nbsp;&nbsp;{ntw_badge}</div>"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # ── 3 colunas de ação rápida ─────────────────────
                qa1, qa2, qa3 = st.columns(3)
                with qa1:
                    tel_wa = ''.join(filter(str.isdigit,
                                            func.get('telefone','')))
                    if tel_wa:
                        msg_wa = (f"Olá {func['nome'].title()}, aqui é o RH do HOVA. "
                                  f"Verificamos que há documentos pendentes no seu dossiê. "
                                  f"Pode confirmar o envio?")
                        url_wa = (f"https://wa.me/55{tel_wa}?"
                                  f"text={urllib.parse.quote(msg_wa)}")
                        st.markdown(
                            f'<a href="{url_wa}" target="_blank" class="wa-btn"'
                            f' style="display:block;text-align:center;">'
                            f'WhatsApp Colaborador</a>',
                            unsafe_allow_html=True)
                with qa2:
                    # WhatsApp Paula
                    msg_paula = (
                        f"Paula, os documentos de admissão de "
                        f"{func['nome'].title()} estão disponíveis no sistema. "
                        f"Cargo: {func.get('cargo_atual','—')}.")
                    url_paula = (f"https://wa.me/5531886000023?"
                                 f"text={urllib.parse.quote(msg_paula)}")
                    st.markdown(
                        f'<a href="{url_paula}" target="_blank"'
                        f' style="display:block;text-align:center;'
                        f'background:#00A884;color:#FFF;padding:13px 18px;'
                        f'border-radius:9px;font-size:11px;font-weight:700;'
                        f'letter-spacing:1px;text-transform:uppercase;'
                        f'text-decoration:none;">'
                        f'WhatsApp Paula</a>',
                        unsafe_allow_html=True)
                with qa3:
                    if st.button("VOLTAR A EQUIPE", key="voltar_grid",
                                 use_container_width=True):
                        st.session_state.perfil_foco = None
                        st.rerun()

                st.write("")

                # ── TABS internas do dossiê ──────────────────────
                eh_eptom_func = func.get('eptom', False) or func.get('setor','') == 'JOVEM APRENDIZ'
                tabs_dossie = ["Dados & RH", "Documentos", "Enviar para Contabilidade"]
                if eh_eptom_func:
                    tabs_dossie.append("📋 Ficha EPTOM")
                tab_lista = st.tabs(tabs_dossie)
                tab_dados = tab_lista[0]
                tab_docs  = tab_lista[1]
                tab_ntw   = tab_lista[2]
                tab_eptom = tab_lista[3] if eh_eptom_func else None

                # ── TAB 1: DADOS ─────────────────────────────────
                with tab_dados:
                    with st.form(f"form_dados_{func['id']}"):
                        st.markdown(
                            "<div style='font-size:10px;font-weight:800;"
                            "color:#004D40;letter-spacing:2px;"
                            "text-transform:uppercase;margin-bottom:16px;'>"
                            "Informações Pessoais e Contratuais</div>",
                            unsafe_allow_html=True)

                        r1c1, r1c2 = st.columns(2)
                        novo_nome   = r1c1.text_input("Nome",  value=func.get('nome',''))
                        novo_cargo  = r1c2.text_input("Cargo", value=func.get('cargo_atual',''))

                        r2c1, r2c2 = st.columns(2)
                        novo_tel    = r2c1.text_input("Telefone/WhatsApp",
                                                      value=func.get('telefone',''))
                        novo_email  = r2c2.text_input("E-mail", value=func.get('email',''))

                        r3c1, r3c2 = st.columns(2)
                        SETORES_LIST = ["TRIAGEM GERAL","RECEPCAO E ATENDIMENTO",
                                        "TECNICO E ENFERMAGEM","ADMINISTRATIVO",
                                        "FATURAMENTO","JOVEM APRENDIZ"]
                        novo_setor  = r3c1.selectbox(
                            "Setor", SETORES_LIST,
                            index=SETORES_LIST.index(func['setor'])
                                  if func.get('setor') in SETORES_LIST else 0)
                        novo_unif   = r3c2.text_input("Nº Uniforme",
                                                      value=func.get('num_uniforme',''))

                        r4c1, r4c2 = st.columns(2)
                        novo_adm    = r4c1.date_input(
                            "Data de Admissão",
                            value=func.get('data_inicio_contrato') or datetime.date.today())
                        novo_exp    = r4c2.date_input(
                            "Início da Experiência",
                            value=func.get('data_inicio_experiencia')
                                  or func.get('data_inicio_contrato')
                                  or datetime.date.today())

                        r5c1, r5c2 = st.columns(2)
                        novo_carga  = r5c1.text_input("Carga Horária",
                                                      value=func.get('carga_horaria',''),
                                                      placeholder="Ex: 6h semanais")
                        novo_vt     = r5c2.checkbox("Vale Transporte",
                                                     value=func.get('vale_transporte',False))
                        novo_linhas = st.text_input("Linhas de Ônibus",
                                                    value=func.get('linhas_onibus',''),
                                                    placeholder="Ex: 201, 405",
                                                    disabled=not novo_vt)

                        # Upload de foto
                        nova_foto_up = st.file_uploader(
                            "Atualizar foto do perfil",
                            type=["jpg","jpeg","png"],
                            key=f"foto_edit_{func['id']}")

                        novo_obs = st.text_area(
                            "Observações de Desempenho",
                            value=func.get('observacoes',''), height=90,
                            placeholder="Feedbacks, advertências, elogios...")

                        st.markdown("<div style='height:8px;'></div>",
                                    unsafe_allow_html=True)
                        sb1, sb2, sb3 = st.columns(3)
                        salvar_ok   = sb1.form_submit_button(
                            "SALVAR", type="primary", use_container_width=True)
                        fechar_ok   = sb2.form_submit_button(
                            "FECHAR", use_container_width=True)
                        desligar_ok = sb3.form_submit_button(
                            "DESLIGAR", use_container_width=True)

                    if salvar_ok:
                        func['nome']               = novo_nome.upper().strip()
                        func['cargo_atual']        = novo_cargo.strip()
                        func['telefone']           = ''.join(filter(str.isdigit, novo_tel))
                        func['email']              = novo_email.lower().strip()
                        func['setor']              = novo_setor
                        func['num_uniforme']       = novo_unif.strip()
                        func['data_inicio_contrato']    = novo_adm
                        func['data_inicio_experiencia'] = novo_exp
                        func['carga_horaria']      = novo_carga.strip()
                        func['vale_transporte']    = novo_vt
                        func['linhas_onibus']      = novo_linhas.strip()
                        func['observacoes']        = novo_obs.strip()
                        if nova_foto_up:
                            func['foto'] = nova_foto_up.read()
                        salvar_json()
                        st.success("Dados salvos.")
                        st.rerun()

                    if fechar_ok:
                        st.session_state.perfil_foco = None
                        st.rerun()

                    if desligar_ok:
                        func['data_desligamento'] = datetime.date.today()
                        st.session_state.ex_funcionarios.append(func)
                        st.session_state.contratados.remove(func)
                        st.session_state.perfil_foco = None
                        salvar_json()
                        st.success(f"{func['nome']} desligado e movido para Histórico.")
                        st.rerun()

                    # ── Reenviar e-mail de admissão ──────────────────
                    st.markdown("<hr style='border:none;border-top:1px solid #E2E6EA;margin:20px 0 14px;'>",
                                unsafe_allow_html=True)

                    # ── PASSO 1: Pedido de documentos ────────────────
                    st.markdown(
                        "<div style='font-size:10px;font-weight:800;color:#004D40;"
                        "letter-spacing:2px;text-transform:uppercase;margin-bottom:10px;'>"
                        "📋 Passo 1 — Pedido de Documentos</div>", unsafe_allow_html=True)

                    _adm_ok      = func.get('email_admissao_enviado', False)
                    _eh_aprendiz = func.get('eptom', False) or func.get('setor','') == 'JOVEM APRENDIZ'

                    if _adm_ok:
                        st.markdown(
                            "<div class='notif notif-ok' style='margin-bottom:10px;'>"
                            "✅ Pedido de documentos já foi enviado.</div>",
                            unsafe_allow_html=True)

                    with st.form(f"form_reenvio_{func['id']}"):
                        re1, re2 = st.columns(2)
                        re_email = re1.text_input(
                            "E-mail:",
                            value=func.get('email',''),
                            placeholder="email@exemplo.com",
                            key=f"re_email_{func['id']}")
                        re_dl = re2.date_input(
                            "Prazo para documentos:",
                            value=datetime.date.today() + datetime.timedelta(days=5),
                            key=f"re_dl_{func['id']}")
                        reenviar_ok = st.form_submit_button(
                            "ENVIAR PEDIDO DE DOCUMENTOS",
                            type="primary", use_container_width=True)

                    if reenviar_ok:
                        if not re_email.strip():
                            st.error("Informe o e-mail.")
                        else:
                            with st.spinner("Enviando..."):
                                ok_re = send_email_admissao(
                                    re_email, func['nome'],
                                    re_dl, None, None,
                                    func.get('id',''),
                                    aprendiz=_eh_aprendiz)
                            if ok_re:
                                func['email']                  = re_email.lower().strip()
                                func['email_admissao_enviado'] = True
                                salvar_json()
                                st.success(f"✅ Pedido de documentos enviado para {re_email}")
                            else:
                                st.error("Falha ao enviar. Tente novamente.")

                    # ── PASSO 2: Confirmação de início ───────────────
                    st.markdown("<hr style='border:none;border-top:1px solid #E2E6EA;margin:20px 0 14px;'>",
                                unsafe_allow_html=True)
                    st.markdown(
                        "<div style='font-size:10px;font-weight:800;color:#004D40;"
                        "letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;'>"
                        "🗓 Passo 2 — Confirmar Data de Início</div>"
                        "<div style='font-size:11px;color:#9AA5B4;margin-bottom:12px;'>"
                        "Enviar após receber os documentos e alinhar com a contabilidade.</div>",
                        unsafe_allow_html=True)

                    _inicio_ok = func.get('email_inicio_enviado', False)
                    if _inicio_ok:
                        st.markdown(
                            "<div class='notif notif-ok' style='margin-bottom:10px;'>"
                            "✅ E-mail de início já foi enviado.</div>",
                            unsafe_allow_html=True)

                    with st.form(f"form_inicio_{func['id']}"):
                        fi1, fi2 = st.columns(2)
                        fi_email = fi1.text_input(
                            "E-mail:",
                            value=func.get('email',''),
                            key=f"fi_email_{func['id']}")
                        fi_di = fi2.date_input(
                            "Data de início:",
                            value=func.get('data_inicio_contrato') or datetime.date.today(),
                            key=f"fi_di_{func['id']}")

                        if not _eh_aprendiz:
                            fi3, fi4 = st.columns(2)
                            fi_hi = fi3.time_input(
                                "Horário de entrada:",
                                value=func.get('hora_inicio_contrato') or datetime.time(8,0),
                                key=f"fi_hi_{func['id']}")
                            fi_carga = fi4.text_input(
                                "Carga horária:",
                                value=func.get('carga_horaria',''),
                                placeholder="Ex: 44h semanais",
                                key=f"fi_carga_{func['id']}")
                        else:
                            fi_hi    = None
                            fi_carga = func.get('carga_horaria','')

                        fi_vt = st.checkbox(
                            "Utiliza Vale Transporte",
                            value=func.get('vale_transporte', False),
                            key=f"fi_vt_{func['id']}")
                        fi_linhas = st.text_input(
                            "Linhas de ônibus (se VT):",
                            value=func.get('linhas_onibus',''),
                            placeholder="Ex: 201, 405",
                            key=f"fi_linhas_{func['id']}")

                        inicio_ok = st.form_submit_button(
                            "ENVIAR CONFIRMAÇÃO DE INÍCIO",
                            type="primary", use_container_width=True)

                    if inicio_ok:
                        if not fi_email.strip():
                            st.error("Informe o e-mail.")
                        else:
                            _hi_str = fi_hi.strftime('%H:%M') if fi_hi else 'a definir'
                            _vt_str = (f"\n\nSobre o Vale Transporte: sim, será fornecido."
                                       + (f" Linhas: {fi_linhas}." if fi_linhas else "")
                                       if fi_vt else
                                       "\n\nSobre o Vale Transporte: não será necessário.")
                            if _eh_aprendiz:
                                corpo_inicio = (
                                    f"Prezada(o) {func['nome'].title()}, bom dia!\n\n"
                                    f"Aqui é a equipe de RH do Hospital de Olhos Vale do Aço.\n\n"
                                    f"Conforme alinhamento com a EPTOM, informamos que seu início "
                                    f"será no dia {fi_di.strftime('%d/%m/%Y')}.\n\n"
                                    f"O horário de trabalho será informado diretamente pela EPTOM.{_vt_str}\n\n"
                                    f"Qualquer dúvida, estamos à disposição!\n\n"
                                    f"Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
                                )
                            else:
                                corpo_inicio = (
                                    f"Prezada(o) {func['nome'].title()}, bom dia!\n\n"
                                    f"Aqui é a equipe de RH do Hospital de Olhos Vale do Aço.\n\n"
                                    f"Confirmamos que seu início será no dia "
                                    f"{fi_di.strftime('%d/%m/%Y')} às {_hi_str}.\n\n"
                                    f"Carga horária: {fi_carga or 'a confirmar'}.{_vt_str}\n\n"
                                    f"Endereço: {ENDERECO_HOVA}\n"
                                    f"Ao chegar, informe na recepção e pergunte por Josi ou Paula.\n\n"
                                    f"Qualquer dúvida, estamos à disposição!\n\n"
                                    f"Atenciosamente,\nEquipe de RH — Hospital de Olhos Vale do Aço"
                                )
                            with st.spinner("Enviando confirmação de início..."):
                                ok_ini = send_email(
                                    fi_email,
                                    "Hospital de Olhos Vale do Aço — Confirmação de Início",
                                    corpo_inicio)
                            if ok_ini:
                                func['data_inicio_contrato']  = fi_di
                                func['hora_inicio_contrato']  = fi_hi
                                func['carga_horaria']         = fi_carga
                                func['vale_transporte']       = fi_vt
                                func['linhas_onibus']         = fi_linhas
                                func['email_inicio_enviado']  = True
                                salvar_json()
                                st.success(f"✅ Confirmação de início enviada para {fi_email}")
                            else:
                                st.error("Falha ao enviar. Tente novamente.")

                # ── TAB 2: DOCUMENTOS ────────────────────────────
                with tab_docs:
                    # ── Botão de varredura automática ─────────────
                    st.markdown(
                        "<div style='font-size:10px;font-weight:800;color:#004D40;"
                        "letter-spacing:2px;text-transform:uppercase;"
                        "margin-bottom:10px;'>Verificar Documentos Recebidos por E-mail</div>",
                        unsafe_allow_html=True)

                    cand_id_8 = func.get('id','')[:8]
                    if st.button("VERIFICAR CAIXA DE ENTRADA",
                                 key=f"scan_docs_{func['id']}",
                                 use_container_width=True):
                        with st.spinner("Varrendo e-mails em busca de documentos..."):
                            encontrados = varrer_documentos_recebidos()

                        # Filtrar os que são deste funcionário
                        docs_func = [e for e in encontrados
                                     if e['cand_id_8'] == cand_id_8
                                     or e['remetente'] == func.get('email','').lower()]

                        if docs_func:
                            func.setdefault('documentos', {})
                            func.setdefault('docs_check', {})
                            novos_salvos = []
                            for doc in docs_func:
                                nd = doc['nome_doc']
                                func['documentos'][nd] = doc['bytes_pdf']
                                # Marcar no checklist se bater com algum item
                                for item in func.get('docs_check', {}):
                                    if item.lower() in nd.lower() or nd.lower() in item.lower():
                                        func['docs_check'][item] = True
                                novos_salvos.append(nd)
                            salvar_json()
                            st.markdown(
                                f"<div class='notif notif-ok'>"
                                f"{len(novos_salvos)} documento(s) recebido(s) e salvo(s) automaticamente: "
                                f"{', '.join(novos_salvos)}</div>",
                                unsafe_allow_html=True)
                        else:
                            st.markdown(
                                "<div class='notif notif-info'>"
                                "Nenhum documento novo encontrado. "
                                "O candidato ainda nao respondeu o e-mail com os PDFs.</div>",
                                unsafe_allow_html=True)

                    st.markdown(
                        "<div style='font-size:10px;font-weight:800;color:#004D40;"
                        "letter-spacing:2px;text-transform:uppercase;"
                        "margin:16px 0 10px;'>Checklist de Documentos</div>",
                        unsafe_allow_html=True)

                    DOCS_LISTA = [
                        "RG","CPF","PIS","Comprovante de Residência",
                        "Diploma","Cartão de Vacina",
                        "Certidão de Casamento",
                        "Certidão de Nascimento dos Filhos","Foto 3x4",
                    ]

                    checks_atuais = func.get('docs_check', {})
                    novos_checks  = {}
                    col_a, col_b = st.columns(2)
                    for idx, doc_nome in enumerate(DOCS_LISTA):
                        col = col_a if idx % 2 == 0 else col_b
                        with col:
                            val = st.checkbox(
                                doc_nome,
                                value=checks_atuais.get(doc_nome, False),
                                key=f"chk_{func['id']}_{doc_nome}")
                            novos_checks[doc_nome] = val

                    # Atualizar checklist em tempo real
                    func['docs_check'] = novos_checks

                    # Resumo
                    total_docs = len(DOCS_LISTA)
                    ok_docs    = sum(novos_checks.values())
                    pct        = int(ok_docs / total_docs * 100)
                    cor_pct    = "#166534" if pct == 100 else ("#004D40" if pct >= 50 else "#92540A")
                    st.markdown(
                        f"<div style='margin:16px 0 20px;padding:12px 18px;"
                        f"background:#F8FAFB;border-radius:10px;border:1px solid #E2E6EA;"
                        f"font-size:13px;color:{cor_pct};font-weight:700;'>"
                        f"{ok_docs} de {total_docs} documentos recebidos ({pct}%)"
                        f"</div>", unsafe_allow_html=True)

                    # Upload de PDFs
                    st.markdown(
                        "<div style='font-size:10px;font-weight:800;color:#004D40;"
                        "letter-spacing:2px;text-transform:uppercase;margin-bottom:10px;'>"
                        "Adicionar PDF ao Dossiê</div>", unsafe_allow_html=True)

                    doc_up_nome = st.selectbox(
                        "Documento:", DOCS_LISTA, key=f"doc_sel_{func['id']}")
                    doc_up_file = st.file_uploader(
                        "Selecionar PDF", type=["pdf"],
                        key=f"doc_up_{func['id']}")

                    if st.button("SALVAR DOCUMENTO",
                                 key=f"salvar_doc_{func['id']}",
                                 type="primary", use_container_width=True):
                        if doc_up_file:
                            func['documentos'][doc_up_nome] = doc_up_file.read()
                            func['docs_check'][doc_up_nome] = True
                            salvar_json()
                            st.success(f"{doc_up_nome} salvo no dossiê.")
                            st.rerun()
                        else:
                            st.warning("Selecione um arquivo PDF primeiro.")

                    # Listar PDFs já enviados
                    if func.get('documentos'):
                        st.markdown(
                            "<div style='font-size:10px;font-weight:800;color:#004D40;"
                            "letter-spacing:2px;text-transform:uppercase;"
                            "margin:16px 0 8px;'>PDFs no Dossiê</div>",
                            unsafe_allow_html=True)
                        for nome_doc, bytes_doc in func['documentos'].items():
                            if bytes_doc:
                                dc1, dc2 = st.columns([3,1])
                                dc1.markdown(
                                    f"<div style='font-size:13px;color:#0D1B2A;"
                                    f"font-weight:600;padding:8px 0;'>"
                                    f"{nome_doc}</div>",
                                    unsafe_allow_html=True)
                                dc2.download_button(
                                    "Baixar", bytes_doc,
                                    file_name=f"{nome_doc}_{func['nome']}.pdf",
                                    mime="application/pdf",
                                    key=f"dwn_{func['id']}_{nome_doc}",
                                    use_container_width=True)

                    salvar_json()  # salvar checklist

                # ── TAB 3: ENVIAR PARA NTW ───────────────────────
                with tab_ntw:
                    st.markdown(
                        "<div class='ntw-box'>"
                        "<div style='font-size:14px;font-weight:800;color:#004D40;"
                        "margin-bottom:4px;'>Encaminhar para Contabilidade</div>"
                        "<div style='font-size:12px;color:#4A5568;'>"
                        f"Envia e-mail para <b>{EMAIL_CONTABILIDADE}</b> "
                        "(MODO TESTE — troque para pessoal.expert@ntwdoctor.com.br quando confirmar) "
                        "com os dados de admissão e todos os PDFs do dossiê anexados."
                        "</div></div>",
                        unsafe_allow_html=True)

                    st.write("")
                    with st.form(f"form_ntw_{func['id']}"):
                        n1, n2 = st.columns(2)
                        ntw_cargo  = n1.text_input(
                            "Cargo (confirmar)",
                            value=func.get('cargo_atual',''))
                        ntw_carga  = n2.text_input(
                            "Carga Horária",
                            value=func.get('carga_horaria',''),
                            placeholder="Ex: 44h semanais")
                        n3, n4 = st.columns(2)
                        ntw_adm = n3.date_input(
                            "Data de início",
                            value=func.get('data_inicio_contrato') or datetime.date.today())
                        ntw_exp = n4.date_input(
                            "Início da Experiência",
                            value=func.get('data_inicio_experiencia')
                                  or func.get('data_inicio_contrato')
                                  or datetime.date.today())
                        ntw_vt  = st.checkbox(
                            "Vale Transporte",
                            value=func.get('vale_transporte', False))
                        ntw_linhas = st.text_input(
                            "Linhas", value=func.get('linhas_onibus',''),
                            disabled=not ntw_vt)

                        # Mostrar PDFs que serão anexados
                        docs_disponiveis = [k for k,v in func.get('documentos',{}).items() if v]
                        if docs_disponiveis:
                            st.markdown(
                                f"<div style='font-size:12px;color:#004D40;"
                                f"font-weight:600;margin-top:8px;'>"
                                f"PDFs que serão anexados: "
                                f"{', '.join(docs_disponiveis)}</div>",
                                unsafe_allow_html=True)
                        else:
                            st.markdown(
                                "<div style='font-size:12px;color:#92540A;"
                                "font-weight:600;margin-top:8px;'>"
                                "Nenhum PDF no dossiê ainda. Adicione na aba Documentos.</div>",
                                unsafe_allow_html=True)

                        enviar_ntw_ok = st.form_submit_button(
                            "ENVIAR PARA NTW DOCTOR",
                            type="primary", use_container_width=True)

                    if enviar_ntw_ok:
                        func['cargo_atual']             = ntw_cargo.strip()
                        func['carga_horaria']           = ntw_carga.strip()
                        func['data_inicio_contrato']    = ntw_adm
                        func['data_inicio_experiencia'] = ntw_exp
                        func['vale_transporte']         = ntw_vt
                        func['linhas_onibus']           = ntw_linhas.strip()
                        with st.spinner("Enviando para NTW Doctor..."):
                            ok_ntw, msg_ntw = enviar_ntw(func)
                        if ok_ntw:
                            func['ntw_enviado'] = True
                            salvar_json()
                            st.markdown(f"<div class='notif notif-ok'>{msg_ntw}</div>", unsafe_allow_html=True)
                        else:
                            st.markdown(f"<div class='notif notif-warn'>{msg_ntw}</div>", unsafe_allow_html=True)

                # ── TAB EPTOM ────────────────────────────────────
                if tab_eptom:
                    with tab_eptom:
                        d = FICHA_EMPRESA_DADOS

                        st.markdown(
                            "<div style='background:#FFF8EC;border:1.5px solid #D69E2E;"
                            "border-radius:14px;padding:16px 20px;margin-bottom:16px;'>"
                            "<div style='font-size:14px;font-weight:800;color:#92540A;margin-bottom:4px;'>"
                            "📋 Ficha da Empresa — EPTOM</div>"
                            "<div style='font-size:12px;color:#4A5568;'>"
                            "Preencha o que varia por contrato, confira a pré-visualização e envie.</div></div>",
                            unsafe_allow_html=True)

                        # ── Dados fixos do hospital (somente leitura) ──
                        st.markdown(
                            f"<div style='background:#F8FAFB;border:1px solid #E2E6EA;"
                            f"border-radius:10px;padding:12px 16px;margin-bottom:14px;"
                            f"font-size:12px;color:#4A5568;line-height:1.9;'>"
                            f"<b style='color:#004D40;'>Empresa:</b> {d['empresa']} &nbsp;·&nbsp; "
                            f"<b style='color:#004D40;'>CNPJ:</b> {d['cnpj']}<br>"
                            f"<b style='color:#004D40;'>Resp. no Setor:</b> {d['resp_setor']} &nbsp;·&nbsp; "
                            f"<b style='color:#004D40;'>Tel:</b> {d['tel_resp_setor']}<br>"
                            f"<b style='color:#004D40;'>Resp. pelo Contrato:</b> {d['resp_contrato']} &nbsp;·&nbsp; "
                            f"<b style='color:#004D40;'>E-mail:</b> {d['email_contrato']}"
                            f"</div>",
                            unsafe_allow_html=True)
                        st.caption("Para alterar dados fixos acima, edite FICHA_EMPRESA_DADOS no código.")

                        # ── Apenas campos variáveis ──
                        with st.form(key=f"form_eptom_{func['id']}"):
                            ec1, ec2 = st.columns(2)
                            ep_empresa  = ec1.text_input("Empresa:", value=func.get('eptom_empresa_edit', d['empresa']), key=f"ep_emp_{func['id']}")
                            ep_cnpj     = ec2.text_input("CNPJ:", value=func.get('eptom_cnpj_edit', d['cnpj']), key=f"ep_cnpj_{func['id']}")
                            ec3, ec4 = st.columns(2)
                            ep_resp_set = ec3.text_input("Resp. no Setor:", value=func.get('eptom_resp_setor_edit', d['resp_setor']), key=f"ep_rset_{func['id']}")
                            ep_tel_set  = ec4.text_input("Tel. Resp. Setor:", value=func.get('eptom_tel_setor_edit', d['tel_resp_setor']), key=f"ep_tset_{func['id']}")
                            ec5, ec6 = st.columns(2)
                            ep_resp_con = ec5.text_input("Resp. pelo Contrato:", value=func.get('eptom_resp_contrato_edit', d['resp_contrato']), key=f"ep_rcon_{func['id']}")
                            ep_email_con= ec6.text_input("E-mail Resp. Contrato:", value=func.get('eptom_email_contrato_edit', d['email_contrato']), key=f"ep_econ_{func['id']}")
                            ec7, ec8, ec9 = st.columns(3)
                            ep_nome     = ec7.text_input("Nome do Aprendiz: *", value=func.get('eptom_nome_edit', func.get('nome','').title()), key=f"ep_nome_{func['id']}")
                            ep_horario  = ec8.text_input("Horário:", value=func.get('eptom_horario_edit', func.get('carga_horaria','')), placeholder="Ex: 13h às 17h", key=f"ep_hor_{func['id']}")
                            ep_salario  = ec9.text_input("Salário: *", value=func.get('eptom_salario_edit', d['salario']), key=f"ep_sal_{func['id']}")

                            fb1, fb2, fb3 = st.columns(3)
                            salvar_ep = fb1.form_submit_button("SALVAR", use_container_width=True)
                            baixar_ep = fb2.form_submit_button("📄 BAIXAR .DOCX", use_container_width=True)
                            enviar_ep = fb3.form_submit_button("✉ ENVIAR PARA EPTOM", type="primary", use_container_width=True)

                        # ── Pré-visualização HTML fiel ──
                        _nome_ep    = func.get('eptom_nome_edit', func.get('nome','').title())
                        _hor_ep     = func.get('eptom_horario_edit', func.get('carga_horaria','—'))
                        _sal_ep     = func.get('eptom_salario_edit', d['salario'])
                        _empresa_ep = func.get('eptom_empresa_edit', d['empresa'])
                        _cnpj_ep    = func.get('eptom_cnpj_edit', d['cnpj'])
                        _rset_ep    = func.get('eptom_resp_setor_edit', d['resp_setor'])
                        _tset_ep    = func.get('eptom_tel_setor_edit', d['tel_resp_setor'])
                        _rcon_ep    = func.get('eptom_resp_contrato_edit', d['resp_contrato'])
                        _econ_ep    = func.get('eptom_email_contrato_edit', d['email_contrato'])
                        _meses   = ['janeiro','fevereiro','março','abril','maio','junho',
                                    'julho','agosto','setembro','outubro','novembro','dezembro']
                        _hj      = datetime.date.today()
                        _hoje_str = f"Ipatinga, {_hj.day} de {_meses[_hj.month-1]} de {_hj.year}"

                        st.markdown(
                            "<div style='font-size:10px;font-weight:800;color:#004D40;"
                            "letter-spacing:2px;text-transform:uppercase;margin:16px 0 8px;'>"
                            "Pré-visualização da Ficha</div>", unsafe_allow_html=True)

                        st.markdown(f"""
<div style="background:#fff;border:1px solid #CBD5E0;border-radius:12px;
            padding:24px 28px;font-family:'Calibri',Arial,sans-serif;font-size:12px;">
  <div style="text-align:center;font-weight:900;font-size:14px;text-transform:uppercase;
              border-bottom:2px solid #003329;padding-bottom:8px;margin-bottom:14px;">
    Formulário para Empresas de Contratação de Aprendizes
  </div>
  <div style="margin-bottom:6px;color:#555;">{_hoje_str}</div>
  <div style="margin-bottom:16px;font-weight:700;color:#003329;">
    EPTOM – Núcleo de Atendimento e Aprendizagem de Adolescentes e Jovens
  </div>
  <div style="overflow-x:auto;">
  <table style="width:100%;border-collapse:collapse;font-size:11px;">
    <thead>
      <tr style="background:#003329;color:#fff;">
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:110px;">Empresa</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:110px;">CNPJ</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:100px;">Resp. no Setor</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:100px;">Tel. Resp. Setor</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:110px;">Resp. pelo Contrato</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:150px;">E-mail Resp. Contrato</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:140px;">Nome do Aprendiz</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:90px;">Horário</th>
        <th style="padding:7px 9px;border:1px solid #CBD5E0;min-width:80px;">Salário</th>
      </tr>
    </thead>
    <tbody>
      <tr style="background:#F0FAF8;">
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_empresa_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_cnpj_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_rset_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_tset_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_rcon_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_econ_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;font-weight:700;color:#003329;">{_nome_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_hor_ep}</td>
        <td style="padding:7px 9px;border:1px solid #CBD5E0;">{_sal_ep}</td>
      </tr>
    </tbody>
  </table>
  </div>
  <div style="margin-top:12px;font-size:10px;color:#9AA5B4;text-align:center;">
    ⚠️ Confira todos os dados antes de enviar. O .docx gerado seguirá o modelo original da EPTOM.
  </div>
</div>
""", unsafe_allow_html=True)

                        def _salvar_eptom_fields():
                            func['eptom_nome_edit']             = ep_nome
                            func['eptom_horario_edit']          = ep_horario
                            func['eptom_salario_edit']          = ep_salario
                            func['eptom_empresa_edit']          = ep_empresa
                            func['eptom_cnpj_edit']             = ep_cnpj
                            func['eptom_resp_setor_edit']       = ep_resp_set
                            func['eptom_tel_setor_edit']        = ep_tel_set
                            func['eptom_resp_contrato_edit']    = ep_resp_con
                            func['eptom_email_contrato_edit']   = ep_email_con

                        if salvar_ep:
                            _salvar_eptom_fields()
                            salvar_json()
                            st.success("Dados salvos.")
                            st.rerun()

                        if baixar_ep:
                            _salvar_eptom_fields()
                            docx_bytes = gerar_ficha_eptom_docx(
                                ep_nome, ep_horario, ep_salario, ep_cnpj,
                                ep_empresa, ep_resp_set, ep_tel_set, ep_resp_con, ep_email_con)
                            if docx_bytes:
                                st.download_button(
                                    "⬇ Baixar Ficha EPTOM preenchida (.docx)",
                                    data=docx_bytes,
                                    file_name=f"Ficha_EPTOM_{ep_nome.replace(' ','_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    key=f"dl_eptom_{func['id']}")
                                st.info("Abra no Word para conferir antes de enviar.")
                            else:
                                st.error("Erro ao gerar o arquivo.")

                        if enviar_ep:
                            _salvar_eptom_fields()
                            with st.spinner("Enviando ficha para a EPTOM..."):
                                ok_ep, msg_ep = enviar_ficha_eptom(
                                    ep_nome, ep_horario, ep_salario, ep_cnpj,
                                    ep_empresa, ep_resp_set, ep_tel_set, ep_resp_con, ep_email_con)
                            if ok_ep:
                                func['eptom_ficha_enviada'] = True
                                salvar_json()
                                st.success(f"✅ {msg_ep}")
                            else:
                                st.error(msg_ep)

                        if func.get('eptom_ficha_enviada'):
                            st.markdown(
                                "<div class='notif notif-ok' style='margin-top:10px;'>"
                                "✅ Ficha já foi enviada para a EPTOM anteriormente. "
                                "Você pode reenviar se necessário.</div>",
                                unsafe_allow_html=True)

        # ── GRID DE CARDS — estilo referência ────────────────────
        else:
            n_ativos = len(ativos)
            st.markdown(
                f"<div style='font-size:12px;color:#8A94A6;margin-bottom:24px;'>"
                f"<b style='color:#004D40;font-size:26px;font-weight:900;letter-spacing:-1px;'>"
                f"{n_ativos}</b> colaborador(es) ativo(s)</div>",
                unsafe_allow_html=True)

            # CSS específico do grid — injetado uma vez
            st.markdown("""
            <style>
            .hova-card {
                background: #FFFFFF;
                border-radius: 18px;
                padding: 28px 16px 20px;
                text-align: center;
                box-shadow: 0 3px 16px rgba(0,0,0,0.08);
                border: 1px solid #E8EAED;
                transition: box-shadow 0.2s, transform 0.2s;
                margin-bottom: 4px;
            }
            .hova-card:hover {
                box-shadow: 0 8px 28px rgba(0,77,64,0.14);
                transform: translateY(-3px);
            }
            .hova-card-foto {
                width: 110px; height: 110px;
                border-radius: 50%;
                object-fit: cover;
                object-position: center top;
                border: 4px solid #E0F2F1;
                display: block;
                margin: 0 auto 16px auto;
                box-shadow: 0 4px 16px rgba(0,0,0,0.14);
            }
            .hova-card-iniciais {
                width: 110px; height: 110px;
                border-radius: 50%;
                background: linear-gradient(145deg, #004D40, #26A69A);
                color: #FFF;
                display: flex; justify-content: center; align-items: center;
                font-size: 36px; font-weight: 900;
                margin: 0 auto 16px auto;
                box-shadow: 0 4px 18px rgba(0,77,64,0.28);
                letter-spacing: 1px;
            }
            .hova-card-nome {
                font-size: 15px; font-weight: 800;
                color: #0D1B2A; text-transform: uppercase;
                letter-spacing: 0.8px; line-height: 1.2;
                margin-bottom: 10px;
            }
            .hova-card-cargo-bar {
                background: #003329;
                color: #FFFFFF;
                font-size: 11px; font-weight: 700;
                letter-spacing: 1.5px; text-transform: uppercase;
                padding: 8px 12px;
                border-radius: 8px;
                margin-bottom: 12px;
            }
            .hova-card-tel {
                font-size: 13px;
                color: #2D3748;
                margin-bottom: 10px;
                font-weight: 600;
            }
            .hova-card-data {
                font-size: 11px;
                color: #718096;
                margin-bottom: 4px;
                font-weight: 500;
            }
            .hova-card-ntw {
                background: #DCFCE7; color: #166534;
                font-size: 9px; font-weight: 700;
                padding: 3px 10px; border-radius: 20px;
                display: inline-block; margin-bottom: 8px;
                letter-spacing: 1px;
            }
            </style>
            """, unsafe_allow_html=True)

            N = 4
            for row_start in range(0, n_ativos, N):
                row  = ativos[row_start:row_start + N]
                cols = st.columns(N)

                for j, f in enumerate(row):
                    with cols[j]:
                        ini_f      = (f['data_inicio_contrato'].strftime('%d/%m/%Y')
                                      if f.get('data_inicio_contrato') else '—')
                        cargo_exib = (f.get('cargo_atual') or f.get('setor','—')).upper()
                        tel_fmt    = f.get('telefone','')
                        # Formatar telefone: 31999990000 → 31 9 9999-0000
                        if len(tel_fmt) == 11:
                            tel_fmt = f"{tel_fmt[:2]} {tel_fmt[2]} {tel_fmt[3:7]}-{tel_fmt[7:]}"
                        elif len(tel_fmt) == 10:
                            tel_fmt = f"{tel_fmt[:2]} {tel_fmt[2:6]}-{tel_fmt[6:]}"

                        # Avatar
                        if f.get('foto'):
                            b64f = base64.b64encode(f['foto']).decode()
                            av   = (f"<img src='data:image/jpeg;base64,{b64f}'"
                                    f" class='hova-card-foto'>")
                        else:
                            av   = (f"<div class='hova-card-iniciais'>"
                                    f"{iniciais(f['nome'])}</div>")

                        # Badges
                        badges_html = ""
                        if f.get('ntw_enviado'):
                            badges_html += "<div class='hova-card-ntw'>NTW ENVIADO</div>"
                        if f.get('eptom_ficha_enviada'):
                            badges_html += "<div class='hova-card-ntw' style='background:#FEF3C7;color:#92540A;'>EPTOM ENVIADO</div>"
                        # Contador docs
                        docs_ok  = sum(1 for v in f.get('docs_check',{}).values() if v)
                        docs_tot = 9
                        cor_doc  = "#166534" if docs_ok==docs_tot else ("#92540A" if docs_ok==0 else "#004D40")
                        docs_badge = (f"<div style='font-size:10px;font-weight:700;color:{cor_doc};"
                                      f"margin-bottom:6px;'>{docs_ok}/{docs_tot} docs</div>")
                        tel_html = (f"<div class='hova-card-tel'>{tel_fmt}</div>"
                                    if tel_fmt else "")

                        st.markdown(
                            f"<div class='hova-card'>"
                            f"{av}"
                            f"<div class='hova-card-nome'>{f['nome']}</div>"
                            f"<div class='hova-card-cargo-bar'>{cargo_exib}</div>"
                            f"{tel_html}"
                            f"{docs_badge}"
                            f"{badges_html}"
                            f"<div class='hova-card-data'>Desde {ini_f}</div>"
                            f"</div>",
                            unsafe_allow_html=True
                        )
                        st.markdown("<div style='height:4px;'></div>",
                                    unsafe_allow_html=True)
                        if st.button("ACESSAR PERFIL",
                                     key=f"perfil_{f['id']}",
                                     use_container_width=True,
                                     type="primary"):
                            st.session_state.perfil_foco = f['id']
                            st.rerun()

                # Preencher colunas vazias
                for j in range(len(row), N):
                    cols[j].empty()
                st.write("")


    # ════════════════════════════════════════════════════════════
    # EX-COLABORADORES
    # ════════════════════════════════════════════════════════════
    else:
        ex_list = _busca(st.session_state.ex_funcionarios, termo)
        ex_list = sorted(ex_list,
                         key=lambda x: x.get('data_desligamento') or datetime.date.min,
                         reverse=True)

        if not ex_list:
            st.markdown('<div class="empty"><div class="e-title">NENHUM REGISTRO</div>'
                        '<div class="e-sub">O histórico de ex-colaboradores aparecerá aqui.</div></div>',
                        unsafe_allow_html=True)
        else:
            st.markdown(
                f"<div style='font-size:12px;color:#8A94A6;margin-bottom:16px;'>"
                f"<b style='color:#9AA5B4;font-size:18px;'>{len(ex_list)}</b>"
                f" ex-colaborador(es)</div>", unsafe_allow_html=True)

            for f in ex_list:
                adm_f = f['data_inicio_contrato'].strftime('%d/%m/%Y') \
                        if f.get('data_inicio_contrato') else '—'
                dem_f = f['data_desligamento'].strftime('%d/%m/%Y') \
                        if f.get('data_desligamento') else '—'
                docs_count = len([v for v in f.get('documentos',{}).values() if v])

                xc1, xc2 = st.columns([4,1])
                with xc1:
                    st.markdown(
                        f"<div class='ex-func-card'>"
                        f"<div style='font-weight:700;font-size:14px;color:#4A5568;'>"
                        f"{f['nome']}</div>"
                        f"<div style='font-size:12px;color:#9AA5B4;margin-top:3px;'>"
                        f"{f.get('cargo_atual') or f.get('setor','—')}"
                        f" &nbsp;·&nbsp; Admissão: {adm_f}"
                        f" &nbsp;·&nbsp; Desligamento: {dem_f}"
                        f" &nbsp;·&nbsp; {docs_count} doc(s) no dossiê"
                        f"</div></div>",
                        unsafe_allow_html=True)
                with xc2:
                    # Botão para reativar
                    if st.button("Reativar", key=f"reativar_{f['id']}",
                                 use_container_width=True):
                        f.pop('data_desligamento', None)
                        st.session_state.contratados.append(f)
                        st.session_state.ex_funcionarios.remove(f)
                        salvar_json()
                        st.success(f"{f['nome']} reativado.")
                        st.rerun()

# ── ABA 9: FAVORITOS ──────────────────────
with abas[9]:
    import streamlit.components.v1 as _comp

    favs  = _busca(st.session_state.favoritos, termo)
    n_fav = len(favs)

    if not favs:
        st.markdown(
            '<div class="empty"><div class="e-title">NENHUM FAVORITO</div>'
            '<div class="e-sub">Durante a triagem, use a estrela para guardar '
            'candidatos de interesse.</div></div>',
            unsafe_allow_html=True)

    else:
        # ── índice de navegação dos favoritos ──
        if 'fav_idx' not in st.session_state:
            st.session_state.fav_idx = 0
        st.session_state.fav_idx = st.session_state.fav_idx % n_fav

        idx_fav = st.session_state.fav_idx
        c = favs[idx_fav]

        # Contador
        st.markdown(
            f"<div style='font-size:12px;color:#8A94A6;margin-bottom:14px;'>"
            f"<b style='color:#F59E0B;font-size:17px;'>★ {n_fav}</b> favorito(s)"
            f" &nbsp;·&nbsp; Exibindo <b style='color:#004D40;'>{idx_fav+1}</b> de {n_fav}"
            f"</div>", unsafe_allow_html=True)

        # ── FORMULÁRIO DE AGENDAMENTO ──
        if st.session_state.candidato_foco == c['id']:
            st.markdown("<div class='form-sched'>", unsafe_allow_html=True)
            st.markdown(
                "<div style='font-size:16px;font-weight:800;color:#004D40;"
                "margin-bottom:20px;'>AGENDAR ENTREVISTA</div>",
                unsafe_allow_html=True)

            fa1, fa2 = st.columns(2)
            with fa1:
                st.caption("NOME DO CANDIDATO")
                ne_fav = st.text_input("", value=c['nome'],
                                       key=f"fav_ne_{c['id']}", label_visibility="collapsed")
            with fa2:
                st.caption("E-MAIL")
                ee_fav = st.text_input("", value=c['email'],
                                       key=f"fav_ee_{c['id']}", label_visibility="collapsed")

            st.caption("DATA E HORÁRIOS (3 OPÇÕES)")
            fd, fh1, fh2, fh3 = st.columns(4)
            da_fav = fd.date_input("",  key=f"fav_da_{c['id']}", label_visibility="collapsed")
            h1_fav = fh1.time_input("", datetime.time(9, 0),  key=f"fav_h1_{c['id']}", label_visibility="collapsed")
            h2_fav = fh2.time_input("", datetime.time(14, 0), key=f"fav_h2_{c['id']}", label_visibility="collapsed")
            h3_fav = fh3.time_input("", datetime.time(16, 0), key=f"fav_h3_{c['id']}", label_visibility="collapsed")

            msg_fav = (
                f"Olá {ne_fav},\n\n"
                f"O Hospital de Olhos Vale do Aço analisou seu perfil e você foi "
                f"selecionada(o) para a próxima fase do Processo Seletivo.\n\n"
                f"Temos disponibilidade para o dia {da_fav.strftime('%d/%m/%Y')}. "
                f"Responda com o NUMERO da sua escolha de horário:\n\n"
                f"1 - {h1_fav.strftime('%H:%M')}\n"
                f"2 - {h2_fav.strftime('%H:%M')}\n"
                f"3 - {h3_fav.strftime('%H:%M')}\n\n"
                f"Endereço: {ENDERECO_HOVA}\n"
                f"Ao chegar, pergunte por Josi ou Paula.\n\n"
                f"Atenciosamente,\nEquipe de RH — HOVA"
            )
            with st.expander("Visualizar e-mail que será enviado"):
                st.code(msg_fav, language=None)

            fbc, fbenv = st.columns(2)
            with fbc:
                if st.button("CANCELAR", key=f"fav_canc_{c['id']}",
                             type="secondary", use_container_width=True):
                    st.session_state.candidato_foco = None
                    st.rerun()
            with fbenv:
                if st.button("ENVIAR CONVITE", type="primary",
                             key=f"fav_conf_{c['id']}", use_container_width=True):
                    with st.spinner("Enviando convite..."):
                        ok = send_email(ee_fav, "HOVA — Convite para Entrevista", msg_fav)
                    if ok:
                        c.update({'nome': ne_fav, 'email': ee_fav,
                                  'data_entrevista': da_fav,
                                  'opcao_1': h1_fav, 'opcao_2': h2_fav, 'opcao_3': h3_fav})
                        st.session_state.aguardando_retorno.append(c)
                        st.session_state.favoritos = [
                            f for f in st.session_state.favoritos if f['id'] != c['id']
                        ]
                        st.session_state.candidato_foco = None
                        st.session_state.fav_idx = 0
                        salvar_json()
                        st.rerun()
                    else:
                        st.error("Falha no envio.")
            st.markdown("</div>", unsafe_allow_html=True)

        # ── CARD TELA CHEIA (igual triagem) ──
        else:
            # Avatar
            if c.get('foto'):
                b64 = base64.b64encode(c['foto']).decode()
                av  = f"<img src='data:image/jpeg;base64,{b64}' class='avatar-img'>"
            else:
                av = f"<div class='avatar'>{iniciais(c['nome'])}</div>"

            cid_b  = f"<span class='tag tag-cinza'>{c['cidade']}</span>" if c.get('cidade') else ""
            dat_b  = f"<span class='tag tag-azul'>{c['data']}</span>"
            tags_b = "".join(f"<span class='tag tag-verde'>{t}</span>" for t in c.get('tags',[]))

            st.markdown(
                f"<div class='card-cand' style='position:relative;'>"
                f"<div style='position:absolute;top:16px;right:20px;"
                f"font-size:26px;color:#F59E0B;'>★</div>"
                f"{av}"
                f"<div class='cand-nome'>{c['nome']}</div>"
                f"<div style='margin:8px 0;'>{cid_b} {dat_b}</div>"
                f"<div class='cand-info'>{c['email']}"
                f"{'  |  '+c['telefone'] if c.get('telefone') else ''}</div>"
                f"<div style='margin:10px 0;'>{tags_b}</div>"
                f"<div class='cv-resumo'>{c.get('preview','')}</div></div>",
                unsafe_allow_html=True)

            # Documento original — tela cheia
            if c.get('arquivo_bytes') and c.get('nome_arquivo'):
                with st.expander("Ver documento original"):
                    if c['nome_arquivo'].lower().endswith('.pdf'):
                        b64p = base64.b64encode(c['arquivo_bytes']).decode()
                        _comp.html(f"""<!DOCTYPE html><html>
<head><meta charset="utf-8"><style>
*{{margin:0;padding:0;box-sizing:border-box;}}
body{{background:#525659;}}
#tb{{background:#323639;padding:8px 14px;display:flex;align-items:center;
     gap:10px;border-radius:8px 8px 0 0;}}
#tb button{{background:#004D40;color:#fff;border:none;padding:6px 14px;
            border-radius:6px;cursor:pointer;font-size:12px;font-weight:700;}}
#tb span{{color:#ccc;font-size:12px;}}
#cw{{height:680px;overflow-y:auto;display:flex;flex-direction:column;
     align-items:center;padding:14px 0;gap:10px;}}
canvas{{box-shadow:0 2px 8px rgba(0,0,0,0.5);}}
</style></head><body>
<div id="tb">
  <button onclick="prev()">&#8592; Anterior</button>
  <span id="info">Carregando...</span>
  <button onclick="next()">Próxima &#8594;</button>
  <button onclick="zm(0.2)">+ Zoom</button>
  <button onclick="zm(-0.2)">- Zoom</button>
</div>
<div id="cw"></div>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<script>
pdfjsLib.GlobalWorkerOptions.workerSrc=
  'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
const arr=new Uint8Array(atob('{b64p}').split('').map(x=>x.charCodeAt(0)));
let pdf,pg=1,sc=1.4;
pdfjsLib.getDocument({{data:arr}}).promise.then(p=>{{pdf=p;render(pg);}});
function render(n){{
  document.getElementById('cw').innerHTML='';
  pdf.getPage(n).then(page=>{{
    const vp=page.getViewport({{scale:sc}});
    const cv=document.createElement('canvas');
    cv.width=vp.width;cv.height=vp.height;
    document.getElementById('cw').appendChild(cv);
    page.render({{canvasContext:cv.getContext('2d'),viewport:vp}});
    document.getElementById('info').textContent='Página '+n+' de '+pdf.numPages;
  }});
}}
function prev(){{if(pg>1){{pg--;render(pg);}}}}
function next(){{if(pg<pdf.numPages){{pg++;render(pg);}}}}
function zm(d){{sc=Math.min(Math.max(sc+d,0.5),3);render(pg);}}
</script></body></html>""", height=760, scrolling=False)

                    st.download_button(
                        f"Baixar — {c['nome_arquivo']}",
                        c['arquivo_bytes'],
                        file_name=c['nome_arquivo'],
                        mime="application/pdf",
                        use_container_width=True,
                        key=f"dl_fav_{c['id']}")

            # ── 4 botões ──
            st.write("")
            bv, bp, brem, bacc = st.columns(4)
            with bv:
                if st.button("← VOLTAR", key=f"fv_vol_{c['id']}", use_container_width=True):
                    st.session_state.fav_idx = (idx_fav - 1) % n_fav
                    st.rerun()
            with bp:
                if st.button("PULAR →", key=f"fv_pul_{c['id']}", use_container_width=True):
                    st.session_state.fav_idx = (idx_fav + 1) % n_fav
                    st.rerun()
            with brem:
                if st.button("REMOVER", key=f"fv_rem_{c['id']}", type="secondary",
                             use_container_width=True):
                    st.session_state.favoritos = [
                        f for f in st.session_state.favoritos if f['id'] != c['id']
                    ]
                    st.session_state.fav_idx = 0
                    salvar_json()
                    st.rerun()
            with bacc:
                if st.button("AGENDAR", key=f"fv_age_{c['id']}", type="primary",
                             use_container_width=True):
                    st.session_state.candidato_foco = c['id']
                    st.rerun()

# ── ABA 10: BANCO ANTIGOS ─────────────────
with abas[10]:
    c_mes, c_bsc = st.columns([1,2])
    with c_mes:
        fm_ant = st.selectbox("", ["Todos"]+list(MESES_NOMES.values()), key="fm_ant", label_visibility="collapsed")
    with c_bsc:
        bsc_ant = st.text_input("", placeholder="Pesquisar...", key="bsc_ant", label_visibility="collapsed")
    st.write("---")
    ant = st.session_state.cvs_antigos
    if fm_ant != "Todos":
        ant = [c for c in ant if c.get('mes_nome')==fm_ant]
    if bsc_ant.strip():
        t2 = bsc_ant.lower()
        ant = [c for c in ant if t2 in c.get('nome','').lower() or t2 in c.get('email','').lower()]
    if not ant:
        st.markdown('<div class="empty"><div class="e-title">BANCO VAZIO</div>'
                    '<div class="e-sub">Nenhum currículo encontrado.</div></div>',
                    unsafe_allow_html=True)
    else:
        st.markdown(f"<div style='font-size:12px;color:#8A94A6;margin-bottom:12px;'><b style='color:#004D40;'>{len(ant)}</b> currículo(s)</div>", unsafe_allow_html=True)
        for c in ant:
            st.markdown(f"""
            <div style="background:#FFF;padding:16px 20px;border-radius:10px;margin-bottom:8px;
                        border:1px solid #E4E7EB;border-left:3px solid #CBD5E0;">
                <div style="font-weight:700;font-size:14px;color:#0A1628;">{c['nome']}</div>
                <div style="color:#8A94A6;font-size:12px;margin-top:2px;">
                    {c.get('setor','—')} &nbsp;·&nbsp; {c.get('email','—')} &nbsp;·&nbsp; {c.get('data','—')}
                </div>
            </div>""", unsafe_allow_html=True)
